#!/usr/bin/env python3
"""
Adam AI — Cryptonary OS
Writing Studio | Data Studio | Creative Studio
- Emails, Social Content, Ad Copy, Landing Pages
- Idea generation with live RSS + X tweet context
- Voice input via Whisper
- Image generation via Gemini + DALL-E
- Visual briefs: storyboards, thumbnails, slide briefs
- Performance data feedback loops
- Brand guidelines + best posts reference library
"""

import os, json, ssl, urllib.request, time, re, traceback

TELEGRAM_TOKEN = "8611455908:AAH2zTch0Nf5tM590-_ouPZO2at-sqDpj_Y"
ANTHROPIC_KEY  = os.environ.get("ANTHROPIC_KEY",  "YOUR_ANTHROPIC_KEY_HERE")
OPENAI_KEY     = os.environ.get("OPENAI_KEY",     "")
X_BEARER_TOKEN = os.environ.get("X_BEARER_TOKEN", "")
GEMINI_KEY     = os.environ.get("GEMINI_KEY",     "")
ssl._create_default_https_context = ssl._create_unverified_context

BOOK_KNOWLEDGE = """

=== COPYWRITING KNOWLEDGE BASE ===

FROM HORMOZI ($100M OFFERS / $100M LEADS):
- Value Equation: Dream Outcome x Likelihood / Time x Effort = Value
- Grand Slam Offer: Stack so much value the price feels embarrassing
- Pain/Dream/Fix: Open on pain, paint the dream, position as mechanism
- Specificity sells: numbers, dates, exact levels beat vague claims
- Urgency must be real and tied to market events

FROM CASHVERTISING (WHITMAN):
- LF8 Life Force Desires: survival/enjoyment of life, enjoyment of food/drink, freedom from fear/pain/danger, sexual companionship, comfortable living, to be superior/win, care for loved ones, social approval
- Secondary Wants: information, curiosity, cleanliness, efficiency, convenience, dependability, expression of beauty/style, economy/profit, bargains
- Fear works when threat is real + reader vulnerable + solution clear
- Bucket brigade technique: keep them reading line by line

FROM CIALDINI (INFLUENCE / PRE-SUASION):
- Reciprocity: Give genuine value first, then the ask
- Commitment/Consistency: Anchor their identity as a serious investor
- Social Proof: Show what members DO not just what they have access to
- Authority: Track record with specific wins, not claims
- Scarcity: Information gap is most powerful for Cryptonary
- Liking: People buy from people they like and relate to
- Unity: Joining Pro = joining a group, not buying a product
- Pre-Suasion: What comes before the message primes the response

FROM OGILVY ON ADVERTISING:
- Headline is everything — 5x more people read the headline than the body
- Promise a benefit or provoke curiosity in the headline
- Specifics always outperform generalities
- P.S. is the second most-read element — use it for the sharpest proof point
- One CTA, one action, zero friction
- Long copy outperforms short copy when the reader is interested

FROM BUILDING A STORYBRAND (MILLER):
- The reader is ALWAYS the hero — the brand is the guide
- Guide characteristics: empathy (I understand your struggle) + authority (I can help)
- The 3 levels of problem: external (the market), internal (the fear/frustration), philosophical (it's not fair)
- A clear plan removes the risk of action — people don't act when confused
- Stakes must be clear: what do they gain if they act? What do they lose if they don't?
- Success should be painted explicitly — never assume they'll imagine it themselves
- One message, one call to action — confusion kills conversion

FROM ALCHEMY (SUTHERLAND):
- Logic is not what drives decisions — perceived value and meaning drive decisions
- Reframing changes everything: the same thing positioned differently creates different desire
- Loss aversion framing is more powerful than gain framing in almost every context
- Signalling: expensive or inconvenient things signal quality and commitment
- Dare to be trivial: small, specific, seemingly irrelevant details make claims feel more credible
- People don't know why they want things — emotion precedes rationalisation

--- WRITING FRAMEWORKS ---

AIDA (Attention → Interest → Desire → Action)
- Attention: stop the scroll
- Interest: hold them with a relevant problem or story
- Desire: build the want through proof, benefits, transformation
- Action: one clear CTA

PAS (Problem → Agitate → Solution)
- Problem: name the exact problem
- Agitate: make them feel how bad it is — the cost of inaction
- Solution: position Cryptonary as the relief
"""


# ══════════════════════════════════════════════════════════════════
# CRYPTONARY BEST POSTS — PROVEN CONTENT REFERENCE LIBRARY
# Extracted from top performing Instagram posts sorted by engagement,
# shares, saves and reach. Used as reference when generating concepts.
# ══════════════════════════════════════════════════════════════════

BEST_POSTS = """

=== CRYPTONARY TOP PERFORMING INSTAGRAM POSTS ===

These are real posts from @Cryptonary that have proven high engagement.
Use these as creative reference — study the hooks, formats, angles and
what made each one work. Do not copy verbatim. Extract the pattern.

--- ENGAGEMENT BAIT (Comments/Shares focused) ---

POST: "Tag a friend that isn't into crypto but should be."
TYPE: Static | ENGAGEMENT: 1.24% | LIKES: 4,956 | COMMENTS: 1,940
WHY IT WORKED: Ultra low friction ask. Peer-to-peer shareability. Identity signal ("I'm in, you should be").

POST: "You can only pick one... [BTC / ETH / XRP / DOGE / ? grid]"
TYPE: Static | ENGAGEMENT: 1.44% | LIKES: 6,538 | COMMENTS: 1,513
WHY IT WORKED: Binary/forced choice triggers debate. Coin tribalism drives comments. Simple visual.

POST: "Comment your biggest mistakes in crypto."
TYPE: Static | ENGAGEMENT: 2.88% | LIKES: 14,892 | COMMENTS: 1,138
WHY IT WORKED: Vulnerability prompt. Collective commiseration. People love sharing their losses.

POST: "What was the first crypto you purchased? [coin grid]"
TYPE: Static | ENGAGEMENT: 2.66% | LIKES: 14,810 | COMMENTS: 340
WHY IT WORKED: Nostalgia + identity. Everyone has an origin story. Grid visual makes it scrollable.

POST: "What coin have you been bag holding for the last 2 years and can't let go?"
TYPE: Static | ENGAGEMENT: 9.96% | LIKES: 55,106 | COMMENTS: 338
WHY IT WORKED: Pain point + humour. Universal crypto experience. Relatable loss = huge response.

POST: "What is your largest bag without telling us your largest bag?"
TYPE: Static | ENGAGEMENT: 0.86% | LIKES: 3,797 | COMMENTS: 986
WHY IT WORKED: Playful secrecy mechanic. Riddle format. Community inside joke energy.

POST: "What made you start crypto and how did you first get started?"
TYPE: Static | ENGAGEMENT: 0.98% | LIKES: 4,912 | COMMENTS: 526
WHY IT WORKED: Origin story prompt. High personal investment in reply. Storytelling trigger.

POST: "If you were given $2,500,000 what is the first thing you would do?"
TYPE: Static | ENGAGEMENT: 3.09% | LIKES: 16,441 | COMMENTS: 787
WHY IT WORKED: Fantasy scenario. Zero barrier to entry. Universal aspiration.

POST: "Without dropping any names, what's your favourite crypto famous for?"
TYPE: Static | ENGAGEMENT: 0.55% | LIKES: 2,658 | COMMENTS: 414
WHY IT WORKED: Inside-joke format. Insider community feel. Coded language = belonging signal.

POST: "Share your dumbest crypto trading story with us, go."
TYPE: Static | ENGAGEMENT: 0.9% | LIKES: 4,280 | COMMENTS: 710
WHY IT WORKED: Confessional format. Humour + pain. Story prompt = high-effort comments.

POST: "$250K sitting in your crypto exchange account. What are you doing with it?"
TYPE: Static | ENGAGEMENT: 0.58% | LIKES: 2,883 | COMMENTS: 357
WHY IT WORKED: Aspirational scenario. Forces real strategy thinking. Debate-generating.

--- BREAKING NEWS / MACRO (Shares + Reach focused) ---

POST: "WARNING! CRYPTO IS ON ROCKY GROUNDS. Here's everything you need to know in 60 seconds."
TYPE: Static | ENGAGEMENT: 0.41% | LIKES: 1,521 | SHARES: 792 | VIEWS: 193,631
WHY IT WORKED: Fear + urgency. WARNING hook. Time promise (60 seconds). Dark aesthetic + explosion visual.

POST: "SOLD BTC AT BOTTOM. BOUGHT OIL AT $118"
TYPE: Reel | ENGAGEMENT: 0.33% | LIKES: 1,217 | SHARES: 969 | VIEWS: 379,861
WHY IT WORKED: Cringe/pain narrative. Highest shares in recent batch. Human failure story = virality.

POST: "MARKETS POST WAR [chart grid showing Buy the War pattern across conflicts]"
TYPE: Carousel | ENGAGEMENT: 0.29% | LIKES: 997 | SHARES: 764
WHY IT WORKED: Data-led historical pattern. Counterintuitive insight. Actionable for investors.

POST: "IRAN JUST NAMED A NEW SUPREME LEADER AND HE'S MORE DANGEROUS THAN THE LAST ONE"
TYPE: Carousel | ENGAGEMENT: 0.29% | LIKES: 1,359 | SHARES: 372 | VIEWS: 234,697
WHY IT WORKED: Geopolitical tension. Real-world macro hook. Crypto implications for informed audience.

POST: "CRYPTO IS NOW BIGGER THAN THE ENTIRE UK STOCK MARKET"
TYPE: Static | ENGAGEMENT: 1.53% | LIKES: 7,403 | SHARES: 49
WHY IT WORKED: Milestone credibility stat. Legitimacy narrative. Shocks the skeptic.

POST: "BRICS TO OFFICIALLY ABANDON US DOLLAR FOR TRADE SETTLEMENTS"
TYPE: Static | ENGAGEMENT: 6.58% | LIKES: 29,018 | SHARES: 2,372 | SAVES: 4,348 | VIEWS: 729,123
WHY IT WORKED: Macro news with direct crypto implication. High saves = reference content. Geopolitical fear.

POST: "US JUDGE RULES RIPPLE XRP IS NOT A SECURITY"
TYPE: Static | ENGAGEMENT: 2.84% | LIKES: 12,597 | SHARES: 2,177
WHY IT WORKED: Major regulatory win. Instant share trigger. XRP community mobilised.

POST: "$1 BILLION WIPED FROM THE CRYPTO MARKET IN 60 MINUTES"
TYPE: Static | ENGAGEMENT: 0.88% | LIKES: 4,009 | SHARES: 2,011 | VIEWS: 855,387
WHY IT WORKED: Shock number. Time specificity. Fear = share. Macro audience crossover.

POST: "EMIRATES TO ACCEPT BITCOIN AND CRYPTO FOR FLIGHTS"
TYPE: Static | ENGAGEMENT: 2.35% | LIKES: 13,607 | SHARES: 1,851 | VIEWS: 2,448,771
WHY IT WORKED: Adoption story. Mainstream brand = credibility. Shareable to non-crypto audience.

POST: "OVER $1 TRILLION WIPED OUT FROM CRYPTO & US STOCK MARKET"
TYPE: Static | ENGAGEMENT: 0.5% | LIKES: 2,305 | SHARES: 957 | VIEWS: 222,153
WHY IT WORKED: Trillion = attention. Stock market crossover = wider audience. Fear share.

POST: "RUSSIA & IRAN LAUNCHING GOLD-BACKED STABLECOIN"
TYPE: Static | ENGAGEMENT: 2.52% | LIKES: 11,206 | SHARES: 709 | SAVES: 1,718 | VIEWS: 161,548
WHY IT WORKED: Geopolitical + crypto convergence. Contrarian macro angle. High saves.

POST: "BINANCE SEIZED ALL PALESTINIAN ASSETS AT ISRAEL'S REQUEST"
TYPE: Static | ENGAGEMENT: 1.54% | LIKES: 7,532 | SHARES: 88 | SAVES: 1,046 | REACH: 353,213
WHY IT WORKED: Controversy + moral debate. Self-custody argument. Very high saves.

POST: "DOJ LAUNCHES INVESTIGATION INTO IRAN USING BINANCE TO EVADE SANCTIONS"
TYPE: Static | ENGAGEMENT: 0.15% | LIKES: 666 | SHARES: 324
WHY IT WORKED: Regulatory fear. Exchange risk narrative. Timely news hook.

POST: "PRICE CHANGES SINCE WAR STARTED [BTC/ETH/SOL/AXIE with % gains]"
TYPE: Static | ENGAGEMENT: 0.14% | LIKES: 612 | SHARES: 287
WHY IT WORKED: Data table + war angle. Counterintuitive "buy the war" thesis. Timely and specific.

--- DATA / COMPARISON / HISTORICAL (Saves focused) ---

POST: "CRYPTO JUST CHANGED FOREVER [16 tokens reclassified as securities]"
TYPE: Static | ENGAGEMENT: 0.22% | LIKES: 741 | COMMENTS: 71 | SHARES: 436 | SAVES: 355
WHY IT WORKED: Regulatory milestone. "Forever" framing. List format = saves magnet.

POST: "2026 FEELS LIKE A TIME MACHINE [BTC/ETH/SEI 2024 vs 2026 prices same]"
TYPE: Static | ENGAGEMENT: 0.2% | LIKES: 959 | SHARES: 349
WHY IT WORKED: Ironic data insight. Cognitive dissonance. "We're back to square one" angle.

POST: "THE CRYPTO YOUR FRIEND BOUGHT vs THE CRYPTO YOU SOLD [wavy lines chart]"
TYPE: Static | ENGAGEMENT: 0.18% | LIKES: 740 | SHARES: 520
WHY IT WORKED: Relatable pain. FOMO + regret combo. Visual simplicity. No text needed.

POST: "BITCOIN AT $70K IN 2024 [Extreme Greed] vs BITCOIN AT $70K IN 2026 [Extreme Fear]"
TYPE: Static | ENGAGEMENT: 0.18% | LIKES: 965 | SHARES: 224
WHY IT WORKED: Same number, opposite sentiment. Contrarian insight. Market psychology angle.

POST: "BITCOIN PRICES ON EID [2010: $0.06 → 2026: $70,600 table]"
TYPE: Carousel | ENGAGEMENT: 0.13% | LIKES: 638 | SHARES: 180 | SAVES: 86
WHY IT WORKED: Cultural hook (Eid). Historical journey. Aspirational wealth narrative.

POST: "BITCOIN SURPASSES $125K [NEW ATH]"
TYPE: Static | ENGAGEMENT: 0.48% | LIKES: 2,327 | SHARES: 855
WHY IT WORKED: Milestone celebration. ATH = instant share. Simple, bold, celebratory.

POST: "BTC ALL-TIME HIGH $100K"
TYPE: Static | ENGAGEMENT: 1.06% | LIKES: 5,221 | SHARES: 47 | SAVES: 1,403
WHY IT WORKED: Historic milestone. Bold visual. High saves as memento.

POST: "ETH Monthly Prices [Jan $1,300 → May ???]"
TYPE: Static | ENGAGEMENT: 1.15% | LIKES: 6,073 | COMMENTS: 348
WHY IT WORKED: Cliffhanger data. Forces prediction in comments. Monthly cadence = repeatable.

POST: "JAMIE DIMON ON BITCOIN [2014-2016 quotes all wrong]"
TYPE: Static | ENGAGEMENT: 1.09% | LIKES: 6,009 | COMMENTS: 90
WHY IT WORKED: Authority figure proven wrong. Vindication for HODLers. Screenshot-worthy.

POST: "ETH has risen +3,870% from March 2020. Why has it risen so much and where will it head?"
TYPE: Carousel | ENGAGEMENT: 2.67% | LIKES: 14,593 | SHARES: 0 | SAVES: 0
WHY IT WORKED: Specific % gain. Education + curiosity gap. Carousel = read-through.

POST: "DOLLAR VS BITCOIN [visual comparison bar]"
TYPE: Static | ENGAGEMENT: 1.72% | LIKES: 9,446 | COMMENTS: 118
WHY IT WORKED: Simple visual contrast. Devaluation narrative. Bitcoin maximalist bait.

POST: "What is Solana [carousel explainer with saves: 1,716]"
TYPE: Carousel | ENGAGEMENT: 1.2% | LIKES: 4,864 | SAVES: 1,716
WHY IT WORKED: Evergreen education. High saves = bookmarked reference. New entrant traffic.

POST: "1 ETH = 1 AIRPOD MAX [December 2020]"
TYPE: Static | ENGAGEMENT: 1.29% | LIKES: 7,087 | COMMENTS: 101
WHY IT WORKED: Relatable price anchor. Consumer goods comparison. Timing: peak cycle nostalgia.

POST: "BTC AT 60K: 'This is going to 100K, strong people buy strength, time to buy'"
TYPE: Static | ENGAGEMENT: 1.36% | LIKES: 7,378 | COMMENTS: 176
WHY IT WORKED: Conviction quote. Bold market call. Tribal "strong hands" identity.

--- MEME / POP CULTURE CROSSOVER ---

POST: "RUG BOYS [Elon Musk + Vitalik Buterin as movie characters]"
TYPE: Static | ENGAGEMENT: 3.58% | LIKES: 19,502 | COMMENTS: 451
WHY IT WORKED: Celebrity faces. Film poster format. Humour + tribalism. Instantly shareable.

POST: "Pick your pill [6 crypto outcomes: send BTC to 100K / send ETH to 5K / etc.]"
TYPE: Static | ENGAGEMENT: 0.95% | LIKES: 4,911 | COMMENTS: 396
WHY IT WORKED: Matrix reference. Forced fantasy choice. Debate in comments. Pop culture hook.

POST: "MY PORTFOLIO 3 MONTHS AGO [Tesla Cybertruck] / MY PORTFOLIO NOW [rusted bin]"
TYPE: Static | ENGAGEMENT: 0.16% | LIKES: 661 | SHARES: 431
WHY IT WORKED: Meme format. Self-deprecating loss humour. Highly shareable. No text needed to get it.

POST: "WHEN MY PORTFOLIO KEEPS GETTING HIT [battered man meme]"
TYPE: Reel | ENGAGEMENT: 0.14% | LIKES: 567 | SHARES: 371 | VIEWS: 203,693
WHY IT WORKED: Pain meme. Universal investor experience. High views = algorithmic push.

POST: "What is a bigger flex? [$10M in cash vs $10M on a ledger]"
TYPE: Static | ENGAGEMENT: 18.05% | LIKES: 99,346 | COMMENTS: 387 | SHARES: 332 | SAVES: 580
WHY IT WORKED: HIGHEST ENGAGEMENT POST. Debate mechanic. Status + crypto conviction. Simple binary.

POST: "Elon went from Robin Hood to Hood Robbin."
TYPE: Static | ENGAGEMENT: 1.76% | LIKES: 9,418 | COMMENTS: 378 | SHARES: 45
WHY IT WORKED: Wordplay pun. Cultural moment. Sharp one-liner. Screenshot/share bait.

POST: "YouTube star KSI buys $2.8 Million of LUNA, which is now worth less than $10"
TYPE: Static | ENGAGEMENT: 1.78% | LIKES: 8,651 | COMMENTS: 253 | SAVES: 993
WHY IT WORKED: Celebrity schadenfreude. LUNA collapse moment. Warning narrative.

POST: "Vitalik Buterin sells $SHIB and $AKITA and donates $1B to India Covid Crypto Fund"
TYPE: Static | ENGAGEMENT: 5.24% | LIKES: 28,485 | COMMENTS: 735
WHY IT WORKED: Celebrity action + charity = viral. Meme coin angle. News + emotion combo.

POST: "My enemies made me short Bitcoin at $42K! I'm at war, please send 50 ETH today."
TYPE: Static | ENGAGEMENT: 1.74% | LIKES: 7,807 | COMMENTS: 207 | SAVES: 504
WHY IT WORKED: Scam awareness humour. Relatable parody. Community in-joke format.

POST: "DID COINBASE MAKE THE BEST CRYPTO AD EVER?"
TYPE: Reel | ENGAGEMENT: 0.15% | LIKES: 779 | COMMENTS: 14 | SHARES: 145
WHY IT WORKED: Engagement question. Industry credibility moment. Debate trigger.

--- EDUCATIONAL / EVERGREEN ---

POST: "Reasons why you should never take crypto advice from a friend."
TYPE: Static | ENGAGEMENT: 0.69% | LIKES: 3,171 | COMMENTS: 444 | SAVES: 193
WHY IT WORKED: Warning format. Relatable mistake. Positions Cryptonary as the better alternative.

POST: "HOW TO BUY A CAR FOR FREE [crypto gains guide]"
TYPE: Carousel | ENGAGEMENT: 2.37% | LIKES: 7,258 | COMMENTS: 189 | SAVES: 5,770
WHY IT WORKED: Aspirational outcome. Specific tangible goal. High saves = practical guide.

POST: "THE LUNA SAGA: EXPLAINED FROM $100 TO $0.003 IN 48 HOURS — BUT HOW?"
TYPE: Carousel | ENGAGEMENT: 2.34% | LIKES: 9,209 | COMMENTS: 229
WHY IT WORKED: Explainer of a disaster. Curiosity gap. "But how?" hook. High shareability.

POST: "2013: Bitcoin is a scam / 2015: Blockchains are a scam / ... / 2021: Web3 is a scam"
TYPE: Static | ENGAGEMENT: 2.8% | LIKES: 12,178 | COMMENTS: 247 | SAVES: 3,117
WHY IT WORKED: Pattern recognition list. Vindicates believers. Screenshot + share magnet.

POST: "DeFi-ing the rules of traditional finance [coin orbit visual]"
TYPE: Static | ENGAGEMENT: 1.28% | LIKES: 7,070 | COMMENTS: 59
WHY IT WORKED: Wordplay headline. Visual identity of the space. Educational positioning.

POST: "MUST WATCH FINANCE MOVIES & DOCUMENTARIES [series]"
TYPE: Carousel | ENGAGEMENT: 1.1% / 1.08% | LIKES: 6,006 / 5,929
WHY IT WORKED: Evergreen saves content. Useful list. Positions as curator not just news.

POST: "Buying crypto now is like buying Dubai land in March 2026 [Tweet format]"
TYPE: Static | ENGAGEMENT: 0.14% | LIKES: 690 | COMMENTS: 36
WHY IT WORKED: Analogy format. Local cultural reference. Conviction statement.

POST: "Crypto Twitter raised $800,000 for a 6-year-old who just beat leukemia"
TYPE: Static | ENGAGEMENT: 1.69% | LIKES: 9,359 | COMMENTS: 77
WHY IT WORKED: Human interest + crypto good news. Emotional story. Contradicts negative narrative.

POST: "SATOSHI NAKAMOTO [bold name reveal energy]"
TYPE: Static | ENGAGEMENT: 1.1% | LIKES: 4,988 | COMMENTS: 809 | SAVES: 254
WHY IT WORKED: Mystery figure. Debate magnet. Name alone triggers massive comment response.

POST: "WE ARE HERE [market cycle chart with arrow at bottom]"
TYPE: Static | ENGAGEMENT: 2.01% | LIKES: 8,247 | SAVES: 4,121 | VIEWS: 508,514
WHY IT WORKED: Market cycle positioning. Accumulation narrative. High saves = reference content.

POST: "WHERE DO YOU THINK WE ARE? [market psychology chart]"
TYPE: Static | ENGAGEMENT: 1.18% | LIKES: 6,173 | COMMENTS: 387
WHY IT WORKED: Participation mechanic. Everyone has an opinion. Psychological model = saves.

POST: "HODL = misspelled HOLD [plain text]"
TYPE: Static | ENGAGEMENT: 1.4% | LIKES: 7,665 | COMMENTS: 129
WHY IT WORKED: Crypto etymology. Inside knowledge signal. Clean minimal design.

POST: "If you know, you know."
TYPE: Static | ENGAGEMENT: 0.49% | LIKES: 1,753 | COMMENTS: 985
WHY IT WORKED: Exclusivity signal. Insider community. Drove 985 comments = massive engagement.

POST: "How are you feeling 24 hours later? [Twitter screenshot format]"
TYPE: Static | ENGAGEMENT: 0.88% | LIKES: 4,578 | COMMENTS: 346
WHY IT WORKED: Follow-up content mechanic. Community check-in. Simple and timely.

POST: "The year is 2025 / The price of ETH is... [cliffhanger]"
TYPE: Carousel | ENGAGEMENT: 1.19% | LIKES: 5,963 | COMMENTS: 653
WHY IT WORKED: Prediction cliffhanger. Forces swipe. Debate in comments about the answer.

POST: "Bitcoin just formed 10 consecutive green candles in July"
TYPE: Static | ENGAGEMENT: 1.02% | LIKES: 5,618 | COMMENTS: 97
WHY IT WORKED: Specific data point. Bullish signal. Timely market observation.

POST: "EVER WONDERED HOW MUCH $100,000 STAKED TODAY CAN GENERATE?"
TYPE: Carousel | ENGAGEMENT: 1.05% | LIKES: 5,706 | COMMENTS: 126
WHY IT WORKED: Financial fantasy. Specific number. Passive income angle. Swipe to reveal.

POST: "I'm proud for not investing in crypto — Charlie Munger [quote card]"
TYPE: Static | ENGAGEMENT: 1.32% | LIKES: 6,527 | COMMENTS: 296 | SAVES: 536
WHY IT WORKED: Authority figure quote. Contrarian vindication. Tribal anger + saves.

POST: "Warren Buffet invests $1 Billion into Bitcoin-friendly Nubank"
TYPE: Static | ENGAGEMENT: 1.26% | LIKES: 5,516 | SAVES: 527 | REACH: 140,242
WHY IT WORKED: Credibility transfer. Buffett + crypto = cognitive dissonance. News moment.

POST: "Stars who bounced back from financial losses [celebrity feature]"
TYPE: Carousel | ENGAGEMENT: 1.01% | LIKES: 4,655 | SAVES: 959
WHY IT WORKED: Celebrity + resilience narrative. Relatable to investors in drawdown. Saves = inspiration.

POST: "WHICH OF THESE IS HAPPENING FIRST? [BTC $100K / $10K / $1K + altcoins]"
TYPE: Static | ENGAGEMENT: 2.02% | LIKES: 10,134 | COMMENTS: 1,121
WHY IT WORKED: Prediction poll format. Bullish vs bearish debate. Grid visual = quick read.

POST: "PENTAGON HEXAGON OCTAGON PORTFOLIOGONE [chart showing crash]"
TYPE: Static | ENGAGEMENT: 5.87% | LIKES: 30,961 | SHARES: 2,285 | SAVES: 4,562 | VIEWS: 1,407,145
WHY IT WORKED: HIGHEST REACH POST. Wordplay punchline reveal. Chart visual. Meme energy.
Shares 2,285 — highest share post. Saves 4,562. Reach 1.36M. Follower change +780.

POST: "It's 2025... What's the price of Bitcoin? [then vs now]"
TYPE: Static | ENGAGEMENT: 4.36% | LIKES: 23,136 | COMMENTS: 1,196 | VIEWS: 136,797
WHY IT WORKED: Prediction mechanic. Aspirational year framing. Comment your answer = high engagement.

POST: "It's 2025... What's the price of Ethereum?"
TYPE: Static | ENGAGEMENT: 3.92% | LIKES: 21,215 | COMMENTS: 625
WHY IT WORKED: Same mechanic as above. Prediction + community input. Series format.

POST: "January 2022: $2,604 / February: $2,629 / March: $3,383 [ETH monthly prices]"
TYPE: Static | ENGAGEMENT: 4.19% | LIKES: 22,318 | COMMENTS: 552 | VIEWS: 0
WHY IT WORKED: Data cadence. Month-by-month narrative. Repeatable format. Bullish momentum chart.

POST: "One has to go... [BTC / ETH / BNB / ADA / XRP grid]"
TYPE: Static | ENGAGEMENT: 0.74% | LIKES: 3,440 | COMMENTS: 678
WHY IT WORKED: Forced elimination game. Tribalism trigger. Repeat engagement mechanic.

POST: "You have $10 to pick your trio [coin grid]"
TYPE: Static | ENGAGEMENT: 0.97% | LIKES: 4,674 | COMMENTS: 732
WHY IT WORKED: Budget constraint + choice = debate. Portfolio strategy in comment form.

POST: "BTC OR ETH? [$1M hardware wallet visual]"
TYPE: Static | ENGAGEMENT: 1.41% | LIKES: 7,135 | COMMENTS: 711
WHY IT WORKED: Simple binary debate. Store of value thesis. Visual anchor = premium feel.

POST: "My Ukrainian credit cards don't work... Crypto is the only money I still have [Tweet]"
TYPE: Static | ENGAGEMENT: 4.65% | LIKES: 15,781 | SHARES: 7,534 | SAVES: 2,453
WHY IT WORKED: Real human story. Crypto as freedom/survival narrative. Emotional share trigger.

POST: "These scam coins are getting crazy: 27 trillion in circulation, unlimited supply..."
TYPE: Static | ENGAGEMENT: 1.06% | LIKES: 5,816 | COMMENTS: 96
WHY IT WORKED: Obvious reference (USD). Subversive angle. Forces the reframe. Insider laugh.

POST: "EVER WONDERED HOW MUCH $100,000 STAKED can generate?"
TYPE: Carousel | ENGAGEMENT: 1.05% | LIKES: 5,706
WHY IT WORKED: Specific aspiration. Passive income hook. Curiosity gap swipe format.

POST: "CRYPTO MINE RAID FOUND THOUSANDS OF PS4s"
TYPE: Carousel | ENGAGEMENT: 0.87% | LIKES: 4,736 | COMMENTS: 103
WHY IT WORKED: Bizarre news story. Visual hook. Unusual = shareworthy.

POST: "Like Bitcoin as a portfolio diversifier — Paul Tudor Jones [quote]"
TYPE: Carousel | ENGAGEMENT: 0.86% | LIKES: 4,776
WHY IT WORKED: TradFi authority validates BTC. Credibility bridge for skeptics.

=== KEY PATTERNS THAT DRIVE CRYPTONARY PERFORMANCE ===

FORMAT INSIGHTS:
- Engagement bait (pick/choose/comment) = highest comment count
- Breaking news with share implication = highest shares
- Data tables + historical prices = highest saves
- Meme crossovers = highest viral reach
- Quote cards from TradFi figures = credibility + debate

HOOK PATTERNS THAT WORK:
- "X years ago [price] → today [price]"
- "WARNING: [thing happening]"
- "You can only pick one..."
- "What was your [first/biggest/worst]..."
- "[Year] feels like a time machine"
- "Without dropping names..."
- "If you know, you know."
- "One has to go..."
- "[Specific stat] in [specific time]"
- "[Celebrity/institution] just [did crypto thing]"
- "[Thing everyone said was a scam] list"
- "What is a bigger flex?"

CONTENT ANGLES:
- Geopolitical events → crypto implications
- TradFi/institution adoption → legitimacy
- Famous person proven wrong about crypto → vindication
- Real human crypto story → emotional share
- Market cycle positioning → saves
- Price prediction → comment debate
- Scam/loss awareness → warning + humour
- Crypto as freedom/sovereignty → identity
- Forced choice/debate → tribal comments

"""


# ══════════════════════════════════════════════════════════════════
# CRYPTONARY BRAND GUIDELINES — used in every image generation prompt
# ══════════════════════════════════════════════════════════════════

BRAND_GUIDELINES = """
=== CRYPTONARY BRAND SYSTEM ===

COLOURS:
- Black #000000 — primary background (default for all dark posts)
- White #FFFFFF — primary text on dark backgrounds
- Cryptonary Blue #005EFF — brand accent, highlights, CTAs, data callouts
- Red #FF0000 — urgency, warnings, bearish signals, breaking news
- Green #0DA500 — bullish signals, gains, positive news
- Bitcoin Orange #F7931A — Bitcoin-specific content, BTC price posts
- No gradients unless explicitly requested. Dark background is always default.

LOGO:
- Primary logomark: geometric C shape inside an octagonal frame, inner square cutout, diagonal depth shadow bottom-left. Architectural and structural.
- Alt logomark: flat C cutout, simpler form
- Wordmark: "cryptonary" lowercase in Inter paired with logomark
- White logo on dark backgrounds. Black logo on light. Blue as accent only.
- Never distort, rotate or recolour the logo outside brand palette.
- Placement: bottom-right corner at ~8% of image width. Never on busy areas.

TYPOGRAPHY:
- Inter: primary font, all body text and standard headers. Clean, geometric. Lowercase preferred for brand name.
- Tungsten Regular: condensed impactful headlines in tight spaces. Use for WARNING, BREAKING, ATH, single-word punches. All caps.
- Termina: wide technical header. Use for data posts, price tables, market updates. Feels authoritative.
- Hierarchy: Headline (Tungsten/Termina, large, white or accent) → Subhead (Inter bold) → Body (Inter regular, #FFFFFF or #CCCCCC)
- All caps: breaking news, price callouts. Sentence case: educational and narrative posts.

POST STYLE PATTERNS:
- BREAKING NEWS: Black bg, red NEWS/BREAKING badge top-left, bold Tungsten headline, relevant image/flag
- PRICE/DATA: Black bg, coin logo, large Termina price figure, % change in green/red
- ENGAGEMENT BAIT: Clean black/white bg, bold Inter question, coin grid or choice visual
- EDUCATIONAL: Dark bg, structured layout, data table or numbered list, logo bottom-right
- MEME/CULTURAL: Cinematic dark image, bold white text overlay, minimal C mark bottom
- MACRO/GEOPOLITICAL: News photo with dark overlay, white headline, news badge

DIMENSIONS:
- Instagram square: 1080x1080px
- Instagram story/reel cover: 1080x1920px
- Instagram carousel slide: 1080x1080px
"""


CTA_OPTIONS = {
    "free": {
        "upgrade": "Upgrade to Pro",
        "read_free": "Read for Free",
        "discount": "Upgrade with Discount/Offer"
    },
    "pro": {
        "read_report": "Read the Full Report",
        "check_levels": "Check the Levels",
        "upgrade_save": "Upgrade & Save"
    }
}

CTA_INSTRUCTIONS = {
    "free": {
        "upgrade": "CTA pushes reader to upgrade to Pro. Sell the transformation not the subscription. Use Hormozi value equation.",
        "read_free": "CTA invites reader to read specific content for free. Low friction. Frame as a gift.",
        "discount": "CTA offers a specific discount on Pro. Make the saving concrete. Use Cialdini scarcity and urgency."
    },
    "pro": {
        "read_report": "CTA directs Pro member to read the full report. Direct, no friction. Frame as time-sensitive.",
        "check_levels": "CTA sends Pro member to check updated market levels. Frame as the specific numbers they need right now.",
        "upgrade_save": "CTA offers a saving on Pro — annual plan, loyalty discount, or limited offer. Make the saving concrete."
    }
}

VOICE_GUIDE = """You are writing emails and social content for Adam, Co-Founder of Cryptonary, a crypto research and education platform with 300K+ newsletter subscribers.

ADAM'S VOICE RULES:
- Open emails with "Gm [Name],"
- Short punchy sentences. Very short. Sometimes one word. Or two.
- Bold key phrases using *asterisks* notation
- Never use em dashes. Use full stops or line breaks instead.
- Bullets use the bullet point symbol
- Free email sign off: Talk soon, / Adam / Co-Founder, Cryptonary
- Pro email sign off: Adam / Co-Founder, Cryptonary
- Always include a punchy P.S. line
- No fluff. Every line earns its place.
- Rhetorical questions work well.
- Numbers and specifics make it real.
- Write like talking to one person.
- 2-3 sentences max per paragraph.

""" + BOOK_KNOWLEDGE

performance_data = {}
user_state = {}
_global_brief_store = {}  # persists visual briefs across state resets
_last_callback_time = {}  # per-user callback cooldown (prevents double-fire)

# ── CRYPTONARY LOGO (white, transparent bg) ───────────────────
# LogomarkAlt_W.png encoded as base64 for watermarking generated images
CRYPTONARY_LOGO_B64 = "iVBORw0KGgoAAAANSUhEUgAAAZgAAAGYCAYAAAB/O/RVAAARPklEQVR4nO3dS44cVd7G4f+JzKpqWQYhC/WoN9BqDIxghNgHatsyuLoZsQEmIMQOGHER1zUwRGLAAPCUNXBp5obMjIhv4C+StMFZzsupiBPxPBLqpqVyl6rs/OWJiLccAQAAAAAAAAAAAAAAAAAAAAAAAExR2vUDvvvuu/aFF17Y+eOIeOutt9o333wz6rqOlFKk5MsIObVtG03TxMnJSaxWq/j888/j/PzcH7xLMt/1A65fvx7ffvtt++KLL/om7ahpmlgul3FychJN00Rd11FVVd+fFozafD6P5XIZVVXFq6++Gk3TtP/973+9fl2CnV/dTk9P4/nnn4+7d++2OT6hsdsMihMM5NfFpaqqWCwWcfv27Xj//fe9fl2CnQNz7969mM1m8c9//lNkdtRdFqvrOiLC6QUya9s2Tk9PI6UUTdNEVVXRtm3cvHkzPvroI69fme38CndychJ1Xcd8Po/r16/H999/75v0mJqmiYg/Ti5t60sHOXVv6B7+s5ZSin//+9/xwQcf+EOY0cFvoZ955hknGaBIt27dEpmMDg5MVVXx3HPPOckAxekul3344YdevzI4ODCz2SxWq1U8++yzIgMUJ6UUN2/edJLJ4ODALJfL9T0FJxmgJE3TRNM00bZt3LhxQ2SO7ODAzOfzSCnFbDaLtm3d+AeKMZ/P10+WpZTi1q1bLpcd0cGB6caC3RNSTdPEc88958Y/MHjdE2ZdZCLCSeaIDg5MSilWq9X6OfOTkxMnGaAIKaX16aVt23VkPF12HEdZ+nUnmM2TTES48Q8M3uYIs9M9XSYyh8k+JbeTAUrlJHOY7IGxkwFKZSdzmOyBsZMBSmYns7/sgbGTAUplJ3OY7IGxkwFKZSdzmOyBsZMBSmUnc5jsgbGTAUplJ3OYS/kbr+xkgFLZyeyv979S0U4GKJWTzHa9B8ZOBiiVncx2vQfGTgYomZ3Mo/UeGDsZoFR2Mtv1Hhg7GaBUdjLb9R4YOxmgVHYy2/UeGDsZoFR2Mtv1HpgIOxmgXHYyjzaIwGxjJwOUauonmcEHxk4GKNXUdzKDD4ydDFCyKe9kBh8YOxmgVFPfyQw+MHYyQKmmvpMZfGDsZIBSTX0nM/jA2MkApZr6TmbwgYmwkwHKNeWdTBGB2cZOBijV2E8yxQfGTgYo1dh3MsUHxk4GKNmYdzLFB8ZOBijV2HcyxQfGTgYo1dh3MsUHxk4GKNXYdzLFB8ZOBijV2HcyxQcmwk4GKNeYdzKjCMw2djJAqUo/yYw+MHYyQKlK38mMPjB2MkDJSt7JjD4wdjJAqUrfyYw+MHYyQKlK38mMPjB2MkCpSt/JjD4wdjJAqUrfyYw+MBF2MkC5St7JTCIw29jJAKUa+klm8oGxkwFKNfSdzOQDYycDlGzIO5nJB8ZOBijV0Hcykw+MnQxQqqHvZCYfGDsZoFRD38lMPjB2MkCphr6TmXxgIuxkgHINeScjMBewkwFK1fdJRmAuYCcDlKrvnYzAXMBOBihZnzsZgbmAnQxQqr53MgJzATsZoFR972QE5gJ2MkCp+t7JCMwF7GSAUvW9kxGYx3CsnUz3je7+++Zz6wA59LmT8Qp3oH13Mm3bRl3XOT4lgMeS+yQjMAfaZyeTUnJ6AXqXeyfjVe5Au+5kuktsKaX15TKAvuTcyQjMgXbZyWzef9m84QbQh9w7GYE50C47GacWYEhy72QE5kC77GS6b+LmPwB9yb2TEZgD7bqT2bw85hIZ0KfcOxmBOYJddzJN0zi9wCXo3vh15vO5N3YPybmTEZjMNm/8z2az9SkHuByz2Wz9n/fu3Xvgninb3b59O95///29v1DzY34y/FnbtvGvf/0r7t6923755Zcxn89juVxGVVV/OvEAx7VarWI+n6/vNfztb3+L5XLZ96dVjNVqFbdv346IaP/zn//sfNll5w/4/fffZX8HmzfPvvnmm/jqq68eOLZ3766A40sprR/E+cc//hGvvvrqA9MCtuteu05PT+ONN96I9957b6cvnBNMZt07qOVyGXVdxzvvvON3NvTgpZdeal955ZX1tODh+zP82Xw+j8ViEYvFIp566qmdP949mMyqqorVahWz2cxpBXqUUlrHpa5rcXkM3ROy3ddtV04wmXX3Wtq2de0XerT5SG73Zs+N/u26n5vYLf535QSTWdM065+a7Dcz9KdpmvXPDuz+ne02H2He5wQjMJfIjUXoT/dOvHtE2Z/Hi3Vfp7quBQaA4RAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGACyEBgAshAYALIQGGAS6rqOlFLUdR0RESmlnj+j4dv8GjVNs/PHz4/5yQAM1Xw+j5RSnJ2dxWKxiNlsJjKPoa7rOD093etrJTDAJCyXy2ia5oGTjMBsV1VVVFUVq9Vqr48XGGASvvnmm3R2dtb3pzEp7sEAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACQhcAAkIXAAJCFwACT8PLLL7c///xz2/fnMSUCA0zGtWvXQmQuj8AAk7FareLatWvxyy+/iMwlEBhgMlJK0bZtXL16NX799VeRyUxggElo2zbato2UUpycnMSVK1ecZDITGGASUkoREdE0TaSUYj6fx9WrV0UmI4EBJiGlFCmlaJom6rqOtm3j7Owsnnzyyfjf//4nMhkIDDAJ3eWx7iQTEbFYLCKl5CSTicAAk9C29/vRBaa7TNY0Tczn83jiiSecZI5MYIBJatt2HZfVahWz2SyuXr1qJ3NEAgNMwualsU5VVbFcLmM2m0XTNHF6ehpPPvmkR5iPRGCAyWrbNmazWUTcD1B3krly5YrIHIHAAJPVbWM2/7GTOR6BASZt89JZd18mIuxkjkBggMlLKUVVVXYyRyYwwGR1l8Wapom2bdeR6djJHEZggMmqqiqq6v7L4OZOpvvHTuYwAgNMWnd5LOKPyGyebOxk9icwwGStVqt1VDZ1Jxg7mcMIDDBZs9nsgceTN3X/u53M/gQG4BHsZA4jMABb2MnsT2AALmAnsx+BAXgEO5nDCAzAI9jJHEZgALawk9mfwAA8gp3MYQQG4BHsZA4jMAB7spPZTmAADmAn82gCA3AgO5m/JjAAe7KT2U5gAPZkJ7OdwAAcwE7m0QQGYE92MtsJDMCe7GS2ExiATKa+kxEYgIymvJMRGIDMprqTERiATKa+kxEYgEymvpMRGICMpryTERiATKa+kxEYgEymvpMRGICejH0nIzAAPRrzTkZgAHo21p2MwAD0ZOw7GYEB6MnYdzICA9CjMe9kBAagJ2PfyQgMQE/GvpMRGICBKn0nIzAAA1byTkZgAAau1J2MwAAMVOk7GYEBGKjSdzICAzBgJe9kBAZgoErfyQgMwECVvpMRGIBCDX0nIzAABRvyTkZgAAo31J2MwAAUaug7GYEBKNTQdzICA1CwIe9kBAagUEPfyQgMQKGGvpMRGICR6nsnIzAAI9bnTkZgAEaur52MwACMVN87GYEBGKm+dzICAzBife5kBAZgpPreyewcmO46XldAtlsul3F6evqXz6kD/ZnC61ffO5m9TjB1Xa8/Iba7cuVK3Lt3LyJi/XggcPlSSlHX9foN8nw+7/tT6t1FO5mffvrpoMjsHJiudps3j3i07vjZHUmBfnTv1pumiZRS/P77731/SoOwbSfz9NNPx48//rh3ZHZOeF3XMZvN1i+cbNd989q2jWvXrsXLL7/cds+id5cbgXxSStG2bVy/fn29A6mqKk5PT11V+H/dG+CmadavTd0ls6effjp++eWX9u9///vO75B3/oCPP/64vXnzZiwWizg5OVlfLuOvbb5r6o6gJycn8dtvv63vzQD5dG+KuxfR5XIZZ2dnsVgsJv8Gb/P1ZzMy3b9H/HEf+d1334233357p2bs/NW9fft2+uCDD9anGC62XC4fuLa5Wq28e4JLcnJyEhH374F2V14Wi4UrMHHxTqZ7zdr3Ev9e+X799dfTZ599Nvn6P47Nb1Dbtuvf2O7JwOXorrJUVRVnZ2frezCuHty3bScTEesfLfPbb7/t/GvvXYg7d+6kL774Yt8Pn4yUUiwWi/Vitrus6NIi5Nc9LbZardZ/Frs3eFy8k4mI9dN2Z2dnO//6Bx1BXnvttfTZZ5898M3qbqB5d3Df5m/wzSdY/AaH/LpHk7s3eF6bHvQ4O5nuNWufS/oHX+M6Pz9Pn376aUT8Ub3uk/WNBJiuo9xEOT8/T59//nlExLp0KSU30QAm7GhT1jt37qS2bdsbN26sI2PnATBdR/1ZCefn5yki2lu3brlEBjBxR/9hPOfn5yml1N64cePYvzQABcly/erOnTt/eroMgGnJ9uNEu8tlN2/ezPV/AcCAZb0Df35+bicDMFHZH/GykwGYpkt5hthOBmB6Lu2vdLOTAZiWS/07Q+1kAKbj0v9SajsZgGno5fqUnQzA+F36CaZjJwMwbr3eYbeTARiv3h/hspMBGKfeAxNhJwMwRr3dg3mYnQzAuAwmMBF2MgBjMqjARNjJAIzFIK8/2ckAlG9wJ5iOnQxA2QZ5gunYyQCUa9CBibCTASjV4AMTYScDUKLB3oN5mJ0MQFmKCUyEnQxASYoKTISdDEApiry+ZCcDMHzFnWA6djIAw1bkCaZjJwMwXEUHJsJOBmCoig9MhJ0MwBAVew/mYXYyAMMymsBE2MkADMmoAhNhJwMwFKO8fmQnA9C/0Z1gOnYyAP0a5QmmYycD0J9RBybCTgagL6MPTISdDEAfRnsP5mF2MgCXazKBibCTAbhMkwpMhJ0MwGWZ5PUhOxmA/CZ3gunYyQDkNckTTMdOBiCfSQcmwk4GIJfJBybCTgYgh8neg3mYnQzAcQnMBjsZgOMRmIfYyQAch+s/f8FOBuBwTjCPYCcDcBgnmC3sZAD2JzAXsJMB2I/APAY7GYDduQfzmOxkAHYjMDuwkwF4fAKzIzsZgMfj+s4e7GQALuYEsyc7GYDtnGAOYCcD8GgCcyA7GYC/JjBHYCcD8GfuwRyJnQzAgwTmiOxkAP4gMEdmJwNwn+s3GdjJADjBZGMnA0ydE0xG3dNlVVVF0zTRNI0ny4CidPOLpml2/liByezOnTvpk08+ifl8HlVVxXK5jIjwdBkweFVVRV3XERExn+9+wcslsktwfn6e6rpuz8/PY7FYrE80m988gKGp6zqqqorVarXXCUZgLsnrr7+efvjhh/app556YIzpJAMMVXdZv2ma+Prrr/v+dAAAAAAAAAAAAAAAAAAAAAAAAIBC/B9Raav64ChYtAAAAABJRU5ErkJggg=="

# SVG wordmark for HTML file embedding
CRYPTONARY_SVG_WORDMARK = """<?xml version="1.0" encoding="utf-8"?>
<!-- Generator: Adobe Illustrator 28.0.0, SVG Export Plug-In . SVG Version: 6.00 Build 0)  -->
<svg version="1.1" id="Layer_1" xmlns="http://www.w3.org/2000/svg" xmlns:xlink="http://www.w3.org/1999/xlink" x="0px" y="0px"
	 viewBox="0 0 563.91 94.4" style="enable-background:new 0 0 563.91 94.4;" xml:space="preserve">
<style type="text/css">
	.st0{clip-path:url(#SVGID_00000044857774131872014620000018353702127037256634_);}
	.st1{fill:#FFFFFF;}
	.st2{fill:#F9F9F9;}
	.st3{fill:#005FFF;}
	.st4{fill:#0043C9;}
</style>
<g>
	<path class="st1" d="M84.29,1.21H7.22v77.06l14.92,14.91H99.2V16.12L84.29,1.21z M89.05,40.55H74.14V25.64h-42.1v42.11h42.1V52.84
		l14.91,14.91v14.91H32.03L17.11,67.75V10.73l14.92,14.91V10.73h42.1l14.91,14.91V40.55z"/>
</g>
<g>
	<g>
		<path class="st1" d="M140.33,30.9c12.48,0,19.13,7.71,19.89,14.75h-9.91c-0.23-0.76-2.87-6.05-10.21-6.05
			c-6.81,0-10.96,4.23-11.04,11.04c0,6.73,4.39,11.57,11.34,11.65c7.18,0,9.91-5.29,10.13-6.13h9.91
			c-1.21,7.41-7.64,14.82-19.74,14.9c-12.48-0.08-20.8-8.02-20.8-20.34C119.91,38.39,128.07,30.9,140.33,30.9z"/>
		<path class="st1" d="M195.38,30.9v8.62c-9.38,0-19.13,3.33-19.13,13.84v16.49h-9.23V32.03h9.23v1.13c0,2.87-1.59,6.35-2.72,9
			l1.59,0.68c1.13-2.65,2.34-6.43,4.69-8.17C183.81,31.58,187.97,30.9,195.38,30.9z"/>
		<path class="st1" d="M221.24,53.21c1.29,2.72,1.29,6.58,1.29,9.6h1.74c0-3.02,0-6.88,1.29-9.6l9.83-21.17h11.34l-18.38,37.74
			c-1.13,2.34-2.19,4.23-3.33,6.65c-3.4,7.34-6.96,12.33-13.31,12.33h-10.29v-9.23h5.97c4.54,0,6.2-2.42,8.39-6.5l2.34-4.54
			l-18.07-36.45h11.34L221.24,53.21z"/>
		<path class="st1" d="M261.69,88.75l-9.3,0.08V32.11l9.3-0.08v2.72c0,2.27-1.29,4.92-2.04,7.11l1.59,0.53
			c0.76-2.19,1.13-5.29,2.65-6.88c2.8-2.95,6.73-4.69,11.8-4.69c11.34,0,19.13,8.77,18.91,19.96c-0.3,10.89-7.94,20.19-18.91,20.19
			c-4.99,0-8.92-1.97-11.72-4.92c-1.59-1.59-1.97-4.69-2.72-6.81l-1.59,0.53c0.68,2.19,2.04,4.84,2.04,7.11V88.75z M263.36,50.64
			c0,6.43,4.01,11.72,11.04,11.72c7.11,0,10.89-5.29,10.89-11.72c0-6.35-3.78-11.19-10.89-11.19
			C267.36,39.44,263.36,44.28,263.36,50.64z"/>
		<path class="st1" d="M306.84,40.12h-7.56v-8.09h7.56V17.74h9.23v14.29h9.83v8.09h-9.83v17.02c0,3.18,1.44,4.08,4.99,4.08h4.46
			v8.62h-9.15c-7.11,0-9.53-6.65-9.53-12.7V40.12z"/>
		<path class="st1" d="M352.97,30.9c13.08,0,21.63,8.32,21.63,19.96c0,11.72-8.55,20.19-21.63,20.19
			c-13.01,0-21.55-8.02-21.55-20.19S339.96,30.9,352.97,30.9z M352.97,62.13c8.39,0,11.8-5.07,11.8-11.27
			c0-6.2-3.4-11.04-11.8-11.04c-8.39,0-11.72,4.84-11.72,11.04C341.24,57.06,344.65,62.13,352.97,62.13z"/>
		<path class="st1" d="M404.84,30.9c7.71,0,13.69,4.39,13.69,13.01v25.94h-9.3V50.56c0-7.26-2.65-10.51-8.17-10.51
			c-6.05,0-11.19,3.25-11.19,10.51v19.28h-9.23V32.03h9.23v0.6c0,3.03-1.51,6.66-2.65,9.53l1.59,0.68c1.13-2.87,2.19-6.96,4.61-8.85
			C396.6,31.5,399.92,30.9,404.84,30.9z"/>
		<path class="st1" d="M444.01,71.05c-11.27,0-19.21-9.07-19.21-20.34c0-11.27,7.94-19.81,19.21-19.81c5.07,0,9,1.74,11.72,4.69
			c1.59,1.66,1.89,4.76,2.65,6.96l1.59-0.53c-0.68-2.19-2.04-4.84-2.04-7.11l-0.08-2.87h9.3v37.81h-9.3l0.08-3.1
			c0-2.27,1.36-4.99,2.04-7.11l-1.66-0.53c-0.68,2.12-0.98,5.22-2.5,6.96C453.01,69.16,449.08,71.05,444.01,71.05z M444.54,62.43
			c6.2,0,10.36-5.52,10.59-11.72c0-6.65-3.71-11.19-10.59-11.19c-6.28,0-10.44,4.76-10.44,11.19
			C434.1,57.06,438.26,62.43,444.54,62.43z"/>
		<path class="st1" d="M504.05,30.9v8.62c-9.38,0-19.13,3.33-19.13,13.84v16.49h-9.23V32.03h9.23v1.13c0,2.87-1.59,6.35-2.72,9
			l1.59,0.68c1.13-2.65,2.34-6.43,4.69-8.17C492.48,31.58,496.64,30.9,504.05,30.9z"/>
		<path class="st1" d="M529.91,53.21c1.29,2.72,1.29,6.58,1.29,9.6h1.74c0-3.02,0-6.88,1.29-9.6l9.83-21.17h11.34l-18.38,37.74
			c-1.13,2.34-2.19,4.23-3.33,6.65c-3.4,7.34-6.96,12.33-13.31,12.33H510.1v-9.23h5.97c4.54,0,6.2-2.42,8.39-6.5l2.34-4.54
			l-18.07-36.45h11.34L529.91,53.21z"/>
	</g>
</g>
</svg>
"""

# ── HELPERS ───────────────────────────────────────────────────────

def tg(method, data):
    url = "https://api.telegram.org/bot" + TELEGRAM_TOKEN + "/" + method
    payload = json.dumps(data).encode()
    req = urllib.request.Request(url, data=payload, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=30) as r:
            result = json.loads(r.read())
            if not result.get("ok"):
                print(f"Telegram API error [{method}]: {result.get('description', 'unknown')}", flush=True)
            elif method == "sendMessage":
                msg_id = result.get("result", {}).get("message_id", "?")
                chat = result.get("result", {}).get("chat", {}).get("id", "?")
                print(f"Message sent ok: chat={chat} msg_id={msg_id}", flush=True)
            return result
    except urllib.error.HTTPError as e:
        body = e.read().decode("utf-8", errors="replace")
        print(f"Telegram HTTP error [{method}] {e.code}: {body}", flush=True)
        return None
    except Exception as e:
        print(f"Telegram error [{method}]: {e}", flush=True)
        return None

def send(chat_id, text, keyboard=None):
    # Telegram max message length is 4096 chars - chunk if needed
    max_len = 4000
    if len(text) <= max_len:
        data = {"chat_id": chat_id, "text": text, "parse_mode": "Markdown"}
        if keyboard:
            data["reply_markup"] = {"inline_keyboard": keyboard}
        tg("sendMessage", data)
    else:
        # Send in chunks, keyboard only on last chunk
        chunks = []
        while len(text) > max_len:
            # Break at last newline before limit
            split_at = text.rfind("\n", 0, max_len)
            if split_at == -1: split_at = max_len
            chunks.append(text[:split_at])
            text = text[split_at:].lstrip("\n")
        chunks.append(text)
        for i, chunk in enumerate(chunks):
            is_last = (i == len(chunks) - 1)
            data = {"chat_id": chat_id, "text": chunk, "parse_mode": "Markdown"}
            if is_last and keyboard:
                data["reply_markup"] = {"inline_keyboard": keyboard}
            tg("sendMessage", data)
            if not is_last:
                time.sleep(0.3)  # avoid hitting rate limits

def send_plain(chat_id, text, keyboard=None):
    """Send without Markdown parsing - safe for raw Claude output."""
    max_len = 4000
    if len(text) <= max_len:
        data = {"chat_id": chat_id, "text": text}
        if keyboard:
            data["reply_markup"] = {"inline_keyboard": keyboard}
        tg("sendMessage", data)
    else:
        chunks = []
        while len(text) > max_len:
            split_at = text.rfind("\n", 0, max_len)
            if split_at == -1: split_at = max_len
            chunks.append(text[:split_at])
            text = text[split_at:].lstrip("\n")
        chunks.append(text)
        for i, chunk in enumerate(chunks):
            is_last = (i == len(chunks) - 1)
            data = {"chat_id": chat_id, "text": chunk}
            if is_last and keyboard:
                data["reply_markup"] = {"inline_keyboard": keyboard}
            tg("sendMessage", data)
            if not is_last:
                time.sleep(0.3)


def claude(prompt, max_tokens=900, system=None):
    payload = json.dumps({
        "model": "claude-sonnet-4-20250514",
        "max_tokens": max_tokens,
        "system": system if system else VOICE_GUIDE,
        "messages": [{"role": "user", "content": prompt}]
    }).encode()
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=payload,
        headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"}
    )
    with urllib.request.urlopen(req, timeout=60) as r:
        data = json.loads(r.read())
        return "".join(c.get("text", "") for c in data["content"])

def clean_json(s):
    s = s.replace("```json", "").replace("```", "").strip()
    start = next((i for i, c in enumerate(s) if c in ('[', '{')), -1)
    if start >= 0: s = s[start:]
    end = next((i for i in range(len(s)-1, -1, -1) if s[i] in (']', '}')), -1)
    if end >= 0: s = s[:end+1]
    return s

def sanitise(text):
    if not text: return ""
    pairs = [
        (u"\u2018", "'"), (u"\u2019", "'"), (u"\u201a", "'"),
        (u"\u201c", '"'), (u"\u201d", '"'), (u"\u201e", '"'),
        (u"\u2013", "-"), (u"\u2014", "-"), (u"\u2015", "-"),
        (u"\u2026", "..."), (u"\u2022", "-"), (u"\u00b7", "-"),
    ]
    for a, b in pairs:
        text = text.replace(a, b)
    out = []
    for ch in text:
        cp = ord(ch)
        if cp == 10 or cp == 9 or cp == 13:
            out.append(ch)
        elif cp < 32:
            continue
        elif 127 <= cp <= 159:
            continue
        else:
            out.append(ch)
    return "".join(out)


def parse_numbered_list(text, count):
    """Parse a numbered list response into a list of strings."""
    lines = text.strip().split("\n")
    results = []
    current = ""
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Check if line starts with a number
        m = re.match(r"^\d+\.\s*(.+)", line)
        if m:
            if current:
                results.append(current.strip())
            current = m.group(1)
        else:
            if current:
                current += " " + line
    if current:
        results.append(current.strip())
    # Pad or trim to expected count
    while len(results) < count:
        results.append("Alternative angle " + str(len(results)+1))
    return results[:count]

def parse_hooks(text):
    """Parse numbered SUBJECT/PREVIEW format into list of dicts."""
    hooks = []
    # Split by numbered entries
    blocks = re.split(r"(?m)^\d+\.\s*", text.strip())
    for block in blocks:
        if not block.strip():
            continue
        subject = ""
        preview = ""
        for line in block.strip().split("\n"):
            line = line.strip()
            if line.upper().startswith("SUBJECT:"):
                subject = line[8:].strip()
            elif line.upper().startswith("PREVIEW:"):
                preview = line[8:].strip()
        if subject:
            hooks.append({"subject": subject, "preview": preview or ""})
    if not hooks:
        # Fallback: try to parse as plain lines
        lines = [l.strip() for l in text.strip().split("\n") if l.strip()]
        for i in range(0, len(lines)-1, 2):
            hooks.append({"subject": lines[i], "preview": lines[i+1] if i+1 < len(lines) else ""})
    while len(hooks) < 4:
        hooks.append({"subject": "Alternative hook " + str(len(hooks)+1), "preview": ""})
    return hooks[:4]

def parse_subject_alternatives(text):
    """Parse numbered PRINCIPLE/SUBJECT/PREVIEW/REASON format."""
    alts = []
    blocks = re.split(r"(?m)^\d+\.\s*", text.strip())
    for block in blocks:
        if not block.strip():
            continue
        principle = ""
        subject = ""
        preview = ""
        reason = ""
        for line in block.strip().split("\n"):
            line = line.strip()
            if line.upper().startswith("PRINCIPLE:"):
                principle = line[10:].strip()
            elif line.upper().startswith("SUBJECT:"):
                subject = line[8:].strip()
            elif line.upper().startswith("PREVIEW:"):
                preview = line[8:].strip()
            elif line.upper().startswith("REASON:"):
                reason = line[7:].strip()
        if subject:
            alts.append({"principle": principle, "subject": subject, "preview": preview, "reason": reason})
    while len(alts) < 3:
        alts.append({"principle": "Alternative", "subject": "Alternative subject " + str(len(alts)+1), "preview": "", "reason": ""})
    return alts[:3]

def parse_enhance_suggestions(text):
    suggestions = []
    idx = 1
    current = {"principle": "", "issue": "", "fix": ""}
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        num_match = re.match(r"^(\d+)\.\s*(.*)", line)
        if num_match:
            if current.get("issue") or current.get("fix"):
                suggestions.append({"id": idx, "principle": current["principle"] or "Copywriting", "issue": current["issue"] or current["fix"][:50], "fix": current["fix"] or current["issue"]})
                idx += 1
            current = {"principle": "", "issue": "", "fix": ""}
            rest = num_match.group(2)
            if rest.upper().startswith("PRINCIPLE:"):
                current["principle"] = rest[10:].strip()
        elif line.upper().startswith("PRINCIPLE:"):
            current["principle"] = line[10:].strip()
        elif line.upper().startswith("ISSUE:"):
            current["issue"] = line[6:].strip()
        elif line.upper().startswith("FIX:"):
            current["fix"] = line[4:].strip()
    if current.get("issue") or current.get("fix"):
        suggestions.append({"id": idx, "principle": current["principle"] or "Copywriting", "issue": current["issue"] or current["fix"][:50], "fix": current["fix"] or current["issue"]})
    if not suggestions:
        suggestions = [{"id": 1, "principle": "General", "issue": "Could not parse", "fix": "Please try again"}]
    return suggestions


def parse_delimited_emails(text):
    """Parse FREE/PRO email delimiter format into dict."""
    emails = {}
    if "===FREE EMAIL START===" in text and "===FREE EMAIL END===" in text:
        start = text.find("===FREE EMAIL START===") + len("===FREE EMAIL START===")
        end = text.find("===FREE EMAIL END===")
        emails["free"] = text[start:end].strip()
    if "===PRO EMAIL START===" in text and "===PRO EMAIL END===" in text:
        start = text.find("===PRO EMAIL START===") + len("===PRO EMAIL START===")
        end = text.find("===PRO EMAIL END===")
        emails["pro"] = text[start:end].strip()
    # Fallback: if delimiters not found, try JSON
    if not emails:
        try:
            emails = json.loads(clean_json(text))
        except:
            emails = {"free": text, "pro": text}
    return emails

def parse_delimited_segments(text):
    """Parse HOT/WARM/COLD segment delimiter format into dict."""
    segments = {}
    for key in ["hot", "warm", "cold"]:
        start_tag = "===" + key.upper() + " EMAIL START==="
        end_tag = "===" + key.upper() + " EMAIL END==="
        if start_tag in text and end_tag in text:
            start = text.find(start_tag) + len(start_tag)
            end = text.find(end_tag)
            segments[key] = text[start:end].strip()
    if not segments:
        try:
            segments = json.loads(clean_json(text))
        except:
            segments = {"hot": text, "warm": text, "cold": text}
    return segments

def extract_text(v):
    if isinstance(v, str): return v
    if isinstance(v, dict):
        for k in ["body", "content", "email"]:
            if k in v: return v[k]
        return "\n\n".join(str(x) for x in v.values())
    return str(v)

def email_action_keyboard():
    return [
        [{"text": "Quick Edit", "callback_data": "quick_edit"},
         {"text": "Enhance", "callback_data": "enhance"},
         {"text": "Approve", "callback_data": "approve_emails"}],
        [{"text": "Subject Lines", "callback_data": "subject_ab"},
         {"text": "Tone", "callback_data": "tone_menu"},
         {"text": "Segments", "callback_data": "segments"}],
        [{"text": "Length", "callback_data": "length_email"},
         {"text": "Add Context", "callback_data": "add_context_email"},
         {"text": "Critique", "callback_data": "critique_email"}]
    ]

def ad_action_keyboard():
    return [
        [{"text": "Quick Edit",           "callback_data": "ad_quick_edit"},
         {"text": "Enhance",              "callback_data": "ad_enhance"},
         {"text": "Approve",              "callback_data": "approve_ad"}],
        [{"text": "Critique",             "callback_data": "critique_ad"},
         {"text": "Generate another set", "callback_data": "ads_again"}],
        [{"text": "🎨 Visual brief",      "callback_data": "vb_type_ad_static"},
         {"text": "Mark Complete",        "callback_data": "mark_complete"}]
    ]

def format_action_keyboard(fmt_key, fmt_label):
    """Action keyboard for a specific format — encodes format in callback."""
    kb = [
        [{"text": "Quick Edit", "callback_data": "sfmt_edit_" + fmt_key},
         {"text": "Enhance",    "callback_data": "sfmt_enhance_" + fmt_key},
         {"text": "Approve",    "callback_data": "sfmt_approve_" + fmt_key}],
        [{"text": "Critique",   "callback_data": "critique_social_" + fmt_key},
         {"text": "Length",     "callback_data": "length_social"}],
        [{"text": "🎨 Visual brief", "callback_data": "vb_auto"}],
    ]
    return kb

def social_action_keyboard():
    return [
        [{"text": "Quick Edit", "callback_data": "social_quick_edit"},
         {"text": "Enhance", "callback_data": "social_enhance"},
         {"text": "Approve", "callback_data": "approve_social"}],
        [{"text": "Length", "callback_data": "length_social"},
         {"text": "Critique", "callback_data": "critique_social"}]
    ]

def mark_complete_keyboard():
    return [
        [{"text": "New session", "callback_data": "start_over"}],
        [{"text": "Log email", "callback_data": "log_email_start"},
         {"text": "Log ad", "callback_data": "log_ad_start"}],
        [{"text": "Email report", "callback_data": "run_email_report"},
         {"text": "Ad report", "callback_data": "run_ad_report"}]
    ]

def get_perf_context(chat_id):
    records = performance_data.get(chat_id, [])
    if not records: return ""
    ctx = "\n\nPERFORMANCE DATA FROM PAST EMAILS:\n"
    for r in records[-10:]:
        ctx += "- Subject: " + r.get("subject","") + " | Open: " + str(r.get("open_rate","?")) + "% | Click: " + str(r.get("click_rate","?")) + "%\n"
    high = [r for r in records if r.get("open_rate", 0) > 30]
    if high:
        ctx += "\nHIGH PERFORMERS (>30% open):\n"
        for r in high:
            ctx += "- " + r.get("subject","") + " (" + str(r.get("open_rate","?")) + "% open)\n"
    return ctx

# ── FLOW ──────────────────────────────────────────────────────────

def ask_context(chat_id):
    user_state[chat_id]["stage"] = "awaiting_context_choice"
    keyboard = [
        [{"text": "Yes — I have extra context", "callback_data": "context_yes"}],
        [{"text": "No — just the report", "callback_data": "context_no"}]
    ]
    send(chat_id, "*Any extra context to factor in?*\n\n_Promos, discounts, Inner Circle open, upcoming events, factoids, PSAs..._", keyboard)

def gen_angles(chat_id):
    """Generate ONE set of angles used for BOTH Free and Pro emails."""
    report = sanitise(user_state[chat_id].get("report", ""))
    context = sanitise(user_state[chat_id].get("context", ""))
    perf = get_perf_context(chat_id)
    prompt = "REPORT:\n" + report
    if context: prompt += "\n\nEXTRA CONTEXT:\n" + context
    prompt += perf
    prompt += "\n\nGenerate exactly 4 distinct content angles for this crypto update. Each angle will be used to write both a Free email (curiosity/tease) and a Pro email (full analysis). Different emotional lenses or hook strategies. Apply copywriting principles.\n1. [angle]\n2. [angle]\n3. [angle]\n4. [angle]\nNothing else."
    send(chat_id, "Finding angles...")
    try:
        raw = claude(prompt)
        angles = parse_numbered_list(raw, 4)
        # Store under both keys for compatibility
        user_state[chat_id]["angles"] = angles
        user_state[chat_id]["free_angles"] = angles
        user_state[chat_id]["pro_angles"] = angles
        user_state[chat_id]["stage"] = "pick_angle"
        text = "*Pick an angle:*\n_(Used for both Free and Pro emails)_\n\n"
        keyboard = []
        for i, a in enumerate(angles):
            text += "*" + str(i+1) + ".* " + a + "\n\n"
            keyboard.append([{"text": str(i+1), "callback_data": "angle_" + str(i)}])
        keyboard.append([{"text": "✏️ Write my own angle", "callback_data": "custom_angle"}])
        keyboard.append([{"text": "Regenerate angles", "callback_data": "regen_angles"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def gen_hooks(chat_id):
    """Generate hooks — one set used for both Free and Pro emails."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    perf = get_perf_context(chat_id)
    angle = state.get("selected_angle", state.get("selected_free_angle", ""))

    base = "REPORT:\n" + report
    if context: base += "\n\nEXTRA CONTEXT:\n" + context
    base += perf

    send(chat_id, "Writing hooks...")
    try:
        raw = claude(base + "\n\nANGLE: " + angle +
            "\n\nWrite 4 subject line + preview text combos for this crypto email.\n" +
            "Each hook should create a strong curiosity gap or bold data-led statement.\n\n" +
            "1. SUBJECT: [subject]\nPREVIEW: [preview]\n\n(repeat for all 4)", max_tokens=400)
        hooks = parse_hooks(raw)
        # Store under all keys for compatibility
        state["free_hooks"] = hooks
        state["pro_hooks"] = hooks
        state["hooks"] = hooks
        state["stage"] = "pick_free_hook"
        text = "*Pick a hook:*\n_(Used for both Free and Pro emails)_\n\n"
        keyboard = []
        for i, h in enumerate(hooks):
            text += "*" + str(i+1) + ".* " + h["subject"] + "\n_" + h["preview"] + "_\n\n"
            keyboard.append([{"text": str(i+1), "callback_data": "free_hook_" + str(i)}])
        keyboard.append([{"text": "✏️ Write my own", "callback_data": "custom_free_hook"}])
        keyboard.append([{"text": "Regenerate", "callback_data": "regen_free_hooks"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def show_pro_angle_picker(chat_id):
    """Show Pro angle picker after Free angle is selected."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    pro_angles = state.get("pro_angles", [])
    state["stage"] = "pick_pro_angle"
    text = "*PRO EMAIL — Pick an angle:*\n_(Full analysis, data-led, complete conviction)_\n\n"
    keyboard = []
    for i, a in enumerate(pro_angles):
        text += "*" + str(i+1) + ".* " + a + "\n\n"
        keyboard.append([{"text": str(i+1), "callback_data": "pro_angle_" + str(i)}])
    keyboard.append([{"text": "✏️ Write my own", "callback_data": "custom_pro_angle"}])
    keyboard.append([{"text": "Regenerate", "callback_data": "regen_pro_angles"}])
    send(chat_id, text, keyboard)

def gen_pro_hooks(chat_id):
    """Generate Pro email hooks — called after Free hook selected."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    perf = get_perf_context(chat_id)
    pro_angle = state.get("selected_pro_angle", state.get("selected_angle", ""))

    base = "REPORT:\n" + report
    if context: base += "\n\nEXTRA CONTEXT:\n" + context
    base += perf

    send(chat_id, "Writing hooks for the Pro email...")
    try:
        raw = claude(base + "\n\nPRO EMAIL ANGLE: " + pro_angle +
            "\n\nWrite 4 subject line + preview text combos for the PRO email.\n" +
            "PRO hooks: Data-led, specific, authoritative. Paying members expect depth and directness.\n\n" +
            "1. SUBJECT: [subject]\nPREVIEW: [preview]\n\n(repeat for all 4)", max_tokens=400)
        hooks = parse_hooks(raw)
        state["pro_hooks"] = hooks
        state["stage"] = "pick_pro_hook"
        text = "*PRO EMAIL — Pick a hook:*\n_(Data-led, authoritative, specific)_\n\n"
        keyboard = []
        for i, h in enumerate(hooks):
            text += "*" + str(i+1) + ".* " + h["subject"] + "\n_" + h["preview"] + "_\n\n"
            keyboard.append([{"text": str(i+1), "callback_data": "pro_hook_" + str(i)}])
        keyboard.append([{"text": "✏️ Write my own", "callback_data": "custom_pro_hook"}])
        keyboard.append([{"text": "Regenerate", "callback_data": "regen_pro_hooks"}])
        keyboard.append([{"text": "Same as Free hook", "callback_data": "pro_hook_same"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))


def ask_free_cta(chat_id):
    user_state[chat_id]["stage"] = "pick_free_cta"
    keyboard = [[{"text": v, "callback_data": "cta_free_" + k}] for k, v in CTA_OPTIONS["free"].items()]
    send(chat_id, "*Free email CTA:*", keyboard)

def ask_pro_cta(chat_id):
    user_state[chat_id]["stage"] = "pick_pro_cta"
    keyboard = [[{"text": v, "callback_data": "cta_pro_" + k}] for k, v in CTA_OPTIONS["pro"].items()]
    send(chat_id, "*Pro email CTA:*", keyboard)

def gen_emails(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "")
    hook = state.get("selected_hook", {})
    free_cta = state.get("free_cta", "upgrade")
    pro_cta = state.get("pro_cta", "read_report")
    voice_examples = get_voice_corpus_context(chat_id)
    prompt = "Write BOTH a Free email and a Pro email for Cryptonary. Apply all copywriting principles.\n\n"
    if voice_examples:
        prompt += voice_examples + "\n\n"
    prompt += "REPORT:\n" + report
    if context: prompt += "\n\nEXTRA CONTEXT (weave in naturally):\n" + context
    prompt += "\n\nAngle: " + angle
    prompt += "\nSubject: " + hook.get("subject", "")
    prompt += "\nPreview: " + hook.get("preview", "")
    prompt += "\n\nFREE EMAIL CTA: " + CTA_INSTRUCTIONS["free"].get(free_cta, "")
    prompt += "\nPRO EMAIL CTA: " + CTA_INSTRUCTIONS["pro"].get(pro_cta, "")
    prompt += "\n\nWrite the two emails separated by this exact delimiter. Do not use JSON. IMPORTANT: Every email MUST begin with Subject Line: and Preview: on the first two lines.\n\n===FREE EMAIL START===\nSubject Line: [subject here]\nPreview: [preview here]\n\n[complete free email body here]\n===FREE EMAIL END===\n\n===PRO EMAIL START===\nSubject Line: [subject here]\nPreview: [preview here]\n\n[complete pro email body here]\n===PRO EMAIL END==="
    send(chat_id, "Writing your emails...")
    try:
        raw = claude(prompt, max_tokens=2500)
        emails = parse_delimited_emails(raw)
        state["current_emails"] = emails
        state["stage"] = "emails_ready"
        # Store for voice learning on approve
        state["current_emails"]["free_raw"] = emails.get("free", "")
        state["current_emails"]["pro_raw"] = emails.get("pro", "")
        state["email_record"] = {
            "subject": hook.get("subject", ""),
            "angle": angle,
            "date": time.strftime("%Y-%m-%d"),
            "timestamp": time.strftime("%Y-%m-%d %H:%M")
        }
        if "free" in emails:
            send_plain(chat_id, "FREE EMAIL\n\n" + extract_text(emails["free"]))
        if "pro" in emails:
            send_plain(chat_id, "PRO EMAIL\n\n" + extract_text(emails["pro"]))
        send(chat_id, "Emails ready. What would you like to do?", email_action_keyboard())
    except Exception as e:
        send(chat_id, "Error: " + str(e))

# ── QUICK EDIT ────────────────────────────────────────────────────

def ask_quick_edit(chat_id, mode="email"):
    user_state[chat_id]["stage"] = "awaiting_quick_edit"
    user_state[chat_id]["quick_edit_mode"] = mode
    send(chat_id, "*Quick Edit*\n\nType your instruction:\n\n_Examples: Make the opening more urgent / Shorten the P.S. / Make the CTA stronger / Add a specific number to paragraph 2_")

def apply_quick_edit(chat_id, instruction):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    mode = state.get("quick_edit_mode", "email")
    if mode == "social":
        content = state.get("current_social", "")
        social_type = state.get("current_social_type", "social content")
        send(chat_id, "Applying edit to " + social_type + "...")
        try:
            result = claude(
                "Edit this Cryptonary " + social_type + " based on this instruction: " + instruction +
                "\n\nKeep Adam's voice intact. Only change what the instruction specifies.\n\nCONTENT:\n" + content +
                "\n\nReturn the edited content as a plain string, not JSON.",
                max_tokens=1500
            )
            state["current_social"] = result
            state["stage"] = "social_ready"
            send_plain(chat_id, "*EDITED " + social_type.upper() + "*\n\n" + result)
            send(chat_id, "Edit applied.", social_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "social_ready"
    else:
        emails = state.get("current_emails", {})
        free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
        pro_email = extract_text(emails.get("pro", "")) if "pro" in emails else ""
        email_text = ""
        if free_email: email_text += "FREE EMAIL:\n" + free_email + "\n\n"
        if pro_email: email_text += "PRO EMAIL:\n" + pro_email
        send(chat_id, "Applying edit...")
        try:
            raw = claude(
                "Edit these Cryptonary emails based on this instruction: " + instruction +
                "\n\nKeep Adam's voice intact. Only change what the instruction specifies.\n\nEMAILS:\n" + email_text +
                "\n\nReturn the edited emails using this exact format. No JSON:\n\n===FREE EMAIL START===\n[edited free email]\n===FREE EMAIL END===\n\n===PRO EMAIL START===\n[edited pro email]\n===PRO EMAIL END===",
                max_tokens=2500
            )
            edited = parse_delimited_emails(raw)
            state["current_emails"] = edited
            state["stage"] = "emails_ready"
            if "free" in edited:
                send_plain(chat_id, "EDITED FREE EMAIL\n\n" + extract_text(edited["free"]))
            if "pro" in edited:
                send_plain(chat_id, "EDITED PRO EMAIL\n\n" + extract_text(edited["pro"]))
            send(chat_id, "Edit applied.", email_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "emails_ready"

# ── ENHANCE ───────────────────────────────────────────────────────

def gen_enhance(chat_id, mode="email"):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    if mode == "social":
        content = state.get("current_social", "")
        social_type = state.get("current_social_type", "social content")
        send(chat_id, "Analysing " + social_type + "...")
        try:
            raw = claude(
                "Analyse this Cryptonary " + social_type + " and generate exactly 6 specific improvement suggestions. Each must reference a copywriting principle and give a concrete fix.\n\nCONTENT:\n" + content +
                "\n\nFormat as a numbered list:\n1. PRINCIPLE: [name]\nISSUE: [what is weak]\nFIX: [specific improvement]\n\n2. PRINCIPLE: ...\nAnd so on for all 6. Nothing else.",
                max_tokens=900
            )
            suggestions = parse_enhance_suggestions(raw)
            state["enhance_suggestions"] = suggestions
            state["selected_enhancements"] = []
            state["enhance_mode"] = "social"
            state["stage"] = "enhance_select"
            text = "*Choose improvements:*\n_(tap to select, then Apply)_\n\n"
            keyboard = []
            for s in suggestions:
                keyboard.append([{"text": "[ ] " + str(s["id"]) + ". [" + s["principle"] + "] " + s["issue"], "callback_data": "enh_" + str(s["id"])}])
            keyboard.append([{"text": "Apply selected", "callback_data": "apply_enhancements"}])
            keyboard.append([{"text": "Skip", "callback_data": "skip_enhance"}])
            send(chat_id, text, keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))
    else:
        emails = state.get("current_emails", {})
        free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
        pro_email = extract_text(emails.get("pro", "")) if "pro" in emails else ""
        email_text = ""
        if free_email: email_text += "FREE EMAIL:\n" + free_email + "\n\n"
        if pro_email: email_text += "PRO EMAIL:\n" + pro_email
        send(chat_id, "Analysing emails against copywriting principles...")
        try:
            raw = claude(
                "Analyse these Cryptonary emails against copywriting principles (Hormozi, Cashvertising LF8, Cialdini, Ogilvy). Generate exactly 8 specific, actionable improvement suggestions.\n\nEmails:\n" + email_text +
                "\n\nFormat as a numbered list:\n1. PRINCIPLE: [name]\nISSUE: [what is weak]\nFIX: [specific improvement]\n\n2. PRINCIPLE: ...\nAnd so on for all 8. Nothing else.",
                max_tokens=1200
            )
            suggestions = parse_enhance_suggestions(raw)
            state["enhance_suggestions"] = suggestions
            state["selected_enhancements"] = []
            state["enhance_mode"] = "email"
            state["stage"] = "enhance_select"
            text = "*Choose improvements:*\n_(tap to select, then Apply)_\n\n"
            keyboard = []
            for s in suggestions:
                keyboard.append([{"text": "[ ] " + str(s["id"]) + ". [" + s["principle"] + "] " + s["issue"], "callback_data": "enh_" + str(s["id"])}])
            keyboard.append([{"text": "Apply selected", "callback_data": "apply_enhancements"}])
            keyboard.append([{"text": "Skip", "callback_data": "skip_enhance"}])
            send(chat_id, text, keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))

def toggle_enhancement(chat_id, enh_id, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_enhancements", [])
    suggestions = state.get("enhance_suggestions", [])
    if enh_id in selected: selected.remove(enh_id)
    else: selected.append(enh_id)
    state["selected_enhancements"] = selected
    keyboard = []
    for s in suggestions:
        is_sel = s["id"] in selected
        label = ("[x] " if is_sel else "[ ] ") + str(s["id"]) + ". [" + s["principle"] + "] " + s["issue"]
        keyboard.append([{"text": label, "callback_data": "enh_" + str(s["id"])}])
    keyboard.append([{"text": "Apply selected (" + str(len(selected)) + ")", "callback_data": "apply_enhancements"}])
    keyboard.append([{"text": "Skip", "callback_data": "skip_enhance"}])
    tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id, "reply_markup": {"inline_keyboard": keyboard}})

def apply_enhancements(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected_ids = state.get("selected_enhancements", [])
    suggestions = state.get("enhance_suggestions", [])
    mode = state.get("enhance_mode", "email")
    if not selected_ids:
        send(chat_id, "No improvements selected. Tap numbers to select first.")
        return
    selected = [s for s in suggestions if s["id"] in selected_ids]
    improvements = "\n".join([str(s["id"]) + ". [" + s["principle"] + "] " + s["fix"] for s in selected])
    send(chat_id, "Applying " + str(len(selected_ids)) + " improvements...")
    if mode == "social":
        content = state.get("current_social", "")
        social_type = state.get("current_social_type", "social content")
        try:
            result = claude(
                "Rewrite this Cryptonary " + social_type + " applying these improvements only:\n\n" + improvements +
                "\n\nORIGINAL:\n" + content +
                "\n\nReturn the improved content as a plain string, not JSON.",
                max_tokens=1500
            )
            state["current_social"] = result
            state["stage"] = "social_ready"
            send_plain(chat_id, "*IMPROVED " + social_type.upper() + "*\n\n" + result)
            send(chat_id, str(len(selected_ids)) + " improvements applied.", social_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "social_ready"
    else:
        emails = state.get("current_emails", {})
        free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
        pro_email = extract_text(emails.get("pro", "")) if "pro" in emails else ""
        email_text = ""
        if free_email: email_text += "FREE EMAIL:\n" + free_email + "\n\n"
        if pro_email: email_text += "PRO EMAIL:\n" + pro_email
        try:
            raw = claude(
                "Rewrite these Cryptonary emails applying these improvements only. Keep Adam's voice intact.\n\nIMPROVEMENTS:\n" + improvements +
                "\n\nORIGINAL EMAILS:\n" + email_text +
                "\n\nReturn the improved emails using this exact format. No JSON:\n\n===FREE EMAIL START===\n[improved free email]\n===FREE EMAIL END===\n\n===PRO EMAIL START===\n[improved pro email]\n===PRO EMAIL END===",
                max_tokens=2500
            )
            improved = parse_delimited_emails(raw)
            prev_free = extract_text(emails.get("free", ""))
            prev_pro = extract_text(emails.get("pro", ""))
            state["prev_emails"] = state.get("current_emails", {})
            state["current_emails"] = improved
            state["stage"] = "emails_ready"
            if "free" in improved:
                send_plain(chat_id, "IMPROVED FREE EMAIL\n\n" + extract_text(improved["free"]))
            if "pro" in improved:
                send_plain(chat_id, "IMPROVED PRO EMAIL\n\n" + extract_text(improved["pro"]))
            # Show diff
            after_free = extract_text(improved.get("free", ""))
            show_diff_with_revert(chat_id, prev_free, after_free, "email")
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "emails_ready"

# ── SOCIAL CONTENT ────────────────────────────────────────────────

def show_social_source_menu(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    emails = state.get("current_emails", {})
    segments = state.get("current_segments", {})
    keyboard = []
    if "free" in emails:
        keyboard.append([{"text": "Free email", "callback_data": "src_free"}])
    if "pro" in emails:
        keyboard.append([{"text": "Pro email", "callback_data": "src_pro"}])
    if "hot" in segments:
        keyboard.append([{"text": "Free — Hot segment", "callback_data": "src_hot"}])
    if "warm" in segments:
        keyboard.append([{"text": "Free — Warm segment", "callback_data": "src_warm"}])
    if "cold" in segments:
        keyboard.append([{"text": "Free — Cold segment", "callback_data": "src_cold"}])
    keyboard.append([{"text": "Cancel", "callback_data": "back_to_done"}])
    send(chat_id, "*Which email should social content be based on?*", keyboard)

def show_social_format_menu(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_social_formats", [])
    formats = [
        ("Reel Script (45-60s)",  "fmt_reel"),
        ("Carousel (5-8 slides)", "fmt_carousel"),
        ("Static Post + Caption", "fmt_static"),
        ("Story — Single slide",  "fmt_story_single"),
        ("Story — Multi slide",   "fmt_story_multi"),
    ]
    keyboard = []
    for label, cb in formats:
        is_sel = cb in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + label, "callback_data": cb}])
    keyboard.append([{"text": "Generate selected (" + str(len(selected)) + ")", "callback_data": "gen_social_selected"}])
    keyboard.append([{"text": "Back", "callback_data": "back_to_done"}])
    send(chat_id, "*Select formats to generate:*\n_(tap to select multiple)_", keyboard)

def toggle_social_format(chat_id, fmt_cb, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_social_formats", [])
    if fmt_cb in selected:
        selected.remove(fmt_cb)
    else:
        selected.append(fmt_cb)
    state["selected_social_formats"] = selected
    # Clear hooks so they regenerate fresh for the new format selection
    state.pop("social_hooks", None)
    state.pop("selected_social_hooks", None)
    formats = [
        ("Reel Script (45-60s)",  "fmt_reel"),
        ("Carousel (5-8 slides)", "fmt_carousel"),
        ("Static Post + Caption", "fmt_static"),
        ("Story — Single slide",  "fmt_story_single"),
        ("Story — Multi slide",   "fmt_story_multi"),
    ]
    keyboard = []
    for label, cb in formats:
        is_sel = cb in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + label, "callback_data": cb}])
    keyboard.append([{"text": "Generate selected (" + str(len(selected)) + ")", "callback_data": "gen_social_selected"}])
    keyboard.append([{"text": "Back", "callback_data": "back_to_done"}])
    tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id, "reply_markup": {"inline_keyboard": keyboard}})

def get_social_source_text(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    source = state.get("social_source", "free")
    emails = state.get("current_emails", {})
    segments = state.get("current_segments", {})
    if source == "free":
        return extract_text(emails.get("free", ""))
    elif source == "pro":
        return extract_text(emails.get("pro", ""))
    elif source in ["hot", "warm", "cold"]:
        return extract_text(segments.get(source, ""))
    return extract_text(emails.get("free", ""))

def gen_social_selected(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_social_formats", [])
    if not selected:
        send(chat_id, "No formats selected. Tap the format names to select them first.")
        return
    # If coming from email flow, use the email angle directly — skip angle/hook pickers
    origin = state.get("social_origin", "")
    if origin == "email" and not state.get("social_angle"):
        # Pull angle from the email flow
        email_angle = state.get("selected_angle", state.get("angle", ""))
        if email_angle:
            state["social_angle"] = email_angle
    # If no angle yet, generate avatar-aware angles first
    if not state.get("social_angle"):
        gen_social_angles(chat_id)
        return
    # Always regenerate hooks when formats change — prevents stale hook state causing freeze
    existing_hooks = state.get("social_hooks", {})
    needs_hooks = any(f not in existing_hooks for f in selected)
    if needs_hooks:
        state["social_hooks"] = {}
        state["selected_social_hooks"] = {}
        gen_social_hooks(chat_id)
        return
    # If reel/video selected and no framework chosen yet, ask PAS or AIDA first
    video_fmts = {"fmt_reel", "fmt_story_single", "fmt_story_multi"}
    if any(f in video_fmts for f in selected) and not state.get("social_framework"):
        keyboard = [
            [{"text": "AIDA — build desire then act", "callback_data": "social_fw_aida"}],
            [{"text": "PAS — problem, agitate, solve", "callback_data": "social_fw_pas"}]
        ]
        send(chat_id, "*Which framework?*\n\n*AIDA* — Attention → Interest → Desire → Action. Builds want, then converts.\n*PAS* — Problem → Agitate → Solution. Opens on pain, intensifies it, then offers the fix.", keyboard)
        return
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "")
    source_email = get_social_source_text(chat_id)[:600]
    source_label = state.get("social_source", "free")
    fmt_map = {
        "fmt_reel": ("Reel Script", gen_reel),
        "fmt_carousel": ("Carousel", gen_carousel),
        "fmt_static": ("Static Post", gen_static),
        "fmt_story_single": ("Story (single)", lambda c: gen_story(c, multi=False)),
        "fmt_story_multi": ("Story (multi)", lambda c: gen_story(c, multi=True)),
    }
    # Only say "email" in the status if we're actually coming from the email flow
    origin = state.get("social_origin", "email")
    if origin == "idea_engine":
        status_msg = "Generating " + str(len(selected)) + " format(s)..."
    elif origin == "voice":
        status_msg = "Generating " + str(len(selected)) + " format(s)..."
    elif origin == "email" and state.get("report"):
        status_msg = "Generating " + str(len(selected)) + " format(s) based on " + source_label + " email..."
    else:
        status_msg = "Generating " + str(len(selected)) + " format(s)..."
    send(chat_id, status_msg)
    for fmt_cb in selected:
        if fmt_cb in fmt_map:
            label, fn = fmt_map[fmt_cb]
            fn(chat_id)
    state["selected_social_formats"] = []

def gen_reel(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "") or state.get("social_angle", "")
    reel_duration = state.get("reel_duration", 52)
    source_email = get_social_source_text(chat_id)[:500]
    hook = state.get("selected_social_hooks", {}).get("fmt_reel", "")
    word_count = int(reel_duration * 2.3)
    voice_examples = get_voice_corpus_context(chat_id)
    framework = state.get("social_framework", "AIDA")
    fw_instruction = {
        "AIDA": "Structure: Attention (hook stops scroll) → Interest (relevant problem/insight) → Desire (show the transformation/benefit) → Action (follow Cryptonary).",
        "PAS": "Structure: Problem (name the exact pain or risk) → Agitate (make them feel the cost of inaction, intensify) → Solution (position Cryptonary as the relief)."
    }.get(framework, "")
    try:
        result = claude(
            (voice_examples + "\n\n" if voice_examples else "") +
            "Write a " + str(reel_duration) + "-second Instagram Reel voiceover script for Cryptonary.\n\n" +
            "FRAMEWORK: " + framework + " — " + fw_instruction + "\n\n" +
            ("OPENING HOOK (use this as the first spoken line): " + hook + "\n\n" if hook else "") +
            "SOURCE:\nReport: " + report + ("\nContext: " + context if context else "") +
            "\nAngle: " + angle + "\nEmail reference: " + source_email +
            "\n\nRULES:\n- Approximately " + str(word_count) + " words (matches " + str(reel_duration) + "s at natural pace)\n- First 3 seconds must stop the scroll\n- Format: voiceover text | [B-roll instruction] for each line\n- CTA at end: follow Cryptonary\n- Adam's voice: punchy, direct, data-led\n\nReturn as plain string.",
            max_tokens=900
        )
        state["current_social"] = result
        state["current_social_type"] = "Reel Script"
        state["stage"] = "social_ready"
        other_fw = "PAS" if framework == "AIDA" else "AIDA"
        send_plain(chat_id, "*REEL SCRIPT (" + str(reel_duration) + "s) — " + framework + "*\n\n" + result)
        reel_kb = format_action_keyboard("fmt_reel", "Reel Script")
        reel_kb.append([{"text": "Switch to " + other_fw, "callback_data": "social_switch_fw"}])
        send(chat_id, "Reel script ready.", reel_kb)
    except Exception as e:
        send(chat_id, "Error generating reel: " + str(e))

def gen_carousel(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "") or state.get("social_angle", "")
    slide_count = state.get("carousel_slides", 6)
    source_email = get_social_source_text(chat_id)[:500]
    hook = state.get("selected_social_hooks", {}).get("fmt_carousel", "")
    voice_examples = get_voice_corpus_context(chat_id)
    try:
        result = claude(
            (voice_examples + "\n\n" if voice_examples else "") +
            "Create a " + str(slide_count) + "-slide Instagram Carousel for Cryptonary.\n\n" +
            ("COVER SLIDE HEADLINE (use this): " + hook + "\n\n" if hook else "") +
            "SOURCE:\nReport: " + report + ("\nContext: " + context if context else "") +
            "\nAngle: " + angle + "\nEmail reference: " + source_email +
            "\n\nRULES:\n- Exactly " + str(slide_count) + " slides including cover and CTA final slide\n- Format: SLIDE N: [headline max 8 words] + [visual direction in brackets]\n- Mix bold text, data slides, list slides\n- Each slide earns the next swipe\n- Final slide: follow for more\n\nReturn as plain string.",
            max_tokens=900
        )
        state["current_social"] = result
        state["current_social_type"] = "Carousel"
        state["stage"] = "social_ready"
        send_plain(chat_id, "*CAROUSEL (" + str(slide_count) + " slides)*\n\n" + result)
        send(chat_id, "Carousel ready.", format_action_keyboard("fmt_carousel", "Carousel"))
    except Exception as e:
        send(chat_id, "Error generating carousel: " + str(e))


def gen_static(chat_id):
    """Generate a static Instagram post — limited on-screen text + full caption."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "") or state.get("social_angle", "")
    hook = state.get("selected_social_hooks", {}).get("fmt_static", "")
    source_email = get_social_source_text(chat_id)[:500]
    voice_examples = get_voice_corpus_context(chat_id)
    try:
        result = claude(
            (voice_examples + "\n\n" if voice_examples else "") +
            "Create an Instagram static post for Cryptonary.\n\n" +
            ("HOOK (use as the main on-screen text): " + hook + "\n\n" if hook else "") +
            "SOURCE:\nReport: " + report + ("\nContext: " + context if context else "") +
            "\nAngle: " + angle + "\nEmail reference: " + source_email +
            """

RULES:
ON-SCREEN TEXT (max 10 words total — this is what appears on the image):
- Headline: [bold, punchy, max 6 words — think newspaper front page]
- Subtext: [supporting line, max 4 words] (optional)
- Keep it minimal — most of the content lives in the caption

CAPTION (this is the full post text under the image):
- Open with a hook that expands on the headline
- 3-5 short punchy paragraphs
- Data, insight or story in the body
- End with a question or CTA to drive comments
- Adam's voice: direct, data-led, no fluff
- 150-250 words

Format EXACTLY:
ON-SCREEN TEXT:
Headline: [max 6 words]
Subtext: [max 4 words or leave blank]

CAPTION:
[full caption here]

Return as plain string.""",
            max_tokens=700
        )
        state["current_social"] = result
        state["current_social_type"] = "Static Post"
        state["stage"] = "social_ready"
        send_plain(chat_id, "*STATIC POST*\n\n" + result)
        send(chat_id, "Static post ready.", social_action_keyboard())
    except Exception as e:
        send(chat_id, "Error generating static post: " + str(e))

def gen_story(chat_id, multi=False):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    context = sanitise(state.get("context", ""))
    angle = state.get("selected_angle", "") or state.get("social_angle", "")
    story_slides = state.get("story_slides", 3) if multi else 1
    source_email = get_social_source_text(chat_id)[:500]
    fmt_key = "fmt_story_multi" if multi else "fmt_story_single"
    hook = state.get("selected_social_hooks", {}).get(fmt_key, "")
    slide_instruction = str(story_slides) + " slides" if multi else "1 single slide"
    voice_examples_story = get_voice_corpus_context(chat_id)
    try:
        result = claude(
            (voice_examples_story + "\n\n" if voice_examples_story else "") +
            "Create an Instagram Story for Cryptonary: " + slide_instruction + "\n\n" +
            ("OPENING SLIDE TEXT (use this as slide 1): " + hook + "\n\n" if hook else "") +
            "SOURCE:\nReport: " + report + ("\nContext: " + context if context else "") +
            "\nAngle: " + angle + "\nEmail reference: " + source_email +
            "\n\nRULES:\n- Format per slide: SLIDE N: [text max 15 words] + [visual/background direction] + [optional sticker]\n- Cliffhanger between slides if multi\n- Final slide: swipe up or link in bio\n- Urgent, direct tone\n\nReturn as plain string.",
            max_tokens=700
        )
        social_type = "Story (" + str(story_slides) + " slides)" if multi else "Story (single)"
        state["current_social"] = result
        state["current_social_type"] = social_type
        state["stage"] = "social_ready"
        send_plain(chat_id, "*" + social_type.upper() + "*\n\n" + result)
        send(chat_id, "Done.", social_action_keyboard())
    except Exception as e:
        send(chat_id, "Error generating story: " + str(e))

# ── LENGTH CONTROL ─────────────────────────────────────────────────

def show_length_menu(chat_id, mode="email"):
    user_state[chat_id]["length_mode"] = mode
    if mode == "email":
        keyboard = [
            [{"text": "Extend (add more detail)", "callback_data": "length_extend"}],
            [{"text": "Shorten (tighten it up)", "callback_data": "length_shorten"}],
            [{"text": "Cancel", "callback_data": "cancel_length"}]
        ]
        send(chat_id, "*Adjust length:*", keyboard)
    else:
        social_type = user_state[chat_id].get("current_social_type", "content")
        is_reel = "Reel" in social_type
        is_carousel = "Carousel" in social_type
        is_story = "Story" in social_type
        keyboard = []
        if is_reel:
            current = user_state[chat_id].get("reel_duration", 52)
            keyboard.append([{"text": "Extend to " + str(current+15) + "s", "callback_data": "reel_extend"}])
            keyboard.append([{"text": "Shorten to " + str(max(20, current-15)) + "s", "callback_data": "reel_shorten"}])
        elif is_carousel:
            current = user_state[chat_id].get("carousel_slides", 6)
            keyboard.append([{"text": "Add 2 more slides (" + str(current+2) + " total)", "callback_data": "carousel_extend"}])
            keyboard.append([{"text": "Remove 2 slides (" + str(max(3, current-2)) + " total)", "callback_data": "carousel_shorten"}])
        elif is_story:
            current = user_state[chat_id].get("story_slides", 3)
            social_type = user_state[chat_id].get("current_social_type", "")
            if "single" in social_type.lower() or current == 1:
                # Single slide — adjust text length not slide count
                keyboard.append([{"text": "Longer (add more detail)", "callback_data": "length_extend"}])
                keyboard.append([{"text": "Shorter (tighten text)", "callback_data": "length_shorten"}])
            else:
                keyboard.append([{"text": "Add a slide (" + str(current+1) + " total)", "callback_data": "story_extend"}])
                keyboard.append([{"text": "Remove a slide (" + str(max(1, current-1)) + " total)", "callback_data": "story_shorten"}])
        else:
            keyboard.append([{"text": "Extend", "callback_data": "length_extend"}])
            keyboard.append([{"text": "Shorten", "callback_data": "length_shorten"}])
        keyboard.append([{"text": "Cancel", "callback_data": "cancel_length"}])
        send(chat_id, "*Adjust length:*", keyboard)

def apply_length(chat_id, direction):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    mode = state.get("length_mode", "email")
    instruction = "Make this significantly longer. Add more detail, context, and depth. Keep Adam's voice." if direction == "extend" else "Make this significantly shorter. Cut anything that doesn't earn its place. Keep the core message and Adam's voice."
    if mode == "social":
        content = state.get("current_social", "")
        social_type = state.get("current_social_type", "content")
        send(chat_id, ("Extending" if direction == "extend" else "Shortening") + " " + social_type + "...")
        try:
            result = claude(instruction + "\n\nCONTENT:\n" + content + "\n\nReturn as plain string.", max_tokens=1200)
            state["current_social"] = result
            state["stage"] = "social_ready"
            send_plain(chat_id, social_type.upper() + " — " + direction.upper() + "ED\n\n" + result)
            send(chat_id, "Done.", social_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "social_ready"
    else:
        emails = state.get("current_emails", {})
        free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
        pro_email = extract_text(emails.get("pro", "")) if "pro" in emails else ""
        email_text = ""
        if free_email: email_text += "FREE EMAIL:\n" + free_email + "\n\n"
        if pro_email: email_text += "PRO EMAIL:\n" + pro_email
        send(chat_id, ("Extending" if direction == "extend" else "Shortening") + " emails...")
        try:
            raw = claude(
                instruction + "\n\nEMAILS:\n" + email_text +
                "\n\nReturn using this format. No JSON:\n\n===FREE EMAIL START===\n[adjusted free email]\n===FREE EMAIL END===\n\n===PRO EMAIL START===\n[adjusted pro email]\n===PRO EMAIL END===",
                max_tokens=2500
            )
            adjusted = parse_delimited_emails(raw)
            state["current_emails"] = adjusted
            state["stage"] = "emails_ready"
            if "free" in adjusted:
                send_plain(chat_id, "FREE EMAIL — " + direction.upper() + "ED\n\n" + extract_text(adjusted["free"]))
            if "pro" in adjusted:
                send_plain(chat_id, "PRO EMAIL — " + direction.upper() + "ED\n\n" + extract_text(adjusted["pro"]))
            send(chat_id, "Done.", email_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e))
            state["stage"] = "emails_ready"


# ── PERFORMANCE TRACKING ──────────────────────────────────────────

def start_log_performance(chat_id):
    user_state[chat_id]["stage"] = "log_subject"
    send(chat_id, "*Log Performance*\n\nEnter the subject line of the email you want to log:")

def log_step(chat_id, text):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    stage = state.get("stage")
    if stage == "log_subject":
        state["log_subject"] = text
        state["stage"] = "log_date"
        send(chat_id, "Date generated? (e.g. 2026-03-21)")
    elif stage == "log_date":
        state["log_date"] = text
        state["stage"] = "log_open"
        send(chat_id, "Open rate? (e.g. 28)")
    elif stage == "log_open":
        try:
            state["log_open"] = float(text)
            state["stage"] = "log_click"
            send(chat_id, "Click rate? (e.g. 4.2)")
        except:
            send(chat_id, "Please enter a number, e.g. 28")
    elif stage == "log_click":
        try:
            click = float(text)
            record = {
                "subject": state.get("log_subject", ""),
                "date": state.get("log_date", ""),
                "open_rate": state.get("log_open", 0),
                "click_rate": click
            }
            if chat_id not in performance_data:
                performance_data[chat_id] = []
            performance_data[chat_id].append(record)
            save_all_data()
            records = performance_data[chat_id]
            avg_open = sum(r.get("open_rate", 0) for r in records) / len(records)
            avg_click = sum(r.get("click_rate", 0) for r in records) / len(records)
            msg = "*Performance logged.*\n\n"
            msg += "Subject: _" + record["subject"] + "_\n"
            msg += "Open: " + str(record["open_rate"]) + "% | Click: " + str(click) + "%\n\n"
            msg += "Your averages (" + str(len(records)) + " emails tracked):\n"
            msg += "Open: " + str(round(avg_open, 1)) + "% | Click: " + str(round(avg_click, 1)) + "%"
            if len(records) >= 3:
                best = max(records, key=lambda r: r.get("open_rate", 0))
                msg += "\n\nBest subject: _" + best.get("subject", "") + "_\n(" + str(best.get("open_rate", 0)) + "% open)"
            state["stage"] = "idle"
            send(chat_id, msg, [[{"text": "New email set", "callback_data": "start_over"}]])
        except:
            send(chat_id, "Please enter a number, e.g. 4.2")

def show_stats(chat_id):
    records = performance_data.get(chat_id, [])
    if not records:
        send(chat_id, "No performance data yet. Use /logperformance after sending an email.")
        return
    avg_open = sum(r.get("open_rate", 0) for r in records) / len(records)
    avg_click = sum(r.get("click_rate", 0) for r in records) / len(records)
    best = max(records, key=lambda r: r.get("open_rate", 0))
    msg = "*Performance Summary*\n\n"
    msg += "Emails tracked: " + str(len(records)) + "\n"
    msg += "Avg open: " + str(round(avg_open, 1)) + "% | Avg click: " + str(round(avg_click, 1)) + "%\n\n"
    msg += "*Best subject:*\n_" + best.get("subject", "") + "_\n"
    msg += str(best.get("open_rate", 0)) + "% open / " + str(best.get("click_rate", 0)) + "% click\n\n"
    msg += "*Recent:*\n"
    for r in records[-5:]:
        msg += "- " + r.get("subject", "")[:45] + "\n  " + str(r.get("open_rate", "?")) + "% open / " + str(r.get("click_rate", "?")) + "% click\n"
    send(chat_id, msg)


# ── SUBJECT LINE A/B (MULTI-SELECT) ─────────────────────────────

def gen_subject_ab(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    hook = state.get("selected_hook", {})
    current_subject = hook.get("subject", "")
    current_preview = hook.get("preview", "")
    report = sanitise(state.get("report", ""))
    angle = state.get("selected_angle", "")
    send(chat_id, "Generating subject line alternatives...")
    try:
        raw = claude(
            "Generate 3 alternative subject lines for this Cryptonary email. Each uses a different copywriting principle.\n\nReport: " + report[:600] +
            "\nAngle: " + angle +
            "\nCurrent subject: " + current_subject +
            "\n\nPrinciples to use (one each): curiosity gap, fear with specific data, contrarian/counterintuitive.\n\nFormat each option as:\n1. PRINCIPLE: [principle name]\nSUBJECT: [subject line]\nPREVIEW: [preview text]\nREASON: [one sentence]\n\n2. PRINCIPLE: ...\nAnd so on for all 3. Nothing else.",
            max_tokens=700
        )
        alternatives = parse_subject_alternatives(raw)
        state["subject_alternatives"] = alternatives
        state["selected_subjects"] = []
        state["stage"] = "subject_select"
        text = "*Select subject lines to include:*\n_(tap to select multiple, then tap Done)_\n\n"
        text += "*Original:* " + current_subject + "\n_" + current_preview + "_\n\n"
        keyboard = []
        for i, a in enumerate(alternatives):
            text += "*" + str(i+1) + ". [" + a.get("principle","") + "]*\n"
            text += a.get("subject","") + "\n"
            text += "_" + a.get("preview","") + "_\n"
            text += a.get("reason","") + "\n\n"
            keyboard.append([{"text": "[ ] " + str(i+1) + ". " + a.get("subject","")[:40], "callback_data": "sel_sub_" + str(i)}])
        keyboard.append([{"text": "Done — apply selected", "callback_data": "apply_subjects"}])
        keyboard.append([{"text": "Keep original only", "callback_data": "keep_subject"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "emails_ready"

def toggle_subject_select(chat_id, idx, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_subjects", [])
    alternatives = state.get("subject_alternatives", [])
    if idx in selected:
        selected.remove(idx)
    else:
        selected.append(idx)
    state["selected_subjects"] = selected
    keyboard = []
    for i, a in enumerate(alternatives):
        is_sel = i in selected
        label = ("[x] " if is_sel else "[ ] ") + str(i+1) + ". " + a.get("subject","")[:40]
        keyboard.append([{"text": label, "callback_data": "sel_sub_" + str(i)}])
    keyboard.append([{"text": "Done — apply selected (" + str(len(selected)) + ")", "callback_data": "apply_subjects"}])
    keyboard.append([{"text": "Keep original only", "callback_data": "keep_subject"}])
    tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id, "reply_markup": {"inline_keyboard": keyboard}})

def apply_subjects(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected_ids = state.get("selected_subjects", [])
    alternatives = state.get("subject_alternatives", [])
    hook = state.get("selected_hook", {})
    original_subject = hook.get("subject", "")
    original_preview = hook.get("preview", "")
    emails = state.get("current_emails", {})

    # Build subject block with A/B/C labels
    subject_block = "Subject A: " + original_subject + "\nPreview A: " + original_preview
    labels = ["B", "C", "D"]
    for li, idx in enumerate(selected_ids):
        if idx < len(alternatives):
            a = alternatives[idx]
            label = labels[li] if li < len(labels) else str(li+2)
            subject_block += "\nSubject " + label + ": " + a.get("subject","")
            subject_block += "\nPreview " + label + ": " + a.get("preview","")

    # Update both emails to have the multi-subject block at the top
    for key in ["free", "pro"]:
        if key in emails:
            email_str = extract_text(emails[key])
            lines = email_str.split("\n")
            new_lines = []
            skip_next = False
            replaced = False
            for line in lines:
                if not replaced and (line.lower().startswith("subject") or line.lower().startswith("preview")):
                    if not replaced:
                        new_lines.append(subject_block)
                        replaced = True
                    if line.lower().startswith("preview"):
                        skip_next = False
                    continue
                new_lines.append(line)
            if not replaced:
                new_lines.insert(0, subject_block)
            emails[key] = "\n".join(new_lines)

    state["current_emails"] = emails
    state["stage"] = "emails_ready"
    count = len(selected_ids) + 1
    label = ["B","C","D"][min(len(selected_ids)-1,2)] if selected_ids else "A"
    send(chat_id, str(count) + " subject line(s) applied.")
    # Show updated emails immediately
    if "free" in emails:
        send_plain(chat_id, "*FREE EMAIL (updated subjects)*\n\n" + extract_text(emails["free"]))
    if "pro" in emails:
        send_plain(chat_id, "*PRO EMAIL (updated subjects)*\n\n" + extract_text(emails["pro"]))
    send(chat_id, "What next?", email_action_keyboard())


# ── TONE VARIANTS ─────────────────────────────────────────────────

TONE_DEFS = {
    "standard": ("Standard", "Adam's default voice. Confident, direct, data-led. Short punchy sentences. Balanced urgency."),
    "aggressive": ("Aggressive", "Bull market mode. High urgency, FOMO-driven. The opportunity is NOW. Stronger social proof. Reader will regret missing this. Shorter sentences. More bold statements. Commands not invitations."),
    "empathetic": ("Empathetic", "Bear market mode. Reader is stressed, possibly down on portfolio. Adam is alongside them. Acknowledges the pain. Positions Cryptonary as steady hand in chaos. Calmer tone. Less urgency, more reassurance. CTA is a lifeline not a push.")
}

def show_tone_menu(chat_id):
    keyboard = [
        [{"text": "Standard", "callback_data": "tone_standard"}],
        [{"text": "Aggressive (bull / FOMO)", "callback_data": "tone_aggressive"}],
        [{"text": "Empathetic (bear / stressed reader)", "callback_data": "tone_empathetic"}],
        [{"text": "Cancel", "callback_data": "cancel_tone"}]
    ]
    send(chat_id, "*Choose tone:*\n\nRewrites both emails in the selected tone. Same content, different emotional delivery.", keyboard)

def apply_tone(chat_id, tone_key):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    tone_label, tone_instruction = TONE_DEFS.get(tone_key, TONE_DEFS["standard"])
    emails = state.get("current_emails", {})
    free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
    pro_email = extract_text(emails.get("pro", "")) if "pro" in emails else ""
    email_text = ""
    if free_email: email_text += "FREE EMAIL:\n" + free_email + "\n\n"
    if pro_email: email_text += "PRO EMAIL:\n" + pro_email
    send(chat_id, "Rewriting in " + tone_label + " tone...")
    try:
        raw = claude(
            "Rewrite these Cryptonary emails in the following tone. Keep same content, facts, structure. Only change emotional delivery and language.\n\nTONE: " + tone_label +
            "\nINSTRUCTION: " + tone_instruction +
            "\n\nEMAILS:\n" + email_text +
            "\n\nRewrite and return the emails using this exact format. No JSON:\n\n===FREE EMAIL START===\n[rewritten free email]\n===FREE EMAIL END===\n\n===PRO EMAIL START===\n[rewritten pro email]\n===PRO EMAIL END===",
            max_tokens=2500
        )
        toned = parse_delimited_emails(raw)
        state["current_emails"] = toned
        state["stage"] = "emails_ready"
        if "free" in toned:
            send_plain(chat_id, "*FREE EMAIL — " + tone_label.upper() + "*\n\n" + extract_text(toned["free"]))
        if "pro" in toned:
            send_plain(chat_id, "*PRO EMAIL — " + tone_label.upper() + "*\n\n" + extract_text(toned["pro"]))
        send(chat_id, tone_label + " tone applied.", email_action_keyboard())
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "emails_ready"
        send(chat_id, "What next?", email_action_keyboard())

# ── AUDIENCE SEGMENTS ─────────────────────────────────────────────

SEG_DEFS = {
    "hot": ("Hot (last 30 days)", "This reader opens every email. Engaged, informed, trusts Cryptonary. Assume shared context. Reference recent calls briefly. Reward loyalty with insider-feeling language. Skip basic explanations. Treat as fellow serious investor. CTA feels like natural next step for someone already bought in."),
    "warm": ("Warm (31-90 days)", "This reader has drifted. Knows Cryptonary but hasn't engaged lately. Re-establish value without being preachy. Reference what they may have missed as opportunity cost not guilt. Make them feel the gap between where they are and where engaged members are. Stronger hook than usual to cut through inertia. CTA feels like re-engaging with something good."),
    "cold": ("Cold (90+ days / never opened)", "Almost a cold lead. May barely remember signing up. No assumed context or knowledge. Re-introduce Cryptonary value proposition clearly. Lead with strongest possible hook: a track record stat, specific win, market call that played out. Feels like a first impression. Lower assumed knowledge. CTA is low-friction and easy to say yes to.")
}

def gen_segments(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    emails = state.get("current_emails", {})
    free_email = extract_text(emails.get("free", "")) if "free" in emails else ""
    if not free_email:
        send(chat_id, "No free email found.", email_action_keyboard())
        return
    free_cta_key = state.get("free_cta", "upgrade")
    free_cta_instruction = CTA_INSTRUCTIONS["free"].get(free_cta_key, "")
    send(chat_id, "Writing 3 segment versions of the Free email...")
    try:
        raw = claude(
            "Rewrite this Cryptonary Free email for 3 different audience segments. Same core content and CTA. Different tone, assumed knowledge, emotional approach.\n\nORIGINAL FREE EMAIL:\n" + free_email +
            "\n\nCTA INSTRUCTION: " + free_cta_instruction +
            "\n\nSEGMENT HOT: " + SEG_DEFS["hot"][1] +
            "\n\nSEGMENT WARM: " + SEG_DEFS["warm"][1] +
            "\n\nSEGMENT COLD: " + SEG_DEFS["cold"][1] +
            "\n\nReturn the 3 segment emails using this exact format. No JSON:\n\n===HOT EMAIL START===\n[hot segment email]\n===HOT EMAIL END===\n\n===WARM EMAIL START===\n[warm segment email]\n===WARM EMAIL END===\n\n===COLD EMAIL START===\n[cold segment email]\n===COLD EMAIL END===",
            max_tokens=3000
        )
        segments = parse_delimited_segments(raw)
        state["current_segments"] = segments
        state["stage"] = "segments_ready"
        for key, (label, _) in SEG_DEFS.items():
            if key in segments:
                send_plain(chat_id, "*FREE EMAIL — " + label.upper() + "*\n\n" + extract_text(segments[key]))
        keyboard = [
            [{"text": "Edit Hot", "callback_data": "seg_edit_hot"},
             {"text": "Edit Warm", "callback_data": "seg_edit_warm"},
             {"text": "Edit Cold", "callback_data": "seg_edit_cold"}],
            [{"text": "Approve segments", "callback_data": "approve_segments"},
             {"text": "Back to emails", "callback_data": "back_to_emails"}]
        ]
        send(chat_id, "3 segment versions ready.", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "emails_ready"
        send(chat_id, "What next?", email_action_keyboard())

def ask_seg_edit(chat_id, segment):
    user_state[chat_id]["stage"] = "awaiting_seg_edit"
    user_state[chat_id]["seg_edit_target"] = segment
    label = SEG_DEFS.get(segment, ("",))[0]
    send(chat_id, "*Quick Edit — " + label + "*\n\nType your instruction:")

def apply_seg_edit(chat_id, instruction):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    segment = state.get("seg_edit_target", "hot")
    segments = state.get("current_segments", {})
    content = extract_text(segments.get(segment, ""))
    label = SEG_DEFS.get(segment, ("",))[0]
    send(chat_id, "Applying edit...")
    try:
        result = claude(
            "Edit this Cryptonary Free email (" + label + " segment). Instruction: " + instruction +
            "\n\nKeep segment tone intact. Only change what instruction specifies.\n\nEMAIL:\n" + content +
            "\n\nReturn edited email as a plain string.",
            max_tokens=1200
        )
        segments[segment] = result
        state["current_segments"] = segments
        state["stage"] = "segments_ready"
        send_plain(chat_id, "*EDITED — " + label.upper() + "*\n\n" + result)
        keyboard = [
            [{"text": "Edit Hot", "callback_data": "seg_edit_hot"},
             {"text": "Edit Warm", "callback_data": "seg_edit_warm"},
             {"text": "Edit Cold", "callback_data": "seg_edit_cold"}],
            [{"text": "Approve segments", "callback_data": "approve_segments"},
             {"text": "Back to emails", "callback_data": "back_to_emails"}]
        ]
        send(chat_id, "Edit applied.", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "segments_ready"

# ── MESSAGE + CALLBACK HANDLERS ───────────────────────────────────

def handle_message(msg):
    chat_id = msg["chat"]["id"]
    text = msg.get("text", "").strip()
    if not text: return

    if text == "/start":
        user_state[chat_id] = {"stage": "idle"}
        show_main_menu(chat_id)
        return

    # URL detection — drop any link, route based on current flow
    url_match = URL_PATTERN.search(text)
    if url_match and len(text.strip()) < 500 and text.strip().startswith("http"):
        url = url_match.group(0)
        user_state.setdefault(chat_id, {"stage": "idle"})
        current_stage = user_state[chat_id].get("stage", "idle")

        # EMAIL FLOW — if in email mode, fetch URL content and treat as report
        email_stages = {"awaiting_email_report", "awaiting_report", "buffering_report"}
        if current_stage in email_stages or user_state[chat_id].get("mode") == "email":
            send(chat_id, "Fetching content from link...")
            fetched = fetch_url_content(url, detect_url_type(url))
            if fetched and len(fetched.strip()) > 50:
                user_state[chat_id]["stage"] = "awaiting_context_choice"
                user_state[chat_id]["report"] = fetched[:4000]
                user_state[chat_id]["mode"] = "email"
                ask_context(chat_id)
            else:
                send(chat_id, "Could not read content from that link. Paste the text directly instead.")
            return

        # SOCIAL FLOW — if in social mode, fetch and use as report
        social_stages = {"awaiting_social_report"}
        if current_stage in social_stages:
            send(chat_id, "Fetching content from link...")
            fetched = fetch_url_content(url, detect_url_type(url))
            if fetched and len(fetched.strip()) > 50:
                user_state[chat_id]["report"] = fetched[:4000]
                show_standalone_social_menu(chat_id)
            else:
                send(chat_id, "Could not read content from that link. Paste the text directly instead.")
            return

        # IDEA ENGINE INSPIRATION FLOW — awaiting_inspiration handles URL directly
        if current_stage == "ie_awaiting_inspiration":
            analyse_url(chat_id, url, mode="ideas")
            return

        # DEFAULT — route to ideas
        mode = "critique" if current_stage == "ie_awaiting_screenshot_critique" else "ideas"
        user_state[chat_id]["stage"] = "idea_engine_idle"
        analyse_url(chat_id, url, mode=mode)
        return
    if text == "/stats":
        show_stats(chat_id)
        return
    if text == "/logperformance":
        user_state.setdefault(chat_id, {})
        start_log_performance(chat_id)
        return
    if text == "/logemail":
        user_state.setdefault(chat_id, {})
        start_log_email(chat_id)
        return
    if text == "/logad":
        user_state.setdefault(chat_id, {})
        start_log_ad(chat_id)
        return
    if text == "/emailreport":
        run_email_analysis(chat_id)
        return
    if text == "/adreport":
        run_ad_analysis(chat_id)
        return
    if text == "/briefing":
        user_state.setdefault(chat_id, {})
        # Check if there's a previous briefing to recap
        last_brief = user_state[chat_id].get("last_briefing_summary", "")
        last_brief_date = user_state[chat_id].get("last_briefing_date", "")
        if last_brief and last_brief_date:
            user_state[chat_id]["include_recap"] = True
        generate_briefing(chat_id)
        return

    if text == "/imageprompt" or text.startswith("/imageprompt "):
        user_state.setdefault(chat_id, {"stage": "idle"})
        brief = text.replace("/imageprompt", "").strip()
        if brief:
            generate_image_prompts(chat_id, brief)
        else:
            user_state[chat_id]["stage"] = "awaiting_image_brief"
            send(chat_id, "*Image Prompt Generator*\n\nDescribe what you need.\n\n_e.g. Hero image for Inner Circle landing page, premium dark crypto wealth theme_\n_e.g. Static ad for Passive Income avatar, warm lifestyle feel_\n_e.g. Carousel cover slide, Bitcoin price breakout energy_")
        return

    if text == "/help":
        send(chat_id, "*Writing Studio V9*\n\nFrom /start choose: Emails, Ads, or Social\n\n*Email commands:*\n/logemail — log open rate + CTR\n/emailreport — analyse all logged emails\n\n*Ad commands:*\n/logad — log video or static ad results\n/adreport — analyse all logged ads\n\n*Legacy:*\n/logperformance — old performance log\n/stats — old stats summary\n\n/start — return to main menu")
        return

    stage = user_state.get(chat_id, {}).get("stage", "idle")

    if stage in ["log_subject", "log_date", "log_open", "log_click"]:
        log_step(chat_id, text)
        return

    if stage == "awaiting_context_text":
        user_state[chat_id]["context"] = text
        user_state[chat_id]["stage"] = "pick_angle"
        send(chat_id, "Context added.")
        gen_angles(chat_id)
        return

    if stage in ("awaiting_quick_edit", "social_quick_edit",
                  "social_ready", "social_approved",
                  "emails_ready", "emails_approved"):
        # Voice/text edit arriving at active content stage — treat as quick edit
        if stage in ("social_ready", "social_approved"):
            user_state[chat_id]["pre_edit_stage"] = stage
            user_state[chat_id]["quick_edit_mode"] = "social"
        elif stage in ("emails_ready", "emails_approved"):
            user_state[chat_id]["pre_edit_stage"] = stage
            user_state[chat_id]["quick_edit_mode"] = "email"
        user_state[chat_id]["stage"] = user_state[chat_id].get("pre_edit_stage", stage)
        apply_quick_edit(chat_id, text)
        return

    if stage == "awaiting_seg_edit":
        apply_seg_edit(chat_id, text)
        return

    if stage == "logging_email":
        handle_email_log_step(chat_id, text)
        return

    if stage == "logging_ad":
        handle_ad_log_step(chat_id, text)
        return

    if stage == "lp_outline_review":
        # User sent numbered feedback for outline
        apply_lp_outline_feedback(chat_id, text)
        return

    if stage == "lp_awaiting_paste_back":
        # User pasted their edited version
        state["lp_pre_edit"] = state.get("current_lp", "")
        state["current_lp"] = text
        state["stage"] = "lp_copy_review"
        send_plain(chat_id, "*UPDATED VERSION RECEIVED*\n\n" + text[:2000])
        keyboard = [
            [{"text": "Quick Edit", "callback_data": "lp_quick_edit"},
             {"text": "Approve — Generate Design Brief", "callback_data": "lp_approve_copy"}],
            [{"text": "Critique", "callback_data": "critique_lp"},
             {"text": "Revert", "callback_data": "lp_revert"}]
        ]
        send(chat_id, "Paste-back received and saved as new version.", keyboard)
        return

    if stage == "lp_awaiting_length_instruction":
        state["stage"] = "lp_copy_review"
        lp_content = state.get("current_lp", "")
        send(chat_id, "Adjusting length...")
        try:
            result = claude(
                "INSTRUCTION: " + text + "\n\nLANDING PAGE:\n" + lp_content[:4000] +
                "\n\nApply the length adjustment. Return the full revised page.",
                max_tokens=3000, system=BRANDSCRIPT_PROMPT
            )
            state["lp_pre_edit"] = lp_content
            state["current_lp"] = result
            send_plain(chat_id, result[:3000])
            keyboard = [
                [{"text": "Approve — Generate Design Brief", "callback_data": "lp_approve_copy"}],
                [{"text": "More adjustments", "callback_data": "lp_length"},
                 {"text": "Revert", "callback_data": "lp_revert"}]
            ]
            send(chat_id, "Length adjusted.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))
        return

    if stage == "awaiting_additional_context":
        existing_context = user_state[chat_id].get("context", "")
        new_context = (existing_context + " " + text).strip() if existing_context else text
        user_state[chat_id]["context"] = new_context
        send(chat_id, "Context added. Regenerating emails with: _" + text[:100] + "_")
        gen_emails(chat_id)
        return

    if stage == "ie_awaiting_custom_concept":
        user_state[chat_id]["stage"] = "ie_angle_review"
        generate_ie_angle_from_concept(chat_id, text)
        return

    if stage == "ie_awaiting_custom_angle_ie":
        user_state[chat_id]["ie_selected_angle"] = text
        generate_ie_hook_from_angle(chat_id, text)
        return

    if stage == "ie_awaiting_custom_hook_ie":
        user_state[chat_id]["ie_selected_hook"] = text
        generate_ie_final_content(chat_id, text)
        return

    if stage == "ie_awaiting_content_edit":
        content_val = user_state[chat_id].get("ie_generated_content", "")
        user_state[chat_id]["ie_pre_edit_content"] = content_val
        send(chat_id, "Applying edit...")
        try:
            result = claude(
                "INSTRUCTION: " + text +
                "\n\nCONTENT:\n" + content_val +
                "\n\nApply the instruction. Return the full revised content only.",
                max_tokens=1500, system=VOICE_GUIDE
            )
            user_state[chat_id]["ie_generated_content"] = result
            user_state[chat_id]["stage"] = "ie_content_ready"
            before = content_val
            send_plain(chat_id, result)
            show_diff_with_revert(chat_id, before, result, "ie")
            keyboard = [
                [{"text": "Another edit", "callback_data": "ie_content_edit"},
                 {"text": "Develop in Content Studio", "callback_data": "ie_to_content_studio"}],
                [{"text": "Start over", "callback_data": "open_idea_engine"}]
            ]
            send(chat_id, "Edit applied.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))
        return

    if stage == "ie_awaiting_link_ideas":
        import re as _re
        url_match = _re.search(r"https?://\S+", text)
        if url_match:
            user_state[chat_id]["stage"] = "idea_engine_idle"
            show_ie_format_menu(chat_id)
            user_state[chat_id]["stage"] = "ie_pending_url"
            user_state[chat_id]["ie_pending_url"] = url_match.group(0)
            show_ie_format_menu(chat_id)
        else:
            send(chat_id, "Please paste a valid URL starting with https://")
        return

    if stage == "ie_awaiting_inspiration":
        # User dropped a brief, article, or URL into the inspiration flow
        import re as _rei
        url_match = _rei.search(r"https?://\S+", text)
        if url_match:
            # It's a link — use analyse_url
            url = url_match.group(0)
            user_state[chat_id]["ie_idea_type"] = "instagram"
            analyse_url(chat_id, url, mode="ideas")
        elif len(text.strip()) < 200:
            # Short message = a brief. Use it as direct creative direction.
            user_state[chat_id]["ie_source_content"] = ""
            user_state[chat_id]["ie_source_label"] = "Brief: " + text.strip()[:60]
            user_state[chat_id]["ie_idea_type"] = "instagram"
            # Inject the brief as source content so generate_ie_concept picks it up
            user_state[chat_id]["ie_source_content"] = "CREATIVE BRIEF: " + text.strip()
            generate_ie_concept(chat_id)
        else:
            # Long text = reference material. Analyse and extract concepts.
            user_state[chat_id]["ie_idea_type"] = "instagram"
            _process_ie_text_content(chat_id, stage, text, "Pasted content", mode="ideas")
        return

    if stage == "ie_awaiting_pasted_text":
        # User pasted text from a platform we couldn't scrape (Twitter, Instagram, TikTok, etc.)
        mode = user_state[chat_id].get("pending_url_mode", "ideas")
        platform = user_state[chat_id].get("pending_url_platform", "Pasted content")
        user_state[chat_id]["stage"] = "idea_engine_idle"
        _process_ie_text_content(chat_id, stage, text, platform, mode=mode)
        return

    if stage == "ie_pending_url":
        # Format was selected, now fetch the URL and generate concepts
        url = user_state[chat_id].get("ie_pending_url", "")
        if url:
            fetched = fetch_url_content(url, detect_url_type(url))
            user_state[chat_id]["ie_source_content"] = fetched[:2000] if fetched else ""
            user_state[chat_id]["ie_source_label"] = url[:60]
        generate_ie_concept(chat_id)
        return

    if stage == "awaiting_critique_apply":
        apply_critique_fix(chat_id, text.strip())
        return

    if stage == "awaiting_custom_free_angle":
        user_state[chat_id]["selected_free_angle"] = text
        user_state[chat_id]["selected_angle"] = text
        show_pro_angle_picker(chat_id)
        return

    if stage == "awaiting_custom_pro_angle":
        user_state[chat_id]["selected_pro_angle"] = text
        gen_hooks(chat_id)
        return

    if stage == "awaiting_custom_free_hook":
        parts = text.split("|")
        hook = {"subject": parts[0].strip(), "preview": parts[1].strip() if len(parts) > 1 else ""}
        user_state[chat_id]["selected_free_hook"] = hook
        user_state[chat_id]["selected_hook"] = hook
        gen_pro_hooks(chat_id)
        return

    if stage == "awaiting_custom_pro_hook":
        parts = text.split("|")
        hook = {"subject": parts[0].strip(), "preview": parts[1].strip() if len(parts) > 1 else ""}
        user_state[chat_id]["selected_pro_hook"] = hook
        user_state[chat_id]["stage"] = "pick_free_cta"
        ask_free_cta(chat_id)
        return

    if stage == "awaiting_custom_angle":
        user_state[chat_id]["selected_angle"] = text
        user_state[chat_id]["selected_free_angle"] = text
        show_pro_angle_picker(chat_id)
        return

    if stage == "awaiting_custom_hook":
        # Parse "subject | preview" or just subject
        parts = text.split("|")
        subject = parts[0].strip()
        preview = parts[1].strip() if len(parts) > 1 else ""
        hook = {"subject": subject, "preview": preview, "hook_a": "", "hook_b": "", "hook_c": ""}
        user_state[chat_id]["selected_hook"] = hook
        user_state[chat_id]["stage"] = "pick_free_cta"
        ask_free_cta(chat_id)
        return

    if stage == "awaiting_custom_social_angle":
        user_state[chat_id]["social_angle"] = text
        user_state[chat_id]["stage"] = "pick_social_formats"
        user_state[chat_id]["selected_social_formats"] = []
        show_standalone_social_menu(chat_id)
        return

    if stage == "awaiting_ad_existing_upload":
        # They pasted existing ad copy as text
        user_state[chat_id]["existing_ad_content"] = text
        user_state[chat_id]["stage"] = "pick_existing_ad_action"
        show_existing_ad_action_menu(chat_id)
        return

    if stage == "awaiting_ad_theme":
        product_context = user_state[chat_id].get("ad_product_context", "")
        theme = text
        if product_context:
            theme = "PRODUCT: " + product_context + "\n\nTHEME/CONTEXT: " + text
        user_state[chat_id]["ad_theme"] = theme
        user_state[chat_id]["stage"] = "pick_ad_avatars"
        show_avatar_menu(chat_id)
        return

    if stage == "awaiting_image_brief":
        user_state[chat_id]["stage"] = "idle"
        generate_image_prompts(chat_id, text)
        return

    if stage.startswith("ie_awaiting_source_"):
        source_type = stage.replace("ie_awaiting_source_", "")
        sources = load_idea_sources()
        key_map = {"instagram": "instagram", "twitter": "twitter",
                   "facebook": "facebook_pages", "telegram": "telegram", "reddit": "reddit"}
        key = key_map.get(source_type)
        if key:
            if key not in sources: sources[key] = []
            val = text.strip().lstrip("@").lstrip("/")
            if val not in sources[key]:
                sources[key].append(val)
                save_idea_sources(sources)
                send(chat_id, "Added " + val + " to " + source_type + " sources.")
            else:
                send(chat_id, val + " is already in your sources.")
        user_state[chat_id]["stage"] = "idea_engine_idle"
        show_ie_manage_sources(chat_id)
        return

    if stage == "ie_awaiting_develop":
        idea_text = text
        last_ideas = user_state[chat_id].get("last_ideas", "")
        user_state[chat_id]["stage"] = "awaiting_social_report"
        user_state[chat_id]["report"] = "IDEA TO DEVELOP:\n" + idea_text + "\n\nFULL IDEAS CONTEXT:\n" + last_ideas[:500]
        user_state[chat_id]["social_angle"] = idea_text[:100]
        user_state[chat_id]["social_origin"] = "idea_engine"
        show_standalone_social_menu(chat_id)
        return

    if stage == "ie_awaiting_develop_ad":
        idea_text = text
        last_ideas = user_state[chat_id].get("last_ideas", "")
        user_state[chat_id]["stage"] = "pick_ad_avatars"
        user_state[chat_id]["ad_theme"] = idea_text
        user_state[chat_id]["report"] = last_ideas[:500]
        show_avatar_menu(chat_id)
        return

    if stage == "ds_awaiting_email_splitvar":
        user_state[chat_id]["ds_split_var"] = text
        user_state[chat_id]["stage"] = "ds_awaiting_email_split_data"
        user_state[chat_id]["ds_images"] = []
        keyboard = [
            [{"text": "Done — analyse now", "callback_data": "ds_analyse_emails_split"}],
            [{"text": "Type numbers manually", "callback_data": "ds_email_type_numbers"}]
        ]
        send(chat_id, "Got it. Upload screenshots/CSV or type raw numbers directly (e.g. \'Variant A: 5000 sent, 1200 opens, 25 clicks\').", keyboard)
        return

    if stage == "ds_awaiting_email_split_numbers":
        # User typed raw numbers — use as CSV text
        user_state[chat_id]["ds_csv_text"] = text
        analyse_emails(chat_id, split_var=user_state[chat_id].get("ds_split_var"))
        return

    if stage == "ds_awaiting_landing_splitvar":
        user_state[chat_id]["ds_split_var"] = text
        user_state[chat_id]["stage"] = "ds_awaiting_landing_split_data"
        user_state[chat_id]["ds_images"] = []
        keyboard = [[{"text": "Done — analyse now", "callback_data": "ds_analyse_landing_split"}]]
        send(chat_id, "Got it. Now upload your landing page screenshots or CSV.", keyboard)
        return

    if stage in ("ds_awaiting_followup", "ds_analysis_done") and user_state.get(chat_id, {}).get("last_ds_analysis"):
        last_analysis = user_state[chat_id].get("last_ds_analysis", "")
        user_state[chat_id]["stage"] = "ds_analysis_done"
        send(chat_id, "Working on it...")
        try:
            result = claude(
                "Previous analysis:\n" + last_analysis[:3000] +
                "\n\nFollow-up question: " + text +
                "\n\nAnswer specifically based on the data already analysed. Be direct and actionable.",
                max_tokens=1500,
                system=DATA_STUDIO_SYSTEM
            )
            send_plain(chat_id, result)
            keyboard = [
                [{"text": "Ask another follow-up", "callback_data": "ds_followup"}],
                [{"text": "Back to Data Studio", "callback_data": "open_data_studio"}]
            ]
            send(chat_id, "Done.", keyboard)
        except Exception as e:
            print("Followup error:", e, flush=True)
            print(traceback.format_exc(), flush=True)
            send(chat_id, "Error in follow-up: " + str(e))
        return

    if stage == "awaiting_lp_context_text":
        state["lp_context"] = text
        generate_lp_outline(chat_id)
        return

    if stage == "awaiting_lp_quick_edit":
        instruction = text
        lp_content = state.get("current_lp", "")
        state["stage"] = "lp_ready"
        send(chat_id, "Applying edit...")
        try:
            result = claude(
                "Edit this Cryptonary landing page section. Instruction: " + instruction +
                "\n\nApply BrandScript framework and copywriting principles. Keep all other sections intact. Only change what the instruction specifies.\n\nFULL PAGE:\n" + lp_content[:3000] +
                "\n\nReturn only the edited section(s) as plain text.",
                max_tokens=1500,
                system=BRANDSCRIPT_PROMPT
            )
            send_plain(chat_id, result)
            keyboard = [
                [{"text": "Quick Edit", "callback_data": "lp_quick_edit"},
                 {"text": "Regenerate section", "callback_data": "lp_regen"}],
                [{"text": "Critique", "callback_data": "critique_lp"},
                 {"text": "Mark Complete", "callback_data": "mark_complete"}],
                [{"text": "Generate another page", "callback_data": "lp_again"}]
            ]
            send(chat_id, "Edit applied.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))
        return

    if stage == "awaiting_lp_regen_instruction":
        instruction = text.strip()
        section = state.get("lp_regen_section", "Hero")
        cta_key = state.get("lp_cta", "pro")
        cta = LP_CTA_DEFS.get(cta_key, LP_CTA_DEFS["pro"])
        avatar_keys = state.get("selected_avatars", [])
        avatar_descs = "\n".join([k + ": " + AVATARS_AD.get(k, ("",))[1] for k in avatar_keys])
        context = sanitise(state.get("lp_context", ""))
        state["stage"] = "lp_ready"
        send(chat_id, "Regenerating " + section + "...")
        try:
            prompt = BRANDSCRIPT_PROMPT
            prompt += "\n\nAVATAR(S): " + avatar_descs
            prompt += "\nCTA: " + cta["label"] + " — " + cta["price"]
            prompt += "\nPOSITIONING: " + cta["positioning"]
            if context: prompt += "\nCONTEXT: " + context
            if instruction: prompt += "\nSPECIFIC DIRECTION: " + instruction
            prompt += "\n\nREGENERATE ONLY: " + section.upper() + " SECTION. Return as plain text with graphic recommendations."
            result = claude(prompt, max_tokens=1200, system=BRANDSCRIPT_PROMPT)
            send_plain(chat_id, result)
            keyboard = [
                [{"text": "Quick Edit", "callback_data": "lp_quick_edit"},
                 {"text": "Regenerate section", "callback_data": "lp_regen"}],
                [{"text": "Critique", "callback_data": "critique_lp"},
                 {"text": "Mark Complete", "callback_data": "mark_complete"}],
                [{"text": "Generate another page", "callback_data": "lp_again"}]
            ]
            send(chat_id, "Section regenerated.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))
        return

    if stage == "awaiting_ad_quick_edit":
        instruction = text
        ad_content = state.get("current_ad", "")
        label = state.get("current_ad_label", "AD")
        state["stage"] = "ads_ready"
        send(chat_id, "Applying edit...")
        try:
            result = claude("Edit this ad copy. Instruction: " + instruction + "\n\nKeep the same structure and format. Only change what specified.\n\nAD:\n" + ad_content + "\n\nReturn as plain string.", max_tokens=1500)
            state["current_ad"] = result
            send_plain(chat_id, label + result, ad_action_keyboard())
        except Exception as e:
            send(chat_id, "Error: " + str(e), ad_action_keyboard())
        return

    if stage == "voice_awaiting_edit":
        # User is adding detail to their voice brief — re-interpret with the extra context
        original = user_state[chat_id].get("voice_transcript", "")
        combined = original + ". Additional detail: " + text
        user_state[chat_id]["voice_transcript"] = combined
        user_state[chat_id]["stage"] = "idle"
        # Re-run interpretation with the enriched brief
        try:
            interpretation = claude(
                "A user sent this voice brief to a crypto content bot:\n\n" +
                "BRIEF: \"" + combined[:400] + "\"\n\n" +
                "Extract their intent. Reply EXACTLY:\n\n" +
                "CONTENT_TYPE: [email/instagram_post/reel/carousel/story/ad/idea/general]\n" +
                "TOPIC: [topic, 1 sentence]\n" +
                "ANGLE: [angle, 1 sentence]\n" +
                "FORMAT_HINT: [format clues]\n" +
                "READY_TO_GENERATE: [yes/no]",
                max_tokens=300
            )
            lines = {l.split(":")[0].strip(): ":".join(l.split(":")[1:]).strip()
                     for l in interpretation.strip().splitlines() if ":" in l}
            user_state[chat_id]["voice_content_type"] = lines.get("CONTENT_TYPE", "general")
            user_state[chat_id]["voice_topic"] = lines.get("TOPIC", combined[:100])
            user_state[chat_id]["voice_angle"] = lines.get("ANGLE", "")
            summary = "*Updated brief:*\n\n"
            summary += "📌 *Content:* " + user_state[chat_id]["voice_content_type"].replace("_", " ").title() + "\n"
            summary += "📰 *Topic:* " + user_state[chat_id]["voice_topic"] + "\n"
            if user_state[chat_id]["voice_angle"]:
                summary += "🎯 *Angle:* " + user_state[chat_id]["voice_angle"] + "\n"
            keyboard = [
                [{"text": "✅ Generate this", "callback_data": "voice_confirm"}],
                [{"text": "✏️ Change more",   "callback_data": "voice_edit"}],
                [{"text": "❌ Discard",        "callback_data": "voice_discard"}],
            ]
            send(chat_id, summary, keyboard)
        except Exception as e:
            send(chat_id, "Error re-interpreting: " + str(e))
        return

    if stage == "img_awaiting_brief":
        user_state[chat_id]["pending_img_concept"] = text
        user_state[chat_id]["pending_img_angle"] = ""
        show_image_type_menu(chat_id)
        return

    if stage == "vb_awaiting_edit":
        # User is editing the visual brief text, not an image
        brief = _global_brief_store.get(chat_id, "")
        vb_type = user_state[chat_id].get("last_visual_type", "static")
        user_state[chat_id]["stage"] = "idle"
        if not brief:
            send(chat_id, "No brief to edit. Generate a visual brief first.")
            return
        send(chat_id, "Updating brief...")
        try:
            updated = claude(
                "CURRENT BRIEF:\n" + brief[:1500] +
                "\n\nUSER FEEDBACK: " + text +
                "\n\nUpdate the brief applying this feedback. Keep the same format and structure. Only change what was specified.",
                max_tokens=900
            )
            _global_brief_store[chat_id] = updated
            user_state[chat_id]["last_visual_brief"] = updated
            user_state[chat_id]["last_visual_type"] = vb_type
            send_plain(chat_id, "*UPDATED BRIEF*\n\n" + updated)
            img_row = [{"text": "🎨 Generate image", "callback_data": "img_from_brief"}] if (OPENAI_KEY or GEMINI_KEY) else []
            keyboard = []
            if img_row:
                keyboard.append(img_row)
            keyboard.append([{"text": "✏️ Adjust more", "callback_data": "vb_edit"}])
            keyboard.append([{"text": "✅ Done", "callback_data": "mark_complete"}])
            send(chat_id, "Brief updated.", keyboard)
        except Exception as e:
            send(chat_id, "Error updating brief: " + str(e))
        return

    if stage == "img_awaiting_direction":
        user_state[chat_id]["stage"] = "idle"
        handle_image_direction(chat_id, text)
        return

    if stage == "awaiting_social_report":
        user_state[chat_id]["report"] = text
        user_state[chat_id]["context"] = ""
        # Format selection first (moved before angles)
        show_standalone_social_menu(chat_id)
        return

    if len(text) > 100:
        current_stage = user_state.get(chat_id, {}).get("stage", "idle")

        # Already buffering — append and extend the window, never trigger twice
        if current_stage == "buffering_report":
            user_state[chat_id]["report_buffer"] += "\n" + text
            user_state[chat_id]["buffer_timer"] = time.time()
            return

        # Suppress stale chunks arriving after buffer already resolved
        if current_stage in ["awaiting_context_choice", "pick_angle", "pick_free_hook",
                              "pick_pro_hook", "social_ready", "emails_ready"]:
            return

        # EMAIL REPORT — buffer and wait for all chunks
        if current_stage == "awaiting_email_report":
            user_state[chat_id] = {"stage": "buffering_report", "report_buffer": text,
                                    "buffer_timer": time.time(), "mode": "email"}
            time.sleep(2)
            if user_state.get(chat_id, {}).get("stage") == "buffering_report":
                full_report = user_state[chat_id].get("report_buffer", text)
                user_state[chat_id] = {"stage": "awaiting_context_choice", "report": full_report, "mode": "email"}
                ask_context(chat_id)
            return

        # SOCIAL REPORT — buffer and wait
        if current_stage == "awaiting_social_report":
            user_state[chat_id]["report"] = (user_state[chat_id].get("report", "") + "\n" + text).strip()
            # Don't trigger again — awaiting_social_report already handled above
            return

        # DEFAULT fallback — treat as email report
        if current_stage in ["idle", "buffering_report"]:
            return
        user_state[chat_id] = {"stage": "buffering_report", "report_buffer": text,
                                "buffer_timer": time.time(), "mode": "email"}
        time.sleep(2)
        if user_state.get(chat_id, {}).get("stage") == "buffering_report":
            full_report = user_state[chat_id].get("report_buffer", text)
            user_state[chat_id] = {"stage": "awaiting_context_choice", "report": full_report, "mode": "email"}
            ask_context(chat_id)
    else:
        stage = user_state.get(chat_id, {}).get("stage", "idle")
        if stage == "idle":
            show_main_menu(chat_id)
        # Silently ignore short text in active flows — don't reset or confuse the user

def handle_callback(cb):
    chat_id = cb["message"]["chat"]["id"]
    message_id = cb["message"]["message_id"]
    data = cb["data"]
    tg("answerCallbackQuery", {"callback_query_id": cb["id"]})
    state = user_state.get(chat_id, {})

    # Cooldown: ignore duplicate callbacks fired within 1.5s (queued updates / double-tap)
    cb_key = str(chat_id) + ":" + data
    now = time.time()
    last = _last_callback_time.get(cb_key, 0)
    # Never cooldown these — they need to fire immediately every time
    no_cooldown = {
        "mark_complete", "start_over", "vb_auto", "img_from_brief",
        "img_regen", "img_direction", "img_restyle",
        "approve_social", "approve_emails", "approve_ad",
        "gen_social_selected", "gen_social_confirmed",
        "voice_confirm", "voice_discard",
    }
    if now - last < 1.5 and data not in no_cooldown and not data.startswith("img_engine_") and not data.startswith("img_style_") and not data.startswith("sfmt_approve_"):
        return
    _last_callback_time[cb_key] = now
    # Clean old entries periodically
    if len(_last_callback_time) > 1000:
        _last_callback_time.clear()

    if data == "start_over":
        show_main_menu(chat_id)

    elif data == "open_content_studio":
        keyboard = [
            [{"text": "Emails", "callback_data": "mode_email"}],
            [{"text": "Ad Copy", "callback_data": "mode_ads"}],
            [{"text": "Social Content", "callback_data": "mode_social"}],
            [{"text": "Landing Page", "callback_data": "mode_landing"}]
        ]
        send(chat_id, "*Writing Studio*\n\nWhat do you want to create?", keyboard)

    elif data == "open_data_studio":
        show_data_studio_menu(chat_id)

    elif data == "mode_email":
        user_state[chat_id] = {"stage": "awaiting_email_report"}
        send(chat_id, "Paste your report:")

    elif data == "mode_ads":
        user_state[chat_id] = {"stage": "awaiting_ad_theme", "selected_avatars": [], "selected_stages": []}
        send(chat_id, "*Ad Creation*\n\nPaste your campaign theme, report, or context:\n\n_Examples: Bitcoin halving setup, inner circle launch, market crash opportunity, weekly market update..._")

    elif data == "mode_social":
        user_state[chat_id] = {"stage": "awaiting_social_report", "selected_social_formats": []}
        send(chat_id, "Paste your report or content to base social posts on:")

    elif data == "mode_landing":
        user_state[chat_id] = {"stage": "lp_idle", "selected_avatars": [], "lp_outline": {}, "lp_full_copy": {}}
        start_landing_page_flow(chat_id)

    elif data.startswith("adavatar_"):
        avatar_key = data.replace("adavatar_", "")
        toggle_avatar(chat_id, avatar_key, message_id)

    elif data == "adavatars_done":
        selected = state.get("selected_avatars", [])
        if not selected:
            send(chat_id, "Please select at least one avatar.")
        else:
            state["selected_stages"] = []
            show_stage_menu(chat_id)

    elif data.startswith("adstage_"):
        stage_key = data.replace("adstage_", "")
        toggle_stage(chat_id, stage_key, message_id)

    elif data == "adstages_done":
        selected = state.get("selected_stages", [])
        if not selected:
            send(chat_id, "Please select at least one funnel stage.")
        else:
            show_adtype_menu(chat_id)

    elif data == "adtype_static":
        state["ad_type"] = "static"
        generate_all_ads(chat_id)

    elif data == "adtype_video":
        state["ad_type"] = "video"
        generate_all_ads(chat_id)

    elif data == "ad_quick_edit":
        state["stage"] = "awaiting_ad_quick_edit"
        send(chat_id, "*Quick Edit — Ad Copy*\n\nType your instruction:")

    elif data == "ad_enhance":
        ad_content = state.get("current_ad", "")
        if not ad_content:
            send(chat_id, "No ad content found.", ad_action_keyboard())
        else:
            send(chat_id, "Analysing ad copy...")
            try:
                prompt = "Analyse this Cryptonary ad copy and give 6 specific improvements. Apply Hormozi, Cashvertising LF8, Cialdini, Ogilvy principles.\n\nAD:\n" + ad_content + "\n\nFormat:\n1. PRINCIPLE: [name]\nISSUE: [what is weak]\nFIX: [specific fix]\n\nAnd so on for all 6."
                raw = claude(prompt, max_tokens=1000)
                send_plain(chat_id, "AD ENHANCEMENT SUGGESTIONS\n\n" + raw, ad_action_keyboard())
            except Exception as e:
                send(chat_id, "Error: " + str(e), ad_action_keyboard())

    elif data == "avatarpage_next":
        page = state.get("avatar_page", 0)
        show_avatar_menu(chat_id, page + 1)

    elif data == "avatarpage_prev":
        page = state.get("avatar_page", 0)
        show_avatar_menu(chat_id, max(0, page - 1))

    elif data == "ads_again":
        state["selected_avatars"] = []
        state["selected_stages"] = []
        show_avatar_menu(chat_id)

    elif data.startswith("lpavatar_"):
        avatar_key = data.replace("lpavatar_", "")
        toggle_lp_avatar(chat_id, avatar_key, message_id)

    elif data == "lpavatarpage_next":
        page = state.get("avatar_page", 0)
        show_lp_avatar_menu(chat_id, page + 1)

    elif data == "lpavatarpage_prev":
        page = state.get("avatar_page", 0)
        show_lp_avatar_menu(chat_id, max(0, page - 1))

    elif data == "lpavatars_done":
        selected = state.get("selected_avatars", [])
        if not selected:
            send(chat_id, "Please select at least one avatar.")
        else:
            state["stage"] = "awaiting_lp_context"
            keyboard = [
                [{"text": "Yes — add context", "callback_data": "lp_context_yes"}],
                [{"text": "No — generate outline now", "callback_data": "lp_context_no"}]
            ]
            send(chat_id, "*Any extra context?*\n\nCurrent offer, promo, urgency mechanic, specific hook, seasonal angle...", keyboard)

    elif data == "lpcta_pro":
        state["lp_cta"] = "pro"
        state["stage"] = "awaiting_lp_context"
        keyboard = [
            [{"text": "Yes — add context", "callback_data": "lp_context_yes"}],
            [{"text": "No — generate outline now", "callback_data": "lp_context_no"}]
        ]
        send(chat_id, "*Any extra context?*\n\nCurrent offer, promo, urgency mechanic, specific hook, seasonal angle...", keyboard)

    elif data == "lpcta_inner_circle":
        state["lp_cta"] = "inner_circle"
        state["stage"] = "awaiting_lp_context"
        keyboard = [
            [{"text": "Yes — add context", "callback_data": "lp_context_yes"}],
            [{"text": "No — generate outline now", "callback_data": "lp_context_no"}]
        ]
        send(chat_id, "*Any extra context?*\n\nCurrent offer, promo, urgency mechanic, specific hook, seasonal angle...", keyboard)

    elif data == "lp_context_yes":
        state["stage"] = "awaiting_lp_context_text"
        send(chat_id, "Type your extra context:")

    elif data == "lp_context_no":
        state["lp_context"] = ""
        generate_lp_outline(chat_id)

    elif data == "lp_quick_edit":
        state["stage"] = "awaiting_lp_quick_edit"
        send(chat_id, "*Quick Edit — Landing Page*\n\nWhich section and what to change?\n_e.g. Rewrite the hero headline / Make the value stack more specific / Shorten the FAQ_")

    elif data == "lp_regen":
        state["stage"] = "awaiting_lp_regen"
        keyboard = []
        for name in ["Hero", "Problem", "Guide", "Plan", "Value Stack", "Social Proof", "CTA Block", "FAQ", "Footer CTA"]:
            keyboard.append([{"text": name, "callback_data": "lp_regen_" + name.lower().replace(" ", "_")}])
        send(chat_id, "Which section to regenerate?", keyboard)

    elif data == "lp_revert":
        pre_edit = state.get("lp_pre_edit", "")
        if pre_edit:
            state["current_lp"] = pre_edit
            send(chat_id, "Reverted to previous version.")
            keyboard = [
                [{"text": "Approve — Generate Design Brief", "callback_data": "lp_approve_copy"}],
                [{"text": "Quick Edit", "callback_data": "lp_quick_edit"}]
            ]
            send(chat_id, "Previous version restored.", keyboard)
        else:
            send(chat_id, "No previous version to revert to.")

    elif data == "lp_now_regen":
        section = state.get("lp_regen_section", "Hero")
        state["stage"] = "lp_ready"
        cta_key = state.get("lp_cta", "pro")
        cta = LP_CTA_DEFS.get(cta_key, LP_CTA_DEFS["pro"])
        avatar_keys = state.get("selected_avatars", [])
        avatar_descs = "\n".join([k + ": " + AVATARS_AD.get(k, ("",))[1] for k in avatar_keys])
        context = sanitise(state.get("lp_context", ""))
        send(chat_id, "Regenerating " + section + "...")
        try:
            user_prompt = "LANDING PAGE BRIEF:\nAVATAR(S): " + avatar_descs
            user_prompt += "\nCTA: " + cta["label"] + " — " + cta["price"]
            user_prompt += "\nPOSITIONING: " + cta["positioning"]
            if context: user_prompt += "\nCONTEXT: " + context
            user_prompt += "\n\nREGENERATE ONLY: " + section.upper() + " SECTION. Return as plain text with graphic recommendations."
            result = claude(user_prompt, max_tokens=1200, system=BRANDSCRIPT_PROMPT)
            send_plain(chat_id, result)
            keyboard = [
                [{"text": "Quick Edit", "callback_data": "lp_quick_edit"},
                 {"text": "Regenerate section", "callback_data": "lp_regen"}],
                [{"text": "Critique", "callback_data": "critique_lp"},
                 {"text": "Mark Complete", "callback_data": "mark_complete"}],
                [{"text": "Generate another page", "callback_data": "lp_again"}]
            ]
            send(chat_id, "Section regenerated.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))


    elif data.startswith("lp_regen_"):
        section = data.replace("lp_regen_", "").replace("_", " ").title()
        state["lp_regen_section"] = section
        state["stage"] = "awaiting_lp_regen_instruction"
        keyboard = [[{"text": "Regenerate as-is", "callback_data": "lp_now_regen"}]]
        send(chat_id, "*Regenerate: " + section + "*\n\nType a specific direction, or tap to regenerate as-is:", keyboard)


    elif data == "lp_again":
        state["selected_avatars"] = []
        state.pop("lp_cta", None)
        state.pop("lp_context", None)
        show_lp_avatar_menu(chat_id)

    elif data == "context_yes":
        user_state[chat_id]["stage"] = "awaiting_context_text"
        send(chat_id, "Type your extra context:")

    elif data == "context_no":
        user_state[chat_id]["context"] = ""
        user_state[chat_id]["stage"] = "pick_angle"
        gen_angles(chat_id)

    elif data == "custom_angle":
        user_state[chat_id]["stage"] = "awaiting_custom_angle"
        send(chat_id, "Type your angle:")

    elif data == "custom_hook":
        user_state[chat_id]["stage"] = "awaiting_custom_hook"
        send(chat_id, "Type your subject line:\n\n_Format: Subject line | Preview text_\n_Or just type the subject line_")

    elif data == "social_custom_angle":
        user_state[chat_id]["stage"] = "awaiting_custom_social_angle"
        send(chat_id, "Type your angle:")

    elif data == "regen_angles":
        gen_angles(chat_id)

    elif data.startswith("free_angle_"):
        idx = int(data.replace("free_angle_", ""))
        free_angles = state.get("free_angles", [])
        if idx < len(free_angles):
            state["selected_free_angle"] = free_angles[idx]
            state["selected_pro_angle"] = free_angles[idx]
            state["selected_angle"] = free_angles[idx]
        gen_hooks(chat_id)

    elif data.startswith("pro_angle_"):
        idx = int(data.replace("pro_angle_", ""))
        pro_angles = state.get("pro_angles", state.get("angles", []))
        if idx < len(pro_angles):
            state["selected_pro_angle"] = pro_angles[idx]
            state["selected_angle"] = pro_angles[idx]
            state["selected_free_angle"] = pro_angles[idx]
        gen_hooks(chat_id)

    elif data == "regen_free_angles":
        gen_angles(chat_id)

    elif data == "regen_pro_angles":
        show_pro_angle_picker(chat_id)

    elif data == "custom_free_angle":
        state["stage"] = "awaiting_custom_free_angle"
        send(chat_id, "Type your Free email angle:")

    elif data == "custom_pro_angle":
        state["stage"] = "awaiting_custom_pro_angle"
        send(chat_id, "Type your Pro email angle:")

    elif data.startswith("free_hook_"):
        idx = int(data.replace("free_hook_", ""))
        hooks = state.get("free_hooks", state.get("hooks", []))
        if idx < len(hooks):
            # One hook serves both Free and Pro
            state["selected_free_hook"] = hooks[idx]
            state["selected_pro_hook"] = hooks[idx]
            state["selected_hook"] = hooks[idx]
        # Skip Pro hook picker — go straight to CTA
        ask_free_cta(chat_id)

    elif data.startswith("pro_hook_") and data != "pro_hook_same":
        idx = int(data.replace("pro_hook_", ""))
        pro_hooks = state.get("pro_hooks", [])
        if idx < len(pro_hooks):
            state["selected_pro_hook"] = pro_hooks[idx]
        state["stage"] = "pick_free_cta"
        ask_free_cta(chat_id)

    elif data == "pro_hook_same":
        state["selected_pro_hook"] = state.get("selected_free_hook", state.get("selected_hook", {}))
        state["stage"] = "pick_free_cta"
        ask_free_cta(chat_id)

    elif data == "regen_free_hooks":
        gen_hooks(chat_id)

    elif data == "regen_pro_hooks":
        gen_pro_hooks(chat_id)

    elif data == "custom_free_hook":
        state["stage"] = "awaiting_custom_free_hook"
        send(chat_id, "Type your Free email hook:\n_Format: Subject | Preview_")

    elif data == "custom_pro_hook":
        state["stage"] = "awaiting_custom_pro_hook"
        send(chat_id, "Type your Pro email hook:\n_Format: Subject | Preview_")

    elif data.startswith("angle_"):
        idx = int(data.split("_")[1])
        angles = state.get("angles", state.get("free_angles", []))
        if idx < len(angles):
            # One angle serves both Free and Pro emails
            state["selected_angle"] = angles[idx]
            state["selected_free_angle"] = angles[idx]
            state["selected_pro_angle"] = angles[idx]
        gen_hooks(chat_id)

    elif data == "regen_hooks":
        gen_hooks(chat_id)

    elif data.startswith("hook_"):
        idx = int(data.split("_")[1])
        hooks = state.get("hooks", [])
        if idx < len(hooks):
            user_state[chat_id]["selected_hook"] = hooks[idx]
            send(chat_id, "Hook: _" + hooks[idx]["subject"] + "_")
            ask_free_cta(chat_id)

    elif data.startswith("cta_free_"):
        key = data.replace("cta_free_", "")
        user_state[chat_id]["free_cta"] = key
        send(chat_id, "Free CTA: _" + CTA_OPTIONS["free"][key] + "_")
        ask_pro_cta(chat_id)

    elif data.startswith("cta_pro_"):
        key = data.replace("cta_pro_", "")
        user_state[chat_id]["pro_cta"] = key
        send(chat_id, "Pro CTA: _" + CTA_OPTIONS["pro"][key] + "_")
        gen_emails(chat_id)

    elif data == "regen_emails":
        gen_emails(chat_id)

    elif data == "quick_edit":
        user_state[chat_id]["pre_edit_stage"] = "emails_ready"
        ask_quick_edit(chat_id, mode="email")

    elif data == "social_quick_edit":
        user_state[chat_id]["pre_edit_stage"] = "social_ready"
        ask_quick_edit(chat_id, mode="social")

    elif data == "enhance":
        gen_enhance(chat_id, mode="email")

    elif data == "social_enhance":
        gen_enhance(chat_id, mode="social")

    elif data.startswith("sfmt_edit_"):
        fmt_key = data.replace("sfmt_edit_", "")
        # Load the right format into current_social
        outputs = state.get("social_outputs", {})
        if fmt_key in outputs:
            state["current_social"] = outputs[fmt_key]["content"]
            state["current_social_type"] = outputs[fmt_key]["type"]
            state["current_social_format"] = fmt_key
        state["stage"] = "social_quick_edit"
        send(chat_id, "What would you like to change in the " + state.get("current_social_type","content") + "?")

    elif data.startswith("sfmt_enhance_"):
        fmt_key = data.replace("sfmt_enhance_", "")
        outputs = state.get("social_outputs", {})
        if fmt_key in outputs:
            state["current_social"] = outputs[fmt_key]["content"]
            state["current_social_type"] = outputs[fmt_key]["type"]
            state["current_social_format"] = fmt_key
        gen_enhance(chat_id, mode="social")

    elif data.startswith("sfmt_approve_"):
        fmt_key = data.replace("sfmt_approve_", "")
        outputs = state.get("social_outputs", {})
        if fmt_key in outputs:
            state["current_social"] = outputs[fmt_key]["content"]
            state["current_social_type"] = outputs[fmt_key]["type"]
            state["current_social_format"] = fmt_key
        # Save to voice corpus
        social_body = state.get("current_social", "")
        social_type = state.get("current_social_type", "social")
        if social_body:
            save_voice_example(chat_id, social_body[:600], "approved_" + social_type.lower().replace(" ", "_"))
        state["stage"] = "social_approved"
        keyboard = [
            [{"text": "Generate another format", "callback_data": "social_yes"}],
            [{"text": "🎨 Visual brief + image",  "callback_data": "vb_auto"}],
            [{"text": "Mark Complete",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, social_type + " approved.", keyboard)

    elif data.startswith("enh_"):
        enh_id = int(data.split("_")[1])
        toggle_enhancement(chat_id, enh_id, message_id)

    elif data == "apply_enhancements":
        apply_enhancements(chat_id)

    elif data == "skip_enhance":
        mode = state.get("enhance_mode", "email")
        if mode == "social":
            state["stage"] = "social_ready"
            send(chat_id, "Enhancement skipped.", social_action_keyboard())
        else:
            state["stage"] = "emails_ready"
            send(chat_id, "Enhancement skipped.", email_action_keyboard())

    elif data == "add_context_email":
        state["stage"] = "awaiting_additional_context"
        send(chat_id, "Add any extra context to factor in — promo, event, angle tweak, anything:\n\n_The emails will be regenerated with this added._")

    elif data == "revert_email_edit":
        pre_edit = state.get("pre_edit_emails", {})
        if pre_edit:
            state["current_emails"] = pre_edit
            free_body = extract_text(pre_edit.get("free", ""))
            pro_body = extract_text(pre_edit.get("pro", ""))
            if free_body: send_plain(chat_id, "FREE EMAIL (reverted)\n\n" + free_body)
            if pro_body: send_plain(chat_id, "PRO EMAIL (reverted)\n\n" + pro_body)
            send(chat_id, "Reverted to previous version.", email_action_keyboard())
        else:
            send(chat_id, "Nothing to revert to.")

    elif data == "accept_email_edit":
        state.pop("pre_edit_emails", None)
        send(chat_id, "Changes accepted.", email_action_keyboard())

    elif data == "approve_emails":
        state["stage"] = "emails_approved"
        # Save content metadata + actual approved copy for voice learning
        free_body = state.get("current_emails", {}).get("free", "")
        pro_body = state.get("current_emails", {}).get("pro", "")
        save_content_metadata(chat_id, "email", {
            "angle": state.get("angle", state.get("selected_angle", "")),
            "hook_subject": state.get("selected_hook", {}).get("subject", ""),
            "hook_preview": state.get("selected_hook", {}).get("preview", ""),
            "free_cta": state.get("free_cta", ""),
            "pro_cta": state.get("pro_cta", ""),
            "tone": state.get("selected_tone", "standard"),
            "free_email_body": free_body[:600],
            "pro_email_body": pro_body[:600]
        })
        # Save to voice corpus for self-learning
        if free_body:
            save_voice_example(chat_id, free_body[:600], "approved_free_email")
        if pro_body:
            save_voice_example(chat_id, pro_body[:600], "approved_pro_email")
        keyboard = [
            [{"text": "📱 Create social content",     "callback_data": "social_yes"}],
            [{"text": "🎨 Email thumbnail brief",     "callback_data": "vb_type_email"}],
            [{"text": "✅ Mark complete",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Emails approved.", keyboard)

    elif data == "social_yes":
        state["selected_social_formats"] = []
        # If we already have report/context from a previous generation, skip source menu
        if state.get("report") or state.get("current_social"):
            state["stage"] = "pick_social_formats"
            show_standalone_social_menu(chat_id)
        else:
            show_social_source_menu(chat_id)

    elif data.startswith("social_angle_"):
        idx = int(data.replace("social_angle_", ""))
        angles = state.get("social_angles", [])
        if idx < len(angles):
            state["social_angle"] = angles[idx]
        # Don't reset formats — go straight to hooks since formats already selected
        gen_social_hooks(chat_id)

    elif data == "social_regen_angles":
        state["stage"] = "awaiting_social_report"
        gen_social_angles(chat_id)

    elif data == "social_fw_aida":
        state["social_framework"] = "AIDA"
        gen_social_selected(chat_id)

    elif data == "social_fw_pas":
        state["social_framework"] = "PAS"
        gen_social_selected(chat_id)

    elif data == "social_switch_fw":
        # Switch framework and regenerate
        current = state.get("social_framework", "AIDA")
        new_fw = "PAS" if current == "AIDA" else "AIDA"
        state["social_framework"] = new_fw
        state["social_hooks"] = None  # force re-generate
        gen_social_selected(chat_id)

    elif data == "gen_social_confirmed":
        gen_social_selected(chat_id)

    elif data.startswith("social_hook_"):
        # Format: social_hook_{fmt_idx}_{hook_idx}
        parts = data.replace("social_hook_", "").split("_")
        if len(parts) == 2:
            fmt_idx = int(parts[0])
            hook_idx = int(parts[1])
            formats = state.get("selected_social_formats", [])
            if fmt_idx < len(formats):
                fmt = formats[fmt_idx]
                hooks = state.get("social_hooks", {}).get(fmt, [])
                if hook_idx < len(hooks):
                    if "selected_social_hooks" not in state:
                        state["selected_social_hooks"] = {}
                    state["selected_social_hooks"][fmt] = hooks[hook_idx]
            show_social_hook_picker(chat_id, fmt_idx + 1)

    elif data == "src_free":
        state["social_source"] = "free"
        show_social_format_menu(chat_id)

    elif data == "src_pro":
        state["social_source"] = "pro"
        show_social_format_menu(chat_id)

    elif data == "src_hot":
        state["social_source"] = "hot"
        show_social_format_menu(chat_id)

    elif data == "src_warm":
        state["social_source"] = "warm"
        show_social_format_menu(chat_id)

    elif data == "src_cold":
        state["social_source"] = "cold"
        show_social_format_menu(chat_id)

    elif data.startswith("fmt_"):
        toggle_social_format(chat_id, data, message_id)

    elif data == "gen_social_selected":
        gen_social_selected(chat_id)



    elif data == "approve_social":
        state["stage"] = "social_approved"
        social_body = state.get("current_social", "")
        social_type = state.get("current_social_type", "social")
        if social_body:
            save_voice_example(chat_id, social_body[:600], "approved_" + social_type.lower().replace(" ", "_"))
        keyboard = [
            [{"text": "Generate another format",   "callback_data": "social_yes"}],
            [{"text": "🎨 Visual brief + image",   "callback_data": "vb_auto"}],
            [{"text": "Mark Complete",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, social_type + " approved.", keyboard)

    elif data == "subject_ab":
        gen_subject_ab(chat_id)

    elif data.startswith("use_subject_"):
        idx = int(data.replace("use_subject_", ""))
        toggle_subject_select(chat_id, idx, message_id)

    elif data == "keep_subject":
        state["stage"] = "emails_ready"
        send(chat_id, "Keeping current subject.", email_action_keyboard())

    elif data == "tone_menu":
        show_tone_menu(chat_id)

    elif data.startswith("tone_") and data != "tone_menu":
        tone_key = data.replace("tone_", "")
        if tone_key in ["standard", "aggressive", "empathetic"]:
            apply_tone(chat_id, tone_key)

    elif data == "cancel_tone":
        state["stage"] = "emails_ready"
        send(chat_id, "Tone change cancelled.", email_action_keyboard())

    elif data == "segments":
        gen_segments(chat_id)

    elif data == "seg_edit_hot":
        ask_seg_edit(chat_id, "hot")

    elif data == "seg_edit_warm":
        ask_seg_edit(chat_id, "warm")

    elif data == "seg_edit_cold":
        ask_seg_edit(chat_id, "cold")

    elif data == "approve_segments":
        state["stage"] = "emails_approved"
        keyboard = [
            [{"text": "Yes — create social content", "callback_data": "social_yes"}],
            [{"text": "No — mark complete", "callback_data": "mark_complete"}]
        ]
        send(chat_id, "Segments approved. Want to create social content?", keyboard)

    elif data == "back_to_emails":
        state["stage"] = "emails_ready"
        send(chat_id, "Back to emails.", email_action_keyboard())

    elif data.startswith("sel_sub_"):
        idx = int(data.replace("sel_sub_", ""))
        toggle_subject_select(chat_id, idx, message_id)

    elif data == "apply_subjects":
        apply_subjects(chat_id)

    elif data == "length_email":
        show_length_menu(chat_id, mode="email")

    elif data == "length_social":
        show_length_menu(chat_id, mode="social")

    elif data == "length_extend":
        apply_length(chat_id, "extend")

    elif data == "length_shorten":
        apply_length(chat_id, "shorten")

    elif data == "cancel_length":
        mode = state.get("length_mode", "email")
        if mode == "social":
            send(chat_id, "Cancelled.", social_action_keyboard())
        else:
            send(chat_id, "Cancelled.", email_action_keyboard())

    elif data == "back_to_done":
        state["stage"] = "emails_ready"
        send(chat_id, "Back to emails.", email_action_keyboard())

    elif data == "reel_extend":
        state["reel_duration"] = state.get("reel_duration", 52) + 15
        gen_reel(chat_id)

    elif data == "reel_shorten":
        state["reel_duration"] = max(20, state.get("reel_duration", 52) - 15)
        gen_reel(chat_id)

    elif data == "carousel_extend":
        state["carousel_slides"] = state.get("carousel_slides", 6) + 2
        gen_carousel(chat_id)

    elif data == "carousel_shorten":
        state["carousel_slides"] = max(3, state.get("carousel_slides", 6) - 2)
        gen_carousel(chat_id)

    elif data == "story_extend":
        state["story_slides"] = state.get("story_slides", 3) + 1
        gen_story(chat_id, multi=True)

    elif data == "story_shorten":
        state["story_slides"] = max(1, state.get("story_slides", 3) - 1)
        is_multi = state.get("story_slides", 1) > 1
        gen_story(chat_id, multi=is_multi)

    elif data == "log_email_start":
        start_log_email(chat_id)

    elif data == "log_ad_start":
        start_log_ad(chat_id)

    elif data == "run_email_report":
        run_email_analysis(chat_id)

    elif data == "run_ad_report":
        run_ad_analysis(chat_id)

    elif data == "email_log_ab_yes":
        state["log_stage"] = "email_ab_variable"
        send(chat_id, "What is the variable being tested? (e.g. 'name in subject line', 'curiosity gap vs data hook', 'short vs long preview')")

    elif data == "email_log_ab_no":
        state.get("log_data", {}).update({"ab_variable": None, "ab_label": None})
        state["log_stage"] = "email_recipients"
        send(chat_id, "How many recipients received this email?")

    elif data.startswith("adlog_type_"):
        ad_type = data.replace("adlog_type_", "")
        state["log_data"] = {"ad_type": ad_type}
        state["log_stage"] = "adlog_headline"
        send(chat_id, "Enter the headline or first line of the ad:")

    elif data.startswith("adlog_avatar_"):
        avatar_key = data.replace("adlog_avatar_", "")
        avatar_name = AVATARS_AD.get(avatar_key, ("Unknown",))[0]
        if "log_data" not in state:
            state["log_data"] = {}
        state["log_data"]["avatar"] = avatar_name
        state["log_stage"] = "adlog_stage"
        keyboard = [[{"text": s.capitalize(), "callback_data": "adlog_stage_" + s}] for s in ["awareness","consideration","conversion"]]
        send(chat_id, "Which funnel stage was this ad?", keyboard)

    elif data.startswith("adlog_stage_"):
        stage_name = data.replace("adlog_stage_", "")
        if "log_data" not in state:
            state["log_data"] = {}
        state["log_data"]["stage"] = stage_name
        state["log_stage"] = "adlog_date"
        send(chat_id, "Date the ad ran? (e.g. 2026-03-22)")

    elif data == "ds_adverts":
        state["ds_images"] = []
        state.pop("ds_csv_text", None)
        start_ds_adverts(chat_id)

    elif data.startswith("ad_filter_"):
        ad_filter = data.replace("ad_filter_", "")
        state["ds_ad_filter"] = ad_filter
        state["stage"] = "ds_awaiting_ad_data"
        state["ds_images"] = []
        filter_label = {"all": "All ads", "video": "Video ads only", "static": "Static ads only"}.get(ad_filter, "All ads")
        keyboard = [[{"text": "Done — analyse now", "callback_data": "ds_analyse_ads"}]]
        send(chat_id, "*" + filter_label + "*\n\nUpload your Meta Ads Manager screenshots or CSV. Send all then tap Done.", keyboard)

    elif data == "ds_social":
        state["ds_images"] = []
        state.pop("ds_csv_text", None)
        start_ds_social(chat_id)

    elif data.startswith("social_filter_"):
        social_filter = data.replace("social_filter_", "")
        state["ds_social_filter"] = social_filter
        state["stage"] = "ds_awaiting_social_data"
        state["ds_images"] = []
        filter_label = {"all": "All formats", "reels": "Reels only", "statics": "Statics only", "carousels": "Carousels only"}.get(social_filter, "All formats")
        keyboard = [[{"text": "Done — analyse now", "callback_data": "ds_analyse_social"}]]
        send(chat_id, "*" + filter_label + "*\n\nUpload your Minter screenshots or CSV. Send all then tap Done.", keyboard)

    elif data == "ds_emails":
        state["ds_images"] = []
        state.pop("ds_csv_text", None)
        start_ds_emails(chat_id)

    elif data == "ds_landing":
        state["ds_images"] = []
        state.pop("ds_csv_text", None)
        start_ds_landing(chat_id)

    elif data == "landing_scope_single":
        state["landing_scope"] = "single"
        state["stage"] = "ds_awaiting_landing_splitvar"
        state["ds_images"] = []
        send(chat_id, "What variable are you testing?\n\n_Subject line (Open rate), Content (CTR), CTA (CTR)_")

    elif data == "landing_scope_multi":
        state["landing_scope"] = "multi"
        state["stage"] = "ds_awaiting_landing_splitvar"
        state["ds_images"] = []
        send(chat_id, "What variable are you testing across pages?\n\n_Then upload screenshots from all page variants_")

    elif data == "ds_analyse_ads":
        analyse_ads(chat_id)

    elif data == "ds_analyse_social":
        analyse_social(chat_id)

    elif data == "ds_analyse_emails":
        analyse_emails(chat_id)

    elif data == "ds_analyse_landing":
        analyse_landing(chat_id)

    elif data == "ds_email_splittest":
        start_ds_email_splittest(chat_id)

    elif data == "ds_landing_splittest":
        start_ds_landing_splittest(chat_id)

    elif data == "ds_analyse_emails_split":
        analyse_emails(chat_id, split_var=state.get("ds_split_var"))

    elif data == "ds_email_type_numbers":
        state["stage"] = "ds_awaiting_email_split_numbers"
        send(chat_id, "Type or paste your numbers in any format:\n\n_e.g. Variant A: 5,000 sent, 1,200 opens, 25 clicks\nVariant B: 4,800 sent, 960 opens, 18 clicks_")

    elif data == "ds_analyse_landing_split":
        analyse_landing(chat_id, split_var=state.get("ds_split_var"))

    elif data.startswith("ds_add_more_"):
        # Keep current stage, just prompt for more
        send(chat_id, "Send your next screenshot:")

    elif data == "briefing_refresh":
        generate_briefing(chat_id)

    # ── AD PRODUCT + EXISTING AD CALLBACKS ───────────────────────
    elif data == "ad_product_pro":
        state["ad_product"] = "pro"
        state["ad_product_label"] = "Cryptonary Pro ($1,197/year)"
        state["ad_product_context"] = "Cryptonary Pro — research platform, 300K+ subscribers, weekly analysis, airdrops, community. Price: $1,197/year."
        show_ad_input_menu(chat_id)

    elif data == "ad_product_ic":
        state["ad_product"] = "ic"
        state["ad_product_label"] = "Inner Circle ($15K–$22K)"
        state["ad_product_context"] = "Cryptonary Inner Circle — dedicated team of 8, personalised portfolio framework, monthly audit reports, 20+ hours call time. Application only. $15,000–$22,000/year. Minimum $200K portfolio."
        show_ad_input_menu(chat_id)

    elif data == "ad_input_new":
        state["ad_existing"] = False
        state["stage"] = "awaiting_ad_theme"
        product_label = state.get("ad_product_label", "")
        send(chat_id, ("*" + product_label + "*\n\n" if product_label else "") + "What\'s the theme or context for this campaign?\n\n_Market event, specific angle, target narrative..._")

    elif data == "ad_input_existing":
        state["ad_existing"] = True
        state["stage"] = "awaiting_ad_existing_upload"
        send(chat_id, "Upload your existing ad — image, screenshot, or paste the copy below.")

    elif data.startswith("existing_ad_") and data != "existing_ad_back":
        action = data.replace("existing_ad_", "")
        if action == "reconceptualise":
            state["pending_ad_action"] = "reconceptualise"
            state["stage"] = "pick_ad_avatars"
            show_avatar_menu(chat_id)
        else:
            generate_existing_ad_action(chat_id, action)

    elif data == "existing_ad_back":
        show_existing_ad_action_menu(chat_id)

    # ── LANDING PAGE NEW FLOW CALLBACKS ──────────────────────────
    elif data == "lp_goal_pro":
        state["lp_goal"] = "pro"
        state["selected_avatars"] = []
        show_lp_avatar_menu(chat_id)

    elif data == "lp_goal_ic":
        state["lp_goal"] = "ic"
        state["selected_avatars"] = []
        show_lp_avatar_menu(chat_id)

    elif data == "lp_outline_approve":
        generate_lp_full_copy(chat_id)

    elif data == "lp_regen_outline":
        generate_lp_outline(chat_id)

    elif data == "lp_more_outline_feedback":
        state["stage"] = "lp_outline_review"
        send(chat_id, "Send feedback by section number (e.g. \'2: Make the villain more specific\'):")

    elif data == "lp_approve_copy":
        lp_content = state.get("current_lp", "")
        if lp_content:
            save_voice_example(chat_id, lp_content[:600], "approved_landing_page")
        generate_lp_design_brief(chat_id)

    elif data == "lp_paste_back":
        state["stage"] = "lp_awaiting_paste_back"
        send(chat_id, "Paste your edited version. It will be saved as the new version.")

    elif data == "lp_enhance":
        gen_enhance(chat_id, mode="lp")

    elif data == "lp_length":
        state["stage"] = "lp_awaiting_length_instruction"
        send(chat_id, "What length adjustment?\n\n_e.g. Make Villain longer, shorten Plan, more concise overall_")

    # ── LANDING PAGE ANALYSIS NEW ─────────────────────────────────
    elif data == "landing_scope_single":
        state["landing_scope"] = "single"
        state["stage"] = "ds_awaiting_landing_splitvar"
        state["ds_images"] = []
        send(chat_id, "What variable are you testing?\n\n_Subject line (Open rate), Content (CTR), or CTA (CTR)_")

    elif data == "landing_scope_multi":
        state["landing_scope"] = "multi"
        state["stage"] = "ds_awaiting_landing_splitvar"
        state["ds_images"] = []
        send(chat_id, "What variable are you testing across pages?\n\n_Then upload screenshots from all variants_")

    elif data == "open_idea_engine":
        # Don't interrupt active content generation flows
        active_stages = {"pick_social_formats", "ie_generating", "social_ready",
                         "pick_angle", "pick_free_hook", "pick_pro_hook",
                         "pick_free_cta", "pick_pro_cta"}
        if state.get("stage") in active_stages:
            return  # ignore queued open_idea_engine during active flow
        preserved = {k: v for k, v in state.items() if k in ("report", "context", "selected_angle", "social_angle")}
        user_state[chat_id] = {"stage": "idea_engine_idle"}
        user_state[chat_id].update(preserved)
        show_idea_engine_menu(chat_id)

    elif data == "ie_from_scratch":
        # Guard against double-tap
        if state.get("stage") == "ie_generating":
            return
        state["stage"] = "ie_generating"
        state["ie_idea_type"] = "instagram"
        state["ie_source_label"] = "Saved library"
        sources = load_idea_sources()
        fetched = fetch_source_content(sources)
        source_text = "\n".join(["[" + i["source"] + "] " + i["content"][:100] for i in fetched[:10]])
        state["ie_source_content"] = source_text
        generate_ie_concept(chat_id)

    elif data == "ie_from_inspiration":
        # Wait for user to drop text, URL, brief, or image
        state["stage"] = "ie_awaiting_inspiration"
        state["ie_idea_type"] = "instagram"
        send(chat_id, "*Generate from Inspiration*\n\nDrop anything:\n\n• A link\n• A screenshot or image\n• Paste text (article, caption, post)\n• Or just type a brief — e.g. \"market manipulation ideas\"")

    elif data == "ie_generate_all":
        # Regenerate concepts using current source content
        generate_ie_concept(chat_id)

    elif data == "ie_generate_social":
        generate_ideas(chat_id, "social")

    elif data == "ie_generate_ads":
        generate_ideas(chat_id, "ads")

    elif data == "ie_type_ad":
        state["ie_idea_type"] = "ad"
        keyboard = [
            [{"text": "Video script", "callback_data": "ie_ad_fmt_video"}],
            [{"text": "Static ad", "callback_data": "ie_ad_fmt_static"}]
        ]
        send(chat_id, "*Ad idea — what format?*", keyboard)

    elif data == "ie_ad_fmt_video":
        state["ie_format"] = "video"
        show_ie_source_menu(chat_id, "ad")

    elif data == "ie_ad_fmt_static":
        state["ie_format"] = "static"
        show_ie_source_menu(chat_id, "ad")

    elif data == "ie_type_instagram":
        state["ie_idea_type"] = "instagram"
        keyboard = [
            [{"text": "Reel", "callback_data": "ie_ig_fmt_reel"},
             {"text": "Carousel", "callback_data": "ie_ig_fmt_carousel"}],
            [{"text": "Static post", "callback_data": "ie_ig_fmt_static"},
             {"text": "Story", "callback_data": "ie_ig_fmt_story"}]
        ]
        send(chat_id, "*Instagram idea — what format?*", keyboard)

    elif data.startswith("ie_ig_fmt_"):
        fmt = data.replace("ie_ig_fmt_", "")
        state["ie_format"] = fmt
        show_ie_source_menu(chat_id, "instagram")

    elif data == "ie_source_database":
        state["ie_using_database"] = True
        # Fetch sources and then show format menu
        show_ie_format_menu(chat_id)

    elif data == "ie_source_upload":
        state["ie_using_database"] = False
        idea_type = state.get("ie_idea_type", "ad")
        state["stage"] = "ie_awaiting_screenshot_ideas"
        send(chat_id, "Upload an image, PDF, or paste content for inspiration:")

    elif data == "ie_source_link":
        state["ie_using_database"] = False
        state["stage"] = "ie_awaiting_link_ideas"
        send(chat_id, "Paste the URL:")

    elif data.startswith("ie_format_"):
        fmt = data.replace("ie_format_", "")
        state["ie_format"] = fmt
        if state.get("ie_using_database", True):
            # Fetch from saved sources first
            fetched = fetch_source_content(load_idea_sources())
            source_text = "\n".join(["[" + i["source"] + "] " + i["content"][:100] for i in fetched[:10]])
            state["ie_source_content"] = source_text
            state["ie_source_label"] = "Saved library"
        generate_ie_concept(chat_id)

    elif data.startswith("ie_develop_concept_"):
        state["stage"] = "idle"  # clear ie_generating to prevent open_idea_engine guard issues
        num = int(data.replace("ie_develop_concept_", "")) - 1
        concepts_raw = state.get("ie_concepts", "")
        # Extract the nth numbered idea block (matches "1." or "2." etc)
        import re as _re
        # New format: numbered blocks "1. CONCEPT: ..."
        blocks = _re.split(r"(?=\n?\d+\.\s+CONCEPT:)", concepts_raw)
        blocks = [b.strip() for b in blocks if b.strip() and _re.match(r"\d+\.", b.strip())]
        if not blocks:
            # Fallback: old CONCEPT [N]: format
            blocks = _re.split(r"CONCEPT \[?\d+\]?:", concepts_raw)
            blocks = [b.strip() for b in blocks if b.strip()]
        if num < len(blocks):
            concept_text = blocks[num][:300]
        else:
            concept_text = concepts_raw[:300]
        # Skip angle picker — concept already contains the angle. Go straight to content flow.
        state["report"] = "INSTAGRAM CONCEPT TO DEVELOP:\n" + concept_text
        state["social_angle"] = concept_text[:120]
        state["social_origin"] = "idea_engine"
        state["stage"] = "pick_social_formats"
        show_standalone_social_menu(chat_id)

    elif data == "ie_regen_concepts":
        generate_ie_concept(chat_id)

    elif data == "ie_custom_concept":
        state["stage"] = "ie_awaiting_custom_concept"
        send(chat_id, "Describe your concept:")

    elif data == "ie_back_to_concepts":
        concepts_raw = state.get("ie_concepts", "")
        if concepts_raw:
            send_plain(chat_id, "*IDEAS*\n\n" + concepts_raw)
            keyboard = [
                [{"text": str(i+1) + " — Expand this", "callback_data": "ie_develop_concept_" + str(i+1)}]
                for i in range(4)
            ] + [[{"text": "Regenerate", "callback_data": "ie_regen_concepts"}]]
            send(chat_id, "Tap a number to expand into full content:", keyboard)

    elif data.startswith("ie_angle_"):
        idx = int(data.replace("ie_angle_", ""))
        angles = state.get("ie_angles", [])
        if idx < len(angles):
            generate_ie_hook_from_angle(chat_id, angles[idx])

    elif data == "ie_custom_angle_ie":
        state["stage"] = "ie_awaiting_custom_angle_ie"
        send(chat_id, "Type your angle:")

    elif data == "ie_back_to_angles":
        concept = state.get("ie_selected_concept", "")
        if concept:
            generate_ie_angle_from_concept(chat_id, concept)

    elif data.startswith("ie_hook_"):
        idx = int(data.replace("ie_hook_", ""))
        hooks = state.get("ie_hooks", [])
        if idx < len(hooks):
            generate_ie_final_content(chat_id, hooks[idx])

    elif data == "ie_custom_hook_ie":
        state["stage"] = "ie_awaiting_custom_hook_ie"
        send(chat_id, "Type your hook:")

    elif data == "ie_back_to_hooks":
        angle = state.get("ie_selected_angle", "")
        if angle:
            generate_ie_hook_from_angle(chat_id, angle)

    elif data == "ie_content_edit":
        state["stage"] = "ie_awaiting_content_edit"
        send(chat_id, "What would you like to change?")

    elif data == "ie_content_regen":
        hook = state.get("ie_selected_hook", "")
        generate_ie_final_content(chat_id, hook)

    elif data == "ie_content_enhance":
        content_to_enh = state.get("ie_generated_content", "")
        ie_format = state.get("ie_format", "social")
        if content_to_enh:
            state["current_social"] = content_to_enh
            state["current_social_type"] = ie_format
            gen_enhance(chat_id, mode="social")
        else:
            send(chat_id, "No content to enhance yet.")

    elif data == "ie_content_critique":
        content_to_crit = state.get("ie_generated_content", "")
        if content_to_crit:
            state["current_social"] = content_to_crit
            state["current_social_type"] = state.get("ie_format", "content")
            state["critique_content_type"] = "social"
            run_critique(chat_id, "social")
        else:
            send(chat_id, "No content to critique yet.")

    elif data == "ie_content_revert":
        pre = state.get("ie_pre_edit_content", "")
        if pre:
            state["ie_generated_content"] = pre
            send_plain(chat_id, "REVERTED:\n\n" + pre)
            keyboard = [
                [{"text": "Quick Edit", "callback_data": "ie_content_edit"},
                 {"text": "Enhance", "callback_data": "ie_content_enhance"}],
                [{"text": "Develop in Studio", "callback_data": "ie_develop"}]
            ]
            send(chat_id, "Reverted.", keyboard)
        else:
            send(chat_id, "Nothing to revert to.")

    elif data == "ie_to_content_studio":
        content_val = state.get("ie_generated_content", "")
        ie_format = state.get("ie_format", "")
        if "reel" in ie_format or "video" in ie_format:
            user_state[chat_id]["report"] = content_val
            user_state[chat_id]["social_angle"] = state.get("ie_selected_angle", "")
            user_state[chat_id]["social_origin"] = "idea_engine"
            user_state[chat_id]["stage"] = "pick_social_formats"
            show_standalone_social_menu(chat_id)
        elif "carousel" in ie_format or "story" in ie_format or "static_post" in ie_format:
            user_state[chat_id]["report"] = content_val
            user_state[chat_id]["social_angle"] = state.get("ie_selected_angle", "")
            user_state[chat_id]["social_origin"] = "idea_engine"
            user_state[chat_id]["stage"] = "pick_social_formats"
            show_standalone_social_menu(chat_id)
        elif "static" in ie_format or "video_script" in ie_format:
            user_state[chat_id]["ad_theme"] = content_val[:500]
            user_state[chat_id]["stage"] = "pick_ad_avatars"
            show_avatar_menu(chat_id)
        else:
            show_main_menu(chat_id)

    elif data == "ie_manage_sources":
        show_ie_manage_sources(chat_id)

    elif data == "ie_develop":
        state["stage"] = "ie_awaiting_develop"
        send(chat_id, "Type the idea number or paste the idea you want to develop into social content:")

    elif data == "ie_develop_ad":
        state["stage"] = "ie_awaiting_develop_ad"
        send(chat_id, "Type the idea number or paste the idea you want to develop into an ad:")

    elif data == "ie_screenshot_ideas":
        state["stage"] = "ie_awaiting_screenshot_ideas"
        send(chat_id, "*Ideas from a Screenshot*\n\nSend an image — Instagram post, ad, competitor content, anything. I\'ll analyse it and generate ideas based on what\'s working in it.")

    elif data == "ie_screenshot_critique":
        state["stage"] = "ie_awaiting_screenshot_critique"
        send(chat_id, "*Critique a Screenshot*\n\nSend an image of any content — your own post, a competitor ad, a carousel. I\'ll critique it against our full knowledge base with specific, numbered fixes.")

    elif data == "ie_imageprompt_again":
        brief = state.get("last_image_brief", "")
        if brief:
            generate_image_prompts(chat_id, brief)
        else:
            send(chat_id, "Use /imageprompt [description] to generate new prompts.")

    elif data.startswith("ie_add_"):
        source_type = data.replace("ie_add_", "")
        state["stage"] = "ie_awaiting_source_" + source_type
        prompts = {
            "instagram": "Enter Instagram handle (without @):",
            "twitter": "Enter Twitter/X handle (without @):",
            "facebook": "Enter Facebook page name:",
            "telegram": "Enter Telegram channel name (without t.me/):",
            "reddit": "Enter subreddit name (without r/):"
        }
        send(chat_id, prompts.get(source_type, "Enter the source name:"))

    elif data == "ds_followup":
        state["stage"] = "ds_awaiting_followup"
        send(chat_id, "Ask your follow-up question:")

    elif data.startswith("critique_") and data != "critique_":
        content_type = data.replace("critique_", "")
        run_critique(chat_id, content_type)

    elif data.startswith("toggle_critique_"):
        num = int(data.replace("toggle_critique_", ""))
        selected = state.get("critique_selected_fixes", [])
        if num in selected:
            selected.remove(num)
        else:
            selected.append(num)
        state["critique_selected_fixes"] = selected
        # Rebuild keyboard showing selections
        total = state.get("critique_fixes_available", 6)
        keyboard = []
        for i in range(1, total + 1):
            tick = "[x]" if i in selected else "[ ]"
            keyboard.append([{"text": tick + " Fix " + str(i), "callback_data": "toggle_critique_" + str(i)}])
        keyboard.append([{"text": "Apply selected (" + str(len(selected)) + ")", "callback_data": "apply_critique_selected"},
                         {"text": "Apply all", "callback_data": "apply_critique_all"}])
        keyboard.append([{"text": "Ignore all", "callback_data": "ignore_critique"}])
        try:
            tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id,
                "reply_markup": json.dumps({"inline_keyboard": keyboard})})
        except:
            send(chat_id, str(len(selected)) + " fix(es) selected.", keyboard)

    elif data == "apply_critique_selected":
        selected = state.get("critique_selected_fixes", [])
        if not selected:
            send(chat_id, "No fixes selected. Tap the fix numbers to select them first.")
        else:
            apply_critique_fixes(chat_id, selected)

    elif data == "apply_critique_all":
        total = state.get("critique_fixes_available", 6)
        apply_critique_fixes(chat_id, list(range(1, total + 1)))

    elif data.startswith("apply_critique_"):
        num = data.replace("apply_critique_", "")
        apply_critique_fixes(chat_id, [int(num)])

    elif data.startswith("revert_to_previous_"):
        content_type = data.replace("revert_to_previous_", "")
        key_map = {"email": ("prev_emails", "current_emails"),
                   "ad": ("prev_ad_output", "current_ad_output"),
                   "social": ("prev_social", "current_social")}
        keys = key_map.get(content_type, ("prev_emails", "current_emails"))
        prev = state.get(keys[0])
        if prev:
            state[keys[1]] = prev
            send(chat_id, "Reverted to previous version.")
            kb_map = {"email": email_action_keyboard, "ad": ad_action_keyboard, "social": social_action_keyboard}
            send(chat_id, "Previous version restored.", kb_map.get(content_type, email_action_keyboard)())
        else:
            send(chat_id, "No previous version saved.")

    elif data.startswith("accept_changes_"):
        content_type = data.replace("accept_changes_", "")
        kb_map = {"email": email_action_keyboard, "ad": ad_action_keyboard, "social": social_action_keyboard}
        send(chat_id, "Changes accepted.", kb_map.get(content_type, email_action_keyboard)())

    elif data.startswith("revert_to_previous_ie"):
        prev = state.get("ie_pre_edit_content", "")
        if prev:
            state["ie_generated_content"] = prev
            send_plain(chat_id, prev)
            keyboard = [
                [{"text": "Quick Edit", "callback_data": "ie_content_edit"},
                 {"text": "Develop in Content Studio", "callback_data": "ie_to_content_studio"}]
            ]
            send(chat_id, "Reverted to previous version.", keyboard)
        else:
            send(chat_id, "No previous version saved.")

    elif data == "revert_critique":
        original = state.get("critique_original", "")
        content_type = state.get("critique_content_type", "email")
        if original:
            if content_type == "email":
                emails = parse_delimited_emails(original) if "===FREE EMAIL" in original else {"free": original[:len(original)//2], "pro": original[len(original)//2:]}
                state["current_emails"] = emails
                send(chat_id, "Reverted to previous version.", email_action_keyboard())
            elif content_type == "ad":
                state["current_ad_output"] = original
                send(chat_id, "Reverted to previous version.", ad_action_keyboard())
            elif content_type == "social":
                state["current_social"] = original
                send(chat_id, "Reverted to previous version.", social_action_keyboard())
        else:
            send(chat_id, "No previous version to revert to.")

    elif data == "ignore_critique":
        # Just dismiss and show content keyboard
        ctype = state.get("critique_content_type", "email")
        kb_map = {"email": email_action_keyboard, "ad": ad_action_keyboard,
                  "social": social_action_keyboard}
        kb_fn = kb_map.get(ctype, email_action_keyboard)
        send(chat_id, "Critique dismissed.", kb_fn())

    elif data == "mark_complete":
        send(chat_id, "Set complete. What would you like to do next?", mark_complete_keyboard())

    elif data == "img_from_brief":
        # Must be BEFORE img_ startswith catch below
        brief = state.get("last_visual_brief", "") or _global_brief_store.get(chat_id, "")
        vb_type = state.get("last_visual_type", "static")
        print(f"img_from_brief fired. chat_id={chat_id}, brief_len={len(brief)}", flush=True)
        if not brief:
            tg("sendMessage", {"chat_id": chat_id, "text": "No brief found. Generate a visual brief first."})
            return
        state["pending_img_concept"] = brief[:600]
        state["last_visual_brief"] = brief
        _global_brief_store[chat_id] = brief
        # Auto-route structured content to Claude, photographic to Gemini
        social_type = state.get("current_social_type", "")
        auto_engine = None
        if any(t in social_type for t in ["Carousel", "Reel", "Story"]):
            auto_engine = "claude"
        elif any(t in social_type for t in ["Static"]):
            auto_engine = "gemini"
        if auto_engine:
            state["img_engine"] = auto_engine
            show_image_style_menu(chat_id, auto_engine)
        else:
            keyboard = [
                [{"text": "Claude (SVG + HTML file)",  "callback_data": "img_engine_claude"}],
                [{"text": "Gemini (graphic/thumbnail)", "callback_data": "img_engine_gemini"}],
                [{"text": "DALL-E (photo/cinematic)",   "callback_data": "img_engine_dalle"}],
            ]
            result = tg("sendMessage", {
                "chat_id": chat_id,
                "text": "Which AI for image generation?",
                "reply_markup": {"inline_keyboard": keyboard}
            })
            print(f"Engine picker result: ok={result.get('ok') if result else 'None'}", flush=True)

    elif data.startswith("img_"):
        handle_image_callbacks(chat_id, data, state)

    elif data == "vb_auto":
        # Auto-detect type from current content — skip picker
        social_type = state.get("current_social_type", "")
        if "Carousel" in social_type:
            state["img_engine"] = "claude"
            generate_visual_brief(chat_id, "carousel")
        elif "Reel" in social_type:
            state["img_engine"] = "claude"
            generate_visual_brief(chat_id, "reel")
        elif "Static" in social_type:
            generate_visual_brief(chat_id, "static")
        elif "Story" in social_type:
            generate_visual_brief(chat_id, "story")
        elif state.get("current_emails"):
            generate_visual_brief(chat_id, "email")
        elif state.get("current_ad"):
            generate_visual_brief(chat_id, "ad_static")
        else:
            show_visual_brief_menu(chat_id)

    elif data.startswith("vb_type_"):
        vb_type = data.replace("vb_type_", "")
        generate_visual_brief(chat_id, vb_type)

    elif data == "vb_edit":
        state["stage"] = "vb_awaiting_edit"
        send(chat_id, "What to change in the brief?")

    elif data == "approve_ad":
        ad_output = state.get("current_ad", "")
        if ad_output:
            save_voice_example(chat_id, ad_output[:600], "approved_ad")
        keyboard = [
            [{"text": "🎨 Visual brief + image", "callback_data": "vb_type_ad_static"}],
            [{"text": "Generate another set",    "callback_data": "ads_again"}],
            [{"text": "Mark Complete",           "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Ad approved.", keyboard)

    elif data == "voice_confirm":
        route_voice_to_flow(chat_id)

    elif data == "voice_edit":
        state["stage"] = "voice_awaiting_edit"
        send(chat_id, "Tell me what to change or add:")

    elif data == "voice_discard":
        state.pop("voice_transcript", None)
        state.pop("voice_content_type", None)
        state.pop("voice_topic", None)
        state.pop("voice_angle", None)
        send(chat_id, "Discarded. Send a new voice note or type your brief.")

    elif data == "voice_as_email":
        state["voice_content_type"] = "email"
        route_voice_to_flow(chat_id)

    elif data == "voice_as_social":
        state["voice_content_type"] = "reel"
        route_voice_to_flow(chat_id)

    elif data == "voice_as_ideas":
        state["voice_content_type"] = "idea"
        route_voice_to_flow(chat_id)

    elif data == "voice_as_ad":
        state["voice_content_type"] = "ad"
        route_voice_to_flow(chat_id)

    elif data == "log_perf_start":
        start_log_performance(chat_id)

# ── POLLING LOOP ──────────────────────────────────────────────────

def poll():
    offset = 0
    processed_updates = set()
    print("Cryptonary Bot V9 running.", flush=True)

    # Step 1: Delete webhook and drop pending updates
    for attempt in range(3):
        try:
            tg("deleteWebhook", {"drop_pending_updates": True})
            print("Webhook cleared.", flush=True)
            break
        except Exception as e:
            print("Webhook clear attempt", attempt+1, "error:", e, flush=True)
            time.sleep(2)

    # Step 2: Close any existing connections held by other instances
    # by calling getUpdates with timeout=0 repeatedly until we get ok=True
    print("Claiming Telegram connection...", flush=True)
    for attempt in range(20):
        try:
            url = "https://api.telegram.org/bot" + TELEGRAM_TOKEN + "/getUpdates?timeout=0&offset=-1"
            req = urllib.request.Request(url)
            with urllib.request.urlopen(req, timeout=10) as r:
                result = json.loads(r.read())
            if result.get("ok"):
                print("Connection claimed successfully.", flush=True)
                break
            else:
                print("Waiting to claim connection, attempt", attempt+1, flush=True)
                time.sleep(3)
        except Exception as e:
            print("Claim attempt", attempt+1, "error:", e, flush=True)
            time.sleep(3)

    while True:
        try:
            url = "https://api.telegram.org/bot" + TELEGRAM_TOKEN + "/getUpdates?timeout=30&offset=" + str(offset)
            req = urllib.request.Request(url)
            with urllib.request.urlopen(req, timeout=35) as r:
                data = json.loads(r.read())
            if not data.get("ok"):
                err_code = data.get("error_code", 0)
                if err_code == 409:
                    print("409 Conflict — retrying claim...", flush=True)
                    time.sleep(5)
                    # Try to claim again
                    try:
                        tg("deleteWebhook", {"drop_pending_updates": False})
                    except:
                        pass
                else:
                    time.sleep(5)
                continue
            for update in data.get("result", []):
                offset = update["update_id"] + 1
                # Deduplication: skip if already processed (handles brief 409 overlaps)
                uid = update["update_id"]
                if uid in processed_updates:
                    continue
                processed_updates.add(uid)
                if len(processed_updates) > 500:
                    processed_updates.clear()
                try:
                    if "message" in update:
                        msg = update["message"]
                        chat_id = msg["chat"]["id"]
                        # Handle photo uploads
                        if "voice" in msg or "audio" in msg:
                            # Voice note — transcribe via Whisper
                            user_state.setdefault(chat_id, {"stage": "idle"})
                            voice_obj = msg.get("voice") or msg.get("audio")
                            handle_voice_message(chat_id, voice_obj)
                        elif "photo" in msg:
                            user_state.setdefault(chat_id, {"stage": "idle"})
                            stage = user_state[chat_id].get("stage", "idle")
                            content_stages = ["awaiting_report","buffering_report",
                                "awaiting_email_report",
                                "awaiting_social_report","awaiting_ad_theme",
                                "awaiting_lp_context_text","awaiting_ad_existing_upload"]
                            ie_stages = ["ie_awaiting_screenshot_ideas",
                                "ie_awaiting_screenshot_critique",
                                "ie_awaiting_pasted_text",
                                "ie_awaiting_inspiration"]
                            if stage in content_stages:
                                handle_content_file(chat_id, msg["photo"][-1], "image")
                            elif stage in ie_stages:
                                handle_ie_screenshot(chat_id, msg["photo"][-1], stage)
                            else:
                                handle_ds_file(chat_id, msg["photo"][-1], "image")
                        elif "document" in msg:
                            user_state.setdefault(chat_id, {"stage": "idle"})
                            stage = user_state[chat_id].get("stage", "idle")
                            doc = msg["document"]
                            mime = doc.get("mime_type", "")
                            content_stages = ["awaiting_report","buffering_report",
                                "awaiting_social_report","awaiting_ad_theme",
                                "awaiting_lp_context_text","awaiting_ad_existing_upload"]
                            ie_stages_doc = ["ie_awaiting_screenshot_ideas",
                                "ie_awaiting_screenshot_critique",
                                "ie_awaiting_pasted_text",
                                "ie_awaiting_inspiration"]
                            if stage in content_stages:
                                if "pdf" in mime:
                                    ftype = "pdf"
                                elif "word" in mime or "docx" in mime or "officedocument" in mime:
                                    ftype = "doc"
                                else:
                                    ftype = "doc"
                                handle_content_file(chat_id, doc, ftype)
                            elif stage in ie_stages_doc:
                                handle_ie_screenshot(chat_id, doc, stage)
                            else:
                                ftype = "image" if mime.startswith("image/") else "csv"
                                handle_ds_file(chat_id, doc, ftype)
                        else:
                            handle_message(msg)
                    elif "callback_query" in update:
                        handle_callback(update["callback_query"])
                except Exception as e:
                    print("Handler error:", e, flush=True)
                    print(traceback.format_exc(), flush=True)
        except KeyboardInterrupt:
            print("Stopped.")
            break
        except Exception as e:
            err = str(e)
            print("Poll error:", err, flush=True)
            if "409" in err:
                print("409 Conflict — another instance running. Waiting 15s...", flush=True)
                time.sleep(15)
            elif "400" in err:
                time.sleep(3)
            else:
                time.sleep(5)

# ── MAIN MENU ─────────────────────────────────────────────────────

def show_main_menu(chat_id):
    user_state[chat_id] = {"stage": "idle"}
    keyboard = [
        [{"text": "✍️ Writing Studio",  "callback_data": "open_content_studio"}],
        [{"text": "📊 Data Studio",     "callback_data": "open_data_studio"}],
        [{"text": "💡 Creative Studio", "callback_data": "open_idea_engine"}],
    ]
    send(chat_id, "*Adam AI*\n\nWhat would you like to do?", keyboard)


# ── AD CREATION FLOW ──────────────────────────────────────────────

AVATARS_AD = {
    "general":         ("General", "Broad appeal. Core fear: missing the move. Tone: direct, confident, data-led."),
    "trader":          ("Trader", "Active, needs an edge. Core fear: missing the setup. Tone: fast, sharp, alpha."),
    "investor":        ("Investor", "Long-term, wants conviction. Core fear: wrong long-term call. Tone: measured, premium."),
    "passive_income":  ("Passive Income Seeker", "Wants yield without complexity. Core fear: losing while trying to earn. Tone: calm, reassuring."),
    "portfolio":       ("Portfolio Builder", "Overwhelmed by token choices. Core fear: bad allocation. Tone: clear, structured."),
    "skeptic":         ("Skeptic", "Has been burned. Only buys on receipts. Core fear: being scammed again. Tone: proof-first."),
    "burned":          ("The Burned", "Lost money, wants redemption. Core fear: losing again. Tone: empathetic, offers a path."),
    "student":         ("College Student", "Smart, low funds, wants escape. Core fear: missing window while broke. Tone: energetic, aspirational."),
    "nine_to_five":    ("Burned-Out 9-5 Worker", "Wants an exit path. Core fear: trapped forever. Tone: emotional, freedom-focused."),
    "boomer":          ("Boomer Near Retirement", "Risk-averse, wants returns. Core fear: losing retirement savings. Tone: conservative, proof-heavy."),
    "side_hustle":     ("Side-Hustle Seeker", "Wants money now, clear steps. Core fear: wasting time. Tone: action-oriented, simple."),
    "beginner":        ("Complete Beginner", "Zero knowledge, overwhelmed. Core fear: getting it wrong. Tone: warm, zero jargon."),
    "chaser":          ("100X Chaser", "Degen, FOMO-driven. Core fear: missing the next 100X. Tone: high energy, exclusive access.")
}

FUNNEL_STAGES = {
    "awareness":      "AWARENESS: Problem/curiosity focus. No price mention. Stop the scroll. Make them feel the problem or the opportunity.",
    "consideration":  "CONSIDERATION: Education and proof. Build authority with track record. Use specific wins. Soft CTA.",
    "conversion":     "CONVERSION: Clear offer. Price $1,197/year. Urgency. Transformation CTA. Make the math obvious."
}

AD_LOGIC_PROMPT = """
CRYPTONARY TRACK RECORD (use these specifics):
- Called SOL at $10 (now $290+), WIF at $0.004 (1,200X), POPCAT (660X), SPX (227X), HYPE airdrop avg $28,500
- 7+ years, 3 full cycles since 2017
- 300,000+ subscribers, 12,000+ Pro members

COPYWRITING PRINCIPLES TO APPLY:
- Hormozi: Sell transformation not product. Stack specific value bullets.
- Cashvertising LF8: Lead with the avatar's dominant desire/fear.
- Cialdini: Social proof (show what members DO), Authority (track record), Scarcity (information gap).
- Ogilvy: Headline is everything. Specifics beat generalities. One CTA.
"""

def show_avatar_menu(chat_id, page=0):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_avatars", [])
    state["avatar_page"] = page
    all_avatars = list(AVATARS_AD.items())
    page_size = 7
    start = page * page_size
    page_avatars = all_avatars[start:start + page_size]
    keyboard = []
    for key, (name, _) in page_avatars:
        is_sel = key in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + name, "callback_data": "adavatar_" + key}])
    nav = []
    if page > 0:
        nav.append({"text": "Previous", "callback_data": "avatarpage_prev"})
    if start + page_size < len(all_avatars):
        nav.append({"text": "Next page", "callback_data": "avatarpage_next"})
    if nav:
        keyboard.append(nav)
    keyboard.append([{"text": "Done — " + str(len(selected)) + " selected", "callback_data": "adavatars_done"}])
    page_label = "(" + str(page+1) + "/2)" if len(all_avatars) > page_size else ""
    send(chat_id, "*Pick avatars:* " + page_label + "\n_(tap to select multiple)_", keyboard)

def show_stage_menu(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_stages", [])
    keyboard = []
    for key in ["awareness", "consideration", "conversion"]:
        is_sel = key in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + key.capitalize(), "callback_data": "adstage_" + key}])
    keyboard.append([{"text": "Done — " + str(len(selected)) + " selected", "callback_data": "adstages_done"}])
    send(chat_id, "*Pick funnel stages:*\n_(tap to select multiple)_", keyboard)

def show_existing_ad_action_menu(chat_id):
    """Show action options when an existing ad is uploaded as base."""
    keyboard = [
        [{"text": "Refine this ad", "callback_data": "existing_ad_refine"}],
        [{"text": "Create a variation", "callback_data": "existing_ad_variation"}],
        [{"text": "Reconceptualise for an avatar", "callback_data": "existing_ad_reconceptualise"}],
        [{"text": "Generate primary text only", "callback_data": "existing_ad_primary_text"}],
        [{"text": "Generate headlines only", "callback_data": "existing_ad_headlines"}]
    ]
    send(chat_id, "*Existing Ad Loaded*\n\nWhat would you like to do with it?", keyboard)

def generate_existing_ad_action(chat_id, action):
    """Generate based on existing ad as creative base."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    existing = state.get("existing_ad_content", "")
    product_context = state.get("ad_product_context", "")
    avatar_keys = state.get("selected_avatars", [])
    avatar_detail = ", ".join([AVATARS_AD.get(k, ("",))[0] for k in avatar_keys]) if avatar_keys else "General"

    action_prompts = {
        "refine": "Refine this ad. Keep the core concept but improve the copy, hook, and CTA. Apply all copywriting principles. Make it stronger.",
        "variation": "Create 3 distinct variations of this ad. Each should take a different angle, hook style, or emotional approach while keeping the same product/offer.",
        "reconceptualise": "Reconceptualise this ad specifically for the " + avatar_detail + " avatar. Rewrite the hook, body, and CTA to speak directly to their specific pain points, desires, and language.",
        "primary_text": "Using this ad as the creative brief, generate 3 variations of the primary text body copy only. Optimise for the " + avatar_detail + " avatar.",
        "headlines": "Using this ad as the creative brief, generate 8 headline variations. Mix: curiosity, fear/loss, data-led, identity, contrarian. Optimise for " + avatar_detail + ".",
    }

    instruction = action_prompts.get(action, "Improve this ad.")
    send(chat_id, "Working on it...")

    try:
        prompt = "EXISTING AD CREATIVE:\n" + existing[:1500]
        if product_context: prompt += "\n\nPRODUCT: " + product_context
        prompt += "\n\nAVATAR TARGET: " + avatar_detail
        prompt += "\n\nINSTRUCTION: " + instruction
        prompt += "\n\nApply all copywriting principles. Return as plain text."

        result = claude(prompt, max_tokens=1800, system=AD_VOICE)
        state["current_ad_output"] = result
        send_plain(chat_id, result)
        keyboard = [
            [{"text": "Quick Edit", "callback_data": "ad_quick_edit"},
             {"text": "Critique", "callback_data": "critique_ad"}],
            [{"text": "Try another action", "callback_data": "existing_ad_back"}],
            [{"text": "Mark Complete", "callback_data": "mark_complete"}]
        ]
        send(chat_id, "Done.", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def show_ad_product_menu(chat_id):
    """Step 1: Which product is this ad for?"""
    keyboard = [
        [{"text": "Cryptonary Pro — $1,197/year", "callback_data": "ad_product_pro"}],
        [{"text": "Inner Circle — $15K–$22K", "callback_data": "ad_product_ic"}]
    ]
    send(chat_id, "*Ad Copy Generator*\n\nStep 1: Which product is this ad for?", keyboard)

def show_ad_input_menu(chat_id):
    """Step 2: New ad or upload existing?"""
    product_label = user_state[chat_id].get("ad_product_label", "the product")
    keyboard = [
        [{"text": "Start from scratch", "callback_data": "ad_input_new"}],
        [{"text": "Upload existing ad as base", "callback_data": "ad_input_existing"}]
    ]
    send(chat_id, "*" + product_label + "*\n\nStart from scratch or use an existing ad as the creative base?", keyboard)

def show_adtype_menu(chat_id):
    keyboard = [
        [{"text": "Static — 3 copy variants", "callback_data": "adtype_static"}],
        [{"text": "Video — AIDA script + 3 hook variants", "callback_data": "adtype_video"}]
    ]
    send(chat_id, "*Static or Video?*", keyboard)

def toggle_avatar(chat_id, avatar_key, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_avatars", [])
    if avatar_key in selected: selected.remove(avatar_key)
    else: selected.append(avatar_key)
    state["selected_avatars"] = selected
    page = state.get("avatar_page", 0)
    all_avatars = list(AVATARS_AD.items())
    page_size = 7
    start = page * page_size
    page_avatars = all_avatars[start:start + page_size]
    keyboard = []
    for key, (name, _) in page_avatars:
        is_sel = key in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + name, "callback_data": "adavatar_" + key}])
    nav = []
    if page > 0:
        nav.append({"text": "Previous", "callback_data": "avatarpage_prev"})
    if start + page_size < len(all_avatars):
        nav.append({"text": "Next page", "callback_data": "avatarpage_next"})
    if nav:
        keyboard.append(nav)
    keyboard.append([{"text": "Done — " + str(len(selected)) + " selected", "callback_data": "adavatars_done"}])
    tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id, "reply_markup": {"inline_keyboard": keyboard}})

def toggle_stage(chat_id, stage_key, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_stages", [])
    if stage_key in selected: selected.remove(stage_key)
    else: selected.append(stage_key)
    state["selected_stages"] = selected
    keyboard = []
    for key in ["awareness", "consideration", "conversion"]:
        is_sel = key in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + key.capitalize(), "callback_data": "adstage_" + key}])
    keyboard.append([{"text": "Done — " + str(len(selected)) + " selected", "callback_data": "adstages_done"}])
    tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id, "reply_markup": {"inline_keyboard": keyboard}})

def generate_all_ads(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    avatars = state.get("selected_avatars", [])
    stages = state.get("selected_stages", [])
    ad_type = state.get("ad_type", "static")
    theme = sanitise(state.get("ad_theme", ""))
    state["current_ad_output"] = ""  # Reset for fresh generation

    if not avatars or not stages:
        send(chat_id, "Please select at least one avatar and one funnel stage.")
        return

    total = len(avatars) * len(stages)
    send(chat_id, "Generating " + str(total) + " ad set(s)... This may take a moment.")

    for avatar_key in avatars:
        avatar_name, avatar_desc = AVATARS_AD.get(avatar_key, ("General", ""))
        for stage_key in stages:
            stage_instruction = FUNNEL_STAGES.get(stage_key, "")
            try:
                if ad_type == "static":
                    prompt = "Write 3 Meta static ad variants for Cryptonary Pro.\n\nAVATAR: " + avatar_name + "\n" + avatar_desc
                    if theme: prompt += "\n\nCAMPAIGN THEME/CONTEXT:\n" + theme
                    prompt += "\n\nFUNNEL STAGE: " + stage_instruction
                    prompt += "\n\n" + AD_LOGIC_PROMPT
                    prompt += "\n\nFor EACH of the 3 variants write:\nVARIANT N:\nHEADLINE: [max 10 words, scroll-stopping]\nPRIMARY TEXT: [150-200 words, leads with avatar emotion, builds to CTA]\nDESCRIPTION: [max 20 words, reinforces headline]\nCTA BUTTON: [e.g. Learn More / Get Started / Join Now]\n\nThen after all 3 variants:\nLOGIC:\nHook mechanism: [what stops the scroll]\nLF8 desire: [which life force desire and why]\nCialdini principle: [which one and why]\nFunnel logic: [why this stage approach]\n\nReturn as plain text."
                    voice_ex = get_voice_corpus_context(chat_id)
                    if voice_ex: prompt = voice_ex + "\n\n" + prompt
                    raw = claude(prompt, max_tokens=1800)
                    header = "*AD SET — " + avatar_name.upper() + " | " + stage_key.upper() + " | STATIC*\n\n"
                    send_plain(chat_id, header + raw)
                    # Accumulate output for critique/voice learning
                    state["current_ad_output"] = state.get("current_ad_output", "") + "\n\n" + header + raw[:600]
                else:
                    # Video: AIDA script + 3 hook variants
                    prompt = "Write a 30-45 second Meta video ad script for Cryptonary Pro using the AIDA formula.\n\nAVATAR: " + avatar_name + "\n" + avatar_desc
                    if theme: prompt += "\n\nCAMPAIGN THEME/CONTEXT:\n" + theme
                    prompt += "\n\nFUNNEL STAGE: " + stage_instruction
                    prompt += "\n\n" + AD_LOGIC_PROMPT
                    prompt += "\n\nSTRUCTURE:\n\nATTENTION (0-3 seconds): Pattern interrupt hook. Must stop the scroll.\n\nINTEREST (3-12 seconds): Agitate the problem or amplify the desire. Make them feel it.\n\nDESIRE (12-28 seconds): The solution. Specific proof points. Transformation.\n\nACTION (28-35 seconds): Clear CTA. What to do next.\n\nFormat each section as:\n[SECTION NAME]\nSPOKEN: [voiceover text]\nON SCREEN: [text overlays]\nVISUAL: [scene direction]\n\nThen write 3 ALTERNATIVE HOOKS (just the Attention section, different each time — pattern interrupt, question, bold claim):\n\nHOOK VARIANT 1:\nSPOKEN: ...\nON SCREEN: ...\nVISUAL: ...\n\nHOOK VARIANT 2: ...\nHOOK VARIANT 3: ...\n\nThen:\nLOGIC:\nHook mechanism: [what stops the scroll]\nLF8 desire: [which life force desire and why]\nCialdini principle: [which one and why]\nFunnel logic: [why this stage approach]\nAIDA breakdown: [one sentence per section on what it does psychologically]\n\nReturn as plain text."
                    raw = claude(prompt, max_tokens=2000)
                    header = "*AD SET — " + avatar_name.upper() + " | " + stage_key.upper() + " | VIDEO*\n\n"
                    send_plain(chat_id, header + raw)
                    # Accumulate output for critique/voice learning
                    state["current_ad_output"] = state.get("current_ad_output", "") + "\n\n" + header + raw[:600]
            except Exception as e:
                send(chat_id, "Error generating " + avatar_name + " / " + stage_key + ": " + str(e))

    state["stage"] = "ads_ready"
    # Silently save content metadata + copy for voice learning
    ad_output = state.get("current_ad_output", "")
    save_content_metadata(chat_id, "ad", {
        "avatars": state.get("selected_avatars", []),
        "stages": state.get("selected_stages", []),
        "ad_type": state.get("ad_type", "static"),
        "theme": state.get("ad_theme", "")[:200],
        "ad_body": ad_output[:600]
    })
    if ad_output:
        save_voice_example(chat_id, ad_output[:600], "approved_ad")
    send(chat_id, "All " + str(total) + " ad set(s) generated.", ad_action_keyboard())


# ── STANDALONE SOCIAL FLOW ────────────────────────────────────────

def gen_social_angles(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    selected_formats = state.get("selected_social_formats", [])
    format_labels = {"fmt_reel": "Reel", "fmt_carousel": "Carousel", "fmt_story_single": "Story (single)", "fmt_story_multi": "Story (multi)"}
    formats_str = ", ".join([format_labels.get(f, f) for f in selected_formats]) if selected_formats else "social content"
    send(chat_id, "Finding angles...")
    try:
        raw = claude(
            "REPORT:\n" + report +
            "\n\nFORMAT(S) SELECTED: " + formats_str +
            "\n\nGenerate exactly 4 distinct content angles for this crypto market update.\n" +
            "Each angle should be different — emotional lens, hook style, or narrative approach.\n" +
            "Make them specific to what works for " + formats_str + " format(s) on social media.\n\n" +
            "Format:\n1. [angle]\n2. [angle]\n3. [angle]\n4. [angle]\nNothing else.",
            max_tokens=500
        )
        angles = parse_numbered_list(raw, 4)
        state["social_angles"] = angles
        state["stage"] = "pick_social_angle"
        text = "*Pick a content angle:*\n\n"
        keyboard = []
        for i, a in enumerate(angles):
            text += str(i+1) + ". " + a + "\n\n"
            keyboard.append([{"text": str(i+1) + ". " + a[:50], "callback_data": "social_angle_" + str(i)}])
        keyboard.append([{"text": "Write my own angle", "callback_data": "social_custom_angle"}])
        keyboard.append([{"text": "Regenerate angles", "callback_data": "social_regen_angles"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "idle"

def gen_social_hooks(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    report = sanitise(state.get("report", ""))
    angle = state.get("social_angle", "")
    formats = state.get("selected_social_formats", [])
    send(chat_id, "Writing hooks for your selected formats...")

    format_hook_instructions = {
        "fmt_reel": "REEL HOOK: Write 4 alternative opening spoken lines (first 3 seconds of a video). Must stop the scroll immediately. Punchy, direct, creates instant curiosity or tension. Max 15 words each.",
        "fmt_carousel": "CAROUSEL COVER: Write 4 alternative cover slide headlines. Bold statement that makes someone swipe. Max 8 words each. Think: headline of a magazine, not a sentence.",
        "fmt_static": "STATIC HEADLINE: Write 4 alternative on-screen headlines for a static Instagram post. Max 6 words each. Bold, punchy, newspaper front page energy. This is the ONLY text on the image.",
        "fmt_story_single": "STORY HOOK: Write 4 alternative single-slide story text options. Immediate, bold, designed for someone who's already engaged. Max 12 words each.",
        "fmt_story_multi": "STORY OPENING: Write 4 alternative opening slide lines for a multi-slide story. Creates a cliffhanger that pulls to the next slide. Max 12 words each."
    }

    format_labels = {
        "fmt_reel": "Reel",
        "fmt_carousel": "Carousel",
        "fmt_static": "Static Post",
        "fmt_story_single": "Story (single)",
        "fmt_story_multi": "Story (multi)"
    }

    state["social_hooks"] = {}
    state["selected_social_hooks"] = {}
    state["stage"] = "pick_social_hooks"

    for fmt in formats:
        if fmt not in format_hook_instructions:
            continue
        instruction = format_hook_instructions[fmt]
        label = format_labels.get(fmt, fmt)
        try:
            raw = claude(
                "REPORT:\n" + report[:500] +
                "\nANGLE: " + angle +
                "\n\n" + instruction +
                "\n\nFormat:\n1. [hook]\n2. [hook]\n3. [hook]\n4. [hook]\nNothing else.",
                max_tokens=400
            )
            hooks = parse_numbered_list(raw, 4)
            state["social_hooks"][fmt] = hooks
        except Exception as e:
            state["social_hooks"][fmt] = ["Hook option 1", "Hook option 2", "Hook option 3", "Hook option 4"]

    # Show hook picker for first format
    show_social_hook_picker(chat_id, 0)

def show_social_hook_picker(chat_id, fmt_idx):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    formats = state.get("selected_social_formats", [])
    if fmt_idx >= len(formats):
        # All hooks selected - show format menu confirmation and generate
        show_standalone_social_menu_confirm(chat_id)
        return

    fmt = formats[fmt_idx]
    state["current_hook_fmt_idx"] = fmt_idx
    format_labels = {
        "fmt_reel":         "Reel",
        "fmt_carousel":     "Carousel",
        "fmt_static":       "Static Post",
        "fmt_story_single": "Story (single)",
        "fmt_story_multi":  "Story (multi)"
    }
    label = format_labels.get(fmt, fmt)
    hooks = state.get("social_hooks", {}).get(fmt, [])

    # If no hooks for this format, skip to next (shouldn\'t happen but prevents freeze)
    if not hooks:
        show_social_hook_picker(chat_id, fmt_idx + 1)
        return

    remaining = len(formats) - fmt_idx
    text = "*Pick a hook for " + label + ":*"
    if remaining > 1:
        text += " (" + str(remaining - 1) + " more after this)"
    text += "\n\n"
    keyboard = []
    for i, h in enumerate(hooks):
        text += str(i+1) + ". " + h + "\n\n"
        keyboard.append([{"text": str(i+1) + ". " + h[:50], "callback_data": "social_hook_" + str(fmt_idx) + "_" + str(i)}])
    send(chat_id, text, keyboard)

def show_standalone_social_menu_confirm(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    formats = state.get("selected_social_formats", [])
    selected_hooks = state.get("selected_social_hooks", {})
    format_labels = {
        "fmt_reel":         "Reel Script",
        "fmt_carousel":     "Carousel",
        "fmt_static":       "Static Post",
        "fmt_story_single": "Story (single)",
        "fmt_story_multi":  "Story (multi)"
    }
    summary = "*Ready to generate:*\n\n"
    for fmt in formats:
        hook = selected_hooks.get(fmt, "")
        summary += format_labels.get(fmt, fmt) + "\n"
        if hook:
            summary += "_Hook: " + hook[:60] + "_\n"
        summary += "\n"
    keyboard = [[{"text": "Generate all", "callback_data": "gen_social_confirmed"}]]
    send(chat_id, summary, keyboard)

def show_standalone_social_menu(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    selected = state.get("selected_social_formats", [])
    formats = [
        ("Reel Script (45-60s)", "fmt_reel"),
        ("Carousel (5-8 slides)", "fmt_carousel"),
        ("Static Post + Caption", "fmt_static"),
        ("Story — Single slide", "fmt_story_single"),
        ("Story — Multi slide", "fmt_story_multi"),
    ]
    keyboard = []
    for label, cb in formats:
        is_sel = cb in selected
        keyboard.append([{"text": ("[x] " if is_sel else "[ ] ") + label, "callback_data": cb}])
    keyboard.append([{"text": "Generate selected (" + str(len(selected)) + ")", "callback_data": "gen_social_selected"}])
    send(chat_id, "*Select formats to generate:*\n_(tap to select multiple)_", keyboard)



# ── FEEDBACK LOOP DATA STORE ──────────────────────────────────────
# email_log: {chat_id: [{subject, preview, angle, cta, email_type, date, open_rate, ctr, ab_variable, ab_label}]}
# ad_log:    {chat_id: [{ad_type, avatar, stage, headline, date, ...metrics}]}

email_log = {}
ad_log = {}


# ── EMAIL LOGGING ─────────────────────────────────────────────────

def start_log_email(chat_id):
    user_state[chat_id]["log_stage"] = "email_subject"
    user_state[chat_id]["stage"] = "logging_email"
    user_state[chat_id]["log_data"] = {}
    send(chat_id, "*Log Email Performance*\n\nEnter the subject line of the email:")

def handle_email_log_step(chat_id, text):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    log = state.get("log_data", {})
    step = state.get("log_stage", "")

    if step == "email_subject":
        log["subject"] = text
        state["log_stage"] = "email_date"
        send(chat_id, "Date sent? (e.g. 2026-03-22)")

    elif step == "email_date":
        log["date"] = text
        state["log_stage"] = "email_open"
        send(chat_id, "Open rate? (enter number, e.g. 28.4)")

    elif step == "email_open":
        try:
            log["open_rate"] = float(text)
            state["log_stage"] = "email_ctr"
            send(chat_id, "Click-through rate? (enter number, e.g. 3.2)")
        except:
            send(chat_id, "Please enter a number, e.g. 28.4")

    elif step == "email_ctr":
        try:
            log["ctr"] = float(text)
            state["log_stage"] = "email_ab"
            keyboard = [
                [{"text": "Yes — part of a split test", "callback_data": "email_log_ab_yes"}],
                [{"text": "No — standalone email", "callback_data": "email_log_ab_no"}]
            ]
            send(chat_id, "Was this part of a split test?", keyboard)
        except:
            send(chat_id, "Please enter a number, e.g. 3.2")

    elif step == "email_ab_variable":
        log["ab_variable"] = text
        state["log_stage"] = "email_ab_label"
        send(chat_id, "What variant is this? (e.g. 'name in subject' or 'no name in subject')")

    elif step == "email_ab_label":
        log["ab_label"] = text
        state["log_stage"] = "email_recipients"
        send(chat_id, "How many recipients received this email? (enter number)")

    elif step == "email_recipients":
        try:
            log["recipients"] = int(text.replace(",",""))
            save_email_log(chat_id, log)
        except:
            send(chat_id, "Please enter a number, e.g. 12500")

def save_email_log(chat_id, log):
    if chat_id not in email_log:
        email_log[chat_id] = []
    email_log[chat_id].append(log)
    save_all_data()
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "idle"
    records = email_log[chat_id]
    opens = [r["open_rate"] for r in records if "open_rate" in r]
    ctrs = [r["ctr"] for r in records if "ctr" in r]
    avg_open = round(sum(opens)/len(opens), 1) if opens else 0
    avg_ctr = round(sum(ctrs)/len(ctrs), 1) if ctrs else 0
    msg = "*Email logged.*\n\n"
    msg += "Subject: _" + log.get("subject","") + "_\n"
    msg += "Open: " + str(log.get("open_rate","?")) + "% | CTR: " + str(log.get("ctr","?")) + "%\n\n"
    msg += "Your averages (" + str(len(records)) + " emails):\n"
    msg += "Open: " + str(avg_open) + "% | CTR: " + str(avg_ctr) + "%"
    if log.get("ab_variable"):
        msg += "\n\nTagged as split test: _" + log["ab_variable"] + "_ — variant: _" + log.get("ab_label","") + "_"
    keyboard = [[{"text": "Back to menu", "callback_data": "start_over"}]]
    send(chat_id, msg, keyboard)


# ── AD LOGGING ────────────────────────────────────────────────────

def start_log_ad(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "logging_ad"
    state["log_data"] = {}
    keyboard = [
        [{"text": "Video ad", "callback_data": "adlog_type_video"}],
        [{"text": "Static ad", "callback_data": "adlog_type_static"}]
    ]
    send(chat_id, "*Log Ad Performance*\n\nVideo or Static?", keyboard)

def handle_ad_log_step(chat_id, text):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    log = state.get("log_data", {})
    step = state.get("log_stage", "")

    if step == "adlog_headline":
        log["headline"] = text
        state["log_stage"] = "adlog_avatar"
        keyboard = [[{"text": name, "callback_data": "adlog_avatar_" + key}] for key, (name, _) in AVATARS_AD.items()]
        send(chat_id, "Which avatar was this ad for?", keyboard)

    elif step == "adlog_date":
        log["date"] = text
        ad_type = log.get("ad_type", "static")
        if ad_type == "video":
            state["log_stage"] = "adlog_v_attention"
            send(chat_id, "3-second view rate? (Attention — 3s views / impressions as %, e.g. 12.5)")
        else:
            state["log_stage"] = "adlog_s_cpc"
            send(chat_id, "Cost per click? (CPC in GBP/USD, e.g. 0.85)")

    # VIDEO METRICS
    elif step == "adlog_v_attention":
        try:
            log["attention"] = float(text)
            state["log_stage"] = "adlog_v_interest"
            send(chat_id, "Average watch time in seconds? (Interest, e.g. 18)")
        except: send(chat_id, "Enter a number, e.g. 12.5")

    elif step == "adlog_v_interest":
        try:
            log["interest"] = float(text)
            state["log_stage"] = "adlog_v_desire"
            send(chat_id, "Outbound CTR? (Desire — %, e.g. 2.4)")
        except: send(chat_id, "Enter a number, e.g. 18")

    elif step == "adlog_v_desire":
        try:
            log["desire"] = float(text)
            state["log_stage"] = "adlog_v_purchases"
            send(chat_id, "Number of purchases (Action)? (e.g. 3)")
        except: send(chat_id, "Enter a number, e.g. 2.4")

    elif step == "adlog_v_purchases":
        try:
            log["purchases"] = int(text)
            state["log_stage"] = "adlog_v_cpp"
            send(chat_id, "Cost per purchase? (e.g. 45.00)")
        except: send(chat_id, "Enter a number, e.g. 3")

    elif step == "adlog_v_cpp":
        try:
            log["cost_per_purchase"] = float(text)
            save_ad_log(chat_id, log)
        except: send(chat_id, "Enter a number, e.g. 45.00")

    # STATIC METRICS
    elif step == "adlog_s_cpc":
        try:
            log["cpc"] = float(text)
            state["log_stage"] = "adlog_s_ctr"
            send(chat_id, "Outbound CTR? (%, e.g. 1.8)")
        except: send(chat_id, "Enter a number, e.g. 0.85")

    elif step == "adlog_s_ctr":
        try:
            log["outbound_ctr"] = float(text)
            state["log_stage"] = "adlog_s_purchases"
            send(chat_id, "Number of purchases? (e.g. 5)")
        except: send(chat_id, "Enter a number, e.g. 1.8")

    elif step == "adlog_s_purchases":
        try:
            log["purchases"] = int(text)
            state["log_stage"] = "adlog_s_cpp"
            send(chat_id, "Cost per purchase? (e.g. 38.50)")
        except: send(chat_id, "Enter a number, e.g. 5")

    elif step == "adlog_s_cpp":
        try:
            log["cost_per_purchase"] = float(text)
            save_ad_log(chat_id, log)
        except: send(chat_id, "Enter a number, e.g. 38.50")

def save_ad_log(chat_id, log):
    if chat_id not in ad_log:
        ad_log[chat_id] = []
    ad_log[chat_id].append(log)
    save_all_data()
    user_state[chat_id]["stage"] = "idle"
    msg = "*Ad logged.*\n\n"
    msg += "Headline: _" + log.get("headline","")[:60] + "_\n"
    msg += "Avatar: " + log.get("avatar","?") + " | Stage: " + log.get("stage","?") + "\n"
    if log.get("ad_type") == "video":
        msg += "Attention: " + str(log.get("attention","?")) + "% | Watch time: " + str(log.get("interest","?")) + "s | CTR: " + str(log.get("desire","?")) + "%\n"
    else:
        msg += "CPC: " + str(log.get("cpc","?")) + " | CTR: " + str(log.get("outbound_ctr","?")) + "%\n"
    msg += "Purchases: " + str(log.get("purchases","?")) + " | CPP: " + str(log.get("cost_per_purchase","?"))
    msg += "\n\nTotal ads logged: " + str(len(ad_log[chat_id]))
    keyboard = [[{"text": "Back to menu", "callback_data": "start_over"}]]
    send(chat_id, msg, keyboard)


# ── ANALYSIS ─────────────────────────────────────────────────────

def run_email_analysis(chat_id):
    records = email_log.get(chat_id, [])
    if len(records) < 2:
        send(chat_id, "Need at least 2 logged emails to analyse. Log more with /logemail")
        return
    send(chat_id, "Analysing " + str(len(records)) + " emails...")

    # Build summary for Claude
    summary = "EMAIL PERFORMANCE DATA:\n\n"
    for i, r in enumerate(records):
        summary += str(i+1) + ". Subject: " + r.get("subject","") + "\n"
        summary += "   Date: " + r.get("date","?") + " | Open: " + str(r.get("open_rate","?")) + "% | CTR: " + str(r.get("ctr","?")) + "%\n"
        if r.get("ab_variable"):
            summary += "   Split test: " + r["ab_variable"] + " — variant: " + r.get("ab_label","") + " | Recipients: " + str(r.get("recipients","?")) + "\n"
        summary += "\n"

    # Split test aggregation
    ab_groups = {}
    for r in records:
        if r.get("ab_variable") and r.get("ab_label") and r.get("recipients"):
            var = r["ab_variable"]
            label = r["ab_label"]
            if var not in ab_groups: ab_groups[var] = {}
            if label not in ab_groups[var]: ab_groups[var][label] = {"recipients": 0, "opens": 0, "clicks": 0, "count": 0}
            rec = r["recipients"]
            ab_groups[var][label]["recipients"] += rec
            ab_groups[var][label]["opens"] += int(rec * r.get("open_rate",0) / 100)
            ab_groups[var][label]["clicks"] += int(rec * r.get("ctr",0) / 100)
            ab_groups[var][label]["count"] += 1

    if ab_groups:
        summary += "\nSPLIT TEST AGGREGATED RESULTS:\n"
        for var, labels in ab_groups.items():
            summary += "\nVariable: " + var + "\n"
            for label, data in labels.items():
                agg_open = round(data["opens"] / data["recipients"] * 100, 1) if data["recipients"] else 0
                agg_ctr = round(data["clicks"] / data["recipients"] * 100, 1) if data["recipients"] else 0
                summary += "  " + label + ": " + str(agg_open) + "% open / " + str(agg_ctr) + "% CTR (" + str(data["count"]) + " emails, " + str(data["recipients"]) + " total recipients)\n"

    prompt = "Analyse this Cryptonary email performance data and provide:\n\n1. TOP PERFORMERS — top 3 emails and exactly what made them work (subject style, angle type, CTA)\n2. WORST PERFORMERS — bottom 3 and what likely caused underperformance\n3. PATTERN RECOGNITION — what patterns emerge across the data (which subject styles, angles, CTAs consistently outperform)\n4. SPLIT TEST RESULTS — if split test data exists, declare the winner with the aggregated numbers and statistical context\n5. ITERATION IDEAS — for the top performers, give 3 specific variations to test next\n6. IMPROVEMENT SUGGESTIONS — for underperformers, give specific fixes based on the copywriting principles you know\n\n" + summary + "\n\nBe specific and actionable. Reference actual subject lines and numbers."
    try:
        analysis = claude(prompt, max_tokens=2000)
        send_plain(chat_id, "EMAIL PERFORMANCE ANALYSIS\n\n" + analysis)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def run_ad_analysis(chat_id):
    records = ad_log.get(chat_id, [])
    if len(records) < 2:
        send(chat_id, "Need at least 2 logged ads to analyse. Log more with /logad")
        return
    send(chat_id, "Analysing " + str(len(records)) + " ads...")

    video_ads = [r for r in records if r.get("ad_type") == "video"]
    static_ads = [r for r in records if r.get("ad_type") == "static"]

    summary = "AD PERFORMANCE DATA:\n\n"

    if video_ads:
        summary += "VIDEO ADS (" + str(len(video_ads)) + "):\n"
        for i, r in enumerate(video_ads):
            summary += str(i+1) + ". Headline: " + r.get("headline","")[:60] + "\n"
            summary += "   Avatar: " + r.get("avatar","?") + " | Stage: " + r.get("stage","?") + " | Date: " + r.get("date","?") + "\n"
            summary += "   ATTENTION (3s view rate): " + str(r.get("attention","?")) + "%\n"
            summary += "   INTEREST (avg watch time): " + str(r.get("interest","?")) + "s\n"
            summary += "   DESIRE (outbound CTR): " + str(r.get("desire","?")) + "%\n"
            summary += "   ACTION: " + str(r.get("purchases","?")) + " purchases @ " + str(r.get("cost_per_purchase","?")) + " CPP\n\n"

    if static_ads:
        summary += "STATIC ADS (" + str(len(static_ads)) + "):\n"
        for i, r in enumerate(static_ads):
            summary += str(i+1) + ". Headline: " + r.get("headline","")[:60] + "\n"
            summary += "   Avatar: " + r.get("avatar","?") + " | Stage: " + r.get("stage","?") + " | Date: " + r.get("date","?") + "\n"
            summary += "   CPC: " + str(r.get("cpc","?")) + " | Outbound CTR: " + str(r.get("outbound_ctr","?")) + "%\n"
            summary += "   ACTION: " + str(r.get("purchases","?")) + " purchases @ " + str(r.get("cost_per_purchase","?")) + " CPP\n\n"

    prompt = "Analyse this Cryptonary Meta ad performance data and provide:\n\n1. TOP PERFORMERS — best 3 ads with exactly what made them work (avatar, stage, hook type, which metrics stood out)\n2. WORST PERFORMERS — bottom 3 and diagnose where they failed (for video: which AIDA stage had the biggest drop-off? for static: which metric was weakest?)\n3. PATTERN RECOGNITION — what patterns emerge? (which avatars convert best, which stages perform, which ad types win)\n4. VIDEO AIDA DIAGNOSIS — for video ads, map the drop-off: high attention but low interest = hook works but body fails. High desire but low action = landing page issue. Give specific diagnosis per video.\n5. ITERATION IDEAS FOR WINNERS — for top performers, give 3 specific variants to test next\n6. IMPROVEMENT SUGGESTIONS FOR LOSERS — specific creative fixes based on the AIDA failure point or static metric weakness\n\n" + summary + "\n\nBe specific. Reference actual headlines and numbers."
    try:
        analysis = claude(prompt, max_tokens=2000)
        send_plain(chat_id, "AD PERFORMANCE ANALYSIS\n\n" + analysis)
    except Exception as e:
        send(chat_id, "Error: " + str(e))


# ── PERSISTENCE ───────────────────────────────────────────────────

ANALYTICS_FILE = "analytics_data.json"

def save_all_data():
    try:
        data = {"email_log": {str(k): v for k, v in email_log.items()},
                "ad_log": {str(k): v for k, v in ad_log.items()},
                "performance_data": {str(k): v for k, v in performance_data.items()}}
        with open(ANALYTICS_FILE, "w") as f:
            json.dump(data, f)
    except Exception as e:
        print("Analytics save error:", e)
    save_content_stores()

def load_all_data():
    global email_log, ad_log, performance_data
    try:
        with open(ANALYTICS_FILE, "r") as f:
            data = json.load(f)
        email_log = {int(k): v for k, v in data.get("email_log", {}).items()}
        ad_log = {int(k): v for k, v in data.get("ad_log", {}).items()}
        performance_data = {int(k): v for k, v in data.get("performance_data", {}).items()}
    except:
        pass

load_all_data()

BRANDSCRIPT_PROMPT = """
You are writing a Cryptonary landing page using the StoryBrand BrandScript framework.

BRANDSCRIPT RULES:
- The READER is the hero, not Cryptonary
- Cryptonary is the GUIDE (wise, empathetic, has a plan)
- Lead with the reader's dominant problem/desire, not our features
- Every section earns the scroll to the next
- One CTA, repeated at key moments

COPY PRINCIPLES:
- Hormozi: Stack specific value bullets before the ask. Make the math obvious. Transformation not subscription.
- Cashvertising LF8: Lead with the avatar's dominant life force desire. Fear works when threat is real + solution is clear.
- Cialdini: Social proof = what members DO (not just what they get). Authority = track record with specific wins. Scarcity = information gap.
- Ogilvy: Headline is a promise. Subheads carry the narrative. Specifics beat generalities. One CTA, zero friction.

CRYPTONARY PROOF POINTS (use these):
- Called SOL at $10 (now $290+), WIF at $0.004 (1,200X), POPCAT (660X), SPX (227X)
- HYPE airdrop average $28,500 for members
- 300,000+ newsletter subscribers, 12,000+ Pro members
- 7+ years, 3 full crypto cycles since 2017
- Pro: $1,197/year | Inner Circle: $15K-$22K (portfolios $200K+)

GRAPHIC RECOMMENDATION FORMAT (for each section):
VISUAL DIRECTION: [What to show, why it works psychologically, where to place it]
AI IMAGE PROMPT: [Ready-to-use Midjourney/DALL-E prompt, photorealistic style]
"""

LP_CTA_DEFS = {
    "pro": {
        "label": "Join Cryptonary Pro",
        "price": "$1,197/year",
        "cta_text": "Get Pro Access Now",
        "alt_cta": "Start Making Better Trades",
        "audience": "crypto investors and traders wanting research-backed guidance",
        "positioning": "The research platform that puts you ahead of the market. Not a signal service — a complete intelligence layer.",
        "urgency": "Markets don't wait. Every day without this costs you positioning."
    },
    "inner_circle": {
        "label": "Apply for Inner Circle",
        "price": "$15,000-$22,000",
        "cta_text": "Apply for Inner Circle",
        "alt_cta": "Apply Now — Limited Spots",
        "audience": "high-net-worth crypto investors with $200K+ portfolios",
        "positioning": "Private advisory for serious capital. Personal access to Co-Founder Asad for portfolio-level guidance.",
        "urgency": "Inner Circle is application-only. Spots are strictly limited to maintain quality of service."
    }
}

# Inner Circle avatars for landing pages
IC_AVATARS = {
    "ic_universal": ("Universal IC", "High-net-worth investor ($200K+ portfolio) seeking institutional-grade research framework and dedicated team support across crypto cycles. The Cryptonary Framework — from exhausted to structured in 14 days."),
    "ic_exhausted": ("Exhausted Believer", "Survived 3+ cycles, made and lost significant money. Doesn't need conviction — needs a system. Tired of 3am chart-watching. Sitting in stables paralysed or still fully exposed waiting for something to change."),
    "ic_timepoor": ("Time-Poor Professional", "Seven-figure portfolio managed between meetings. Delegates accountant, lawyer, financial advisor — but still managing crypto alone. Information without execution. Reactive not systematic."),
    "ic_liquidity": ("Liquidity Event Winner", "New serious capital from business exit, inheritance, stock options, or ballooned early gains. Doesn't know what they don't know. This money is too important to learn on. Traditional advisors can't help."),
    "ic_family": ("Family Wealth", "Managing family trust or generational capital. Every dip is a boardroom question. Every loss is a personal failure. Needs institutional-grade reporting and a framework they can present at family meetings."),
    "ic_uhnw": ("Crypto-Curious UHNW", "Built serious wealth in traditional markets — equities, real estate, private markets. Understands due diligence and risk. The crypto industry doesn't speak their language. Wealth manager can't help. Opportunity cost compounding."),
}

# Pro avatars for LP flow
PRO_AVATARS_LP = {
    "universal": "Universal — broad crypto investor wanting research and guidance",
    "trader": "Trader — active, wants daily structure, setups, and confirmation",
    "investor": "Investor — long-term, wants conviction and clean long-term thesis",
    "passive": "Passive Income Seeker — wants yield without complexity",
    "portfolio": "Portfolio Builder — overwhelmed by tokens, wants sector picks",
    "chaser": "100X Chaser — degen energy, FOMO-driven, loves early narratives",
    "skeptic": "Skeptic — needs proof, track record, receipts before buying",
    "burned": "The Burned — lost money, wants redemption and trusted guide",
    "student": "College Student — smart, ambitious, price-sensitive",
    "worker": "Burned-Out 9-5 Worker — wants side income and eventual exit path",
    "boomer": "Boomer Near Retirement — safety and yield, risk-averse",
    "sidehustle": "Side-Hustle Seeker — wants clear steps and fast action",
    "beginner": "Complete Beginner — overwhelmed, needs step-by-step structure",
}

# ══════════════════════════════════════════════════════════════════
# LANDING PAGE FLOW — 8-STEP REBUILD
# ══════════════════════════════════════════════════════════════════

def start_landing_page_flow(chat_id):
    """Step 1: Goal selection — CPRO or Inner Circle."""
    user_state[chat_id]["stage"] = "lp_idle"
    user_state[chat_id]["selected_avatars"] = []
    user_state[chat_id]["lp_outline"] = {}
    user_state[chat_id]["lp_full_copy"] = {}
    keyboard = [
        [{"text": "CPRO Sales Page", "callback_data": "lp_goal_pro"}],
        [{"text": "Inner Circle Sales Page", "callback_data": "lp_goal_ic"}]
    ]
    send(chat_id, "*Landing Page Builder*\n\nStep 1: What is the goal of this page?", keyboard)

def show_lp_avatar_menu(chat_id, page=0):
    """Step 2: Avatar selection — different options for CPRO vs IC."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    goal = state.get("lp_goal", "pro")
    selected = state.get("selected_avatars", [])
    state["avatar_page"] = page

    if goal == "ic":
        avatars = list(IC_AVATARS.items())
        title = "*Step 2: Which Inner Circle avatar is this page for?*"
    else:
        avatars = list(PRO_AVATARS_LP.items())
        title = "*Step 2: Which avatar is this page for?*"

    page_size = 7
    start = page * page_size
    page_avatars = avatars[start:start + page_size]
    keyboard = []
    for key, val in page_avatars:
        name = val[0] if isinstance(val, tuple) else val.split(" — ")[0]
        is_sel = key in selected
        keyboard.append([{"text": ("✓ " if is_sel else "○ ") + name, "callback_data": "lpavatar_" + key}])
    nav = []
    if page > 0: nav.append({"text": "← Previous", "callback_data": "lpavatarpage_prev"})
    if start + page_size < len(avatars): nav.append({"text": "Next →", "callback_data": "lpavatarpage_next"})
    if nav: keyboard.append(nav)
    keyboard.append([{"text": "Continue with " + str(len(selected)) + " selected →", "callback_data": "lpavatars_done"}])
    send(chat_id, title, keyboard)

def toggle_lp_avatar(chat_id, avatar_key, message_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    goal = state.get("lp_goal", "pro")
    selected = state.get("selected_avatars", [])
    if avatar_key in selected: selected.remove(avatar_key)
    else: selected.append(avatar_key)
    state["selected_avatars"] = selected
    page = state.get("avatar_page", 0)

    if goal == "ic":
        avatars = list(IC_AVATARS.items())
    else:
        avatars = list(PRO_AVATARS_LP.items())

    page_size = 7
    start = page * page_size
    page_avatars = avatars[start:start + page_size]
    keyboard = []
    for key, val in page_avatars:
        name = val[0] if isinstance(val, tuple) else val.split(" — ")[0]
        is_sel = key in selected
        keyboard.append([{"text": ("✓ " if is_sel else "○ ") + name, "callback_data": "lpavatar_" + key}])
    nav = []
    if page > 0: nav.append({"text": "← Previous", "callback_data": "lpavatarpage_prev"})
    if start + page_size < len(avatars): nav.append({"text": "Next →", "callback_data": "lpavatarpage_next"})
    if nav: keyboard.append(nav)
    keyboard.append([{"text": "Continue with " + str(len(selected)) + " selected →", "callback_data": "lpavatars_done"}])
    try:
        tg("editMessageReplyMarkup", {"chat_id": chat_id, "message_id": message_id,
            "reply_markup": json.dumps({"inline_keyboard": keyboard})})
    except:
        show_lp_avatar_menu(chat_id, page)

def generate_lp_outline(chat_id):
    """Step 3: Generate short BrandScript outline (1-2 sentences per section)."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    goal = state.get("lp_goal", "pro")
    avatar_keys = state.get("selected_avatars", [])

    # Build avatar descriptions
    if goal == "ic":
        avatar_descs = []
        for key in avatar_keys:
            val = IC_AVATARS.get(key, ("Universal IC", ""))
            avatar_descs.append(val[0] + ": " + val[1])
        cta_info = "Apply for Inner Circle — $15,000-$22,000/year — application only, limited spots"
        page_type = "Inner Circle Sales Page"
    else:
        avatar_descs = []
        for key in avatar_keys:
            val = PRO_AVATARS_LP.get(key, "General crypto investor")
            avatar_descs.append(val)
        cta_info = "Join Cryptonary Pro — $1,197/year"
        page_type = "CPRO Sales Page"

    avatar_detail = "\n".join(avatar_descs) if avatar_descs else "General investor"
    send(chat_id, "Generating BrandScript outline...")

    try:
        prompt = """Create a SHORT BrandScript outline for a """ + page_type + """ targeting:
""" + avatar_detail + """

CTA: """ + cta_info + """

Write 1-2 sentences for EACH of these 7 sections. This is an outline, not the final copy.
Be specific to the avatar — not generic.

Format EXACTLY as:
1. HERO: [1-2 sentences — what the reader wants and what this page promises]
2. VILLAIN: [1-2 sentences — the specific problem/enemy this avatar faces]
3. GUIDE EMPATHY: [1-2 sentences — how Cryptonary understands their struggle]
4. GUIDE AUTHORITY: [1-2 sentences — specific proof points that earn trust]
5. THE PLAN: [1-2 sentences — how simple the path to the solution is]
6. TRANSFORMATION: [1-2 sentences — who they become / what changes after joining]
7. CTA BLOCK: [1-2 sentences — the ask and why now]

Nothing else. Just the 7 sections."""

        raw = claude(prompt, max_tokens=800, system=BRANDSCRIPT_PROMPT)
        state["lp_outline_raw"] = raw
        state["stage"] = "lp_outline_review"

        send_plain(chat_id, "*BRANDSCRIPT OUTLINE*\n\nReview each section. Send feedback on any numbered section (e.g. \'2: Make the villain more specific to managing family capital\') or approve to continue.\n\n" + raw)

        keyboard = [
            [{"text": "Looks good — continue", "callback_data": "lp_outline_approve"}],
            [{"text": "Regenerate outline", "callback_data": "lp_regen_outline"}]
        ]
        send(chat_id, "Review the outline. Approve or send feedback by section number.", keyboard)

    except Exception as e:
        send(chat_id, "Error: " + str(e))

def apply_lp_outline_feedback(chat_id, feedback_text):
    """Step 4: Apply numbered feedback to outline."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    current_outline = state.get("lp_outline_raw", "")
    send(chat_id, "Applying feedback to outline...")
    try:
        result = claude(
            "CURRENT OUTLINE:\n" + current_outline +
            "\n\nFEEDBACK TO APPLY:\n" + feedback_text +
            "\n\nApply the feedback. Return the full revised outline in the same numbered format. Only change what was specified.",
            max_tokens=800, system=BRANDSCRIPT_PROMPT
        )
        state["lp_outline_raw"] = result
        send_plain(chat_id, "*REVISED OUTLINE*\n\n" + result)
        keyboard = [
            [{"text": "Approve outline — generate full copy", "callback_data": "lp_outline_approve"}],
            [{"text": "More feedback", "callback_data": "lp_more_outline_feedback"}]
        ]
        send(chat_id, "How does the revised outline look?", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def generate_lp_full_copy(chat_id):
    """Step 5: Generate comprehensive BrandScript copy from approved outline."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    goal = state.get("lp_goal", "pro")
    avatar_keys = state.get("selected_avatars", [])
    outline = state.get("lp_outline_raw", "")
    context = sanitise(state.get("lp_context", ""))

    if goal == "ic":
        avatar_descs = [IC_AVATARS.get(k, ("Universal IC", ""))[0] + ": " + IC_AVATARS.get(k, ("", ""))[1] for k in avatar_keys]
        cta = LP_CTA_DEFS.get("inner_circle", LP_CTA_DEFS["pro"])
    else:
        avatar_descs = [PRO_AVATARS_LP.get(k, "General investor") for k in avatar_keys]
        cta = LP_CTA_DEFS.get("pro", LP_CTA_DEFS["pro"])

    avatar_detail = "\n".join(avatar_descs)
    send(chat_id, "Writing comprehensive landing page copy... This will take a moment.")

    # IC-specific copy brief
    ic_copy_brief = ""
    if goal == "ic":
        ic_copy_brief = """
INNER CIRCLE SPECIFIC CONTEXT:
- Product: The Cryptonary Framework — dedicated team of 8, personalised portfolio research
- 8 critical regime calls since 2019 | 129 documented posture shifts | 7,300+ published articles
- Jan 2025: RISK-OFF called at $104,855 — 43% drawdown avoided
- Monthly portfolio audit reports | 20+ hours call time/year | 24hr direct access
- Private dashboard | Security assessment | Airdrop guidance | Yield analysis
- Annual in-person meeting | Starting $15,000/year | $200K+ minimum portfolio
- Tag line: "You maintain full custody. You make every decision. What changes is this: you're no longer deciding alone."
- The 74% Gap framework: BTC $110K → $80K → $110K = 0% vs framework user = 37% gain
"""

    sections = [
        ("HERO", "Write the complete Hero section based on the outline. Include H1 (max 10 words), H2 (max 20 words), hero body (2-3 sentences), primary CTA button, secondary soft CTA, and trust bar stats."),
        ("VILLAIN", "Write the complete Villain/Problem section. Three levels: external problem, internal fear/frustration, philosophical injustice. Make the reader feel deeply understood. This is the longest section — earn the scroll."),
        ("GUIDE — EMPATHY", "Write the Guide Empathy section. Show Cryptonary has been where the reader is. Specific, vulnerable, credible. Not corporate."),
        ("GUIDE — AUTHORITY", "Write the Guide Authority section. Specific proof points only — track record, regime calls, published research, years in market. The guide demonstrates, never brags."),
        ("THE PLAN", "Write the Plan section. 3-4 simple numbered steps. Remove every obstacle. Dead simple."),
        ("TRANSFORMATION", "Write the Transformation section. Before/after table. Who they become. What their life looks like after joining. Paint it vividly."),
        ("VALUE STACK & PRICING", "Write the Value Stack section and pricing reveal. List every deliverable with its standalone value. Make the price look like the obvious decision."),
        ("SOCIAL PROOF", "Write the Social Proof section. Track record wins, member testimonials structure, specific numbers."),
        ("CTA BLOCK", "Write the main CTA block. Restate the transformation. Handle the last objection. Single clear CTA."),
        ("FAQ", "Write 5 FAQs addressing the real objections for this specific avatar. Concise, punchy answers."),
        ("FOOTER CTA + DISCLAIMER", "Write the final footer CTA and include the appropriate disclaimer/risk disclosure for Cryptonary."),
    ]

    all_output = []
    for section_name, section_instruction in sections:
        try:
            user_prompt = "APPROVED OUTLINE:\n" + outline
            user_prompt += "\n\nAVATAR(S):\n" + avatar_detail
            user_prompt += "\n\nCTA: " + cta["label"] + " — " + cta["price"]
            user_prompt += "\n\nPOSITIONING: " + cta["positioning"]
            if ic_copy_brief: user_prompt += ic_copy_brief
            if context: user_prompt += "\n\nEXTRA CONTEXT: " + context
            user_prompt += "\n\nWRITE: " + section_name
            user_prompt += "\n\n" + section_instruction
            user_prompt += "\n\nReturn as plain text. Section name in CAPS as header. Full copy first. Then: GRAPHIC RECOMMENDATIONS: [visual direction] | AI IMAGE PROMPT: [Midjourney/DALL-E prompt]"

            raw = claude(user_prompt, max_tokens=1500, system=BRANDSCRIPT_PROMPT)
            send_plain(chat_id, raw)
            all_output.append(raw)
            time.sleep(0.5)
        except Exception as e:
            send(chat_id, "Error on " + section_name + ": " + str(e))

    state["current_lp"] = "\n\n---\n\n".join(all_output)
    state["lp_pre_edit"] = state["current_lp"]
    state["stage"] = "lp_copy_review"

    keyboard = [
        [{"text": "Quick Edit", "callback_data": "lp_quick_edit"},
         {"text": "Suggest Enhancements", "callback_data": "lp_enhance"}],
        [{"text": "Adjust Length", "callback_data": "lp_length"},
         {"text": "Paste-Back Edit", "callback_data": "lp_paste_back"}],
        [{"text": "Approve — Generate Design Brief", "callback_data": "lp_approve_copy"}],
        [{"text": "Critique", "callback_data": "critique_lp"},
         {"text": "Regenerate a section", "callback_data": "lp_regen"}]
    ]
    send(chat_id, "Full landing page copy complete. Review and refine.", keyboard)

def generate_lp_design_brief(chat_id):
    """Step 8: Generate design brief for UI team based on approved copy."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    lp_content = state.get("current_lp", "")
    goal = state.get("lp_goal", "pro")
    send(chat_id, "Generating design brief for your UI team...")

    try:
        result = claude(
            "Based on this approved landing page copy, create a comprehensive DESIGN BRIEF for the UI/design team.\n\n" +
            "LANDING PAGE COPY:\n" + lp_content[:4000] +
            "\n\nCREATE A DESIGN BRIEF that covers:\n\n" +
            "1. PAGE OVERVIEW: Layout type, scroll depth, overall visual direction\n" +
            "2. SECTION-BY-SECTION DESIGN SPECS: For each section — graphic type (photo/illustration/chart/video), exact content to show, placement, size recommendation, and why it works psychologically\n" +
            "3. TYPOGRAPHY HIERARCHY: Which headlines need special treatment\n" +
            "4. COLOR & MOOD: Per section — dark/light, accent usage, emotional tone\n" +
            "5. MOBILE CONSIDERATIONS: Any sections that need different mobile treatment\n" +
            "6. AI IMAGE PROMPTS: One ready-to-use Midjourney prompt per major section (photorealistic, professional, dark Cryptonary aesthetic)\n" +
            "7. ASSETS NEEDED: Complete list of what the design team needs to source or create\n\n" +
            "Format clearly. This document should be passable directly to a UI designer.",
            max_tokens=2500, system=BRANDSCRIPT_PROMPT
        )
        send_plain(chat_id, "*DESIGN BRIEF — PASS DIRECTLY TO UI TEAM*\n\n" + result)
        keyboard = [
            [{"text": "Mark Complete", "callback_data": "mark_complete"}],
            [{"text": "Build another page", "callback_data": "lp_again"}]
        ]
        send(chat_id, "Design brief complete. Ready to pass to your UI team.", keyboard)

    except Exception as e:
        send(chat_id, "Error generating design brief: " + str(e))

def show_lp_cta_menu(chat_id):
    """Legacy — now handled by goal selection. Kept for backward compat."""
    keyboard = [
        [{"text": "Join Cryptonary Pro ($1,197/year)", "callback_data": "lpcta_pro"}],
        [{"text": "Apply for Inner Circle ($15K-$22K)", "callback_data": "lpcta_inner_circle"}]
    ]
    send(chat_id, "*What is the CTA for this landing page?*", keyboard)

def generate_landing_page(chat_id):
    """Legacy function — now routes to generate_lp_full_copy."""
    generate_lp_full_copy(chat_id)


# ══════════════════════════════════════════════════════════════════
# CRYPTONARY DATA STUDIO
# ══════════════════════════════════════════════════════════════════

DATA_STUDIO_SYSTEM = """
You are a performance analyst for Cryptonary, a crypto research and education platform.
Your job is to analyse marketing performance data extracted from screenshots or CSV files.

ANALYSIS PHILOSOPHY:
- Never just report numbers. Explain WHY something performed well or poorly.
- Look for patterns across the data — topic, format, hook style, avatar, funnel stage.
- Generate specific, actionable ideas based on what the patterns suggest.
- Use the A/B/C/D grading system (A=top 25%, B=above average, C=below average, D=bottom 25%).
- For metrics where lower is better (CPC, CPP), invert the grading.
- Always pool raw numbers for split tests — never average percentages.
- Compare like-for-like: ads vs ads of same avatar/stage, posts vs posts of same format.

GRADING SYSTEM:
A = Top 25% of cohort — Scale / Keep running
B = 26-50% — Above average, monitor
C = 51-75% — Below average, test changes
D = Bottom 25% — Kill or overhaul

OVERALL AD RATINGS:
Mostly A/B = SCALE
Mix A/B/C = KEEP AND MONITOR  
Mix B/C/D = TEST ONE CHANGE
Mostly C/D = KILL

VIDEO AIDA METRICS (correct formula):
A (Attention) = 3-second video views / impressions × 100 = ThruPlay rate
I (Interest) = average watch time (seconds)
D (Desire) = outbound click-through rate %
A (Action) = purchases

ADDITIONAL METRICS: checkouts initiated, cost per purchase, cost per checkout initiated

STATIC AD METRICS: cost per click (replaces 3SVV and watch time), outbound CTR, purchases, CPP, cost per checkout

VIDEO AIDA DIAGNOSIS:
- D on Attention (low 3SVV/impressions) = Hook failing. Change opening 3 seconds.
- A on Attention, D on Interest (low watch time) = Hook works, body fails. Fix middle section.
- A/B on Interest, D on Desire (low outbound CTR) = Watching but not clicking. Weak CTA or offer.
- A on Desire, D on Action (low purchases) = Clicks but no purchase. Landing page problem not the ad.

GRADING RULE: Grade each ad against its own cohort ONLY.
- Videos graded against other videos
- Statics graded against other statics
- NEVER compare video metrics against static metrics — they are fundamentally different

META AD NAMING CONVENTION:
Format: IMG (static) or VID (video)
Stage: AWA (awareness), CDR (consideration), CNV (conversion)
Avatars: UNIVERSAL, TRADER, INVESTOR, PASSIVE_INCOME, PORTFOLIO, SKEPTIC, BURNED, STUDENT, 9_5, BOOMER, SIDE_HUSTLE, BEGINNER, CHASER
Angle: after Msg_

INSTAGRAM ENGAGEMENT RATE:
Engagement Rate = (Likes + Comments + Saves + Shares) / Reach × 100
Grade engagement rate within format cohort (Reels vs Reels etc).
"""

def show_data_studio_menu(chat_id):
    user_state[chat_id] = {"stage": "data_studio_idle"}
    keyboard = [
        [{"text": "Adverts", "callback_data": "ds_adverts"}],
        [{"text": "Social (Instagram)", "callback_data": "ds_social"}],
        [{"text": "Emails", "callback_data": "ds_emails"}],
        [{"text": "Landing Pages", "callback_data": "ds_landing"}]
    ]
    send(chat_id, "*Data Studio*\n\nWhich data would you like to analyse?", keyboard)

# ── AD ANALYSIS ───────────────────────────────────────────────────

def start_ds_adverts(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_ad_format_filter"
    keyboard = [
        [{"text": "Start Analysis — All Ads", "callback_data": "ad_filter_all"}],
        [{"text": "Start Analysis — Video Ads", "callback_data": "ad_filter_video"}],
        [{"text": "Start Analysis — Static Ads", "callback_data": "ad_filter_static"}]
    ]
    msg = "*Ad Performance Analysis*\n\nData sources: Meta Ads Manager, Mixpanel, or any format showing ad metrics.\n\nSelect which ads to analyse:"
    send(chat_id, msg, keyboard)

def analyse_ads(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    images = state.get("ds_images", [])
    csv_text = state.get("ds_csv_text", "")
    if not images and not csv_text:
        send(chat_id, "No data uploaded. Please send screenshots or a CSV file first.")
        return
    send(chat_id, "Analysing your ad data...")
    # Get content context and history
    ctx = get_content_context(chat_id)
    try:
        if images:
            content_blocks = []
            for img_data in images:
                content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": img_data["type"], "data": img_data["data"]}})
            content_blocks.append({"type": "text", "text": """Extract and analyse all ad performance data from these screenshots.

META AD NAMING: IMG=static, VID=video | AWA/CDR/CNV=funnel stage | Avatar in name | Msg_=angle

FILTER: Only analyse """ + {"all": "all ad types", "video": "VIDEO (VID) ads only — ignore any IMG/static ads", "static": "STATIC (IMG) ads only — ignore any VID/video ads"}.get(state.get("ds_ad_filter","all"), "all ad types") + """

IMPORTANT: First count how many qualifying ads are visible. State total count.

STEP 1 — DATA EXTRACTION:
List every qualifying ad with: Ad Name, Type (IMG/VID), Stage, Avatar, Angle, all metrics visible.

STEP 2 — GRADING (A/B/C/D quartiles within cohort):
Grade each ad on every metric. For CPP/CPC lower=better so invert grading.
Give each ad an overall rating: SCALE / KEEP / TEST / KILL

STEP 3 — COHORT ANALYSIS:
Grade each ad against its own cohort (same avatar + stage) AND overall.

STEP 4 — VIDEO AIDA DIAGNOSIS (video ads only):
Map where each video's funnel breaks: Attention → Interest → Desire → Action.

STEP 5 — PATTERN RECOGNITION:
What patterns emerge across avatar, stage, angle, ad type? Be specific.

STEP 6 — IDEAS:
Generate 5 specific new ad ideas based on what the patterns suggest. Not generic — specific angles, avatars, hooks derived from the data.

Format clearly with headers for each step.
""" + ctx})
            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 3000,
                "system": DATA_STUDIO_SYSTEM,
                "messages": [{"role": "user", "content": content_blocks}]
            }).encode()
        else:
            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 3000,
                "system": DATA_STUDIO_SYSTEM,
                "messages": [{"role": "user", "content": "Analyse this Meta Ads Manager CSV data:\n\n" + csv_text + "\n\nApply full grading, cohort analysis, AIDA diagnosis for videos, pattern recognition, and generate 5 specific new ad ideas."}]
            }).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
            headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
        with urllib.request.urlopen(req, timeout=90) as r:
            result = json.loads(r.read())["content"][0]["text"]
        state["last_ds_analysis"] = result
        state["stage"] = "ds_analysis_done"
        send_plain(chat_id, result)
        # Extract and save patterns for future sessions
        extract_and_save_insights(chat_id, result, "ads")
        keyboard = [
            [{"text": "Create ads from insights", "callback_data": "mode_ads"}],
            [{"text": "Ask a follow-up", "callback_data": "ds_followup"}],
            [{"text": "Upload more data", "callback_data": "ds_adverts"}],
            [{"text": "Back to Data Studio", "callback_data": "open_data_studio"}]
        ]
        send(chat_id, "Analysis complete. Create new ads based on what\'s working?", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e), flush=True)
        state["stage"] = "ds_awaiting_ad_data"

# ── SOCIAL ANALYSIS ───────────────────────────────────────────────

def start_ds_social(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_social_format_filter"
    keyboard = [
        [{"text": "All formats", "callback_data": "social_filter_all"}],
        [{"text": "Reels only", "callback_data": "social_filter_reels"}],
        [{"text": "Statics only", "callback_data": "social_filter_statics"}],
        [{"text": "Carousels only", "callback_data": "social_filter_carousels"}]
    ]
    send(chat_id, "*Instagram Performance Analysis*\n\nWhich formats do you want to analyse?", keyboard)

def analyse_social(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    images = state.get("ds_images", [])
    csv_text = state.get("ds_csv_text", "")
    if not images and not csv_text:
        send(chat_id, "No data uploaded. Please send screenshots or a CSV first.")
        return
    send(chat_id, "Analysing your Instagram data...")
    try:
        analysis_prompt = """You are analysing Instagram performance data for Cryptonary.

FORMAT FILTER: Only analyse """ + {"all": "all formats (Reels, Statics, and Carousels)", "reels": "REELS only — ignore Statics and Carousels", "statics": "STATICS only — ignore Reels and Carousels", "carousels": "CAROUSELS only — ignore Reels and Statics"}.get(state.get("ds_social_filter","all"), "all formats") + """

FORMAT DETECTION — look for these visual cues in the screenshots:
- REELS: show a Views metric (in addition to likes/saves etc) — video posts
- STATICS: single image posts — thumbnail has NO overlay indicator
- CAROUSELS: multi-image posts — thumbnail has a small white square tile indicator in the corner

If Minter shows a post type column, use that. Otherwise use the visual cues above.

PRIMARY SUCCESS METRIC: Engagement Rate = (Likes + Comments + Saves + Shares) / Reach × 100
Metric priority: Engagement Rate > Saves > Reach > Follows gained > Comments > Likes > Shares
Views are Reels-only and graded separately.

STEP 1 — DATA EXTRACTION:
State total posts found across ALL images combined.
For each post extract: caption snippet, format (Reel/Static/Carousel), date if shown, Reach, Likes, Comments, Saves, Shares, Views (Reels only), Follows gained if shown.
Calculate Engagement Rate for each post.

STEP 2 — PERFORMANCE RANKING (no A/B/C/D grading needed):
Rank posts by Likes (primary metric).
For Reels, also note Views separately.
Identify: top performers, solid performers, underperformers — with brief reasoning for each.
No grading system required.

STEP 3 — FORMAT BREAKDOWN:
For each format present:
- Count of posts
- Average engagement rate
- Average reach
- Average saves
- Best performer with ER%
- Worst performer with ER%
- What this format is best used for based on this data

STEP 4 — CROSS-FORMAT COMPARISON:
Which format wins on: Engagement Rate / Reach / Saves / Follows gained?
Overall format verdict: which delivers most value right now and for what goal?

STEP 5 — CONTENT PATTERN ANALYSIS:
From captions and topics, identify:
- Which topics drive most saves?
- Which drive most comments?
- Which drive most reach?
- Any pricing/data posts vs opinion posts difference?
- Any BTC vs altcoin vs airdrop vs general crypto difference?

STEP 6 — PATTERN RECOGNITION:
3 strongest patterns with specific post evidence.

STEP 7 — IDEAS:
5 specific content ideas. For each: format, topic, hook, why the data supports it.

STEP 8 — TOTALS SUMMARY:
- Total posts analysed: N
- Overall average engagement rate: X%
- Best format by ER: [format] at [avg ER%]
- Best format by reach: [format]
- Best single post: [caption] — [ER%]
- Worst single post: [caption] — [ER%]

Format clearly with headers."""

        if images:
            content_blocks = []
            for img_data in images:
                content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": img_data["type"], "data": img_data["data"]}})
            content_blocks.append({"type": "text", "text": analysis_prompt})
            messages = [{"role": "user", "content": content_blocks}]
        else:
            messages = [{"role": "user", "content": "Analyse this Instagram CSV data:\n\n" + csv_text + "\n\n" + analysis_prompt}]

        payload = json.dumps({"model": "claude-sonnet-4-20250514", "max_tokens": 3000,
            "system": DATA_STUDIO_SYSTEM, "messages": messages}).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
            headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
        with urllib.request.urlopen(req, timeout=90) as r:
            result = json.loads(r.read())["content"][0]["text"]
        state["last_ds_analysis"] = result
        state["stage"] = "ds_analysis_done"
        send_plain(chat_id, result)
        extract_and_save_insights(chat_id, result, "social")
        keyboard = [
            [{"text": "Create social content from insights", "callback_data": "mode_social"}],
            [{"text": "Ask a follow-up", "callback_data": "ds_followup"}],
            [{"text": "Upload more data", "callback_data": "ds_social"}],
            [{"text": "Back to Data Studio", "callback_data": "open_data_studio"}]
        ]
        send(chat_id, "Analysis complete. Create social content based on what\'s working?", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "ds_awaiting_social_data"

# ── EMAIL ANALYSIS ────────────────────────────────────────────────

def start_ds_emails(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_awaiting_email_splitvar"
    state["ds_images"] = []
    send(chat_id, "*Email Split Test Analysis*\n\nWhat variable are you testing?\n\n_e.g. Image vs No Image, Name in subject vs No name, Short subject vs Long subject_\n\nAfter entering the variable, you can upload screenshots, CSVs, paste raw numbers — any format.")

def start_ds_email_splittest(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_awaiting_email_splitvar"
    send(chat_id, "*Split Test Analysis*\n\nWhat variable are you testing? (e.g. 'name in subject line', 'curiosity gap vs data hook', 'short vs long preview')")

def analyse_emails(chat_id, split_var=None):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    images = state.get("ds_images", [])
    csv_text = state.get("ds_csv_text", "")
    if not images and not csv_text:
        send(chat_id, "No data uploaded. Please send screenshots or a CSV first.")
        return
    ctx = get_content_context(chat_id)
    send(chat_id, "Analysing your email data...")
    split_instruction = ""
    if split_var:
        split_instruction = """YOU ARE RUNNING A SPLIT TEST CALCULATION. DO NOT DO A FULL EMAIL ANALYSIS.
IF THE USER HAS TYPED RAW NUMBERS directly (e.g. "Variant A: 5000 sent, 1200 opens" or a table of numbers), use those directly — do not try to read screenshots.
DO NOT grade emails. DO NOT identify top/bottom performers. DO NOT generate ideas.
ONLY do the following steps and nothing else.

VARIABLE TESTED: """ + split_var + """
Content A = Image version | Content B = No Image version

THESE ARE BREVO A/B TEST SCREENSHOTS. Each screenshot is one campaign.
The TIMELINE section at the bottom of each screenshot contains the real test data.

STEP 1 — FROM EACH CAMPAIGN TIMELINE, EXTRACT:
- Campaign name and audience segment
- Total A/B sample: the number X from "have been sent to X random subscribers"
- Winner: which version (A or B) won and how many openers it had
- Format: | Campaign | Segment | Total A/B Sample | Winner | Winner Openers |

STEP 2 — WEIGHTED AGGREGATION:
For Content A wins: sum their total A/B sample sizes = A weighted sample
For Content B wins: sum their total A/B sample sizes = B weighted sample
Content A evidence % = A weighted / (A + B weighted) × 100
Content B evidence % = B weighted / (A + B weighted) × 100

STEP 3 — POOLED TOTALS ACROSS ALL CAMPAIGNS:
Sum all delivered numbers. Sum all raw opens (delivered × open_rate / 100). Sum all raw clicks.
Combined open rate = total opens / total delivered × 100
Combined CTR = total clicks / total delivered × 100
Show as a TOTALS row.

STEP 4 — VERDICT (5 lines maximum):
WINNER: [A or B] — won campaigns representing [X]% of total test recipients
SEGMENT PATTERN: [one line — did winner hold across all segments?]
POOLED RATE: [combined open rate across all campaigns]
CONFIDENCE: [INCONCLUSIVE <5,000 / DIRECTIONAL 5,000-20,000 / SIGNIFICANT 20,000+]
RECOMMENDATION: [one sentence]

NOTHING ELSE. No performance grades. No insights sections. No suggestions. Just the table, the maths, and the verdict.
"""
    try:
        if split_var:
            # For split tests, skip standard analysis entirely - just run the split calculation
            analysis_prompt = split_instruction
        else:
            analysis_prompt = """Extract and analyse all email performance data.

METRICS TO EXTRACT per email:
- Subject line
- Send date
- Recipients
- Open rate (%) — convert to raw opens = recipients × open_rate / 100
- CTR (%) — convert to raw clicks = recipients × ctr / 100
- Email type if identifiable (Free/Pro)

STEP 1 — DATA EXTRACTION:
List every email with all metrics including calculated raw opens and clicks.

STEP 2 — GRADING (A/B/C/D quartiles):
Grade each email on open rate and CTR.
A = top 25%, B = 26-50%, C = 51-75%, D = bottom 25%.

STEP 3 — TOP AND BOTTOM PERFORMERS:
Top 3 emails — what made them work? (subject style, topic, length, hook type)
Bottom 3 — what likely caused underperformance?

STEP 4 — PATTERN RECOGNITION:
Look for patterns in subject lines:
- Curiosity gap vs data-led vs fear-based vs contrarian
- With name vs without name
- Short vs long subjects
- Question vs statement
- Specific numbers vs general claims
Which patterns correlate with higher open rates and CTR?

STEP 5 — IDEAS:
Generate 5 specific subject line and angle ideas based on what the patterns suggest.
""" + split_instruction

        if split_var:
            send_msg = "Running split test calculation..."
        else:
            send_msg = "Running full analysis..."

        if images:
            content_blocks = []
            for img_data in images:
                content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": img_data["type"], "data": img_data["data"]}})
            content_blocks.append({"type": "text", "text": analysis_prompt})
            messages = [{"role": "user", "content": content_blocks}]
        else:
            messages = [{"role": "user", "content": "Analyse this email performance CSV:\n\n" + csv_text + "\n\n" + analysis_prompt}]

        payload = json.dumps({"model": "claude-sonnet-4-20250514", "max_tokens": 3000,
            "system": DATA_STUDIO_SYSTEM, "messages": messages}).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
            headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
        with urllib.request.urlopen(req, timeout=90) as r:
            result = json.loads(r.read())["content"][0]["text"]
        state["last_ds_analysis"] = result
        state["stage"] = "ds_analysis_done"
        send_plain(chat_id, result)
        extract_and_save_insights(chat_id, result, "emails")
        keyboard = [
            [{"text": "Ask a follow-up", "callback_data": "ds_followup"}],
            [{"text": "Upload more data", "callback_data": "ds_emails"}],
            [{"text": "Back to Data Studio", "callback_data": "open_data_studio"}]
        ]
        send(chat_id, "Analysis complete.", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "ds_awaiting_email_data"

# ── LANDING PAGE ANALYSIS ─────────────────────────────────────────

def start_ds_landing(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_landing_scope"
    keyboard = [
        [{"text": "Single page test", "callback_data": "landing_scope_single"}],
        [{"text": "Multi-page test", "callback_data": "landing_scope_multi"}]
    ]
    send(chat_id, "*Landing Page Analysis*\n\nData source: VWO screenshots or any analytics platform.\n\nTesting a single page variant or comparing multiple pages?", keyboard)

def start_ds_landing_splittest(chat_id):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["stage"] = "ds_awaiting_landing_splitvar"
    send(chat_id, "*Landing Page Split Test*\n\nWhat variable are you testing? (e.g. 'hero headline A vs B', 'Pro vs Inner Circle CTA', 'with guarantee vs without')")

def analyse_landing(chat_id, split_var=None):
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    images = state.get("ds_images", [])
    csv_text = state.get("ds_csv_text", "")
    if not images and not csv_text:
        send(chat_id, "No data uploaded. Please send screenshots or a CSV first.")
        return
    ctx = get_content_context(chat_id)
    send(chat_id, "Analysing your landing page data...")
    split_instruction = ""
    if split_var:
        split_instruction = f"""
SPLIT TEST ANALYSIS FOR: {split_var}

POOL RAW NUMBERS:
For each variant: total sessions, total conversions = sessions × CVR / 100
Pooled CVR = total conversions / total sessions × 100

Statistical confidence:
- Under 500 combined sessions = INCONCLUSIVE
- 500–2,000 = DIRECTIONAL
- 2,000+ = SIGNIFICANT

Declare winner with pooled numbers and confidence.
"""
    try:
        scope = state.get("landing_scope", "single")
        scope_note = "Multiple page variants compared." if scope == "multi" else "Single page variant test."
        analysis_prompt = """Extract and analyse this landing page test. Data source: VWO or any analytics platform.

""" + scope_note + """

METRICS TO EXTRACT (priority order):
- Variant name
- Total traffic / sessions
- Initiate Checkouts (raw number)
- Purchases (raw number)
- CVR = Purchases / Traffic × 100
- Cost per Purchase (if shown)
- Cost per Checkout Initiated (if shown)

SUCCESS PRIORITY: Purchases first, then Checkout Initiation rate. Ignore bounce rate.

STEP 1 — DATA EXTRACTION: Extract all metrics per variant.
STEP 2 — PERFORMANCE TABLE: Each variant with all metrics. Calculate CVR if not shown.
STEP 3 — POOLED TOTALS: Sum raw numbers across cohorts. Never average percentages.
STEP 4 — WINNER: Highest purchase CVR wins. Note if checkout rate tells a different story.
STEP 5 — VERDICT: Winner, confidence, one recommendation. Keep it simple.
""" + ctx

        if images:
            content_blocks = []
            for img_data in images:
                content_blocks.append({"type": "image", "source": {"type": "base64", "media_type": img_data["type"], "data": img_data["data"]}})
            content_blocks.append({"type": "text", "text": analysis_prompt})
            messages = [{"role": "user", "content": content_blocks}]
        else:
            messages = [{"role": "user", "content": "Analyse this landing page CSV:\n\n" + csv_text + "\n\n" + analysis_prompt}]

        payload = json.dumps({"model": "claude-sonnet-4-20250514", "max_tokens": 3000,
            "system": DATA_STUDIO_SYSTEM, "messages": messages}).encode()
        req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
            headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
        with urllib.request.urlopen(req, timeout=90) as r:
            result = json.loads(r.read())["content"][0]["text"]
        state["last_ds_analysis"] = result
        state["stage"] = "ds_analysis_done"
        send_plain(chat_id, result)
        extract_and_save_insights(chat_id, result, "landing")
        keyboard = [
            [{"text": "Ask a follow-up", "callback_data": "ds_followup"}],
            [{"text": "Upload more data", "callback_data": "ds_landing"}],
            [{"text": "Back to Data Studio", "callback_data": "open_data_studio"}]
        ]
        send(chat_id, "Analysis complete.", keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))
        state["stage"] = "ds_awaiting_landing_data"

# ── IMAGE/CSV HANDLER ─────────────────────────────────────────────

def handle_ds_file(chat_id, file_info, file_type="image"):
    """Download and store uploaded file for analysis."""
    state = user_state.get(chat_id, {})
    stage = state.get("stage", "")
    if not stage.startswith("ds_awaiting"):
        return False
    try:
        file_id = file_info.get("file_id")
        # Get file path from Telegram
        path_data = tg("getFile", {"file_id": file_id})
        file_path = path_data.get("result", {}).get("file_path", "")
        if not file_path:
            return False
        url = "https://api.telegram.org/file/bot" + TELEGRAM_TOKEN + "/" + file_path
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=30) as r:
            file_bytes = r.read()
        import base64
        if file_type == "image":
            # Detect image type
            mime = "image/jpeg"
            if file_path.endswith(".png"): mime = "image/png"
            elif file_path.endswith(".gif"): mime = "image/gif"
            elif file_path.endswith(".webp"): mime = "image/webp"
            encoded = base64.b64encode(file_bytes).decode()
            if "ds_images" not in state: state["ds_images"] = []
            state["ds_images"].append({"type": mime, "data": encoded})
            count = len(state["ds_images"])

            # Quick validation - ask Claude what it can see in this image
            stage = state.get("stage", "")
            context_hint = "Meta Ads Manager data" if "ad" in stage else "Instagram analytics data" if "social" in stage else "email performance data" if "email" in stage else "landing page analytics data"
            # For split test flows, skip preview to avoid partial data commentary
            is_split = "split" in stage
            if is_split:
                preview = "Screenshot received."
            else:
                try:
                    check_payload = json.dumps({
                        "model": "claude-sonnet-4-20250514",
                        "max_tokens": 150,
                        "messages": [{"role": "user", "content": [
                            {"type": "image", "source": {"type": "base64", "media_type": mime, "data": encoded}},
                            {"type": "text", "text": "This should contain " + context_hint + ". In one short sentence: what data can you see? If there are rows/ads/posts, how many? If the data is cut off or unclear, say so. Be specific and brief."}
                        ]}]
                    }).encode()
                    check_req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=check_payload,
                        headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
                    with urllib.request.urlopen(check_req, timeout=30) as r:
                        preview = json.loads(r.read())["content"][0]["text"].strip()
                except:
                    preview = "Image received."

            # Determine button label based on count
            ds_stage_map = {
                "ds_awaiting_ad_data": "ds_analyse_ads",
                "ds_awaiting_social_data": "ds_analyse_social",
                "ds_awaiting_email_data": "ds_analyse_emails",
                "ds_awaiting_email_split_data": "ds_analyse_emails_split",
                "ds_awaiting_landing_data": "ds_analyse_landing",
                "ds_awaiting_landing_split_data": "ds_analyse_landing_split"
            }
            analyse_cb = ds_stage_map.get(stage, "ds_analyse_ads")

            keyboard = [
                [{"text": "Analyse this", "callback_data": analyse_cb}],
                [{"text": "Add another screenshot", "callback_data": "ds_add_more_" + stage}]
            ]
            msg = "Got it. " + preview
            if count > 1:
                msg += "\n\n" + str(count) + " images queued."
            send(chat_id, msg, keyboard)
        else:
            # CSV or document — decode as text
            try:
                text = file_bytes.decode("utf-8")
            except:
                text = file_bytes.decode("latin-1")
            state["ds_csv_text"] = text[:15000]  # limit to 15k chars
            send(chat_id, "File received (" + str(len(text)) + " characters). Tap Done to analyse.")
        return True
    except Exception as e:
        send(chat_id, "Error reading file: " + str(e))
        return False

# ══════════════════════════════════════════════════════════════════
# CONTENT METADATA + PATTERN MEMORY ENGINE
# ══════════════════════════════════════════════════════════════════

import datetime as _dt

CONTENT_STORE_FILE = "content_metadata.json"
INSIGHTS_FILE = "pattern_insights.json"

# In-memory stores
content_metadata = {}   # {chat_id: [{type, date, metadata...}]}
pattern_insights = {}   # {chat_id: [{date, insight, source}]}

VOICE_CORPUS_FILE = "voice_corpus.json"
voice_corpus = {}  # {chat_id: [{date, text, source_type}]}

def load_content_stores():
    global content_metadata, pattern_insights, voice_corpus
    try:
        with open(CONTENT_STORE_FILE, "r") as f:
            raw = json.load(f)
        content_metadata = {int(k): v for k, v in raw.items()}
    except:
        pass
    try:
        with open(INSIGHTS_FILE, "r") as f:
            raw = json.load(f)
        pattern_insights = {int(k): v for k, v in raw.items()}
    except:
        pass
    try:
        with open(VOICE_CORPUS_FILE, "r") as f:
            raw = json.load(f)
        voice_corpus = {int(k): v for k, v in raw.items()}
    except:
        pass

def save_content_stores():
    try:
        with open(CONTENT_STORE_FILE, "w") as f:
            json.dump({str(k): v for k, v in content_metadata.items()}, f)
    except Exception as e:
        print("Content store save error:", e, flush=True)
    try:
        with open(INSIGHTS_FILE, "w") as f:
            json.dump({str(k): v for k, v in pattern_insights.items()}, f)
    except Exception as e:
        print("Insights save error:", e, flush=True)
    try:
        with open(VOICE_CORPUS_FILE, "w") as f:
            json.dump({str(k): v for k, v in voice_corpus.items()}, f)
    except Exception as e:
        print("Voice corpus save error:", e, flush=True)


def generate_diff(before, after):
    """Ask Claude to summarise exactly what changed between two versions."""
    try:
        result = claude(
            "Compare these two versions and list EXACTLY what changed.\n\n" +
            "VERSION BEFORE:\n" + before[:2000] +
            "\n\n---\n\nVERSION AFTER:\n" + after[:2000] +
            "\n\nList every change as:\n" +
            "ADDED: [exact sentence or phrase added]\n" +
            "REMOVED: [exact sentence or phrase removed]\n" +
            "EDITED: [what it was → what it became]\n\n" +
            "Be specific and literal. List every single change. Nothing else.",
            max_tokens=600
        )
        return result
    except:
        return ""

def show_diff_with_revert(chat_id, before, after, content_type="email"):
    """Show what changed and offer Revert or Accept."""
    if not before or before == after:
        return
    diff = generate_diff(before, after)
    if diff:
        send_plain(chat_id, "CHANGES MADE:\n\n" + diff)
    kb_map = {"email": email_action_keyboard, "ad": ad_action_keyboard, "social": social_action_keyboard}
    kb_fn = kb_map.get(content_type, email_action_keyboard)
    keyboard = [
        [{"text": "Revert to previous", "callback_data": "revert_to_previous_" + content_type},
         {"text": "Accept", "callback_data": "accept_changes_" + content_type}]
    ] + kb_fn()
    send(chat_id, "Changes shown above.", keyboard)

def save_voice_example(chat_id, text, source_type="approved"):
    """Save approved copy to voice corpus for self-learning."""
    if not text or len(text.strip()) < 50:
        return
    if chat_id not in voice_corpus:
        voice_corpus[chat_id] = []
    voice_corpus[chat_id].append({
        "date": _dt.datetime.now().strftime("%Y-%m-%d"),
        "text": text.strip(),
        "source": source_type
    })
    # Keep last 100 examples
    if len(voice_corpus[chat_id]) > 100:
        voice_corpus[chat_id] = voice_corpus[chat_id][-100:]
    save_content_stores()

def get_voice_corpus_context(chat_id):
    """Return recent approved examples for injection into generation prompts."""
    examples = voice_corpus.get(chat_id, [])
    if not examples:
        return ""
    recent = examples[-5:]
    ctx = "\n\n=== APPROVED WRITING EXAMPLES (Adam's actual voice — match this style) ===\n"
    for ex in recent:
        ctx += "\n[" + ex.get("source","") + " — " + ex.get("date","") + "]\n"
        ctx += ex.get("text","") + "\n"
    return ctx

def save_content_metadata(chat_id, content_type, metadata):
    """Called silently when content is approved. Stores what was created."""
    if chat_id not in content_metadata:
        content_metadata[chat_id] = []
    record = {
        "type": content_type,
        "date": _dt.datetime.now().strftime("%Y-%m-%d %H:%M"),
        **metadata
    }
    content_metadata[chat_id].append(record)
    # Keep last 200 records per user
    if len(content_metadata[chat_id]) > 200:
        content_metadata[chat_id] = content_metadata[chat_id][-200:]
    save_content_stores()

def save_pattern_insight(chat_id, insight, source="analysis"):
    """Save a pattern discovered during analysis for future reference."""
    if chat_id not in pattern_insights:
        pattern_insights[chat_id] = []
    pattern_insights[chat_id].append({
        "date": _dt.datetime.now().strftime("%Y-%m-%d"),
        "insight": insight,
        "source": source
    })
    if len(pattern_insights[chat_id]) > 100:
        pattern_insights[chat_id] = pattern_insights[chat_id][-100:]
    save_content_stores()

def get_content_context(chat_id):
    """Build context string from stored metadata for use in analysis prompts."""
    records = content_metadata.get(chat_id, [])
    insights = pattern_insights.get(chat_id, [])
    if not records and not insights:
        return ""

    ctx = "\n\n=== CONTENT HISTORY & KNOWN PATTERNS ===\n"

    if records:
        ctx += "\nRECENT CONTENT CREATED (last 20):\n"
        for r in records[-20:]:
            line = "- [" + r.get("date","") + "] " + r.get("type","").upper()
            if r.get("angle"): line += " | Angle: " + r["angle"][:60]
            if r.get("hook_subject"): line += " | Subject: " + r["hook_subject"][:50]
            if r.get("free_cta"): line += " | CTA: " + r["free_cta"]
            if r.get("tone") and r["tone"] != "standard": line += " | Tone: " + r["tone"]
            if r.get("avatars"): line += " | Avatars: " + ", ".join(r["avatars"])
            if r.get("ad_type"): line += " | Type: " + r["ad_type"]
            ctx += line + "\n"

    if insights:
        ctx += "\nPREVIOUSLY DISCOVERED PATTERNS:\n"
        for ins in insights[-15:]:
            ctx += "- [" + ins.get("date","") + "] " + ins.get("insight","") + "\n"

    return ctx

def extract_and_save_insights(chat_id, analysis_text, source_type):
    """After analysis, extract key patterns and save them for future sessions."""
    try:
        raw = claude(
            "From this marketing analysis, extract 3-5 specific, memorable patterns or insights that should be remembered for future campaigns. Focus on what consistently works or fails.\n\nAnalysis:\n" + analysis_text[:3000] +
            "\n\nFormat each insight as one clear sentence starting with the content type and specific finding. Example: 'PASSIVE_INCOME ads consistently outperform at CDR stage with 50% lower CPP than average.' Return as a numbered list only.",
            max_tokens=400,
            system=DATA_STUDIO_SYSTEM
        )
        lines = [l.strip() for l in raw.split('\n') if l.strip() and l.strip()[0].isdigit()]
        for line in lines:
            import re as _re
            m = _re.match(r'^\d+\.\s*(.+)', line)
            if m:
                save_pattern_insight(chat_id, m.group(1), source_type)
    except Exception as e:
        print("Insight extraction error:", e, flush=True)

def get_preanalysis_checklist(data_type, item_count):
    """Returns warning if sample size is too small."""
    thresholds = {
        "ads": {"min": 5, "warn": 10},
        "emails": {"min": 3, "warn": 8},
        "social": {"min": 5, "warn": 10},
        "landing": {"min": 3, "warn": 6}
    }
    t = thresholds.get(data_type, {"min": 3, "warn": 5})
    if item_count < t["min"]:
        return ("TOO_SMALL", "Only " + str(item_count) + " items found. Minimum " + str(t["min"]) + " needed for meaningful analysis. Results may be misleading — consider gathering more data first.")
    elif item_count < t["warn"]:
        return ("SMALL", "Small sample (" + str(item_count) + " items). Analysis will run but treat conclusions as directional, not definitive.")
    return ("OK", "")

# Load on startup
load_content_stores()

# ══════════════════════════════════════════════════════════════════
# WEEKLY BRIEFING ENGINE
# ══════════════════════════════════════════════════════════════════

BRIEFING_CHAT_IDS_FILE = "briefing_subscribers.json"

def load_briefing_subscribers():
    try:
        with open(BRIEFING_CHAT_IDS_FILE, "r") as f:
            return json.load(f)
    except:
        return []

def save_briefing_subscribers(ids):
    try:
        with open(BRIEFING_CHAT_IDS_FILE, "w") as f:
            json.dump(ids, f)
    except Exception as e:
        print("Briefing subscriber save error:", e, flush=True)

# ══════════════════════════════════════════════════════════════════
# WHISPER — VOICE TO TEXT
# ══════════════════════════════════════════════════════════════════

def transcribe_voice(file_id):
    """Download a Telegram voice message and transcribe it via OpenAI Whisper.
    Retries once on 429 rate limit errors."""
    if not OPENAI_KEY:
        return None, "OpenAI key not configured. Add OPENAI_KEY to Render."
    try:
        # Step 1: Get file path from Telegram
        path_data = tg("getFile", {"file_id": file_id})
        file_path = path_data.get("result", {}).get("file_path", "")
        if not file_path:
            return None, "Could not get voice file from Telegram."

        # Step 2: Download the audio file
        download_url = "https://api.telegram.org/file/bot" + TELEGRAM_TOKEN + "/" + file_path
        req = urllib.request.Request(download_url)
        with urllib.request.urlopen(req, timeout=30) as r:
            audio_data = r.read()

        # Step 3: Send to Whisper API — retry once on 429
        boundary = "----WhisperBoundary7MA4YWx"
        body = (
            ("--" + boundary + "\r\n").encode() +
            b"Content-Disposition: form-data; name=\"file\"; filename=\"voice.ogg\"\r\n" +
            b"Content-Type: audio/ogg\r\n\r\n" +
            audio_data + b"\r\n" +
            ("--" + boundary + "\r\n").encode() +
            b"Content-Disposition: form-data; name=\"model\"\r\n\r\nwhisper-1\r\n" +
            ("--" + boundary + "--\r\n").encode()
        )
        for attempt in range(2):
            try:
                req = urllib.request.Request(
                    "https://api.openai.com/v1/audio/transcriptions",
                    data=body,
                    headers={
                        "Authorization": "Bearer " + OPENAI_KEY,
                        "Content-Type": "multipart/form-data; boundary=" + boundary
                    }
                )
                with urllib.request.urlopen(req, timeout=60) as r:
                    result = json.loads(r.read())
                    transcript = result.get("text", "").strip()
                    return transcript, None
            except Exception as e:
                err_str = str(e)
                if "429" in err_str and attempt == 0:
                    time.sleep(5)  # wait 5s then retry once
                    continue
                if "429" in err_str:
                    return None, "Whisper rate limit hit. Wait a moment and try again, or upgrade your OpenAI account at platform.openai.com for higher limits."
                return None, "Transcription failed: " + err_str
    except Exception as e:
        print("Whisper error:", e, flush=True)
        return None, "Transcription failed: " + str(e)
    return None, "Transcription failed after retry"


def handle_voice_message(chat_id, voice):
    """Handle a voice note — transcribe and route contextually based on current stage."""
    if not OPENAI_KEY:
        send(chat_id, "⚠️ Voice input needs an OpenAI key. Add OPENAI_KEY to Render.")
        return

    send(chat_id, "🎙️ Transcribing...")
    file_id = voice.get("file_id", "")
    transcript, error = transcribe_voice(file_id)

    if error or not transcript:
        send(chat_id, "❌ Could not transcribe. Try again or just type it.\n\n_Error: " + (error or "empty") + "_")
        return

    current_stage = user_state.get(chat_id, {}).get("stage", "idle")

    # INPUT STAGES — bot is waiting for text, voice is transparent replacement
    # Just transcribe and route exactly as if the user typed it
    input_stages = {
        # Email flow
        "awaiting_email_report", "awaiting_report", "buffering_report",
        "awaiting_context_choice", "awaiting_context_text",
        # Social flow
        "awaiting_social_report",
        # ALL quick edit stages
        "awaiting_quick_edit", "social_quick_edit",
        "awaiting_lp_quick_edit", "awaiting_ad_quick_edit", "awaiting_social_quick_edit",
        # Angle/hook custom inputs
        "awaiting_custom_free_angle", "awaiting_custom_pro_angle",
        "awaiting_custom_free_hook", "awaiting_custom_pro_hook",
        "awaiting_custom_social_angle", "awaiting_custom_angle", "awaiting_custom_hook",
        # Data studio follow-up
        "ds_awaiting_followup", "ds_awaiting_email_split_numbers",
        # Idea engine
        "ie_awaiting_inspiration", "ie_awaiting_pasted_text",
        "ie_awaiting_custom_concept", "ie_awaiting_develop",
        # Image generation
        "img_awaiting_brief", "img_awaiting_direction",
        # Landing page context
        "awaiting_lp_context_text", "lp_outline_review",
        # Brief editing
        "vb_awaiting_edit",
        # Voice edit
        "voice_awaiting_edit",
        # Segment edit
        "awaiting_seg_edit",
        # Logging
        "logging_email", "logging_ad",
        # Subject line
        "awaiting_custom_pro_hook",
        # DS followup
        "ds_followup",
        # Active content stages — voice = quick edit instruction
        "social_ready", "social_approved",
        "emails_ready", "emails_approved",
    }

    if current_stage in input_stages:
        # Show what was heard, then route as typed text
        send(chat_id, "🎙️ _\"" + transcript[:200] + "\"_")
        fake_msg = {"chat": {"id": chat_id}, "text": transcript, "from": {"id": chat_id}}
        handle_message(fake_msg)
        return

    # IDLE — bot is not waiting for anything specific
    # Use interpretation flow so we know what to do with the voice brief
    send(chat_id, "Heard: _\"" + transcript[:300] + "\"_\n\nInterpreting...")
    try:
        interpretation = claude(
            "A user sent this voice brief to a crypto content creation bot:\n\n" +
            "TRANSCRIPT: \"" + transcript + "\"\n\n" +
            "Extract their intent. Reply in this EXACT format and nothing else:\n\n" +
            "CONTENT_TYPE: [one of: email / instagram_post / reel / carousel / story / ad / idea / general]\n" +
            "TOPIC: [the specific crypto topic or event they want to cover, 1 sentence]\n" +
            "ANGLE: [the emotional or argumentative approach, 1 sentence]\n" +
            "READY_TO_GENERATE: [yes if intent is clear enough / no if more info needed]",
            max_tokens=250
        )
        lines = {l.split(":")[0].strip(): ":".join(l.split(":")[1:]).strip()
                 for l in interpretation.strip().splitlines() if ":" in l}

        content_type = lines.get("CONTENT_TYPE", "general").lower()
        topic = lines.get("TOPIC", transcript[:100])
        angle = lines.get("ANGLE", "")
        ready = lines.get("READY_TO_GENERATE", "no").lower() == "yes"

        user_state.setdefault(chat_id, {})
        user_state[chat_id]["voice_transcript"] = transcript
        user_state[chat_id]["voice_content_type"] = content_type
        user_state[chat_id]["voice_topic"] = topic
        user_state[chat_id]["voice_angle"] = angle

        summary = "*Here\'s what I heard:*\n\n"
        summary += "📌 *Content:* " + content_type.replace("_", " ").title() + "\n"
        summary += "📰 *Topic:* " + topic + "\n"
        if angle:
            summary += "🎯 *Angle:* " + angle + "\n"

        keyboard = [
            [{"text": "✅ Generate this",    "callback_data": "voice_confirm"}],
            [{"text": "✏️ Change something", "callback_data": "voice_edit"}],
            [{"text": "❌ Start over",        "callback_data": "voice_discard"}],
        ]
        if not ready:
            summary += "\n_Brief is a bit vague — generate anyway or add detail?_"
        send(chat_id, summary, keyboard)

    except Exception as e:
        # Interpretation failed — just route as text
        send(chat_id, "🎙️ _\"" + transcript[:200] + "\"_")
        fake_msg = {"chat": {"id": chat_id}, "text": transcript, "from": {"id": chat_id}}
        handle_message(fake_msg)


def route_voice_to_flow(chat_id):
    """Route a confirmed voice brief into the right content flow."""
    state = user_state.get(chat_id, {})
    content_type = state.get("voice_content_type", "general")
    topic = state.get("voice_topic", "")
    angle = state.get("voice_angle", "")

    combined = topic + (". Angle: " + angle if angle else "")

    if content_type in ("reel", "instagram_post", "carousel", "story"):
        # Route to social flow
        user_state[chat_id]["report"] = combined
        user_state[chat_id]["social_angle"] = angle or topic
        user_state[chat_id]["social_origin"] = "voice"
        user_state[chat_id]["stage"] = "pick_social_formats"
        show_standalone_social_menu(chat_id)

    elif content_type == "email":
        user_state[chat_id]["report"] = combined
        user_state[chat_id]["mode"] = "email"
        user_state[chat_id]["stage"] = "awaiting_context_choice"
        ask_context(chat_id)

    elif content_type == "ad":
        user_state[chat_id]["ad_theme"] = combined
        user_state[chat_id]["stage"] = "pick_ad_avatars"
        show_avatar_menu(chat_id)

    elif content_type == "idea":
        user_state[chat_id]["ie_source_content"] = "VOICE BRIEF: " + combined
        user_state[chat_id]["ie_source_label"] = "Voice brief"
        user_state[chat_id]["ie_idea_type"] = "instagram"
        generate_ie_concept(chat_id)

    else:
        # General — show them what we got and let them pick
        keyboard = [
            [{"text": "📧 Make it an email",     "callback_data": "voice_as_email"}],
            [{"text": "📱 Make it social content","callback_data": "voice_as_social"}],
            [{"text": "💡 Generate ideas",        "callback_data": "voice_as_ideas"}],
            [{"text": "📢 Make it an ad",         "callback_data": "voice_as_ad"}],
        ]
        send(chat_id, "What do you want to make with this?", keyboard)


# ══════════════════════════════════════════════════════════════════
# X (TWITTER) API — LIVE TWEET FETCHING
# ══════════════════════════════════════════════════════════════════

# Accounts to pull from — these are the high-signal crypto accounts
X_ACCOUNTS = [
    "WatcherGuru",
    "lookonchain",
    "DocumentingBTC",
    "caprioleio",
    "trendingbitcoin",
    "BitcoinMagazine",
    "CoinDesk",
    "Blockworks_",
    "glassnode",
    "woonomic",
]

# Simple cache so we don't hit the API on every "generate from scratch"
_x_cache = {"tweets": [], "fetched_at": 0}
X_CACHE_TTL = 1800  # 30 minutes


def fetch_x_tweets(max_per_account=3, max_total=20):
    """Fetch recent tweets from key crypto accounts via X API v2 pay-per-use."""
    if not X_BEARER_TOKEN:
        return []

    # Return cached results if fresh
    if time.time() - _x_cache["fetched_at"] < X_CACHE_TTL and _x_cache["tweets"]:
        return _x_cache["tweets"]

    all_tweets = []
    headers = {
        "Authorization": "Bearer " + X_BEARER_TOKEN,
        "User-Agent": "CryptonaryBot/1.0"
    }

    for handle in X_ACCOUNTS:
        if len(all_tweets) >= max_total:
            break
        try:
            # Get user ID first
            user_url = "https://api.twitter.com/2/users/by/username/" + handle
            req = urllib.request.Request(user_url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as r:
                user_data = json.loads(r.read())
            user_id = user_data.get("data", {}).get("id")
            if not user_id:
                continue

            # Get recent tweets
            tweets_url = (
                "https://api.twitter.com/2/users/" + user_id +
                "/tweets?max_results=" + str(max_per_account) +
                "&tweet.fields=text,created_at,public_metrics" +
                "&exclude=retweets,replies"
            )
            req = urllib.request.Request(tweets_url, headers=headers)
            with urllib.request.urlopen(req, timeout=10) as r:
                tweets_data = json.loads(r.read())

            for tweet in tweets_data.get("data", []):
                text = tweet.get("text", "").strip()
                metrics = tweet.get("public_metrics", {})
                if text and len(text) > 20:
                    all_tweets.append({
                        "source": "@" + handle,
                        "text": text[:280],
                        "likes": metrics.get("like_count", 0),
                        "retweets": metrics.get("retweet_count", 0),
                    })
        except Exception as e:
            print("X fetch error @" + handle + ":", e, flush=True)
            continue

    # Cache results
    _x_cache["tweets"] = all_tweets[:max_total]
    _x_cache["fetched_at"] = time.time()
    return _x_cache["tweets"]


def format_x_context(tweets):
    """Format tweets into a clean context string for prompts."""
    if not tweets:
        return ""
    lines = ["\n=== LIVE CRYPTO TWITTER (pulled now) ==="]
    for t in tweets:
        engagement = ""
        if t.get("likes", 0) > 100:
            engagement = " [" + str(t["likes"]) + " likes]"
        lines.append("• " + t["source"] + engagement + ": " + t["text"][:200])
    lines.append("===")
    return "\n".join(lines)


# ══════════════════════════════════════════════════════════════════
# IMAGE GENERATION — DALL-E 3 via OpenAI API
# ══════════════════════════════════════════════════════════════════

# Post type → style instruction mapping
IMAGE_STYLE_MAP = {
    "breaking_news":  "Breaking news style. Black background. Bold red 'NEWS' or 'BREAKING' badge top-left. Dramatic relevant photo or flag. Bold white Tungsten-style headline overlay. @Cryptonary wordmark bottom-right.",
    "price_data":     "Clean data post. Black background. Large bold price figure in centre. Coin logo prominent. Percentage change in green (gains) or red (losses). Minimalist, numbers-first layout. @Cryptonary mark bottom-right.",
    "engagement":     "Bold engagement post. Pure black or white background. Single punchy question or statement in large Inter-style font. Simple coin grid or choice visual if relevant. Clean and high contrast.",
    "educational":    "Educational carousel style. Dark background. Structured numbered layout or data table. Clear typographic hierarchy. Blue accent (#005EFF) for section headers. @Cryptonary mark bottom-right.",
    "meme_cultural":  "Cinematic dark tone. Dramatic background image with dark overlay. Bold white text overlay. Minimal Cryptonary C mark bottom corner. Feels shareable and culturally relevant.",
    "macro_geo":      "Geopolitical news style. Real-world photo with dark gradient overlay. White bold headline. Small NEWS badge. Serious, editorial tone. @Cryptonary bottom-right.",
    "background":     "Clean background image for compositing. No text overlays. Dark cinematic mood. Relevant to crypto, finance, technology or the specific topic provided. High quality, editorial.",
}

def build_image_prompt(concept, angle, post_type="breaking_news", extra_direction=""):
    """Build a clean image prompt. Keep under 800 chars, avoid policy-triggering terms."""
    style = IMAGE_STYLE_MAP.get(post_type, IMAGE_STYLE_MAP["background"])
    # Truncate concept/angle tightly — long prompts cause 400 errors
    concept_short = concept[:150].strip()
    angle_short = angle[:80].strip()
    prompt = (
        "Professional digital graphic design. " +
        style + " " +
        "Topic: " + concept_short + ". " +
        ("Tone: " + angle_short + ". " if angle_short else "") +
        "Colour palette: black background, white text, blue and red accents. " +
        "Square format 1:1. High contrast. Editorial quality. No logos."
    )
    if extra_direction:
        prompt += " " + extra_direction[:100]
    # Hard cap at 900 chars to stay well under DALL-E 4000 char limit
    return prompt[:900]


def generate_dalle_image(prompt, size="1024x1024", quality="standard"):
    """Call DALL-E 3 API and return the image URL."""
    if not OPENAI_KEY:
        return None, "OpenAI key not configured. Add OPENAI_KEY to Render."
    try:
        body = json.dumps({
            "model": "dall-e-3",
            "prompt": prompt,
            "n": 1,
            "size": size,
            "quality": quality,
            "response_format": "url"
        }).encode()
        req = urllib.request.Request(
            "https://api.openai.com/v1/images/generations",
            data=body,
            headers={
                "Authorization": "Bearer " + OPENAI_KEY,
                "Content-Type": "application/json"
            }
        )
        with urllib.request.urlopen(req, timeout=60) as r:
            result = json.loads(r.read())
        url = result["data"][0]["url"]
        revised = result["data"][0].get("revised_prompt", "")
        return url, revised
    except Exception as e:
        print("DALL-E error:", e, flush=True)
        return None, "Image generation failed: " + str(e)


def generate_gemini_image(prompt):
    """Call Gemini image generation (gemini-3.1-flash-image-preview).
    Returns (image_bytes, mime_type, error)."""
    if not GEMINI_KEY:
        return None, None, "Gemini key not configured. Add GEMINI_KEY to Render."
    try:
        import base64
        body = json.dumps({
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {"responseModalities": ["IMAGE", "TEXT"]}
        }).encode()
        req = urllib.request.Request(
            "https://generativelanguage.googleapis.com/v1beta/models/gemini-3.1-flash-image-preview:generateContent?key=" + GEMINI_KEY,
            data=body,
            headers={"Content-Type": "application/json"}
        )
        with urllib.request.urlopen(req, timeout=90) as r:
            result = json.loads(r.read())
        parts = result.get("candidates", [{}])[0].get("content", {}).get("parts", [])
        for part in parts:
            if "inlineData" in part:
                mime = part["inlineData"].get("mimeType", "image/png")
                data = part["inlineData"].get("data", "")
                return base64.b64decode(data), mime, None
        return None, None, "No image returned by Gemini"
    except Exception as e:
        print("Gemini image error:", e, flush=True)
        return None, None, "Gemini image failed: " + str(e)


def send_image_bytes(chat_id, image_bytes, mime_type="image/png", caption=""):
    """Upload image bytes directly to Telegram via multipart POST."""
    try:
        boundary = "----TgBoundary7MA4YWxk"
        ext = "jpg" if "jpeg" in mime_type else "png"
        parts = []
        parts.append(("--" + boundary + "\r\n").encode())
        parts.append(("Content-Disposition: form-data; name=\"chat_id\"\r\n\r\n" + str(chat_id) + "\r\n").encode())
        if caption:
            parts.append(("--" + boundary + "\r\n").encode())
            parts.append(("Content-Disposition: form-data; name=\"caption\"\r\n\r\n" + caption[:1024] + "\r\n").encode())
        parts.append(("--" + boundary + "\r\n").encode())
        parts.append(("Content-Disposition: form-data; name=\"photo\"; filename=\"image." + ext + "\"\r\n").encode())
        parts.append(("Content-Type: " + mime_type + "\r\n\r\n").encode())
        parts.append(image_bytes)
        parts.append(b"\r\n")
        parts.append(("--" + boundary + "--\r\n").encode())
        body = b"".join(parts)
        req = urllib.request.Request(
            "https://api.telegram.org/bot" + TELEGRAM_TOKEN + "/sendPhoto",
            data=body,
            headers={"Content-Type": "multipart/form-data; boundary=" + boundary}
        )
        with urllib.request.urlopen(req, timeout=60) as r:
            return json.loads(r.read())
    except Exception as e:
        print("send_image_bytes error:", e, flush=True)
        send(chat_id, "Image generated but failed to send. Error: " + str(e))



# ══════════════════════════════════════════════════════════════════
# CLAUDE VISUAL GENERATION — SVG + HTML
# Claude generates brand-accurate visuals as SVG (sent as image)
# and HTML (sent as downloadable file for editing)
# ══════════════════════════════════════════════════════════════════

def generate_claude_svg(brief, post_type, angle=""):
    """Ask Claude to generate an SVG graphic from a visual brief."""
    type_instructions = {
        "storyboard": "Create a STORYBOARD SVG showing scene frames in a grid layout. Each frame is a labelled rectangle with scene description inside. Clean, minimal, black background, white text, blue accent borders.",
        "static":     "Create a STATIC POST SVG (1080x1080). Bold headline text, dark background, brand colours. Minimal text on the graphic itself. Make it look like a real Instagram post.",
        "carousel":   "Create a CAROUSEL SLIDE SVG (1080x1080) showing the cover slide design. Bold headline, dark background, clear typography hierarchy.",
        "story":      "Create a STORY FRAME SVG (1080x1920). Full screen vertical. Bold text, dark bg, designed for mobile.",
        "thumbnail":  "Create an EMAIL THUMBNAIL SVG (600x300). Bold headline, relevant visual element, dark background. Optimised for email preview.",
        "data":       "Create a DATA VISUALISATION SVG. Charts, numbers, clean layout. Dark background, green/red for up/down, white text.",
    }
    instruction = type_instructions.get(post_type, type_instructions["static"])

    prompt = ("""Generate a complete, valid SVG graphic for Cryptonary.

BRIEF:
""" + brief[:800] + """

VISUAL TYPE: """ + instruction + """

BRAND:
- Background: #000000 (black)
- Primary text: #FFFFFF (white)  
- Accent: #005EFF (blue)
- Bullish: #0DA500 (green)
- Bearish: #FF0000 (red)
- Bitcoin: #F7931A (orange)
- Font family: Arial, sans-serif (use font-weight bold for headlines)
- Add "@Cryptonary" in small white text bottom-right

RULES:
- Return ONLY the complete SVG code, nothing else
- No markdown, no explanation, just the raw SVG starting with <svg
- Must be valid, renderable SVG
- Keep text concise — max 8 words per headline
- Use viewBox for proper scaling
- Make it look professional and on-brand""")

    try:
        result = claude(prompt, max_tokens=2000)
        # Extract SVG - handle markdown wrapping, code blocks, preamble text
        import re as _re
        # Strip markdown code blocks
        cleaned = _re.sub(r'```(?:svg|xml)?\s*', '', result)
        cleaned = _re.sub(r'```', '', cleaned).strip()
        svg_match = _re.search(r'(<svg[\s\S]*?</svg>)', cleaned, _re.IGNORECASE)
        if svg_match:
            return svg_match.group(1), None
        elif cleaned.lower().startswith('<svg'):
            return cleaned, None
        else:
            # Try to find any SVG-like content
            if '<svg' in result.lower():
                start = result.lower().find('<svg')
                end = result.lower().rfind('</svg>') + 6
                if end > start:
                    return result[start:end], None
            return None, "Claude did not return valid SVG — try again"
    except Exception as e:
        return None, str(e)


def generate_claude_html(brief, post_type, angle=""):
    """Ask Claude to generate a styled HTML visual from a visual brief."""
    type_instructions = {
        "storyboard": "A storyboard layout with scene frames in a CSS grid. Each frame has a number, title, and description. Dark theme, professional.",
        "static":     "An Instagram post mockup (1080x1080px equivalent). Bold headline, dark background, brand colours. Looks like a real post.",
        "carousel":   "A FULL Instagram carousel — ALL slides from the brief, displayed as scrollable vertical pages. Each slide is 1080x1080px, dark background, with its own headline and visual direction. Include navigation dots or slide numbers. Show every slide from the brief, not just the cover.",
        "story":      "A vertical story frame (9:16 ratio). Full-screen design, bold text, mobile-optimised.",
        "thumbnail":  "An email header thumbnail (600x200px). Bold headline, clean layout.",
        "data":       "A data visualisation with styled numbers, charts using CSS, clean layout.",
    }
    instruction = type_instructions.get(post_type, type_instructions["static"])

    prompt = ("""Generate a complete, self-contained HTML file for a Cryptonary visual.

BRIEF:
""" + brief[:800] + """

TYPE: """ + instruction + """

BRAND COLOURS: #000000 bg, #FFFFFF text, #005EFF blue accent, #0DA500 green, #FF0000 red, #F7931A bitcoin orange
FONTS: system-ui, Arial, sans-serif — bold for headlines
LOGO: Place the Cryptonary wordmark SVG in the bottom-right corner. Use this exact SVG inline: <svg style="position:absolute;bottom:16px;right:16px;width:120px;opacity:0.9" viewBox="0 0 500 100"><text x="0" y="80" font-family="Arial,sans-serif" font-weight="bold" font-size="80" fill="white">cryptonary</text></svg>

RULES:
- Return ONLY the complete HTML starting with <!DOCTYPE html>
- Fully self-contained — all CSS inline in <style> tags
- No external resources or CDN links
- Professional, on-brand, clean design
- Include a print/screenshot note: <!-- Open in browser and screenshot for best results -->""")

    try:
        result = claude(prompt, max_tokens=3000)
        import re as _re
        html_match = _re.search(r'(<!DOCTYPE[\s\S]*?</html>)', result, _re.IGNORECASE)
        if html_match:
            return html_match.group(1), None
        elif result.strip().lower().startswith('<!doctype') or result.strip().startswith('<html'):
            return result.strip(), None
        else:
            return None, "Claude did not return valid HTML"
    except Exception as e:
        return None, str(e)


def send_svg_as_image(chat_id, svg_content):
    """Convert SVG to PNG using cairosvg if available, else send as file."""
    try:
        import cairosvg, tempfile, os
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
            png_path = tmp.name
        cairosvg.svg2png(bytestring=svg_content.encode(), write_to=png_path, scale=2.0)
        with open(png_path, "rb") as f:
            png_bytes = f.read()
        os.unlink(png_path)
        send_image_bytes(chat_id, png_bytes, "image/png")
        return True
    except ImportError:
        # cairosvg not on Render — send as SVG file (opens in browser)
        send_file_bytes(chat_id, svg_content.encode(), "cryptonary_visual.svg", "image/svg+xml")
        send(chat_id, "📎 Open the SVG file in your browser to view the full-quality image.")
        return True
    except Exception as e:
        print("SVG conversion error:", e, flush=True)
        # Last resort — send as file anyway
        try:
            send_file_bytes(chat_id, svg_content.encode(), "cryptonary_visual.svg", "image/svg+xml")
            send(chat_id, "📎 SVG sent as file — open in browser.")
            return True
        except:
            return False


def send_file_bytes(chat_id, file_bytes, filename, mime_type="application/octet-stream"):
    """Send a file to Telegram."""
    try:
        boundary = "----TgFileBoundary"
        body = (
            ("--" + boundary + "\r\n").encode() +
            ("Content-Disposition: form-data; name=\"chat_id\"\r\n\r\n" + str(chat_id) + "\r\n").encode() +
            ("--" + boundary + "\r\n").encode() +
            ("Content-Disposition: form-data; name=\"document\"; filename=\"" + filename + "\"\r\n").encode() +
            ("Content-Type: " + mime_type + "\r\n\r\n").encode() +
            file_bytes + b"\r\n" +
            ("--" + boundary + "--\r\n").encode()
        )
        req = urllib.request.Request(
            "https://api.telegram.org/bot" + TELEGRAM_TOKEN + "/sendDocument",
            data=body,
            headers={"Content-Type": "multipart/form-data; boundary=" + boundary}
        )
        with urllib.request.urlopen(req, timeout=60) as r:
            return json.loads(r.read())
    except Exception as e:
        print("send_file_bytes error:", e, flush=True)
        send(chat_id, "File generated but failed to send: " + str(e))


def generate_claude_visual(chat_id, brief, post_type, send_html=True):
    """Generate SVG + HTML visual using Claude and send both."""
    vtype_label = post_type.replace("_", " ").title()
    # Store so Give direction / Regenerate work after Claude visual
    user_state.setdefault(chat_id, {})
    user_state[chat_id]["last_img_prompt"] = brief[:600]
    user_state[chat_id]["last_img_type"] = "claude_" + post_type
    user_state[chat_id]["last_claude_brief"] = brief[:600]
    user_state[chat_id]["last_claude_type"] = post_type

    # Step 1: Generate SVG (send as image)
    send(chat_id, "🤖 Claude generating " + vtype_label + " visual...")
    svg, svg_err = generate_claude_svg(brief, post_type)

    if svg:
        print(f"SVG generated: {len(svg)} chars", flush=True)
        ok = send_svg_as_image(chat_id, svg)
        if not ok:
            send(chat_id, "SVG generated but couldn't send it.")
    else:
        print(f"SVG generation failed: {svg_err}", flush=True)
        send(chat_id, "SVG generation failed: " + (svg_err or "unknown"))

    # Step 2: Generate HTML (send as downloadable file)
    if send_html:
        send(chat_id, "📄 Generating HTML version (editable)...")
        html, html_err = generate_claude_html(brief, post_type)
        if html:
            filename = "cryptonary_" + post_type + ".html"
            send_file_bytes(chat_id, html.encode(), filename, "text/html")
            send(chat_id, "Open the HTML file in your browser for an editable, full-quality version.")
        else:
            send(chat_id, "HTML generation failed: " + (html_err or "unknown"))


def apply_logo_watermark(image_bytes, mime_type="image/png"):
    """Overlay Cryptonary logomark on bottom-right of image."""
    try:
        from PIL import Image
        import io, base64
        # Load the generated image
        img = Image.open(io.BytesIO(image_bytes)).convert("RGBA")
        w, h = img.size
        # Load logo from base64
        logo_bytes = base64.b64decode(CRYPTONARY_LOGO_B64)
        logo = Image.open(io.BytesIO(logo_bytes)).convert("RGBA")
        # Scale logo to ~8% of image width
        logo_w = max(60, int(w * 0.08))
        ratio = logo_w / logo.width
        logo_h = int(logo.height * ratio)
        logo = logo.resize((logo_w, logo_h), Image.LANCZOS)
        # Position: bottom-right with 16px padding
        padding = int(w * 0.015)
        pos = (w - logo_w - padding, h - logo_h - padding)
        # Composite onto image
        img.paste(logo, pos, logo)
        # Convert back to bytes
        output = io.BytesIO()
        img.convert("RGB").save(output, format="PNG")
        return output.getvalue(), "image/png"
    except Exception as e:
        print(f"Watermark error: {e}", flush=True)
        return image_bytes, mime_type  # return original if watermark fails


def generate_and_send_image(chat_id, prompt, post_type="static"):
    """Route to best model for post type, generate and send image."""
    # Gemini: better for graphics with text, thumbnails, statics, branded assets
    # DALL-E: better for photorealistic backgrounds, cinematic/macro scenes
    gemini_types = {"static", "email", "story", "carousel", "ad_static",
                    "ad_video", "engagement", "price_data", "educational"}
    use_gemini = post_type in gemini_types and GEMINI_KEY

    if use_gemini:
        send(chat_id, "🎨 Generating with Gemini...")
        img_bytes, mime, error = generate_gemini_image(prompt)
        if img_bytes:
            img_bytes, mime = apply_logo_watermark(img_bytes, mime)
            send_image_bytes(chat_id, img_bytes, mime)
            return True
        else:
            print("Gemini failed:", error, flush=True)
            # 429 = billing not enabled or quota exhausted
            if error and "429" in str(error):
                send(chat_id, "⚠️ Gemini image quota exhausted or billing not enabled.\nGo to console.cloud.google.com → Billing to enable paid access.")
            if OPENAI_KEY:
                send(chat_id, "Trying DALL-E instead...")
            else:
                send(chat_id, "Image generation failed. Enable billing on Google Cloud or add OPENAI_KEY.")
                return False

    if OPENAI_KEY:
        send(chat_id, "🎨 Generating with DALL-E...")
        url, info = generate_dalle_image(prompt)
        if url:
            # Download and watermark before sending
            try:
                import urllib.request as _req
                with _req.urlopen(url, timeout=30) as _r:
                    dalle_bytes = _r.read()
                dalle_bytes, _ = apply_logo_watermark(dalle_bytes, "image/png")
                send_image_bytes(chat_id, dalle_bytes, "image/png")
            except Exception as _e:
                print(f"DALL-E watermark error: {_e}", flush=True)
                send_image_url(chat_id, url)  # fallback to direct URL
            return True
        else:
            send(chat_id, "Image generation failed: " + info)
            return False

    # Final fallback — Claude SVG requires no external API
    send(chat_id, "⚠️ No image API keys configured. Using Claude SVG instead...")
    brief = _global_brief_store.get(chat_id, prompt[:300])
    generate_claude_visual(chat_id, brief, post_type, send_html=True)
    return True


def send_image_url(chat_id, image_url, caption=""):
    """Send a generated image to Telegram by URL."""
    try:
        data = {"chat_id": chat_id, "photo": image_url}
        if caption:
            data["caption"] = caption
        tg("sendPhoto", data)
    except Exception as e:
        send(chat_id, "Image generated but couldn't send it directly. URL: " + image_url)


def show_image_type_menu(chat_id):
    """Ask which AI to use for image generation."""
    print(f"show_image_type_menu called for {chat_id}", flush=True)
    keyboard = [
        [{"text": "🤖 Claude — Storyboard / Layout / Data", "callback_data": "img_engine_claude"}],
        [{"text": "🎨 Gemini — Thumbnail / Graphic / Static", "callback_data": "img_engine_gemini"}],
        [{"text": "🖼️ DALL-E — Cinematic / Photo / Macro",   "callback_data": "img_engine_dalle"}],
    ]
    result = send(chat_id, "Which AI for image generation?\n\nClaude — storyboards, layouts, data visuals\nGemini — thumbnails, static posts, graphics\nDALL-E — cinematic backgrounds, editorial photos", keyboard)
    print(f"Engine picker sent: {result}", flush=True)


def show_image_style_menu(chat_id, engine):
    """After engine selected, ask what style/type."""
    user_state.setdefault(chat_id, {})
    user_state[chat_id]["img_engine"] = engine
    state = user_state[chat_id]

    if engine == "claude":
        keyboard = [
            [{"text": "📋 Storyboard",       "callback_data": "img_style_storyboard"}],
            [{"text": "🖼️ Static post",      "callback_data": "img_style_static"}],
            [{"text": "📖 Carousel slide",   "callback_data": "img_style_carousel"}],
            [{"text": "📱 Story frame",      "callback_data": "img_style_story"}],
            [{"text": "📊 Data visual",      "callback_data": "img_style_data"}],
            [{"text": "📧 Email thumbnail",  "callback_data": "img_style_thumbnail"}],
        ]
        send(chat_id, "*What type of Claude visual?*", keyboard)
    else:
        keyboard = [
            [{"text": "📰 Breaking news",    "callback_data": "img_style_breaking_news"}],
            [{"text": "📊 Price / data",     "callback_data": "img_style_price_data"}],
            [{"text": "🗳️ Engagement post",  "callback_data": "img_style_engagement"}],
            [{"text": "📚 Educational",       "callback_data": "img_style_educational"}],
            [{"text": "🎭 Meme / cultural",  "callback_data": "img_style_meme_cultural"}],
            [{"text": "🌍 Macro / geo",      "callback_data": "img_style_macro_geo"}],
            [{"text": "🖼️ Background only",  "callback_data": "img_style_background"}],
        ]
        send(chat_id, "*What style?*", keyboard)


def handle_image_generation(chat_id, post_type=None):
    """Generate an image based on current state context."""
    state = user_state.get(chat_id, {})

    # Pull concept and angle from whatever was last generated
    concept = (
        state.get("current_social", "") or
        state.get("report", "") or
        state.get("ie_selected_concept", "") or
        state.get("ad_theme", "")
    )[:300]

    angle = (
        state.get("social_angle", "") or
        state.get("selected_angle", "") or
        state.get("ie_selected_angle", "")
    )[:150]

    if not concept:
        state["stage"] = "img_awaiting_brief"
        send(chat_id, "Describe what you want the image to show:")
        return

    if not post_type:
        state["pending_img_concept"] = concept
        state["pending_img_angle"] = angle
        show_image_type_menu(chat_id)
        return

    prompt = build_image_prompt(concept, angle, post_type)
    state["last_img_prompt"] = prompt
    state["last_img_type"] = post_type
    success = generate_and_send_image(chat_id, prompt, post_type)
    if success:
        keyboard = [
            [{"text": "🔄 Regenerate",        "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",     "callback_data": "img_direction"}],
            [{"text": "🎨 Different style",    "callback_data": "img_restyle"}],
            [{"text": "📋 Show prompt used",   "callback_data": "img_show_prompt"}],
            [{"text": "✅ Done",               "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Image ready.", keyboard)



# ══════════════════════════════════════════════════════════════════
# VISUAL BRIEF GENERATOR
# Context-aware visual output per content type
# ══════════════════════════════════════════════════════════════════

def generate_visual_brief(chat_id, content_type):
    """Generate a visual brief appropriate to the content type."""
    state = user_state.get(chat_id, {})
    current_content = state.get("current_social", "") or state.get("current_emails", {}) or state.get("current_ad", "")
    angle = state.get("selected_angle", "") or state.get("social_angle", "")
    report = sanitise(state.get("report", ""))[:500]

    # Build type-specific brief prompt
    brief_prompts = {
        "reel": """Create a STORYBOARD BRIEF for this Reel script.

Format:
SCENE 1: [0-3s] HOOK
Visual: [what's on screen]
Text overlay: [any text shown]
B-roll direction: [footage type]

SCENE 2: [3-10s] SETUP
(repeat format for each scene)

FINAL SCENE: CTA
Visual direction: [closing shot]
Text: [final text overlay]

Also write:
THUMBNAIL OPTION 1: [bold text + visual direction for cover frame]
THUMBNAIL OPTION 2: [alternative]""",

        "carousel": """Create a SLIDE-BY-SLIDE BRIEF for this carousel.

For each slide:
SLIDE [N]: [slide purpose]
Headline text: [max 8 words — this goes ON the image]
Supporting text: [max 15 words] (optional)
Visual direction: [background image / graphic / data viz]
Colour: [which brand colour dominates this slide]

Cover slide must be the strongest hook.
Final slide: CTA with link in bio direction.""",

        "static": """Create a STATIC IMAGE BRIEF for this post.

ON-IMAGE TEXT (this is all that appears on the image):
Headline: [max 6 words — bold, Tungsten-style]
Subtext: [max 4 words] (optional)

VISUAL DIRECTION:
Background: [photo type / graphic style / colour]
Colour palette: [from Cryptonary brand — black/white/red/green/blue/bitcoin orange]
Logo placement: [bottom-right @Cryptonary mark]
Overall feel: [e.g. breaking news / editorial / data / meme]

AI IMAGE PROMPT (for DALL-E/Midjourney):
[Full detailed prompt ready to paste]""",

        "story": """Create a STORY FRAME BRIEF for each slide.

For each slide:
SLIDE [N]:
Screen text: [max 12 words — this appears on screen]
Background: [image direction or colour]
Sticker/element: [poll / countdown / emoji / swipe up]
Animation feel: [fade in / slide / pop]

Keep it urgent and swipeable.""",

        "email": """Create a THUMBNAIL BRIEF for this email.

Email thumbnails appear in the preview pane and drive open rates.

THUMBNAIL OPTION 1:
Style: [news / data / meme / quote]
Headline shown: [max 6 words]
Visual: [image direction]
Colour: [dominant colour from brand palette]

THUMBNAIL OPTION 2: [alternative angle]

THUMBNAIL OPTION 3: [alternative — try a different style]

Note: thumbnails are 600x300px, text must be readable at small size.""",

        "ad_static": """Create a STATIC AD CREATIVE BRIEF.

PRIMARY HEADLINE (on image, max 6 words):
[bold, stops the scroll]

SUPPORTING TEXT (on image, max 10 words):
[reinforces the hook]

PRIMARY TEXT (ad copy — shown below image):
[already written in the ad copy above]

VISUAL DIRECTION:
Background: [photo / graphic / colour]
Style: [AIDA awareness / consideration / conversion]
CTA button text: [Learn More / Sign Up / Get Started]

CREATIVE VARIANTS:
Variant A: [version 1 concept]
Variant B: [version 2 concept — different visual angle]""",

        "ad_video": """Create a VIDEO AD STORYBOARD BRIEF.

HOOK FRAME [0-3s]:
Visual: [what's on screen]
Text overlay: [opening hook text]
Audio direction: [music / silence / SFX]

BODY [3-15s]:
Scene breakdown: [2-3 key scenes with visual + text]

CTA [final 3s]:
Visual: [closing shot]
Text: [CTA overlay]
Button: [Sign Up / Learn More]

THUMBNAIL: [best frame to use as cover — describe it]"""
    }

    prompt_template = brief_prompts.get(content_type, brief_prompts["static"])

    send(chat_id, "Creating visual brief...")
    try:
        raw_content = ""
        if isinstance(current_content, dict):
            raw_content = current_content.get("free", "") or str(current_content)[:600]
        else:
            raw_content = str(current_content)[:600]

        result = claude(
            "CONTENT:\n" + raw_content +
            "\nANGLE: " + angle +
            "\nSOURCE CONTEXT: " + report +
            "\n\n" + prompt_template +
            "\n\nApply Cryptonary brand guidelines: dark backgrounds preferred, "
            "Inter/Tungsten/Termina fonts, colours: black #000000, white #FFFFFF, "
            "blue #005EFF, red #FF0000, green #0DA500, bitcoin orange #F7931A. "
            "Logo @Cryptonary always bottom-right.",
            max_tokens=900
        )
        state["last_visual_brief"] = result
        state["last_visual_type"] = content_type
        _global_brief_store[chat_id] = result  # persist across state changes
        send_plain(chat_id, "*VISUAL BRIEF — " + content_type.upper().replace("_", " ") + "*\n\n" + result)

        img_row = [{"text": "🎨 Generate image", "callback_data": "img_from_brief"}] if (OPENAI_KEY or GEMINI_KEY) else []
        keyboard = [
            [{"text": "✏️ Adjust brief", "callback_data": "vb_edit"}],
        ]
        if img_row:
            keyboard.insert(0, img_row)
        keyboard.append([{"text": "✅ Done", "callback_data": "mark_complete"}])
        send(chat_id, "Brief ready. Use this to brief your designer, or generate an AI image draft.", keyboard)
    except Exception as e:
        send(chat_id, "Error generating brief: " + str(e))


def show_visual_brief_menu(chat_id, auto_type=None):
    """Show visual brief type options or auto-detect from current content."""
    state = user_state.get(chat_id, {})

    # Auto-detect from current content type if possible
    if auto_type:
        generate_visual_brief(chat_id, auto_type)
        return

    social_type = state.get("current_social_type", "")
    if "Reel" in social_type:
        generate_visual_brief(chat_id, "reel")
    elif "Carousel" in social_type:
        generate_visual_brief(chat_id, "carousel")
    elif "Static" in social_type:
        generate_visual_brief(chat_id, "static")
    elif "Story" in social_type:
        generate_visual_brief(chat_id, "story")
    elif state.get("current_emails"):
        generate_visual_brief(chat_id, "email")
    elif state.get("current_ad"):
        ad_type = state.get("ad_type", "static")
        generate_visual_brief(chat_id, "ad_" + ad_type)
    else:
        keyboard = [
            [{"text": "📧 Email thumbnail",     "callback_data": "vb_type_email"}],
            [{"text": "🎬 Reel storyboard",     "callback_data": "vb_type_reel"}],
            [{"text": "🖼️ Static image brief",  "callback_data": "vb_type_static"}],
            [{"text": "📖 Carousel brief",       "callback_data": "vb_type_carousel"}],
            [{"text": "📱 Story brief",          "callback_data": "vb_type_story"}],
            [{"text": "📢 Ad creative brief",    "callback_data": "vb_type_ad_static"}],
        ]
        send(chat_id, "*What visual brief do you need?*", keyboard)


def handle_image_callbacks(chat_id, data, state):
    """Handle all image-related callbacks."""
    print(f"handle_image_callbacks: {data}", flush=True)
    if data.startswith("img_engine_"):
        engine = data.replace("img_engine_", "")
        print(f"Engine selected: {engine}", flush=True)
        show_image_style_menu(chat_id, engine)

    elif data.startswith("img_style_"):
        style = data.replace("img_style_", "")
        engine = state.get("img_engine", "gemini")
        brief = (_global_brief_store.get(chat_id, "") or 
                 state.get("pending_img_concept", "") or
                 state.get("last_visual_brief", "") or
                 state.get("report", ""))[:600]
        print(f"img_style_ fired. style={style}, engine={engine}, brief_len={len(brief)}", flush=True)
        angle = state.get("pending_img_angle", state.get("social_angle", ""))[:150]
        state["last_img_type"] = style

        img_action_kb = [
            [{"text": "🔄 Regenerate",      "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",   "callback_data": "img_direction"}],
            [{"text": "🎨 Different engine", "callback_data": "img_restyle"}],
            [{"text": "✅ Done",             "callback_data": "mark_complete"}],
        ]

        if engine == "claude":
            generate_claude_visual(chat_id, brief, style, send_html=True)
            send(chat_id, "Claude visual ready.", img_action_kb)
        else:
            post_type = style
            prompt = build_image_prompt(brief, angle, post_type)
            state["last_img_prompt"] = prompt
            prefer_gemini = (engine == "gemini")
            # Override routing — force the chosen engine
            if engine == "gemini" and GEMINI_KEY:
                send(chat_id, "🎨 Generating with Gemini...")
                img_bytes, mime, err = generate_gemini_image(prompt)
                if img_bytes:
                    img_bytes, mime = apply_logo_watermark(img_bytes, mime)
                    send_image_bytes(chat_id, img_bytes, mime)
                    send(chat_id, "Image ready.", img_action_kb)
                else:
                    send(chat_id, "Gemini failed: " + (err or "unknown") + "\nTrying DALL-E...")
                    safe_prompt = build_image_prompt(brief[:150], "", post_type)
                    url, info = generate_dalle_image(safe_prompt) if OPENAI_KEY else (None, "No DALL-E key")
                    if url:
                        send_image_url(chat_id, url)
                        send(chat_id, "Image ready (DALL-E fallback).", img_action_kb)
                    else:
                        send(chat_id, "🤖 Falling back to Claude SVG...")
                        generate_claude_visual(chat_id, brief[:400], post_type, send_html=False)
                        send(chat_id, "Claude SVG used as final fallback.", img_action_kb)
            elif OPENAI_KEY:
                send(chat_id, "🖼️ Generating with DALL-E...")
                url, info = generate_dalle_image(prompt)
                if url:
                    try:
                        import urllib.request as _rq
                        with _rq.urlopen(url, timeout=30) as _r:
                            _b = _r.read()
                        _b, _ = apply_logo_watermark(_b, "image/png")
                        send_image_bytes(chat_id, _b, "image/png")
                    except:
                        send_image_url(chat_id, url)
                    send(chat_id, "Image ready.", img_action_kb)
                else:
                    # DALL-E failed — try stripped prompt (content policy fix)
                    send(chat_id, "Retrying with simplified prompt...")
                    safe_prompt = build_image_prompt(brief[:150], "", post_type)
                    url2, info2 = generate_dalle_image(safe_prompt)
                    if url2:
                        send_image_url(chat_id, url2)
                        send(chat_id, "Image ready (simplified prompt).", img_action_kb)
                    elif GEMINI_KEY:
                        send(chat_id, "Trying Gemini instead...")
                        img_bytes, mime, err = generate_gemini_image(safe_prompt)
                        if img_bytes:
                            send_image_bytes(chat_id, img_bytes, mime)
                            send(chat_id, "Image ready (Gemini fallback).", img_action_kb)
                        else:
                            send(chat_id, "🤖 Falling back to Claude SVG...")
                            generate_claude_visual(chat_id, brief[:400], post_type, send_html=False)
                            send(chat_id, "Claude SVG used as final fallback.", img_action_kb)
                    else:
                        send(chat_id, "🤖 Falling back to Claude SVG...")
                        generate_claude_visual(chat_id, brief[:400], post_type, send_html=False)
                        send(chat_id, "Claude SVG used as final fallback.", img_action_kb)
            else:
                # No paid APIs — use Claude SVG
                send(chat_id, "🤖 Using Claude SVG (no paid image API configured)...")
                generate_claude_visual(chat_id, brief[:400], post_type, send_html=False)
                send(chat_id, "Claude SVG generated.", img_action_kb)

    elif data.startswith("img_type_"):
        # Legacy fallback
        post_type = data.replace("img_type_", "")
        concept = state.get("pending_img_concept", state.get("report", ""))[:300]
        angle = state.get("pending_img_angle", state.get("social_angle", ""))[:150]
        prompt = build_image_prompt(concept, angle, post_type)
        state["last_img_prompt"] = prompt
        state["last_img_type"] = post_type
        success = generate_and_send_image(chat_id, prompt, post_type)
        keyboard = [
            [{"text": "🔄 Regenerate",      "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",   "callback_data": "img_direction"}],
            [{"text": "🎨 Different style",  "callback_data": "img_restyle"}],
            [{"text": "✅ Done",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Image ready." if success else "Generation failed.", keyboard)

    elif data == "img_regen":
        prompt = state.get("last_img_prompt", "")
        if not prompt:
            send(chat_id, "No previous prompt found.")
            return
        post_type = state.get("last_img_type", "static")
        img_kb = [
            [{"text": "🔄 Try again",           "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",       "callback_data": "img_direction"}],
            [{"text": "🎨 Different engine",     "callback_data": "img_restyle"}],
            [{"text": "✅ Done",                 "callback_data": "mark_complete"}],
        ]
        if post_type.startswith("claude_"):
            actual_type = post_type.replace("claude_", "")
            brief = state.get("last_claude_brief", prompt[:400])
            generate_claude_visual(chat_id, brief, actual_type, send_html=False)
            send(chat_id, "Regenerated.", img_kb)
        else:
            success = generate_and_send_image(chat_id, prompt, post_type)
            send(chat_id, "New version ready." if success else "Generation failed — try again.", img_kb)

    elif data == "img_direction":
        state["stage"] = "img_awaiting_direction"
        send(chat_id, "Describe what to change or add:")

    elif data == "img_restyle":
        # Show engine picker so user can switch between Claude/Gemini/DALL-E
        show_image_type_menu(chat_id)

    elif data == "img_show_prompt":
        prompt = state.get("last_img_prompt", "No prompt saved.")
        send_plain(chat_id, "PROMPT USED:\n\n" + prompt[:1000])

    elif data == "img_open_studio":
        handle_image_generation(chat_id)


def handle_image_direction(chat_id, text):
    """Apply user direction to regenerate image."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    base_prompt = state.get("last_img_prompt", "") or state.get("last_claude_brief", "")
    post_type = state.get("last_img_type", "background")
    if not base_prompt:
        send(chat_id, "No previous image to refine. Start a new one.")
        return
    refined = base_prompt + "\n\nREFINEMENT: " + text
    state["last_img_prompt"] = refined

    # Route back to correct engine
    if post_type.startswith("claude_"):
        actual_type = post_type.replace("claude_", "")
        state["last_claude_brief"] = refined
        generate_claude_visual(chat_id, refined, actual_type, send_html=False)
        keyboard = [
            [{"text": "🔄 Regenerate",      "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",   "callback_data": "img_direction"}],
            [{"text": "🎨 Different engine", "callback_data": "img_restyle"}],
            [{"text": "✅ Done",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Updated.", keyboard)
        return
    success = generate_and_send_image(chat_id, refined, post_type)
    if success:
        keyboard = [
            [{"text": "🔄 Regenerate",      "callback_data": "img_regen"}],
            [{"text": "✏️ Give direction",   "callback_data": "img_direction"}],
            [{"text": "🎨 Different style",  "callback_data": "img_restyle"}],
            [{"text": "✅ Done",             "callback_data": "mark_complete"}],
        ]
        send(chat_id, "Updated.", keyboard)


def fetch_market_data():
    """Fetch BTC price, Fear & Greed index, and top crypto news."""
    results = {}
    try:
        # BTC price from CoinGecko public API
        req = urllib.request.Request(
            "https://api.coingecko.com/api/v3/simple/price?ids=bitcoin&vs_currencies=usd&include_24hr_change=true&include_7d_change=true",
            headers={"Accept": "application/json", "User-Agent": "Mozilla/5.0"}
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read())
            btc = data.get("bitcoin", {})
            results["btc_price"] = btc.get("usd", 0)
            results["btc_24h"] = round(btc.get("usd_24h_change", 0), 1)
            results["btc_7d"] = round(btc.get("usd_7d_change", 0), 1)
    except Exception as e:
        results["btc_price"] = 0
        results["btc_24h"] = 0
        results["btc_7d"] = 0
        print("BTC price fetch error:", e, flush=True)

    try:
        # Fear & Greed from alternative.me
        req = urllib.request.Request(
            "https://api.alternative.me/fng/?limit=1",
            headers={"Accept": "application/json", "User-Agent": "Mozilla/5.0"}
        )
        with urllib.request.urlopen(req, timeout=10) as r:
            data = json.loads(r.read())
            fng = data.get("data", [{}])[0]
            results["fng_value"] = fng.get("value", "N/A")
            results["fng_label"] = fng.get("value_classification", "N/A")
    except Exception as e:
        results["fng_value"] = "N/A"
        results["fng_label"] = "N/A"
        print("Fear & Greed fetch error:", e, flush=True)

    return results

def get_last_week_performance(chat_id):
    """Summarise last week's content and performance from stored data."""
    summary = []
    records = content_metadata.get(chat_id, [])
    insights = pattern_insights.get(chat_id, [])
    email_records = email_log.get(chat_id, [])
    ad_records = ad_log.get(chat_id, [])

    # Content created in last 7 days
    from datetime import datetime as _datetime, timedelta as _timedelta
    week_ago = (_datetime.now() - _timedelta(days=7)).strftime("%Y-%m-%d")
    recent = [r for r in records if r.get("date","") >= week_ago]
    if recent:
        email_count = len([r for r in recent if r.get("type") == "email"])
        ad_count = len([r for r in recent if r.get("type") == "ad"])
        summary.append(str(len(recent)) + " pieces created last week (" + str(email_count) + " emails, " + str(ad_count) + " ads)")

    # Best logged email
    if email_records:
        best = max(email_records, key=lambda r: r.get("open_rate", 0))
        summary.append("Best email: " + best.get("subject","")[:50] + " — " + str(best.get("open_rate","?")) + "% open / " + str(best.get("ctr","?")) + "% CTR")

    # Recent insights
    recent_insights = [i for i in insights if i.get("date","") >= week_ago]
    if recent_insights:
        summary.append("Recent pattern: " + recent_insights[-1].get("insight",""))

    return summary

def generate_briefing(chat_id):
    """Generate and send the full weekly briefing."""
    send(chat_id, "Generating your weekly briefing...")

    # Register this chat for auto-briefing
    subscribers = load_briefing_subscribers()
    if chat_id not in subscribers:
        subscribers.append(chat_id)
        save_briefing_subscribers(subscribers)

    # Fetch live data
    market = fetch_market_data()
    perf_summary = get_last_week_performance(chat_id)

    # Build context for Claude
    market_context = ""
    if market.get("btc_price"):
        price = "${:,.0f}".format(market["btc_price"])
        market_context += "BTC: " + price + " (" + str(market["btc_7d"]) + "% 7d, " + str(market["btc_24h"]) + "% 24h)\n"
    if market.get("fng_value") != "N/A":
        market_context += "Fear & Greed: " + str(market["fng_value"]) + " (" + market["fng_label"] + ")\n"

    perf_context = "\n".join(perf_summary) if perf_summary else "No performance data logged yet."

    try:
        # Use web search for news via Claude
        news_prompt = "Search for the top 3 most important crypto/Bitcoin news stories from the past 7 days. For each give: headline, one sentence summary, and why it matters for crypto investors and marketers. Be specific with dates and numbers."

        include_recap = user_state.get(chat_id, {}).pop("include_recap", False)
        last_brief = user_state.get(chat_id, {}).get("last_briefing_summary", "")
        last_brief_date = user_state.get(chat_id, {}).get("last_briefing_date", "")

        recap_section = ""
        if include_recap and last_brief:
            recap_section = """
RECAP SECTION (include this FIRST if on-command):
SINCE LAST BRIEF (""" + last_brief_date + """):
[Brief 2-3 sentence recap of what was in the last briefing, then what's changed in the market and performance since then]
---
"""

        briefing_prompt = """You are the chief strategist for Cryptonary, a crypto research platform with 300K+ subscribers.

Generate a """ + ("full briefing with recap" if include_recap else "weekly Monday morning briefing") + """ for Adam (Co-Founder). Be direct, specific, and actionable.

MARKET DATA:
""" + market_context + """

LAST WEEK'S PERFORMANCE:
""" + perf_context + """

""" + recap_section + """

STRUCTURE THE BRIEFING EXACTLY AS:

MARKET PULSE
[BTC price, 7d change, Fear & Greed with one line on what it means for content tone this week]

TOP STORIES THIS WEEK
[3 most important crypto stories from the past 7 days with specific numbers and why each matters for Cryptonary's content]

LAST WEEK RECAP
[Summary of what was created and best performance metrics. If no data: note that logging performance in Data Studio will make this section more useful over time]

THIS WEEK'S CONTENT PLAN
[Specific recommendations based on market conditions and news:]
- Monday email angle: [specific suggestion based on market mood]
- Mid-week content: [social or email suggestion]
- Ad focus: [which avatar/stage to push given current market sentiment]
- Social: [1-2 specific post ideas tied to the week's news]

TONE RECOMMENDATION
[Based on Fear & Greed + price action: Standard / Aggressive / Empathetic and why]

Keep it tight. Every line should be actionable. No fluff."""

        result = claude(briefing_prompt, max_tokens=1500, system=DATA_STUDIO_SYSTEM)

        # Format header
        from datetime import datetime as _dt2
        date_str = _dt2.now().strftime("%A %d %B %Y")
        header = "CRYPTONARY WEEKLY BRIEFING\n" + date_str + "\n\n"

        send_plain(chat_id, header + result)

        # Save summary for next on-command recap
        from datetime import datetime as _dt_now
        user_state[chat_id]["last_briefing_summary"] = result[:500]
        user_state[chat_id]["last_briefing_date"] = _dt_now.now().strftime("%d %b %Y")

        keyboard = [
            [{"text": "Generate Monday email", "callback_data": "mode_email"},
             {"text": "Generate social content", "callback_data": "mode_social"}],
            [{"text": "Generate ads", "callback_data": "mode_ads"},
             {"text": "Refresh briefing", "callback_data": "briefing_refresh"}]
        ]
        send(chat_id, "Briefing complete. What do you want to create first?", keyboard)

    except Exception as e:
        print("Briefing error:", e, flush=True)
        send(chat_id, "Error generating briefing: " + str(e))


# ── AUTO BRIEFING SCHEDULER ───────────────────────────────────────
# Runs inside the poll loop - checks if it's Monday 7-8am GMT and sends briefing

_last_briefing_date = None

def check_and_send_briefing():
    """Called from poll loop. Sends briefing to all subscribers on Monday 7-8am GMT."""
    global _last_briefing_date
    from datetime import datetime as _dt3
    now = _dt3.utcnow()
    today = now.strftime("%Y-%m-%d")

    # Monday = 0, hour 7
    if now.weekday() == 0 and now.hour == 9 and _last_briefing_date != today:
        _last_briefing_date = today
        subscribers = load_briefing_subscribers()
        print("Sending weekly briefing to", len(subscribers), "subscribers", flush=True)
        for chat_id in subscribers:
            try:
                generate_briefing(chat_id)
                time.sleep(2)  # avoid rate limits between sends
            except Exception as e:
                print("Briefing send error for", chat_id, ":", e, flush=True)

# ══════════════════════════════════════════════════════════════════
# CRYPTONARY IDEA ENGINE
# ══════════════════════════════════════════════════════════════════

IDEA_ENGINE_SOURCES_FILE = "idea_sources.json"

# ── RSS FEED LIBRARY ─────────────────────────────────────────────
# All feeds verified live. Grouped by priority tier.
RSS_FEEDS = [
    # TIER 1 — Mainstream crypto news
    {"name": "CoinDesk",        "url": "https://coindesk.com/arc/outboundfeeds/rss/",  "tier": 1},
    {"name": "CoinTelegraph",   "url": "https://cointelegraph.com/rss",                "tier": 1},
    {"name": "Decrypt",         "url": "https://decrypt.co/feed",                      "tier": 1},
    {"name": "Bitcoin Magazine","url": "https://bitcoinmagazine.com/feed",              "tier": 1},
    {"name": "NewsBTC",         "url": "https://newsbtc.com/feed",                     "tier": 1},
    # TIER 2 — Analysis & data
    {"name": "Messari",         "url": "https://messari.io/rss",                       "tier": 2},
    {"name": "Glassnode",       "url": "https://insights.glassnode.com/rss",           "tier": 2},
    {"name": "The Defiant",     "url": "https://thedefiant.io/api/feed",               "tier": 2},
    {"name": "Blockworks",      "url": "https://blockworks.co/feed",                   "tier": 2},
    # TIER 3 — Market & price focused
    {"name": "U.Today",         "url": "https://u.today/rss",                          "tier": 3},
    {"name": "Crypto.news",     "url": "https://crypto.news/feed",                     "tier": 3},
    {"name": "Investing.com",   "url": "https://investing.com/rss/news_301.rss",       "tier": 3},
    {"name": "CoinJournal",     "url": "https://coinjournal.net/feed",                 "tier": 3},
    # TIER 4 — Broader web3
    {"name": "Watcher Guru",    "url": "https://watcherguru.com/feed",                 "tier": 4},
    {"name": "Milk Road",       "url": "https://milkroad.substack.com/feed",           "tier": 4},
    {"name": "Glassnode Blog",  "url": "https://glassnode.com/feed",                   "tier": 4},
    {"name": "Santiment",       "url": "https://santiment.net/blog/feed",              "tier": 4},
]

# Reddit still works reliably — keep as supplement
REDDIT_SOURCES = ["cryptocurrency", "bitcoin", "CryptoMarkets"]

DEFAULT_SOURCES = {
    "reddit": REDDIT_SOURCES,
    "telegram": [],
}


def fetch_rss_feed(feed, max_items=3):
    """Fetch and parse a single RSS feed. Returns list of headline dicts."""
    import xml.etree.ElementTree as ET
    import re as _re
    results = []
    try:
        req = urllib.request.Request(
            feed["url"],
            headers={"User-Agent": "Mozilla/5.0 (compatible; CryptonaryBot/1.0)"}
        )
        with urllib.request.urlopen(req, timeout=8) as r:
            raw = r.read()
        root = ET.fromstring(raw)
        # Handle both RSS <item> and Atom <entry> formats
        items = root.findall(".//item") or root.findall(".//{http://www.w3.org/2005/Atom}entry")
        for item in items[:max_items]:
            title = (
                item.findtext("title") or
                item.findtext("{http://www.w3.org/2005/Atom}title") or ""
            ).strip()
            # Strip HTML from description
            desc_raw = (
                item.findtext("description") or
                item.findtext("{http://www.w3.org/2005/Atom}summary") or ""
            )
            desc = _re.sub(r"<[^>]+>", " ", desc_raw).strip()[:150]
            if title:
                results.append({
                    "source": feed["name"],
                    "tier": feed["tier"],
                    "title": title,
                    "desc": desc,
                })
    except Exception as e:
        print(f"RSS fetch error [{feed['name']}]: {e}", flush=True)
    return results


def fetch_rss_headlines(max_per_feed=2, max_total=20):
    """Pull fresh headlines from all RSS feeds. Tier 1 first, then 2, 3, 4."""
    all_headlines = []
    # Sort by tier so best sources come first
    sorted_feeds = sorted(RSS_FEEDS, key=lambda f: f["tier"])
    for feed in sorted_feeds:
        if len(all_headlines) >= max_total:
            break
        items = fetch_rss_feed(feed, max_items=max_per_feed)
        all_headlines.extend(items)
    return all_headlines[:max_total]


def format_rss_context(headlines):
    """Format RSS headlines into a clean context string for the prompt."""
    if not headlines:
        return ""
    lines = ["\n=== LIVE CRYPTO NEWS & TRENDS (pulled now) ==="]
    current_tier = None
    tier_labels = {1: "TOP STORIES", 2: "ANALYSIS", 3: "MARKET", 4: "COMMUNITY"}
    for h in headlines:
        if h["tier"] != current_tier:
            current_tier = h["tier"]
            lines.append("\n[" + tier_labels.get(current_tier, "NEWS") + "]")
        line = "• " + h["source"] + ": " + h["title"]
        if h.get("desc"):
            line += " — " + h["desc"][:80]
        lines.append(line)
    lines.append("===")
    return "\n".join(lines)

def load_idea_sources():
    try:
        with open(IDEA_ENGINE_SOURCES_FILE, "r") as f:
            return json.load(f)
    except:
        return DEFAULT_SOURCES.copy()

def save_idea_sources(sources):
    try:
        with open(IDEA_ENGINE_SOURCES_FILE, "w") as f:
            json.dump(sources, f)
    except Exception as e:
        print("Sources save error:", e, flush=True)

def show_idea_engine_menu(chat_id):
    user_state[chat_id] = {"stage": "idea_engine_idle"}
    feed_count = len(RSS_FEEDS)
    keyboard = [
        [{"text": "💡 Generate ideas from scratch", "callback_data": "ie_from_scratch"}],
        [{"text": "🔗 Generate from inspiration", "callback_data": "ie_from_inspiration"}],
        [{"text": "Critique a screenshot", "callback_data": "ie_screenshot_critique"}],
        [{"text": "📡 News sources (" + str(feed_count) + " feeds)", "callback_data": "ie_manage_sources"}]
    ]
    send(chat_id, "*Creative Studio*\n\nHow do you want to generate ideas?", keyboard)

def show_ie_source_menu(chat_id, idea_type):
    """After type selection: pull from database or upload new inspiration."""
    user_state[chat_id]["ie_idea_type"] = idea_type
    type_label = {"ad": "Ad", "instagram": "Instagram"}.get(idea_type, idea_type)
    keyboard = [
        [{"text": "Pull from saved sources", "callback_data": "ie_source_database"}],
        [{"text": "Upload / paste new inspiration", "callback_data": "ie_source_upload"}],
        [{"text": "Drop a link", "callback_data": "ie_source_link"}]
    ]
    send(chat_id, "*" + type_label + " Idea*\n\nWhere should I pull inspiration from?", keyboard)

def show_ie_format_menu(chat_id):
    """After source selection: pick specific format."""
    idea_type = user_state[chat_id].get("ie_idea_type", "ad")
    if idea_type == "ad":
        keyboard = [
            [{"text": "Video script", "callback_data": "ie_format_video_script"}],
            [{"text": "Static ad", "callback_data": "ie_format_static"}]
        ]
        send(chat_id, "*What type of ad?*", keyboard)
    else:
        keyboard = [
            [{"text": "Reel", "callback_data": "ie_format_reel"}],
            [{"text": "Carousel", "callback_data": "ie_format_carousel"}],
            [{"text": "Static post", "callback_data": "ie_format_static_post"}],
            [{"text": "Story", "callback_data": "ie_format_story"}]
        ]
        send(chat_id, "*What format?*", keyboard)

def generate_ie_concept(chat_id):
    """Generate concept ideas — Instagram post focused. Concept + Angle + Source per idea."""
    user_state.setdefault(chat_id, {"stage": "idea_engine_idle"})
    state = user_state[chat_id]
    source_content = state.get("ie_source_content", "")
    source_label = state.get("ie_source_label", "Cryptonary library")
    ctx = get_content_context(chat_id)

    market = fetch_market_data()
    market_line = ""
    if market.get("btc_price"):
        market_line = "BTC: ${:,.0f} | Fear & Greed: {} ({})".format(
            market["btc_price"], market.get("fng_value","?"), market.get("fng_label","?"))

    # If no manual source content, pull live RSS + X tweets
    if not source_content:
        send(chat_id, "Pulling live trends...")
        headlines = fetch_rss_headlines(max_per_feed=2, max_total=16)
        rss_context = format_rss_context(headlines)
        # Layer X tweets on top if key is configured
        x_tweets = fetch_x_tweets(max_per_account=2, max_total=10) if X_BEARER_TOKEN else []
        x_context = format_x_context(x_tweets) if x_tweets else ""
        source_content = rss_context + x_context
        label = "Live news + Twitter" if x_tweets else "Live news feeds"
        if not user_state[chat_id].get("ie_source_label"):
            user_state[chat_id]["ie_source_label"] = label

    send(chat_id, "Generating concepts...")

    try:
        prompt = "Generate 4 Instagram content concepts for Cryptonary.\n\n"
        prompt += "MARKET: " + market_line + "\n"
        if source_content:
            prompt += "\nINSPIRATION SOURCE (" + source_label + "):\n" + source_content[:2000] + "\n"
        prompt += "\nCRYPTONARY PROVEN POST PATTERNS (use as creative reference):\n" + BEST_POSTS[:1500] + "\n"
        prompt += ctx
        prompt += """

Each concept is a distinct Instagram post idea — different angle, emotion, or narrative.
Keep each entry SHORT. No fluff.

Format EXACTLY:
1. CONCEPT: [One punchy line — the big idea]
   ANGLE: [The emotional or argumentative approach — why this hits]
   SOURCE: [What inspired it — e.g. market data, competitor post, Cryptonary track record, user pain point]

2. CONCEPT: ...
   ANGLE: ...
   SOURCE: ...

3. CONCEPT: ...
   ANGLE: ...
   SOURCE: ...

4. CONCEPT: ...
   ANGLE: ...
   SOURCE: ...

Nothing else."""

        raw = claude(prompt, max_tokens=800, system=VOICE_GUIDE)
        state["ie_concepts"] = raw
        state["stage"] = "ie_concept_review"
        send_plain(chat_id, "*IDEAS*\n\n" + raw)

        keyboard = [
            [{"text": "1 — Expand this", "callback_data": "ie_develop_concept_1"}],
            [{"text": "2 — Expand this", "callback_data": "ie_develop_concept_2"}],
            [{"text": "3 — Expand this", "callback_data": "ie_develop_concept_3"}],
            [{"text": "4 — Expand this", "callback_data": "ie_develop_concept_4"}],
            [{"text": "Regenerate", "callback_data": "ie_regen_concepts"}],
            [{"text": "✏️ Write my own", "callback_data": "ie_custom_concept"}]
        ]
        send(chat_id, "Tap a number to expand into full content:", keyboard)

    except Exception as e:
        send(chat_id, "Error: " + str(e))

def generate_ie_angle_from_concept(chat_id, concept_text):
    """Generate angles for a selected concept."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    ie_format = state.get("ie_format", "")
    state["ie_selected_concept"] = concept_text
    state["stage"] = "ie_angle_review"
    send(chat_id, "Generating angles for this concept...")

    try:
        raw = claude(
            "CONCEPT: " + concept_text +
            "\nFORMAT: " + ie_format +
            "\n\nGenerate 4 distinct ANGLES for this concept. An angle is the specific emotional or argumentative approach — not the hook.\n\n" +
            "1. [angle]\n2. [angle]\n3. [angle]\n4. [angle]\nNothing else.",
            max_tokens=400
        )
        angles = parse_numbered_list(raw, 4)
        state["ie_angles"] = angles
        text = "*ANGLES for: _" + concept_text[:60] + "_*\n\n"
        keyboard = []
        for i, a in enumerate(angles):
            text += str(i+1) + ". " + a + "\n\n"
            keyboard.append([{"text": str(i+1), "callback_data": "ie_angle_" + str(i)}])
        keyboard.append([{"text": "✏️ Write my own", "callback_data": "ie_custom_angle_ie"}])
        keyboard.append([{"text": "← Back to concepts", "callback_data": "ie_back_to_concepts"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def generate_ie_hook_from_angle(chat_id, angle_text):
    """Generate hooks for a selected angle."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    ie_format = state.get("ie_format", "")
    concept = state.get("ie_selected_concept", "")
    state["ie_selected_angle"] = angle_text
    state["stage"] = "ie_hook_review"
    send(chat_id, "Writing hooks...")

    format_hook_map = {
        "video_script": "Opening 3-second spoken line that stops the scroll",
        "static": "Headline (max 8 words) for a Meta static ad",
        "reel": "First spoken line of a Reel (max 10 words)",
        "carousel": "Cover slide headline (max 8 words, bold statement)",
        "static_post": "First line of the caption (max 10 words)",
        "story": "First slide text (max 8 words, full-screen impact)"
    }
    hook_instruction = format_hook_map.get(ie_format, "Opening hook")

    try:
        raw = claude(
            "CONCEPT: " + concept + "\nANGLE: " + angle_text +
            "\nFORMAT: " + ie_format +
            "\n\nWrite 4 distinct " + hook_instruction + "s.\n\n" +
            "1. [hook]\n2. [hook]\n3. [hook]\n4. [hook]\nNothing else.",
            max_tokens=300
        )
        hooks = parse_numbered_list(raw, 4)
        state["ie_hooks"] = hooks
        text = "*HOOKS — " + hook_instruction + "*\n\n"
        keyboard = []
        for i, h in enumerate(hooks):
            text += str(i+1) + ". " + h + "\n\n"
            keyboard.append([{"text": str(i+1), "callback_data": "ie_hook_" + str(i)}])
        keyboard.append([{"text": "✏️ Write my own", "callback_data": "ie_custom_hook_ie"}])
        keyboard.append([{"text": "← Back to angles", "callback_data": "ie_back_to_angles"}])
        send(chat_id, text, keyboard)
    except Exception as e:
        send(chat_id, "Error: " + str(e))

def generate_ie_final_content(chat_id, hook_text):
    """Generate the final content from concept → angle → hook."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    ie_format = state.get("ie_format", "")
    concept = state.get("ie_selected_concept", "")
    angle = state.get("ie_selected_angle", "")
    state["ie_selected_hook"] = hook_text
    state["stage"] = "ie_content_ready"

    format_instructions = {
        "video_script": "Write a 30-45 second Meta video ad script using AIDA structure. Include SPOKEN, ON SCREEN, and VISUAL for each section.",
        "static": "Write a complete Meta static ad: HEADLINE (max 8 words), PRIMARY TEXT (150-200 words), DESCRIPTION (max 20 words), CTA BUTTON.",
        "reel": "Write a complete Instagram Reel voiceover script (45-60 seconds). Format: text | [B-roll direction] for each line.",
        "carousel": "Write a complete 6-slide Instagram Carousel. Format: SLIDE N: [headline] | [body text] for each slide.",
        "static_post": "Write a complete Instagram static post caption. Opening hook, 3-4 lines of value, CTA, hashtags.",
        "story": "Write a complete 2-3 slide Instagram Story. Format: SLIDE N: [text overlay] | [visual direction] for each."
    }
    instruction = format_instructions.get(ie_format, "Write the content in full.")

    send(chat_id, "Generating final content...")
    try:
        result = claude(
            "CONCEPT: " + concept +
            "\nANGLE: " + angle +
            "\nOPENING HOOK: " + hook_text +
            "\n\n" + instruction +
            "\n\nApply all copywriting principles. Return as plain text.",
            max_tokens=1500, system=VOICE_GUIDE
        )
        state["ie_generated_content"] = result
        state["ie_pre_edit_content"] = result
        send_plain(chat_id, result)

        keyboard = [
            [{"text": "Quick Edit", "callback_data": "ie_content_edit"},
             {"text": "Regenerate", "callback_data": "ie_content_regen"}],
            [{"text": "Develop in Content Studio", "callback_data": "ie_to_content_studio"}],
            [{"text": "Start over", "callback_data": "open_idea_engine"}]
        ]
        send(chat_id, "Content generated. Edit, develop in Content Studio, or start over.", keyboard)

    except Exception as e:
        send(chat_id, "Error: " + str(e))

def show_ie_manage_sources(chat_id):
    sources = load_idea_sources()
    # Show RSS feed count
    x_status = "✅ Connected (" + str(len(X_ACCOUNTS)) + " accounts)" if X_BEARER_TOKEN else "❌ Not configured — add X_BEARER_TOKEN to Render"
    whisper_status = "✅ Connected" if OPENAI_KEY else "❌ Not configured — add OPENAI_KEY to Render"
    msg = "*News Sources & Integrations*\n\n"
    msg += "🐦 X/Twitter: " + x_status + "\n"
    msg += "🎙️ Voice input: " + whisper_status + "\n\n"
    msg += str(len(RSS_FEEDS)) + " RSS feeds active:\n"
    tier_labels = {1: "Tier 1 (Mainstream)", 2: "Tier 2 (Analysis)", 3: "Tier 3 (Market)", 4: "Tier 4 (Community)"}
    for tier in [1, 2, 3, 4]:
        feeds_in_tier = [f["name"] for f in RSS_FEEDS if f["tier"] == tier]
        if feeds_in_tier:
            msg += "\n" + tier_labels[tier] + ":\n"
            msg += ", ".join(feeds_in_tier) + "\n"
    if sources.get("telegram"):
        msg += "\nTelegram: " + ", ".join("t.me/" + s for s in sources["telegram"]) + "\n"
    if sources.get("reddit"):
        msg += "\nReddit: " + ", ".join("r/" + s for s in sources["reddit"]) + "\n"
    keyboard = [
        [{"text": "Add Telegram channel", "callback_data": "ie_add_telegram"}],
        [{"text": "Add Reddit community", "callback_data": "ie_add_reddit"}],
        [{"text": "Back", "callback_data": "open_idea_engine"}]
    ]
    send(chat_id, msg, keyboard)

def fetch_source_content(sources):
    """Fetch live content — now powered by RSS feeds + Reddit."""
    fetched = []

    # RSS feeds — primary source, verified live
    headlines = fetch_rss_headlines(max_per_feed=2, max_total=20)
    for h in headlines:
        fetched.append({
            "source": h["source"],
            "type": "rss",
            "tier": h["tier"],
            "content": h["title"] + (" — " + h["desc"] if h.get("desc") else "")
        })

    # Reddit removed — returns 403 Blocked when fetched from server IPs
    # Users can still add Reddit as a source but it won't fetch live

    # Telegram public channels
    for channel in sources.get("telegram", [])[:3]:
        try:
            url = "https://t.me/s/" + channel
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=10) as r:
                html = r.read().decode("utf-8", errors="ignore")
                import re as _re
                msgs = _re.findall(r'<div class="tgme_widget_message_text[^"]*"[^>]*>(.*?)</div>', html, _re.DOTALL)
                for msg in msgs[:3]:
                    clean = _re.sub(r'<[^>]+>', '', msg).strip()[:200]
                    if clean:
                        fetched.append({"source": "t.me/" + channel, "type": "telegram", "tier": 3, "content": clean})
        except Exception as e:
            print("Telegram fetch error " + channel + ":", e, flush=True)

    return fetched

def generate_ideas(chat_id, idea_type="both"):
    """Generate ideas from all sources."""
    send(chat_id, "Scanning sources and generating ideas...")
    sources = load_idea_sources()
    ctx = get_content_context(chat_id)

    try:
        # Fetch live content
        fetched = fetch_source_content(sources)

        # Build source summary for Claude
        source_summary = ""
        if fetched:
            source_summary = "LIVE SOURCE DATA:\n"
            # Prioritise news sources and ad inspiration
            news_items = [i for i in fetched if i["type"] in ("twitter", "reddit")]
            ad_items = [i for i in fetched if i["type"] in ("twitter_ads", "facebook_ad", "instagram_marketing")]
            other_items = [i for i in fetched if i["type"] not in ("twitter", "reddit", "twitter_ads", "facebook_ad")]
            for item in (news_items[:6] + ad_items[:3] + other_items[:3]):
                source_summary += "- [" + item["source"] + "] " + item["content"][:150] + "\n"
            if ad_items:
                source_summary += "\nAD INSPIRATION SOURCES (use for ad hook/angle ideas):\n"
                for item in ad_items:
                    source_summary += "- [" + item["source"] + "] " + item["content"][:200] + "\n"
        else:
            source_summary = "Note: Live source fetching returned limited data. Generate ideas based on current crypto market context and trends.\n"

        # Also fetch trending crypto news
        market = fetch_market_data()
        market_line = ""
        if market.get("btc_price"):
            market_line = "Current BTC: ${:,.0f} | Fear & Greed: {} ({})".format(
                market["btc_price"], market.get("fng_value","?"), market.get("fng_label","?"))

        type_instruction = {
            "social": "Generate SOCIAL CONTENT ideas only (Reels, Carousels, Stories).",
            "ads": "Generate AD COPY ideas only (Meta static and video ads).",
            "both": "Generate a mix of SOCIAL CONTENT ideas and AD ideas."
        }.get(idea_type, "Generate a mix of social and ad ideas.")

        prompt = """You are the creative strategist for Cryptonary, a crypto research platform.

MARKET CONTEXT: """ + market_line + """

""" + source_summary + """

""" + ctx + """

""" + type_instruction + """

Generate exactly 8 specific, actionable content ideas. For each idea:

IDEA [N]: [Format] — [Hook/Concept in one punchy line]
AVATAR: [Which of these avatars: Trader / Investor / Passive Income / Portfolio Builder / 100X Chaser / Skeptic / Burned / Student / 9-5 Worker / Boomer / Side Hustle / Beginner / Universal]
OBJECTIVE: [Awareness / Consideration / Conversion] — [what action it drives]
INSPIRED BY: [which source or trend inspired this]
WHY IT WORKS: [one sentence — specific psychological principle or copywriting mechanic]
---

Rules:
- Ideas must be specific to Cryptonary's content — crypto research, market analysis, portfolio strategy, airdrops, passive income
- No generic ideas. Every idea should have a specific angle, hook, or narrative
- Mix formats: include Reels, Carousels, Static ads, and Video ads
- Ideas should reflect current market conditions (Fear & Greed: """ + str(market.get("fng_label","unknown")) + """)
- Draw from what's working in the source data above"""

        result = claude(prompt, max_tokens=2000, system=VOICE_GUIDE)

        # Parse into individual ideas and send with action buttons
        send_plain(chat_id, "*CRYPTONARY IDEA ENGINE*\n\n" + result)

        keyboard = [
            [{"text": "Generate more ideas", "callback_data": "ie_generate_all"}],
            [{"text": "Social ideas only", "callback_data": "ie_generate_social"},
             {"text": "Ad ideas only", "callback_data": "ie_generate_ads"}],
            [{"text": "Develop an idea → Content Studio", "callback_data": "ie_develop"}],
            [{"text": "Develop an idea → Ad Copy", "callback_data": "ie_develop_ad"}],
            [{"text": "Back to Creative Studio", "callback_data": "open_idea_engine"}]
        ]
        send(chat_id, "8 ideas generated. Tap Develop to take any idea straight into the Content Studio.", keyboard)

        # Save ideas to state for development
        user_state[chat_id]["last_ideas"] = result
        user_state[chat_id]["stage"] = "idea_engine_idle"

    except Exception as e:
        print("Idea generation error:", e, flush=True)
        send(chat_id, "Error generating ideas: " + str(e))

# ── IMAGE PROMPT GENERATOR ────────────────────────────────────────

def generate_image_prompts(chat_id, brief):
    """Generate Midjourney, DALL-E, and universal image prompts from a brief."""
    send(chat_id, "Generating image prompts...")
    try:
        prompt = """Generate 3 image prompt variants for this brief: """ + brief + """

You are creating prompts for Cryptonary — a premium crypto research brand. Dark, professional, aspirational aesthetic. The brand uses dark backgrounds, gold/amber accents, clean typography.

OUTPUT FORMAT:

MIDJOURNEY PROMPT:
[Full prompt optimised for Midjourney v6. Include: subject, lighting, mood, style reference, color palette, composition. End with: --ar 4:5 --v 6 --style raw --q 2]

DALL-E PROMPT:
[Plain English scene description optimised for DALL-E 3. Be explicit about style, what to include, what to avoid. No special syntax needed.]

RUNWAY / VIDEO THUMBNAIL:
[Opening frame description for video content. Include: subject position, background, lighting mood, any motion direction, emotional tone.]

DESIGN NOTES:
[2-3 specific notes for the designer: suggested text overlay placement, color hex codes that fit Cryptonary brand, any stock photo search terms as alternatives]

Keep each prompt specific and visual. Avoid abstract descriptions — describe exactly what should be in the frame."""

        result = claude(prompt, max_tokens=800)
        send_plain(chat_id, result)

        keyboard = [
            [{"text": "Generate another variation", "callback_data": "ie_imageprompt_again"}],
            [{"text": "Back to main menu", "callback_data": "start_over"}]
        ]
        user_state[chat_id]["last_image_brief"] = brief
        send(chat_id, "Prompts ready. Copy and paste into Midjourney, DALL-E, or Runway.", keyboard)

    except Exception as e:
        send(chat_id, "Error: " + str(e))

import urllib.parse as _urlparse

# ══════════════════════════════════════════════════════════════════
# CRITIQUE ENGINE
# ══════════════════════════════════════════════════════════════════

CRITIQUE_SYSTEM = """You are a senior copy editor for Cryptonary. Your job is to critique marketing content against Adam's voice rules, copywriting principles, and the full knowledge base.

CRITIQUE PHILOSOPHY:
- Only flag issues that genuinely matter. Do not manufacture feedback.
- Every issue must argue WHY it matters and what the cost of not fixing it is.
- Assign severity honestly based on impact on performance.
- If the content is strong, say so clearly. Don't invent problems.

SEVERITY SCALE:
🔴 RED — Must fix. This is actively hurting the content's performance or breaking a non-negotiable rule.
🟡 YELLOW — Worth fixing. This is a meaningful improvement that would noticeably lift results.
🟢 GREEN — Optional polish. Minor refinement that's nice to have but not essential.

ADAM'S VOICE NON-NEGOTIABLES (violation = 🔴):
- Opens with "Gm [Name]," for emails
- Short punchy sentences. No em dashes. Bullets use •
- No weasel words: very, quite, rather, really — delete them
- Bold key phrases with *asterisks*
- P.S. is mandatory on emails
- Sign off: "Talk soon, / Adam"
- Free emails: tease, curiosity gap, CTA to upgrade
- Pro emails: full analysis, exact levels, complete strategy
- No passive voice. Active, direct, confident.

COPYWRITING STANDARDS:
- Headline/subject must promise a benefit or provoke curiosity
- Every line must earn its place
- Specifics beat generalities — numbers, dates, names, exact levels
- CTA must be clear, single, transformation-led not feature-led
- Fear, curiosity, or desire must be present in the opening
- P.S. should contain the sharpest proof point

FORMAT YOUR CRITIQUE AS:

[SEVERITY EMOJI] [Issue title] — [Specific problem]
WHY IT MATTERS: [Argue the case — what is the cost of this issue on performance?]
FIX: [Exact replacement copy or specific instruction]

---

Maximum 6 issues. Order by severity (🔴 first).
If fewer than 3 genuine issues exist, only list those. Do not pad.
If the content is strong, open with: "STRONG COPY — [one sentence on what's working]" then list any refinements."""

def run_critique(chat_id, content_type):
    """Run a critique on the current piece of content."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    state["critique_content_type"] = content_type

    # Get the right content from state
    if content_type == "email":
        free_email = state.get("current_emails", {}).get("free", "")
        pro_email = state.get("current_emails", {}).get("pro", "")
        # Use 1500 chars per email to ensure P.S. and sign-off are included
        content_to_critique = "FREE EMAIL:\n" + free_email[:1500] + "\n\nPRO EMAIL:\n" + pro_email[:1500]
        label = "emails"
    elif content_type == "ad":
        content_to_critique = state.get("current_ad_output", "")[:1500]
        label = "ad copy"
    elif content_type == "social":
        content_to_critique = state.get("current_social", "")[:1500]
        social_type = state.get("current_social_type", "social content")
        label = social_type
    elif content_type == "lp":
        content_to_critique = state.get("lp_content", "")[:1500]
        label = "landing page"
    else:
        content_to_critique = ""
        label = "content"

    if not content_to_critique or len(content_to_critique.strip()) < 50:
        send(chat_id, "Nothing to critique yet — generate content first.")
        return

    send(chat_id, "Critiquing your " + label + "...")

    try:
        result = claude(
            "Critique this " + label + ". IMPORTANT: Only flag issues that actually exist in the content below. "
            "Read the ENTIRE content carefully before critiquing — do not flag missing P.S., sign-off, or CTA if they are present.\n\n" + content_to_critique,
            max_tokens=1200,
            system=CRITIQUE_SYSTEM
        )

        # Store critique and content for applying fixes
        state["current_critique"] = result
        state["critique_original"] = content_to_critique
        state["stage"] = "awaiting_critique_apply"

        send_plain(chat_id, "CRITIQUE\n\n" + result)

        # Build apply buttons for each numbered fix (up to 6)
        import re as _re
        nums = _re.findall(r'^\d+\.', result, _re.MULTILINE)
        keyboard = []
        row = []
        for i, _ in enumerate(nums[:6]):
            row.append({"text": "Apply " + str(i+1), "callback_data": "apply_critique_" + str(i+1)})
            if len(row) == 3:
                keyboard.append(row)
                row = []
        if row:
            keyboard.append(row)
        keyboard.append([{"text": "Ignore all", "callback_data": "ignore_critique"}])

        send(chat_id, "Type a number to apply a fix, or tap a button:", keyboard)

    except Exception as e:
        send(chat_id, "Critique error: " + str(e))
        print("Critique error:", e, flush=True)

def apply_critique_fixes(chat_id, fix_numbers):
    """Apply one or more critique fixes, show diff, offer revert."""
    user_state.setdefault(chat_id, {"stage": "idle"})
    state = user_state[chat_id]
    critique = state.get("current_critique", "")
    original = state.get("critique_original", "")
    content_type = state.get("critique_content_type", "email")

    if not critique or not original:
        send(chat_id, "No active critique to apply.")
        return

    fix_list = ", ".join(str(n) for n in fix_numbers)
    send(chat_id, "Applying fix" + ("es " if len(fix_numbers) > 1 else " ") + fix_list + "...")

    try:
        fix_instruction = ("fixes " + fix_list) if len(fix_numbers) > 1 else ("fix " + fix_list)
        result = claude(
            "ORIGINAL CONTENT:\n" + original[:1500] +
            "\n\nCRITIQUE:\n" + critique[:1000] +
            "\n\nApply ONLY " + fix_instruction + " from the critique above. " +
            "Return the full revised content with those specific changes applied. " +
            "Do not apply other changes. Do not explain what you changed.",
            max_tokens=1500
        )

        # Store previous version for revert
        state["critique_pre_fix"] = original

        # Update state
        if content_type == "email":
            improved = parse_delimited_emails(result)
            if improved:
                state["current_emails"] = improved
            send_plain(chat_id, str(len(fix_numbers)) + " fix(es) applied:\n\n" + result[:2000])
        elif content_type == "ad":
            state["current_ad_output"] = result
            send_plain(chat_id, str(len(fix_numbers)) + " fix(es) applied:\n\n" + result[:2000])
        elif content_type == "social":
            state["current_social"] = result
            send_plain(chat_id, str(len(fix_numbers)) + " fix(es) applied:\n\n" + result[:2000])

        # Save as voice example
        save_voice_example(chat_id, result[:600], "critique_fix_" + content_type)

        kb_map = {"email": email_action_keyboard, "ad": ad_action_keyboard, "social": social_action_keyboard}
        kb_fn = kb_map.get(content_type, email_action_keyboard)
        keyboard = kb_fn()
        keyboard.insert(0, [{"text": "Revert changes", "callback_data": "revert_critique"},
                             {"text": "Critique again", "callback_data": "critique_" + content_type}])
        state["stage"] = "emails_ready" if content_type == "email" else "social_ready" if content_type == "social" else "ads_ready"
        send(chat_id, "Fix applied. Revert if needed or proceed:", keyboard)

    except Exception as e:
        send(chat_id, "Error applying fix: " + str(e))
        print("Apply fix error:", e, flush=True)

# Keep old single-fix function as alias for backward compat
def apply_critique_fix(chat_id, fix_number):
    apply_critique_fixes(chat_id, [int(fix_number)])

# ══════════════════════════════════════════════════════════════════
# CONTENT STUDIO FILE HANDLER
# ══════════════════════════════════════════════════════════════════

def handle_content_file(chat_id, file_info, file_type="image"):
    """Handle file uploads in Content Studio flows (PDF, image, doc)."""
    state = user_state.get(chat_id, {})
    stage = state.get("stage", "idle")

    try:
        # Download the file
        file_id = file_info.get("file_id")
        path_data = tg("getFile", {"file_id": file_id})
        file_path = path_data.get("result", {}).get("file_path", "")
        if not file_path:
            send(chat_id, "Could not download file.")
            return

        url = "https://api.telegram.org/file/bot" + TELEGRAM_TOKEN + "/" + file_path
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=30) as r:
            file_bytes = r.read()

        import base64
        send(chat_id, "Reading your file...")

        if file_type == "image":
            # Use Claude Vision to extract content
            mime = "image/jpeg"
            if file_path.endswith(".png"): mime = "image/png"
            elif file_path.endswith(".webp"): mime = "image/webp"
            encoded = base64.b64encode(file_bytes).decode()

            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 2000,
                "messages": [{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": mime, "data": encoded}},
                    {"type": "text", "text": "Extract all the text and data from this image. This is likely a research report, market analysis, or data table. Return everything you can read as plain text, preserving structure. Include all numbers, dates, percentages, and key points."}
                ]}]
            }).encode()

            req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
                headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
            with urllib.request.urlopen(req, timeout=60) as r:
                extracted = json.loads(r.read())["content"][0]["text"]

        elif file_type == "pdf":
            # Send PDF directly to Claude API as document
            encoded = base64.b64encode(file_bytes).decode()
            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 3000,
                "messages": [{"role": "user", "content": [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": encoded}},
                    {"type": "text", "text": "Extract all the content from this document. This is likely a research report, market analysis, newsletter, or data report. Return the full content as plain text preserving all key information, numbers, analysis, and structure."}
                ]}]
            }).encode()
            req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
                headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
            with urllib.request.urlopen(req, timeout=60) as r:
                extracted = json.loads(r.read())["content"][0]["text"]

        elif file_type == "doc":
            # Try as docx first, fallback to plain text
            try:
                import io
                from docx import Document as _Document
                doc = _Document(io.BytesIO(file_bytes))
                extracted = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])[:8000]
                if not extracted.strip():
                    raise ValueError("Empty docx")
            except Exception:
                try:
                    extracted = file_bytes.decode("utf-8")[:8000]
                except:
                    extracted = file_bytes.decode("latin-1")[:8000]
        else:
            # Plain text / CSV
            try:
                extracted = file_bytes.decode("utf-8")[:8000]
            except:
                extracted = file_bytes.decode("latin-1")[:8000]

        if not extracted or len(extracted.strip()) < 20:
            send(chat_id, "Could not extract readable content from this file.\n\nSupported: PDF, image/screenshot, CSV, .docx, plain text.\n\nTry pasting the text directly instead.")
            return

        # Show preview and proceed
        preview = extracted[:300] + ("..." if len(extracted) > 300 else "")
        send(chat_id, "Got it. Extracted " + str(len(extracted)) + " characters:\n\n_" + preview + "_")

        # Route based on current stage — same as if text was pasted
        if stage in ("awaiting_report", "buffering_report"):
            user_state[chat_id]["report"] = sanitise(extracted)
            user_state[chat_id]["stage"] = "awaiting_context_choice"
            keyboard = [
                [{"text": "Yes — I have extra context", "callback_data": "context_yes"}],
                [{"text": "No — just the report", "callback_data": "context_no"}]
            ]
            send(chat_id, "Any extra context to factor in?\n\n_Promos, discounts, Inner Circle open, upcoming events, factoids, PSAs..._", keyboard)

        elif stage == "awaiting_social_report":
            user_state[chat_id]["report"] = sanitise(extracted)
            user_state[chat_id]["context"] = ""
            gen_social_angles(chat_id)

        elif stage == "awaiting_ad_existing_upload":
            user_state[chat_id]["existing_ad_content"] = sanitise(extracted[:2000])
            user_state[chat_id]["stage"] = "pick_existing_ad_action"
            show_existing_ad_action_menu(chat_id)

        elif stage == "awaiting_ad_theme":
            product_context = user_state[chat_id].get("ad_product_context", "")
            theme = sanitise(extracted[:500])
            if product_context:
                theme = "PRODUCT: " + product_context + "\n\nTHEME/CONTEXT: " + theme
            user_state[chat_id]["ad_theme"] = theme
            user_state[chat_id]["stage"] = "pick_ad_avatars"
            show_avatar_menu(chat_id)

        elif stage == "awaiting_lp_context_text":
            user_state[chat_id]["lp_context"] = sanitise(extracted[:500])
            generate_lp_outline(chat_id)

        else:
            send(chat_id, "File received but not sure where to use it. Try again from the right step.")

    except Exception as e:
        print("Content file handler error:", e, flush=True)
        send(chat_id, "Error reading file: " + str(e))

# ══════════════════════════════════════════════════════════════════
# IDEA ENGINE — SCREENSHOT ANALYSIS
# ══════════════════════════════════════════════════════════════════

def handle_ie_screenshot(chat_id, file_info, stage):
    """Handle file uploads for Idea Engine — images, PDFs, CSVs."""
    try:
        import base64
        file_id = file_info.get("file_id")
        path_data = tg("getFile", {"file_id": file_id})
        file_path = path_data.get("result", {}).get("file_path", "")
        if not file_path:
            send(chat_id, "Could not download the file.")
            return

        url = "https://api.telegram.org/file/bot" + TELEGRAM_TOKEN + "/" + file_path
        req = urllib.request.Request(url)
        with urllib.request.urlopen(req, timeout=30) as r:
            file_bytes = r.read()

        user_state[chat_id]["stage"] = "idea_engine_idle"

        # Detect file type and extract content
        mime_type = file_info.get("mime_type", "")
        is_pdf = "pdf" in mime_type or file_path.endswith(".pdf")
        is_csv = "csv" in mime_type or file_path.endswith(".csv") or file_path.endswith(".txt")
        is_doc = "word" in mime_type or "docx" in mime_type or "officedocument" in mime_type

        if is_pdf:
            # Use Claude API document extraction
            encoded = base64.b64encode(file_bytes).decode()
            send(chat_id, "Reading PDF...")
            payload = json.dumps({
                "model": "claude-sonnet-4-20250514", "max_tokens": 2000,
                "messages": [{"role": "user", "content": [
                    {"type": "document", "source": {"type": "base64", "media_type": "application/pdf", "data": encoded}},
                    {"type": "text", "text": "Extract the key content from this document — headlines, copy, data, analysis. Return as plain text."}
                ]}]
            }).encode()
            req2 = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
                headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
            with urllib.request.urlopen(req2, timeout=60) as r:
                text_content = json.loads(r.read())["content"][0]["text"]
            _process_ie_text_content(chat_id, stage, text_content, "PDF")
            return

        elif is_csv or is_doc:
            if is_doc:
                try:
                    import io
                    from docx import Document as _Doc
                    doc = _Doc(io.BytesIO(file_bytes))
                    text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])[:6000]
                except:
                    text_content = file_bytes.decode("utf-8", errors="ignore")[:6000]
            else:
                text_content = file_bytes.decode("utf-8", errors="ignore")[:6000]
            _process_ie_text_content(chat_id, stage, text_content, "document")
            return

        # Default: treat as image
        mime = "image/jpeg"
        if file_path.endswith(".png"): mime = "image/png"
        elif file_path.endswith(".webp"): mime = "image/webp"
        encoded = base64.b64encode(file_bytes).decode()

        if stage == "ie_awaiting_screenshot_ideas":
            send(chat_id, "Analysing and generating ideas...")
            prompt_text = """Look at this image carefully. It could be an Instagram post, a Facebook ad, a carousel, a reel thumbnail, or any piece of marketing content.

STEP 1 — READ THE CONTENT:
What format is it? What is the hook or opening line? What topic? What's the visual approach? Who is the target audience?

STEP 2 — IDENTIFY WHAT'S WORKING:
What copywriting or creative principles are at play? Why would this stop a scroll or drive engagement?

STEP 3 — GENERATE 6 CONTENT IDEAS INSPIRED BY THIS:
Use this as creative inspiration for Cryptonary content. Each idea should be in a different format or take a different angle — don't just replicate it.

For each idea:
IDEA [N]: [Format] — [Hook/Concept]
AVATAR: [which of Cryptonary's 13 avatars: Trader/Investor/Passive Income/Portfolio Builder/100X Chaser/Skeptic/Burned/Student/9-5 Worker/Boomer/Side Hustle/Beginner/Universal]
OBJECTIVE: [Awareness/Consideration/Conversion]
INSPIRED BY: [what specifically in the screenshot sparked this]
WHY IT WORKS: [one sentence — specific principle or mechanic]
---"""

            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1500,
                "system": VOICE_GUIDE,
                "messages": [{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": mime, "data": encoded}},
                    {"type": "text", "text": prompt_text}
                ]}]
            }).encode()

            req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
                headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
            with urllib.request.urlopen(req, timeout=60) as r:
                result = json.loads(r.read())["content"][0]["text"]

            user_state[chat_id]["last_ideas"] = result
            send_plain(chat_id, "*IDEAS FROM SCREENSHOT*\n\n" + result)
            keyboard = [
                [{"text": "Develop into social content", "callback_data": "ie_develop"}],
                [{"text": "Develop into ad", "callback_data": "ie_develop_ad"}],
                [{"text": "Generate more ideas", "callback_data": "ie_generate_all"}],
                [{"text": "Back to Creative Studio", "callback_data": "open_idea_engine"}]
            ]
            send(chat_id, "6 ideas generated from your screenshot.", keyboard)

        elif stage == "ie_awaiting_screenshot_critique":
            send(chat_id, "Critiquing...")
            prompt_text = """Critique this piece of content against professional copywriting and marketing standards.

This could be a Cryptonary post, a competitor's content, or any marketing material.

Apply these standards:
- Hook strength (does it stop the scroll in 3 seconds?)
- Clarity (is the message instantly understood?)
- Copy quality (voice, specificity, no weasel words, active not passive)
- Avatar targeting (is it speaking to one specific person or everyone and no one?)
- CTA effectiveness (is there a clear next action?)
- Visual/text balance (if visible)
- Copywriting principles from Ogilvy, Cashvertising, Cialdini, Hormozi, Made to Stick

FORMAT:
WHAT'S WORKING:
[1-2 things that are genuinely strong]

ISSUES:
1. [Issue title] — [Specific problem]
   FIX: [Exact suggested improvement]

2. [Issue title] — [Specific problem]
   FIX: [Exact suggested improvement]

(up to 6 issues, ordered by severity)

OVERALL RATING: [A/B/C/D] — [one sentence verdict]"""

            payload = json.dumps({
                "model": "claude-sonnet-4-20250514",
                "max_tokens": 1200,
                "system": CRITIQUE_SYSTEM,
                "messages": [{"role": "user", "content": [
                    {"type": "image", "source": {"type": "base64", "media_type": mime, "data": encoded}},
                    {"type": "text", "text": prompt_text}
                ]}]
            }).encode()

            req = urllib.request.Request("https://api.anthropic.com/v1/messages", data=payload,
                headers={"Content-Type": "application/json", "x-api-key": ANTHROPIC_KEY, "anthropic-version": "2023-06-01"})
            with urllib.request.urlopen(req, timeout=60) as r:
                result = json.loads(r.read())["content"][0]["text"]

            send_plain(chat_id, "SCREENSHOT CRITIQUE\n\n" + result)
            keyboard = [
                [{"text": "Ideas from this screenshot", "callback_data": "ie_screenshot_ideas"}],
                [{"text": "Back to Creative Studio", "callback_data": "open_idea_engine"}]
            ]
            send(chat_id, "Critique complete.", keyboard)

    except Exception as e:
        print("IE screenshot error:", e, flush=True)
        send(chat_id, "Error analysing screenshot: " + str(e))
        user_state[chat_id]["stage"] = "idea_engine_idle"

# ══════════════════════════════════════════════════════════════════
# LINK / URL ANALYSIS ENGINE
# ══════════════════════════════════════════════════════════════════

import re as _url_re

URL_PATTERN = _url_re.compile(
    r'https?://[^\s<>"{}|\\^`\[\]]+'
)

def detect_url_type(url):
    """Classify a URL so we know how to handle it."""
    url_lower = url.lower()
    if "twitter.com" in url_lower or "x.com" in url_lower:
        return "tweet"
    elif "instagram.com" in url_lower:
        return "instagram"
    elif "facebook.com/ads/library" in url_lower:
        return "fb_ads"
    elif "facebook.com" in url_lower:
        return "facebook"
    elif "tiktok.com" in url_lower:
        return "tiktok"
    elif "youtube.com" in url_lower or "youtu.be" in url_lower:
        return "youtube"
    elif "reddit.com" in url_lower:
        return "reddit"
    elif "t.me" in url_lower:
        return "telegram"
    else:
        return "webpage"

def fetch_url_content(url, url_type):
    """Fetch content from a URL and return text + any screenshot data."""
    headers = {"User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36"}
    try:
        req = urllib.request.Request(url, headers=headers)
        with urllib.request.urlopen(req, timeout=15) as r:
            raw = r.read()
            encoding = r.headers.get_content_charset() or "utf-8"
            html = raw.decode(encoding, errors="ignore")

        import re as _re
        # Strip HTML tags for readable text
        text = _re.sub(r'<script[^>]*>.*?</script>', '', html, flags=_re.DOTALL)
        text = _re.sub(r'<style[^>]*>.*?</style>', '', text, flags=_re.DOTALL)
        text = _re.sub(r'<[^>]+>', ' ', text)
        text = _re.sub(r'\s+', ' ', text).strip()
        return text[:6000]
    except Exception as e:
        return ""

def analyse_url(chat_id, url, mode="ideas"):
    """Fetch a URL and analyse it for ideas or critique."""
    state = user_state.get(chat_id, {})
    url_type = detect_url_type(url)

    # Platform-specific fallbacks — these all block scrapers reliably
    PASTE_TEXT_PLATFORMS = {
        "tweet":     ("Twitter/X", "Twitter/X blocks automated access.\n\nPaste the tweet text below and I'll analyse it 👇"),
        "instagram": ("Instagram", "Instagram blocks automated access.\n\nScreenshot the post and upload the image instead, or paste the caption text below 👇"),
        "tiktok":    ("TikTok",    "TikTok blocks automated access.\n\nPaste the caption or script text below and I'll analyse it 👇"),
        "facebook":  ("Facebook",  "Facebook blocks automated access.\n\nScreenshot the post and upload the image, or paste the post text below 👇"),
        "fb_ads":    ("Facebook Ads Library", "The Ads Library is JS-rendered and can't be scraped.\n\nScreenshot the ad and upload the image instead 👇"),
        "reddit":    ("Reddit",    "Reddit's API changes mean most posts can't be scraped.\n\nPaste the post text below and I'll analyse it 👇"),
        "youtube":   ("YouTube",   "YouTube content is JS-rendered and can't be read directly.\n\nPaste the video transcript or description below and I'll analyse it 👇"),
    }

    if url_type in PASTE_TEXT_PLATFORMS:
        platform_name, fallback_msg = PASTE_TEXT_PLATFORMS[url_type]
        user_state[chat_id]["pending_url_mode"] = mode
        user_state[chat_id]["pending_url_platform"] = platform_name
        # Screenshot-only platforms go back to idle; paste platforms await text
        if url_type in ("fb_ads",):
            user_state[chat_id]["stage"] = "idea_engine_idle"
        else:
            user_state[chat_id]["stage"] = "ie_awaiting_pasted_text"
        send(chat_id, fallback_msg)
        return

    type_labels = {
        "telegram": "Telegram post",
        "webpage":  "webpage"
    }
    label = type_labels.get(url_type, "page")
    send(chat_id, "Fetching " + label + "...")

    text_content = fetch_url_content(url, url_type)

    if not text_content or len(text_content.strip()) < 50:
        send(chat_id, "Could not read content from that link. The page may require login or block bots.\n\nTry screenshotting it and uploading the image instead.")
        return

    _process_ie_text_content(chat_id, state.get("stage", "idea_engine_idle"), text_content,
                              label + " (" + url + ")", mode=mode)

def _process_ie_text_content(chat_id, stage, text_content, source_label, mode=None):
    """Process extracted text content for ideas or critique."""
    # Determine mode from stage if not specified
    if mode is None:
        mode = "critique" if "critique" in stage else "ideas"

    if mode == "ideas" or stage == "ie_awaiting_screenshot_ideas":
        send(chat_id, "Generating ideas from " + source_label + "...")
        ctx = get_content_context(chat_id)
        try:
            result = claude(
                "CONTENT FROM " + source_label.upper() + ":\\n" + text_content[:3000] +
                "\\n\\nGenerate 4 Instagram content concepts for Cryptonary inspired by this material.\\n\\n" +
                "Format EXACTLY:\\n" +
                "1. CONCEPT: [One punchy line — the big idea]\\n" +
                "   ANGLE: [The emotional or argumentative approach]\\n" +
                "   SOURCE: [What specifically in this material inspired it]\\n\\n" +
                "2. CONCEPT: ...\\n   ANGLE: ...\\n   SOURCE: ...\\n\\n" +
                "3. CONCEPT: ...\\n   ANGLE: ...\\n   SOURCE: ...\\n\\n" +
                "4. CONCEPT: ...\\n   ANGLE: ...\\n   SOURCE: ...\\n\\n" +
                "Nothing else." +
                ctx,
                max_tokens=800,
                system=VOICE_GUIDE
            )
            user_state[chat_id]["ie_concepts"] = result
            user_state[chat_id]["ie_source_label"] = source_label
            user_state[chat_id]["stage"] = "ie_concept_review"
            send_plain(chat_id, "*IDEAS FROM " + source_label.upper() + "*\\n\\n" + result)
            keyboard = [
                [{"text": "1 — Expand this", "callback_data": "ie_develop_concept_1"}],
                [{"text": "2 — Expand this", "callback_data": "ie_develop_concept_2"}],
                [{"text": "3 — Expand this", "callback_data": "ie_develop_concept_3"}],
                [{"text": "4 — Expand this", "callback_data": "ie_develop_concept_4"}],
                [{"text": "Regenerate", "callback_data": "ie_regen_concepts"}],
                [{"text": "Back to Creative Studio", "callback_data": "open_idea_engine"}]
            ]
            send(chat_id, "Tap a number to expand into full content:", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))

    else:  # critique mode
        send(chat_id, "Critiquing " + source_label + "...")
        try:
            result = claude(
                "Critique this content from " + source_label + ":\\n\\n" + text_content[:3000],
                max_tokens=1200,
                system=CRITIQUE_SYSTEM
            )
            user_state[chat_id]["stage"] = "idea_engine_idle"
            send_plain(chat_id, "CRITIQUE OF " + source_label.upper() + "\\n\\n" + result)
            keyboard = [
                [{"text": "Get ideas from this", "callback_data": "ie_screenshot_ideas"}],
                [{"text": "Back to Creative Studio", "callback_data": "open_idea_engine"}]
            ]
            send(chat_id, "Critique complete.", keyboard)
        except Exception as e:
            send(chat_id, "Error: " + str(e))

if __name__ == "__main__":
    if ANTHROPIC_KEY == "YOUR_ANTHROPIC_KEY_HERE":
        print("ERROR: Set ANTHROPIC_KEY environment variable.", flush=True)
    else:
        poll()
